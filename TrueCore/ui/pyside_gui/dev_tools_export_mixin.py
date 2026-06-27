from __future__ import annotations

import io
import os
import tempfile

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.shared import Inches, Pt
from PyPDF2 import PdfReader, PdfWriter
from PySide6.QtCore import QMarginsF, QRectF, QSizeF, Qt
from PySide6.QtGui import QColor, QFont, QGuiApplication, QPainter, QPageLayout, QPageSize, QPdfWriter, QPixmap, QTextDocument
from PySide6.QtPrintSupport import QPrinter

from TrueCore.ui.pyside_gui.dev_tools_export_utils import convert_docx_to_pdf_via_word, find_word_executable
from TrueCore.ui.pyside_gui.dev_tools_preview_markup import (
    build_packet_builder_document_markup,
    build_packet_builder_sections,
    render_packet_builder_document_export,
)
from TrueCore.ui.pyside_gui.dev_tools_profiles import (
    is_clinical_documentation_profile,
    is_consult_request_profile,
    is_lomn_profile,
    is_referral_request_profile,
    is_seoc_request_profile,
    is_submission_cover_profile,
    is_va_10172_profile,
    is_virtual_consent_profile,
    sanitize_builder_filename,
)
from TrueCore.ui.pyside_gui.dev_tools_signature_logic import signature_image_bytes


def export_checkbox_marker(enabled):
    return "[X]" if enabled else "[ ]"


def _signature_image_bytes(payload, field_name):
    return signature_image_bytes(payload, field_name)


class PacketBuilderExportMixin:
    def _merge_pdf_files(self, source_paths, target_path, add_continuous_page_numbers=False):
        merged = PdfWriter()
        for source_path in source_paths or []:
            reader = PdfReader(source_path)
            for page in reader.pages:
                merged.add_page(page)
        if add_continuous_page_numbers and merged.pages:
            self._overlay_compiled_pdf_page_numbers(merged)
        with open(target_path, "wb") as handle:
            merged.write(handle)

    def _overlay_compiled_pdf_page_numbers(self, writer):
        total_pages = len(writer.pages or [])
        if total_pages <= 0:
            return

        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_handle:
            overlay_path = temp_handle.name

        pdf_writer = QPdfWriter(overlay_path)
        pdf_writer.setResolution(72)
        pdf_writer.setPageMargins(QMarginsF(0, 0, 0, 0))
        pdf_writer.setPageSize(QPageSize(QSizeF(612, 792), QPageSize.Point))

        painter = QPainter(pdf_writer)
        for index in range(total_pages):
            if index > 0:
                pdf_writer.newPage()
            painter.fillRect(QRectF(196, 756, 220, 22), QColor("#FFFFFF"))
            painter.setPen(QColor("#374151"))
            font = QFont("Arial", 9)
            painter.setFont(font)
            painter.drawText(
                QRectF(170, 755, 272, 18),
                Qt.AlignCenter,
                f"Page {index + 1} of {total_pages}",
            )
        painter.end()

        overlay_reader = PdfReader(overlay_path)
        for index, page in enumerate(writer.pages):
            if index < len(overlay_reader.pages):
                page.merge_page(overlay_reader.pages[index])

        try:
            os.remove(overlay_path)
        except OSError:
            pass

    def _load_pdf_widget_maps(self, template_path):
        reader = PdfReader(template_path)
        page_maps = []
        for page in reader.pages:
            page_map = {"fields": {}, "groups": {}}
            annots = page.get("/Annots")
            if hasattr(annots, "get_object"):
                annots = annots.get_object()
            for annot_ref in annots or []:
                annot = annot_ref.get_object()
                if annot.get("/Subtype") != "/Widget":
                    continue
                rect = annot.get("/Rect")
                rect = [float(value) for value in rect] if rect else None
                field_name = annot.get("/T")
                field_type = annot.get("/FT")
                parent = annot.get("/Parent")
                parent_name = None
                if parent:
                    parent_name = str(parent.get_object().get("/T") or "")
                if field_name and field_type in {"/Tx", "/Sig"} and rect:
                    page_map["fields"][str(field_name)] = {"rect": rect, "type": str(field_type)}
                    continue
                if parent_name and rect:
                    appearance = annot.get("/AP", {}).get("/N")
                    option_names = []
                    if appearance:
                        appearance_obj = appearance.get_object() if hasattr(appearance, "get_object") else appearance
                        if hasattr(appearance_obj, "keys"):
                            option_names = [str(key) for key in appearance_obj.keys()]
                    page_map["groups"].setdefault(parent_name, []).append(
                        {"rect": rect, "option": option_names[0] if option_names else ""}
                    )
            for group_name, widgets in page_map["groups"].items():
                page_map["groups"][group_name] = sorted(
                    widgets,
                    key=lambda item: (-item["rect"][1], item["rect"][0]),
                )
            page_maps.append(page_map)
        return page_maps

    def _pdf_rect_to_qt_rect(self, rect, page_height=792.0):
        x0, y0, x1, y1 = rect
        return QRectF(float(x0), float(page_height) - float(y1), float(x1 - x0), float(y1 - y0))

    def _draw_text_in_pdf_rect(self, painter, rect, text, font_size=9, bold=False, multiline=False):
        if not rect or not str(text or "").strip():
            return
        painter.save()
        font = QFont("Arial", font_size)
        font.setBold(bool(bold))
        painter.setFont(font)
        flags = Qt.AlignLeft | Qt.AlignTop
        if multiline:
            flags |= Qt.TextWordWrap
        else:
            flags = Qt.AlignLeft | Qt.AlignVCenter
        painter.drawText(self._pdf_rect_to_qt_rect(rect), flags, str(text).strip())
        painter.restore()

    def _draw_signature_image_in_pdf_rect(self, painter, rect, payload, field_name):
        image_bytes = _signature_image_bytes(payload, field_name)
        if not rect or not image_bytes:
            return False
        pixmap = QPixmap()
        if not pixmap.loadFromData(image_bytes, "PNG"):
            return False
        target_rect = self._pdf_rect_to_qt_rect(rect)
        scaled = pixmap.scaled(
            int(target_rect.width()),
            int(target_rect.height()),
            Qt.KeepAspectRatio,
            Qt.SmoothTransformation,
        )
        painter.save()
        x = target_rect.x()
        y = target_rect.y() + max(0.0, (target_rect.height() - scaled.height()) / 2.0)
        painter.drawPixmap(int(x), int(y), scaled)
        painter.restore()
        return True

    def _draw_checkbox_mark(self, painter, rect):
        if not rect:
            return
        painter.save()
        font = QFont("Arial", 12)
        font.setBold(True)
        painter.setFont(font)
        painter.drawText(self._pdf_rect_to_qt_rect(rect), Qt.AlignCenter, "X")
        painter.restore()

    def _draw_group_choice(self, painter, page_map, group_name, choice_label, option_map):
        widgets = page_map["groups"].get(group_name) or []
        option_name = option_map.get(str(choice_label or "").strip(), "")
        if not option_name:
            return
        for widget in widgets:
            if widget.get("option") == option_name:
                self._draw_checkbox_mark(painter, widget.get("rect"))
                return

    def _draw_va10172_page_overlay(self, painter, page_map, payload, page_index):
        patient_name = payload.get("patient_name") or ""
        dob = payload.get("date_of_birth") or ""
        auth_number = payload.get("authorization_number") or ""
        va_facility_address = payload.get("va10172_va_facility_address") or ""
        provider_office_address = payload.get("va10172_ordering_provider_office_address") or ""
        provider_phone = payload.get("va10172_ordering_provider_phone") or ""
        provider_fax = payload.get("va10172_ordering_provider_fax") or ""
        provider_email = payload.get("va10172_ordering_provider_secure_email") or ""
        provider_name_printed = payload.get("va10172_ordering_provider_name_printed") or payload.get("provider") or payload.get("ordering_doctor") or ""
        provider_npi = payload.get("va10172_ordering_provider_npi") or payload.get("provider_npi") or ""
        today_date = payload.get("va10172_today_date") or ""
        signature_text = payload.get("va10172_signature_text") or ""

        if page_index == 0:
            field_values = {
                "VETERANSNAME[0]": patient_name,
                "DOB[0]": dob,
                "VAFACILITYADDRESS[0]": va_facility_address,
                "VAAUTHORIZATIONNUMBER[0]": auth_number,
                "ORDERINGPROVIDEROFFICENAMEADDRESS[0]": provider_office_address,
                "ORDERINGPROVIDERPHONENUMBER[0]": provider_phone,
                "ORDERINGPROVIDERFAXNUMBER[0]": provider_fax,
                "ORDERINGPROVIDERSECUREEMAILADDRESS[0]": provider_email,
                "SPECIALTY[0]": payload.get("va10172_referral_specialty_text") or "",
                "DIAGNOSISCODES[0]": payload.get("icd_codes") or "",
                "DIAGNOSISDESCRIPTION[0]": payload.get("va10172_diagnosis_description") or payload.get("diagnosis") or "",
                "REQUESTEDCPTHCPCSCODE[0]": payload.get("va10172_requested_cpt_hcpcs_code") or "",
                "DESCRIPTIONCPTHCPCSCODE[0]": payload.get("va10172_description_cpt_hcpcs_code") or "",
                "TextField1[0]": payload.get("va10172_reason_for_request") or "",
                "ORDERINGPROVIDERSNAMEPRINTED[0]": provider_name_printed,
                "ORDERINGPROVIDERSNPI[0]": provider_npi,
                "Date[0]": today_date,
                "SignatureField11[0]": signature_text,
            }
            multiline_fields = {
                "VAFACILITYADDRESS[0]",
                "ORDERINGPROVIDEROFFICENAMEADDRESS[0]",
                "DIAGNOSISDESCRIPTION[0]",
                "DESCRIPTIONCPTHCPCSCODE[0]",
                "TextField1[0]",
            }
            small_font_fields = {
                "VAFACILITYADDRESS[0]": 8,
                "ORDERINGPROVIDEROFFICENAMEADDRESS[0]": 8,
                "ORDERINGPROVIDERSECUREEMAILADDRESS[0]": 8,
                "SPECIALTY[0]": 8,
                "DIAGNOSISCODES[0]": 8,
                "DIAGNOSISDESCRIPTION[0]": 7,
                "REQUESTEDCPTHCPCSCODE[0]": 8,
                "DESCRIPTIONCPTHCPCSCODE[0]": 7,
                "TextField1[0]": 8,
            }
            for field_name, value in field_values.items():
                field_info = page_map["fields"].get(field_name) or {}
                if field_name == "SignatureField11[0]" and self._draw_signature_image_in_pdf_rect(
                    painter, field_info.get("rect"), payload, "va10172_signature_text"
                ):
                    continue
                self._draw_text_in_pdf_rect(
                    painter,
                    field_info.get("rect"),
                    value,
                    font_size=small_font_fields.get(field_name, 9),
                    bold=False,
                    multiline=field_name in multiline_fields,
                )

            yes_no_map = {"No": "/0", "Yes": "/1"}
            self._draw_group_choice(painter, page_map, "HISTHP[0]", payload.get("va10172_is_ihs_provider"), yes_no_map)
            self._draw_group_choice(painter, page_map, "RadioButtonList[0]", payload.get("va10172_care_needed_within_48_hours"), yes_no_map)
            self._draw_group_choice(painter, page_map, "RadioButtonList[1]", payload.get("va10172_is_continuation_of_care"), yes_no_map)
            self._draw_group_choice(painter, page_map, "RadioButtonList[2]", payload.get("va10172_referral_to_specialty"), yes_no_map)

            geriatric_map = {
                "Community Nursing Home": "/0",
                "Home Infusion": "/1",
                "Hospice/Palliative Care": "/2",
                "Skilled Home Health Care": "/3",
                "Community Adult Day Health Care": "/4",
                "Home Homemaker/Home Health Aide": "/5",
                "Respite": "/6",
            }
            self._draw_group_choice(
                painter,
                page_map,
                "RadioButtonList[3]",
                payload.get("va10172_geriatric_care_option"),
                geriatric_map,
            )
            return

        if page_index == 1:
            field_values = {
                "VETERANSNAME[1]": patient_name,
                "DOB[1]": dob,
                "VAFACILITYADDRESS[1]": va_facility_address,
                "VAAUTHORIZATIONNUMBER[1]": auth_number,
                "ORDERINGPROVIDEROFFICENAMEADDRESS[1]": provider_office_address,
                "ORDERINGPROVIDERPHONENUMBER[1]": provider_phone,
                "ORDERINGPROVIDERFAXNUMBER[1]": provider_fax,
                "ORDERINGPROVIDERSECUREEMAILADDRESS[1]": provider_email,
                "ORDERINGPROVIDERSNAMEPRINTED2[0]": provider_name_printed,
                "ORDERINGPROVIDERSNPI2[0]": provider_npi,
                "Date2[0]": today_date,
                "SignatureField2[0]": signature_text,
            }
            for field_name, value in field_values.items():
                field_info = page_map["fields"].get(field_name) or {}
                if field_name == "SignatureField2[0]" and self._draw_signature_image_in_pdf_rect(
                    painter, field_info.get("rect"), payload, "va10172_signature_text"
                ):
                    continue
                self._draw_text_in_pdf_rect(
                    painter,
                    field_info.get("rect"),
                    value,
                    font_size=8 if field_name in {"VAFACILITYADDRESS[1]", "ORDERINGPROVIDEROFFICENAMEADDRESS[1]"} else 9,
                    bold=False,
                    multiline=field_name in {"VAFACILITYADDRESS[1]", "ORDERINGPROVIDEROFFICENAMEADDRESS[1]"},
                )

    def _write_va_10172_pdf(self, path, payload):
        template_path = self._resolve_va10172_template_path()
        if not template_path or not os.path.exists(template_path):
            raise FileNotFoundError("VA Form 10-10172 template PDF not found.")

        page_maps = self._load_pdf_widget_maps(template_path)
        app = QGuiApplication.instance() or QGuiApplication([])
        del app
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_handle:
            overlay_path = temp_handle.name

        pdf_writer = QPdfWriter(overlay_path)
        pdf_writer.setResolution(72)
        pdf_writer.setPageMargins(QMarginsF(0, 0, 0, 0))
        pdf_writer.setPageSize(QPageSize(QSizeF(612, 792), QPageSize.Point))

        painter = QPainter(pdf_writer)
        for index, page_map in enumerate(page_maps):
            if index > 0:
                pdf_writer.newPage()
            self._draw_va10172_page_overlay(painter, page_map, payload, index)
        painter.end()

        base_reader = PdfReader(template_path)
        overlay_reader = PdfReader(overlay_path)
        merged = PdfWriter()
        for index, base_page in enumerate(base_reader.pages):
            if index < len(overlay_reader.pages):
                base_page.merge_page(overlay_reader.pages[index])
            merged.add_page(base_page)
        with open(path, "wb") as handle:
            merged.write(handle)

        try:
            os.remove(overlay_path)
        except OSError:
            pass

    def _write_word_doc(self, path):
        payload = self.collect_payload()
        self._write_word_doc_for_payload(path, payload)

    def _write_pdf_doc(self, path):
        payload = self.collect_payload()
        self._write_pdf_doc_for_payload(path, payload)

    def _write_word_doc_for_payload(self, path, payload):
        if is_referral_request_profile(payload.get("packet_profile")):
            self._write_referral_request_doc(path, payload)
            return
        if is_virtual_consent_profile(payload.get("packet_profile")):
            self._write_virtual_consent_doc_v2(path, payload)
            return
        if is_submission_cover_profile(payload.get("packet_profile")):
            self._write_submission_cover_doc_v2(path, payload)
            return
        if is_seoc_request_profile(payload.get("packet_profile")):
            self._write_seoc_request_doc_v2(path, payload)
            return
        if is_lomn_profile(payload.get("packet_profile")):
            self._write_lomn_doc_v2(path, payload)
            return
        if is_consult_request_profile(payload.get("packet_profile")):
            self._write_consult_request_doc_v2(path, payload)
            return
        if is_clinical_documentation_profile(payload.get("packet_profile")):
            self._write_clinical_documentation_doc_v2(path, payload)
            return
        if is_va_10172_profile(payload.get("packet_profile")):
            self._write_va_10172_summary_doc(path, payload)
            return
        document = Document()
        document.add_heading(payload.get("packet_title") or "TrueCore Packet Draft", level=0)
        document.add_paragraph(f"Packet Profile: {payload.get('packet_profile') or 'Pending'}")
        for section_title, rows in build_packet_builder_sections(payload):
            document.add_heading(section_title, level=1)
            table = document.add_table(rows=0, cols=2)
            table.style = "Table Grid"
            for label, value in rows:
                cells = table.add_row().cells
                cells[0].text = label
                cells[1].text = str(value or "Pending")
        document.save(path)

    def _write_pdf_doc_for_payload(self, path, payload):
        if is_va_10172_profile(payload.get("packet_profile")):
            self._write_va_10172_pdf(path, payload)
            return
        if find_word_executable():
            with tempfile.TemporaryDirectory() as temp_dir:
                docx_name = sanitize_builder_filename(payload.get("packet_title") or payload.get("packet_profile") or "packet_preview")
                docx_path = os.path.join(temp_dir, f"{docx_name}.docx")
                self._write_word_doc_for_payload(docx_path, payload)
                convert_docx_to_pdf_via_word(docx_path, path)
            return
        document = QTextDocument()
        document.setHtml(render_packet_builder_document_export(build_packet_builder_document_markup(payload)))
        printer = QPrinter(QPrinter.HighResolution)
        printer.setOutputFormat(QPrinter.PdfFormat)
        printer.setOutputFileName(path)
        printer.setPageMargins(QMarginsF(14, 14, 14, 14), QPageLayout.Millimeter)
        document.print_(printer)

    def _write_referral_request_doc(self, path, payload):
        document = Document()
        title = document.add_paragraph()
        title.alignment = 1
        title_run = title.add_run(payload.get("packet_title") or "Community Care Referral Request")
        title_run.bold = True
        title_run.font.size = Pt(20)

        subtitle = document.add_paragraph()
        subtitle.alignment = 1
        subtitle_run = subtitle.add_run(payload.get("referral_subtitle") or "Spine & Pain Evaluation")
        subtitle_run.bold = True

        intro = document.add_paragraph()
        intro.alignment = 1
        intro.add_run(
            "This document supports your request for a Community Care Referral for specialty spine and pain evaluation. "
            "Please share this with your Primary Care VA provider during your appointment."
        )

        emphasis = document.add_paragraph()
        emphasis_run = emphasis.add_run("You do not need to request a specific procedure to be referred.")
        emphasis_run.bold = True
        document.add_paragraph(
            "We provide a comprehensive evaluation in a non-surgical spine care setting and determine the most appropriate treatment based on your condition."
        )
        document.add_paragraph(
            "Community Care allows you to receive specialty care outside the VA when clinically appropriate."
        )

        document.add_heading("What to Do Next", level=1)
        steps = [
            "Schedule a visit with your Primary Care VA Provider (PCP).",
            "Make this request clearly:",
            f"“{payload.get('pcp_request_text') or 'Pending'}”",
            "Referral should be entered as:",
            payload.get("referral_entry_text") or "Pending",
            "Include all relevant diagnoses and areas of concern:",
            payload.get("areas_of_concern") or "Pending",
        ]
        for item in steps:
            document.add_paragraph(item, style="List Bullet")

        document.add_heading("Important", level=1)
        for item in [
            "This request is for:",
            "“Specialty evaluation and consultation – not a request for a specific procedure.”",
        ]:
            document.add_paragraph(item, style="List Bullet")

        document.add_heading("Referral Routing", level=1)
        document.add_paragraph(f"Group NPI: {payload.get('group_npi') or 'Pending'}")
        document.add_paragraph(f"Fax: {payload.get('fax_number') or 'Pending'}")
        document.add_paragraph("This ensures proper routing through the HSRM system.")

        document.add_heading("Once Authorized", level=1)
        for item in [
            "Our team will coordinate your evaluation",
            "Imaging and medical history will be reviewed",
            "A personalized treatment plan will be developed",
        ]:
            document.add_paragraph(item, style="List Bullet")

        document.add_heading("Send Records", level=1)
        document.add_paragraph("Please ensure all prior imaging and relevant medical records are included.")

        document.add_heading("Need Help?", level=1)
        document.add_paragraph("Our Veteran liaison team is here to assist you.")
        document.add_paragraph(payload.get("liaison_contact_info") or "Pending")

        document.save(path)

    def _write_virtual_consent_doc(self, path, payload):
        document = Document()

        title = document.add_paragraph()
        title.alignment = 1
        title_run = title.add_run(payload.get("packet_title") or payload.get("consent_form_title") or "TELEHEALTH VIRTUAL CONSENT FORM")
        title_run.bold = True
        title_run.font.size = Pt(18)

        document.add_heading("Demographic Information", level=1)
        demographic_rows = [
            ("Full Name", payload.get("patient_name") or "Pending"),
            ("Date of Birth", payload.get("date_of_birth") or "Pending"),
            ("Street Address", payload.get("street_address") or "Pending"),
            ("City", payload.get("city") or "Pending"),
            ("State", payload.get("state") or "Pending"),
            ("Zip", payload.get("zip_code") or "Pending"),
            ("Home Phone", payload.get("home_phone") or "Pending"),
            ("Mobile Phone", payload.get("mobile_phone") or "Pending"),
            ("Work Phone", payload.get("work_phone") or "Pending"),
            ("Email Address", payload.get("email_address") or "Pending"),
            ("SSN", payload.get("ssn") or "Pending"),
            ("Driver License", payload.get("drivers_license") or "Pending"),
            ("DL State", payload.get("drivers_license_state") or "Pending"),
            ("Appointment Confirmation Method", payload.get("appointment_confirmation_method") or "Pending"),
        ]
        self._write_key_value_table(document, demographic_rows)

        document.add_heading("Eligibility And Legal Questions", level=1)
        legal_rows = [
            ("Filed for Disability", payload.get("filed_for_disability") or "Pending"),
            ("Condition Work Related", payload.get("condition_work_related") or "Pending"),
            ("Condition Due to Accident", payload.get("condition_due_to_accident") or "Pending"),
            ("Yes Response Explanation", payload.get("yes_response_explanation") or "Pending"),
            ("Has Attorney", payload.get("has_attorney") or "Pending"),
            ("Attorney Name", payload.get("attorney_name") or "Pending"),
            ("Attorney Phone", payload.get("attorney_phone") or "Pending"),
            ("Emergency Contact", payload.get("emergency_contact_name") or "Pending"),
            ("Relationship", payload.get("emergency_contact_relationship") or "Pending"),
            ("Emergency Contact Phone", payload.get("emergency_contact_phone") or "Pending"),
        ]
        self._write_key_value_table(document, legal_rows)

        document.add_heading("Insurance Information", level=1)
        insurance_rows = [
            ("Primary Insurance Carrier", payload.get("primary_insurance_carrier") or "Pending"),
            ("Primary Insurance ID", payload.get("primary_insurance_id") or "Pending"),
            ("Primary Insurance Phone", payload.get("primary_insurance_phone") or "Pending"),
            ("Secondary Insurance Carrier", payload.get("secondary_insurance_carrier") or "Pending"),
            ("Secondary Insurance ID", payload.get("secondary_insurance_id") or "Pending"),
            ("Secondary Insurance Phone", payload.get("secondary_insurance_phone") or "Pending"),
            ("PCP / PCM", payload.get("pcp_pcm_name") or "Pending"),
            ("PCP / PCM Phone", payload.get("pcp_pcm_phone") or "Pending"),
            ("PCP / PCM Fax", payload.get("pcp_pcm_fax") or "Pending"),
        ]
        self._write_key_value_table(document, insurance_rows)

        document.add_heading("Consent for Medical Care and Treatment", level=1)
        document.add_paragraph(
            f"I, {payload.get('patient_name') or 'Pending'}, hereby agree and give my consent for "
            f"{payload.get('consent_provider_name') or 'Pending'} to furnish medical care and treatment considered necessary "
            f"and proper in evaluating or treating my physical condition. Initials: {payload.get('consent_initials') or 'Pending'}"
        )
        document.add_paragraph("FOR MINORS ONLY CONSENT FOR CARE:")
        document.add_paragraph(
            f"As parent and/or legal guardian, I authorize {payload.get('minor_doctor_name') or 'Pending'} to treat the minor patient "
            f"named in the attached form while I am not present. Initials: {payload.get('minor_consent_initials') or 'Pending'}"
        )

        document.add_heading("Telehealth Consent", level=1)
        for bullet in [
            "I understand that through an interactive video connection, telehealth might be used to perform my consultation and that we may use telehealth to connect while working together.",
            "I understand the benefits and risks involved with telehealth technology.",
            "I do not need to travel to the consultation location.",
            "I have access to a specialist through this consultation.",
            "There may be interruptions, unauthorized access and technical difficulties and the telemedicine consult/visit can be discontinued if needed.",
            "My healthcare information may be shared with other individuals for scheduling and billing purposes.",
            "Reasonable steps will be taken to maintain confidentiality.",
            "I understand the initial virtual meet-and-greet discussion is primarily intended to determine any and all treatment I might pursue, in addition to the best approach for scheduling these treatments.",
            "I have read this document in its entirety and understand the risks and benefits of telehealth sessions.",
        ]:
            document.add_paragraph(bullet, style="List Bullet")

        document.add_heading("Final Authorization", level=1)
        document.add_paragraph(
            f"By signing below, I agree that all the above information is correct, and that I authorize "
            f"{payload.get('service_authorization_name') or 'Pending'} to provide me with medical services and to furnish my physician, insurance company or attorney information concerning my injury and/or treatment."
        )
        self._add_signature_paragraph(
            document,
            "Veteran / Patient Signature (Parent/Guardian if applicable)",
            payload.get("patient_signature_name") or payload.get("patient_name") or "Pending",
            "patient_signature_name",
            payload,
        )
        document.add_paragraph(f"Date: {payload.get('patient_signature_date') or 'Pending'}")
        document.save(path)

    def _write_virtual_consent_doc_v2(self, path, payload):
        document = Document()
        section = document.sections[0]
        section.top_margin = Pt(48)
        section.bottom_margin = Pt(48)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)

        normal_style = document.styles["Normal"]
        normal_style.font.name = "Calibri"
        normal_style.font.size = Pt(11)

        title = document.add_paragraph()
        title.alignment = 1
        title_run = title.add_run(payload.get("packet_title") or payload.get("consent_form_title") or "TELEHEALTH VIRTUAL CONSENT FORM")
        title_run.bold = True
        title_run.font.size = Pt(16)
        title.paragraph_format.space_after = Pt(12)

        def fill_text(value, blank="_____________"):
            text = str(value or "").strip()
            return text or blank

        def add_line(text, bold=False, space_after=Pt(2)):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = space_after
            run = paragraph.add_run(text)
            run.bold = bold
            return paragraph

        def add_tabbed_line(parts, tab_stops, bold=False, space_after=Pt(2)):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = space_after
            for stop in tab_stops:
                paragraph.paragraph_format.tab_stops.add_tab_stop(Pt(stop), WD_TAB_ALIGNMENT.LEFT)
            run = paragraph.add_run("\t".join(parts))
            run.bold = bold
            return paragraph

        def add_two_line_row(labels, values, tab_stops):
            add_tabbed_line(labels, tab_stops, bold=True, space_after=Pt(0))
            add_tabbed_line(values, tab_stops, space_after=Pt(5))

        add_line("Demographic Information:", bold=True, space_after=Pt(4))
        add_two_line_row(
            ["Full Name (Shown on your Ins Card):", "Date of Birth:", "Phone:"],
            [
                fill_text(payload.get("patient_name"), "____________________________"),
                fill_text(payload.get("date_of_birth")),
                fill_text(payload.get("mobile_phone") or payload.get("phone")),
            ],
            [270, 405],
        )
        add_two_line_row(
            ["Street Address:", "City:", "State:", "Zip:"],
            [
                fill_text(payload.get("street_address"), "______________________"),
                fill_text(payload.get("city"), "_____"),
                fill_text(payload.get("state"), "___"),
                fill_text(payload.get("zip_code"), "_________"),
            ],
            [180, 280, 360],
        )
        add_two_line_row(
            ["Home Phone:", "Mobile Phone:", "Work Phone:"],
            [
                fill_text(payload.get("home_phone")),
                fill_text(payload.get("mobile_phone") or payload.get("phone"), "______________"),
                fill_text(payload.get("work_phone"), "_____________"),
            ],
            [170, 320],
        )
        add_two_line_row(
            ["Email Address:", "SSN:", "DL:", "State:"],
            [
                fill_text(payload.get("email_address"), "____________________"),
                fill_text(payload.get("ssn"), "______________"),
                fill_text(payload.get("drivers_license"), "_____________"),
                fill_text(payload.get("drivers_license_state"), "__________"),
            ],
            [180, 270, 360],
        )

        appointment = str(payload.get("appointment_confirmation_method") or "").strip().lower()
        add_line(
            "Appointment Confirmation Method:\t"
            + export_checkbox_marker(appointment == "phone")
            + " Phone\t"
            + export_checkbox_marker(appointment == "text")
            + " Text\t"
            + export_checkbox_marker(appointment == "email")
            + " Email",
            bold=True,
            space_after=Pt(6),
        )

        def yes_no_line(label, value):
            selected = str(value or "").strip().lower()
            add_line(
                f"{label}\t"
                + export_checkbox_marker(selected == "yes")
                + " Yes\t"
                + export_checkbox_marker(selected == "no")
                + " No",
                bold=True,
                space_after=Pt(4),
            )

        yes_no_line("Have you filed for Disability?", payload.get("filed_for_disability"))
        yes_no_line("Is your condition work related?", payload.get("condition_work_related"))
        yes_no_line("Is your condition due to an accident?", payload.get("condition_due_to_accident"))
        add_line("Please explain any responses you answered 'Yes' to:", bold=True, space_after=Pt(0))
        add_line(fill_text(payload.get("yes_response_explanation"), "________________________________________________________"), space_after=Pt(6))
        yes_no_line("Do you have an Attorney?", payload.get("has_attorney"))
        add_two_line_row(
            ["Attorney Full Name:", "Attorney Phone Number:"],
            [
                fill_text(payload.get("attorney_name"), "________________________"),
                fill_text(payload.get("attorney_phone"), "______________________"),
            ],
            [240],
        )
        add_two_line_row(
            ["Emergency Contact:", "Relationship:", "Phone:"],
            [
                fill_text(payload.get("emergency_contact_name"), "__________________"),
                fill_text(payload.get("emergency_contact_relationship")),
                fill_text(payload.get("emergency_contact_phone"), "_______________"),
            ],
            [190, 310],
        )

        add_line(
            "Insurance Information - If your condition is due to a work injury or other Personal Injury/Accident, there is a separate form you will need to complete. See accident information form.",
            bold=True,
            space_after=Pt(6),
        )
        add_two_line_row(
            ["Primary Insurance Carrier:", "ID Number:", "Insurance Phone:"],
            [
                fill_text(payload.get("primary_insurance_carrier"), "______________________"),
                fill_text(payload.get("primary_insurance_id"), "____________"),
                fill_text(payload.get("primary_insurance_phone"), "________________"),
            ],
            [250, 365],
        )
        add_two_line_row(
            ["Secondary Insurance Carrier:", "ID Number:", "Insurance Phone:"],
            [
                fill_text(payload.get("secondary_insurance_carrier"), "_______________________"),
                fill_text(payload.get("secondary_insurance_id"), "___________"),
                fill_text(payload.get("secondary_insurance_phone"), "_______________"),
            ],
            [250, 365],
        )
        add_two_line_row(
            ["PCP/PCM:", "Phone:", "Fax:"],
            [
                fill_text(payload.get("pcp_pcm_name"), "_______________________"),
                fill_text(payload.get("pcp_pcm_phone"), "___________"),
                fill_text(payload.get("pcp_pcm_fax"), "_______________"),
            ],
            [170, 300],
        )

        document.add_page_break()
        add_line("CONSENT FOR MEDICAL CARE AND TREATMENT", bold=True, space_after=Pt(8))
        add_line(
            "I, "
            + fill_text(payload.get("patient_name"), "________________")
            + " hereby agree and give my consent for "
            + fill_text(payload.get("consent_provider_name"), "_______________")
            + " to furnish medical care and treatment considered necessary and proper in evaluating or treating my physical condition. "
            + fill_text(payload.get("consent_initials"), "______")
            + " (initial)",
            space_after=Pt(8),
        )
        add_line("FOR MINORS ONLY CONSENT FOR CARE:", bold=True, space_after=Pt(4))
        add_line(
            "As parent and/or legal guardian, I authorize "
            + fill_text(payload.get("minor_doctor_name"), "_________________")
            + " (Doctors Name) to treat the minor patient named in the attached form while I am not present. "
            + fill_text(payload.get("minor_consent_initials"), "_____")
            + " (Parent/Guardian initial)",
            space_after=Pt(10),
        )
        add_line("TELEHEALTH CONSENT", bold=True, space_after=Pt(6))

        telehealth_lines = [
            "I understand that through an interactive video connection, telehealth might be used to perform my consultation and that we may use telehealth to connect while working together.",
            "I understand the benefits and risks involved with telehealth technology:",
            "I do not need to travel to the consultation location.",
            "I have access to a specialist through this consultation.",
            "There may be interruptions, unauthorized access and technical difficulties and my health care provider(s) or myself can discontinue the telemedicine consult/visit if it is felt that the videoconferencing connections are not adequate for the situation.",
            "My healthcare information may be shared with other individuals for scheduling and billing purposes.",
            "I also understand that other individuals may need to use the telehealth platform and that reasonable steps will be taken to maintain confidentiality of the information obtained.",
            "I understand the initial virtual meet and greet discussion is primarily intended to determine any and all treatment I might pursue, in addition to the best approach for scheduling these treatments. The virtual discussion will be followed by individual cost estimates (if applicable).",
            "I have read this document in its entirety and understand the risks and benefits of telehealth consultation/post op visits and have had my questions explained. I hereby consent to participate in telehealth sessions under the conditions described in this document.",
        ]
        for line in telehealth_lines:
            add_line(line, space_after=Pt(4))

        add_line(
            "By signing below, I agree that all the above information is correct, and that I authorize "
            + fill_text(payload.get("service_authorization_name"), "_____________")
            + " to provide me with medical services and to furnish my physician, insurance company or attorney information concerning my injury and/or treatment.",
            space_after=Pt(8),
        )
        self._add_signature_paragraph(
            document,
            "Veteran / Patient Signature (Parent/Guardian if applicable)",
            payload.get("patient_signature_name") or payload.get("patient_name"),
            "patient_signature_name",
            payload,
        )
        add_line(f"Date: {fill_text(payload.get('patient_signature_date'), '___________')}", space_after=Pt(2))
        document.save(path)

    def _write_key_value_table(self, document, rows):
        table = document.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        for label, value in rows:
            cells = table.add_row().cells
            cells[0].text = str(label or "")
            cells[1].text = str(value or "Pending")

    def _add_signature_paragraph(self, document, label, text_value, image_field_name, payload, left_indent=None):
        image_bytes = _signature_image_bytes(payload, image_field_name)
        paragraph = document.add_paragraph()
        if left_indent is not None:
            paragraph.paragraph_format.left_indent = left_indent
        paragraph.paragraph_format.space_after = Pt(2 if image_bytes else 4)
        label_run = paragraph.add_run(f"{label}: ")
        label_run.bold = True
        if image_bytes:
            signature_paragraph = document.add_paragraph()
            if left_indent is not None:
                signature_paragraph.paragraph_format.left_indent = left_indent + Pt(18)
            signature_paragraph.paragraph_format.space_after = Pt(4)
            run = signature_paragraph.add_run()
            run.add_picture(io.BytesIO(image_bytes), width=Inches(2.15))
        else:
            value_text = str(text_value or "").strip() or "_____________"
            value_run = paragraph.add_run(value_text)
            if str(text_value or "").strip():
                value_run.underline = True

    def _write_va_10172_summary_doc(self, path, payload):
        document = Document()
        title = document.add_paragraph()
        title.alignment = 1
        title_run = title.add_run(payload.get("packet_title") or "VA Form 10-10172")
        title_run.bold = True
        title_run.font.size = Pt(18)

        document.add_paragraph(
            "This Word export is a drafting summary for the VA Form 10-10172 profile. "
            "Use the PDF export to generate the actual filled VA form."
        )

        rows = [
            ("Veteran Legal Full Name", payload.get("patient_name") or "Pending"),
            ("Date of Birth", payload.get("date_of_birth") or "Pending"),
            ("VA Authorization Number", payload.get("authorization_number") or "Pending"),
            ("VA Facility & Address", payload.get("va10172_va_facility_address") or "Pending"),
            ("Ordering Provider Office Name & Address", payload.get("va10172_ordering_provider_office_address") or "Pending"),
            ("Indian Health Services / THP Provider", payload.get("va10172_is_ihs_provider") or "Pending"),
            ("Ordering Provider Phone", payload.get("va10172_ordering_provider_phone") or "Pending"),
            ("Ordering Provider Fax", payload.get("va10172_ordering_provider_fax") or "Pending"),
            ("Ordering Provider Secure Email", payload.get("va10172_ordering_provider_secure_email") or "Pending"),
            ("Care Needed Within 48 Hours", payload.get("va10172_care_needed_within_48_hours") or "Pending"),
            ("Continuation of Care", payload.get("va10172_is_continuation_of_care") or "Pending"),
            ("Referral to Another Specialty", payload.get("va10172_referral_to_specialty") or "Pending"),
            ("Specialty", payload.get("va10172_referral_specialty_text") or "Pending"),
            ("Diagnosis Codes", payload.get("icd_codes") or "Pending"),
            ("Diagnosis Description", payload.get("va10172_diagnosis_description") or "Pending"),
            ("Requested CPT/HCPCS Code", payload.get("va10172_requested_cpt_hcpcs_code") or "Pending"),
            ("Description CPT/HCPCS Code", payload.get("va10172_description_cpt_hcpcs_code") or "Pending"),
            ("Geriatric / Extended Care", payload.get("va10172_geriatric_care_option") or "Pending"),
            ("Reason for Request", payload.get("va10172_reason_for_request") or "Pending"),
            ("CCN Ordering Provider Name (Printed)", payload.get("va10172_ordering_provider_name_printed") or "Pending"),
            ("CCN Ordering Provider NPI", payload.get("va10172_ordering_provider_npi") or "Pending"),
            ("CCN Ordering Provider Signature", payload.get("va10172_signature_text") or "Pending"),
            ("Today's Date", payload.get("va10172_today_date") or "Pending"),
            ("Template PDF Path", payload.get("va10172_template_path") or "Pending"),
        ]
        self._write_key_value_table(document, rows)
        document.save(path)

    def _write_submission_cover_doc(self, path, payload):
        document = Document()
        section = document.sections[0]
        section.top_margin = Pt(54)
        section.bottom_margin = Pt(54)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)

        normal_style = document.styles["Normal"]
        normal_style.font.name = "Calibri"
        normal_style.font.size = Pt(12)

        title = document.add_paragraph()
        title.alignment = 1
        title_run = title.add_run(payload.get("packet_title") or payload.get("submission_cover_title") or "VA Submission Cover Sheet")
        title_run.bold = True
        title_run.font.size = Pt(18)
        title.paragraph_format.space_after = Pt(18)

        def add_fill_line(label, value):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.add_run(f"{label}: ").bold = True
            value_text = str(value or "").strip()
            value_run = paragraph.add_run(value_text or "_____________")
            if value_text:
                value_run.underline = True

        add_fill_line("Patient Name", payload.get("patient_name"))
        add_fill_line("DOB", payload.get("date_of_birth"))
        add_fill_line("VA Facility", payload.get("facility"))
        add_fill_line("Ordering Provider", payload.get("ordering_doctor") or payload.get("provider"))
        add_fill_line("Date of Submission", payload.get("submission_date"))
        add_fill_line("Primary Diagnosis Code", payload.get("primary_diagnosis_code") or payload.get("icd_codes"))

        checklist_heading = document.add_paragraph()
        checklist_heading.paragraph_format.space_before = Pt(12)
        checklist_heading.paragraph_format.space_after = Pt(6)
        checklist_heading.add_run("Documents Included (check all):").bold = True
        checklist_rows = [
            ("Virtual Consent Form completed and signed", payload.get("included_virtual_consent_form")),
            ("VA Form 10-10172 completed and signed", payload.get("included_va_form_10_10172")),
            ("SEOC request signed by CCN ordering provider", payload.get("included_seoc_request")),
            ("Consultation & Treatment Request completed", payload.get("included_consult_request")),
            ("Letter of Medical Necessity completed", payload.get("included_lomn")),
            ("Clinical Notes included", payload.get("included_clinical_notes")),
            ("MRI Report included", payload.get("included_mri_report")),
        ]
        for label, checked in checklist_rows:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(2)
            marker = "[X]" if checked else "[ ]"
            paragraph.add_run(f"{marker} {label}")

        review_heading = document.add_paragraph()
        review_heading.paragraph_format.space_before = Pt(12)
        review_heading.paragraph_format.space_after = Pt(6)
        review_heading.add_run("Submitting Office Review").bold = True
        add_fill_line("Submitting Office", payload.get("submitting_office"))
        add_fill_line("Office Staff Name", payload.get("office_staff_name"))
        signature_image = _signature_image_bytes(payload, "office_staff_signature")
        if signature_image:
            self._add_signature_paragraph(document, "Program User / Office Staff Signature", payload.get("office_staff_signature"), "office_staff_signature", payload, left_indent=body_indent)
        else:
            add_fill_line("Program User / Office Staff Signature", payload.get("office_staff_signature"))
        add_fill_line("Date Reviewed", payload.get("date_reviewed"))
        document.save(path)

    def _write_submission_cover_doc_v2(self, path, payload):
        document = Document()
        section = document.sections[0]
        section.top_margin = Pt(48)
        section.bottom_margin = Pt(48)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)

        normal_style = document.styles["Normal"]
        normal_style.font.name = "Calibri"
        normal_style.font.size = Pt(11)
        body_indent = Pt(18)

        title = document.add_paragraph()
        title.alignment = 1
        title_run = title.add_run(payload.get("packet_title") or payload.get("submission_cover_title") or "VA Submission Cover Sheet")
        title_run.bold = True
        title_run.font.size = Pt(18)
        title.paragraph_format.space_after = Pt(14)

        def add_fill_line(label, value):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = body_indent
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.add_run(f"{label}: ").bold = True
            value_text = str(value or "").strip()
            value_run = paragraph.add_run(value_text or "_____________")
            if value_text:
                value_run.underline = True

        add_fill_line("Patient Name", payload.get("patient_name"))
        add_fill_line("DOB", payload.get("date_of_birth"))
        add_fill_line("VA Facility", payload.get("facility"))
        add_fill_line("Ordering Provider", payload.get("ordering_doctor") or payload.get("provider"))
        add_fill_line("Date of Submission", payload.get("submission_date"))
        add_fill_line("Primary Diagnosis Code", payload.get("primary_diagnosis_code") or payload.get("icd_codes"))

        checklist_heading = document.add_paragraph()
        checklist_heading.paragraph_format.left_indent = body_indent
        checklist_heading.paragraph_format.space_before = Pt(10)
        checklist_heading.paragraph_format.space_after = Pt(5)
        checklist_heading.add_run("Documents Included (check all):").bold = True

        checklist_rows = [
            ("Virtual Consent Form completed and signed", payload.get("included_virtual_consent_form")),
            ("VA Form 10-10172 completed and signed", payload.get("included_va_form_10_10172")),
            ("SEOC request signed by CCN ordering provider", payload.get("included_seoc_request")),
            ("Consultation & Treatment Request completed", payload.get("included_consult_request")),
            ("Letter of Medical Necessity completed", payload.get("included_lomn")),
            ("Clinical Notes included", payload.get("included_clinical_notes")),
            ("MRI Report included", payload.get("included_mri_report")),
        ]
        for label, checked in checklist_rows:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = body_indent
            paragraph.paragraph_format.space_after = Pt(1)
            marker = export_checkbox_marker(checked)
            paragraph.add_run(f"{marker} {label}")

        spacer = document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)
        add_fill_line("Submitting Office", payload.get("submitting_office"))
        add_fill_line("Office Staff Name", payload.get("office_staff_name"))
        signature_image = _signature_image_bytes(payload, "office_staff_signature")
        if signature_image:
            self._add_signature_paragraph(document, "Program User / Office Staff Signature", payload.get("office_staff_signature"), "office_staff_signature", payload)
        else:
            add_fill_line("Program User / Office Staff Signature", payload.get("office_staff_signature"))
        add_fill_line("Date Reviewed", payload.get("date_reviewed"))
        document.save(path)

    def _write_seoc_request_doc(self, path, payload):
        document = Document()
        document.add_paragraph(payload.get("seoc_request_date") or "Pending")
        document.add_paragraph("Department of Veterans Affairs\nCommunity Care Office\n" + (payload.get("va_medical_center_name") or "Pending"))

        heading = document.add_paragraph()
        heading_run = heading.add_run("RE: Single Episode of Care (SEOC) Request")
        heading_run.bold = True
        document.add_paragraph(f"Veteran Name: {payload.get('patient_name') or 'Pending'}")
        document.add_paragraph(f"DOB: {payload.get('date_of_birth') or 'Pending'}")
        document.add_paragraph(f"Last Four SSN: {payload.get('last_four_ssn') or 'Pending'}")

        document.add_paragraph(
            "This request is for authorization of a defined, time-limited Single Episode of Care (SEOC) for treatment of lumbar disc pathology."
        )
        document.add_heading("Episode Diagnosis", level=1)
        diagnosis_line = payload.get("episode_diagnosis") or "Pending"
        if payload.get("episode_icd_code"):
            diagnosis_line = f"{diagnosis_line} – {payload.get('episode_icd_code')}"
        document.add_paragraph(diagnosis_line)

        document.add_heading("Scope of Requested Episode", level=1)
        document.add_paragraph(
            "This SEOC includes only the following services directly related to treatment of the above diagnosis:"
        )
        scope_items = [
            (payload.get("seoc_include_preprocedure_eval"), "Pre-procedure evaluation and procedural planning"),
            (payload.get("seoc_include_annulargram"), "Diagnostic Annulargram"),
            (payload.get("seoc_include_fibrin_injection"), "Inter Annular Fibrin Injections if indicated"),
            (payload.get("seoc_include_follow_up"), "Standard post-procedure follow-up visit(s) within the global postoperative period"),
        ]
        for enabled, label in scope_items:
            if enabled:
                document.add_paragraph(label, style="List Bullet")
        document.add_paragraph(
            payload.get("seoc_scope_text")
            or "No additional pain management services, long-term medication management, or unrelated spine care are requested under this episode."
        )

        document.add_heading("Estimated Duration of Episode", level=1)
        document.add_paragraph(
            "The episode is expected to begin upon authorization and conclude within approximately "
            + (payload.get("estimated_duration_text") or "Pending")
        )

        document.add_heading("Clinical Objective", level=1)
        objective_lines = [line.strip() for line in str(payload.get("clinical_objectives") or "").splitlines() if line.strip()]
        for line in objective_lines or ["Pending clinical objective"]:
            document.add_paragraph(line, style="List Bullet")

        document.add_heading("Continuity of Care", level=1)
        document.add_paragraph(
            payload.get("seoc_continuity_text")
            or "Upon completion of the episode, a summary of treatment rendered and clinical outcome will be forwarded to the referring VA provider. Any need for further care beyond this defined episode will require separate evaluation and authorization."
        )
        document.add_paragraph(
            "This request represents a discrete, procedure-based intervention and meets criteria for authorization under the Single Episode of Care (SEOC) framework."
        )

        document.add_paragraph("Sincerely,")
        document.add_paragraph(
            "\n".join(
                [
                    f"{payload.get('provider') or payload.get('ordering_doctor') or 'Pending'}"
                    + (f", {payload.get('provider_credentials')}" if payload.get("provider_credentials") else ""),
                    payload.get("provider_specialty") or "Pending",
                    f"NPI: {payload.get('provider_npi') or 'Pending'}",
                    payload.get("practice_name") or "Pending",
                    payload.get("provider_phone") or "Pending",
                    payload.get("provider_fax") or "Pending",
                ]
            )
        )
        document.save(path)

    def _write_lomn_doc(self, path, payload):
        document = Document()

        title = document.add_paragraph()
        title.alignment = 1
        title_run = title.add_run(payload.get("packet_title") or "LETTER OF MEDICAL NECESSITY")
        title_run.bold = True
        title_run.font.size = Pt(18)

        subtitle = document.add_paragraph()
        subtitle.alignment = 1
        subtitle_run = subtitle.add_run("(Pain Management - Interventional Spine)")
        subtitle_run.bold = True

        document.add_paragraph(payload.get("lmn_request_date") or "Pending")
        document.add_paragraph(
            "Department of Veterans Affairs\nCommunity Care Office\n"
            + (payload.get("va_medical_center_name") or "Pending")
            + ("\n" + (payload.get("facility") or "") if payload.get("facility") else "")
        )

        heading = document.add_paragraph()
        heading_run = heading.add_run("RE: Letter of Medical Necessity")
        heading_run.bold = True
        document.add_paragraph(f"Veteran Name: {payload.get('patient_name') or 'Pending'}")
        document.add_paragraph(f"DOB: {payload.get('date_of_birth') or 'Pending'}")
        document.add_paragraph(f"Last Four SSN: {payload.get('last_four_ssn') or 'Pending'}")
        document.add_paragraph(f"VA Claim Number: {payload.get('lmn_va_claim_number') or 'If Applicable'}")

        document.add_paragraph("To Whom It May Concern:")
        document.add_paragraph(
            "I am writing to formally document the medical necessity of interventional spine treatment for the above-named Veteran."
        )

        document.add_heading("Diagnosis", level=1)
        document.add_paragraph(f"Primary: {payload.get('lmn_primary_diagnosis') or 'Pending'}")
        document.add_paragraph(f"Secondary: {payload.get('lmn_secondary_diagnosis') or 'Pending'}")

        document.add_heading("Clinical Summary", level=1)
        document.add_paragraph(payload.get("lmn_clinical_summary") or "Pending")
        document.add_paragraph(f"MRI Date: {payload.get('lmn_mri_date') or 'Pending'}")
        document.add_paragraph(f"MRI Findings: {payload.get('lmn_mri_findings') or 'Pending'}")
        document.add_paragraph("The Veteran has completed extensive conservative management, including:")

        conservative_items = [
            (payload.get("lmn_include_physical_therapy"), "Structured physical therapy"),
            (payload.get("lmn_include_nsaids"), "NSAIDs and non-opioid analgesics"),
            (payload.get("lmn_include_activity_modification"), "Activity modification"),
            (payload.get("lmn_include_home_exercise"), "Home exercise program"),
            (payload.get("lmn_include_epidural_steroid_injections"), "Epidural Steroid injections (if applicable)"),
        ]
        for enabled, label in conservative_items:
            if enabled:
                document.add_paragraph(label, style="List Bullet")
        document.add_paragraph(
            "Despite appropriate and guideline-based conservative treatment for greater than "
            + (payload.get("lmn_conservative_duration") or "Pending")
            + ", the Veteran continues to experience persistent pain and functional limitation."
        )

        document.add_heading("Medical Necessity", level=1)
        document.add_paragraph(payload.get("lmn_medical_necessity_statement") or "Pending")
        document.add_paragraph("This intervention is indicated to:")

        indication_items = [
            (payload.get("lmn_indication_reduce_pain"), "Reduce pain severity"),
            (payload.get("lmn_indication_improve_function"), "Improve functional capacity"),
            (payload.get("lmn_indication_prevent_degeneration"), "Prevent further disc degeneration"),
            (payload.get("lmn_indication_reduce_opioid_reliance"), "Decrease reliance on long-term opioid therapy"),
            (payload.get("lmn_indication_prevent_surgery"), "Potentially prevent need for more invasive surgical intervention"),
        ]
        for enabled, label in indication_items:
            if enabled:
                document.add_paragraph(label, style="List Bullet")

        document.add_paragraph(payload.get("lmn_risk_statement") or "Pending")
        document.add_paragraph(payload.get("lmn_reasonable_necessary_statement") or "Pending")
        document.add_paragraph(payload.get("lmn_contact_statement") or "Pending")

        document.add_paragraph("Sincerely,")
        document.add_paragraph(
            "\n".join(
                [
                    f"{payload.get('provider') or payload.get('ordering_doctor') or 'Pending'}"
                    + (f", {payload.get('provider_credentials')}" if payload.get("provider_credentials") else ""),
                    payload.get("provider_specialty") or "Pending",
                    f"NPI: {payload.get('provider_npi') or 'Pending'}",
                    payload.get("practice_name") or "Pending",
                    payload.get("provider_phone") or "Pending",
                    payload.get("provider_fax") or "Pending",
                ]
            )
        )
        document.save(path)

    def _write_consult_request_doc(self, path, payload):
        document = Document()

        title = document.add_paragraph()
        title.alignment = 1
        title_run = title.add_run(payload.get("packet_title") or "CONSULTATION AND TREATMENT REQUEST")
        title_run.bold = True
        title_run.font.size = Pt(18)

        subtitle = document.add_paragraph()
        subtitle.alignment = 1
        subtitle_run = subtitle.add_run("(Pain Management - Interventional Spine)")
        subtitle_run.bold = True

        document.add_paragraph(payload.get("consult_request_date") or "Pending")
        document.add_paragraph(
            "Department of Veterans Affairs\nCommunity Care Office\n"
            + (payload.get("va_medical_center_name") or "Pending")
            + ("\n" + (payload.get("facility") or "") if payload.get("facility") else "")
        )

        heading = document.add_paragraph()
        heading_run = heading.add_run("RE: Consultation and Treatment Request")
        heading_run.bold = True
        document.add_paragraph(f"Veteran Name: {payload.get('patient_name') or 'Pending'}")
        document.add_paragraph(f"DOB: {payload.get('date_of_birth') or 'Pending'}")
        document.add_paragraph(f"Last Four SSN: {payload.get('last_four_ssn') or 'Pending'}")
        document.add_paragraph(f"VA Claim Number: {payload.get('consult_va_claim_number') or 'If Applicable'}")
        document.add_paragraph(f"Referring VA Provider: {payload.get('consult_referring_va_provider') or 'If Applicable'}")

        document.add_heading("Reason for Consultation", level=1)
        document.add_paragraph(payload.get("consult_reason_text") or "Pending")

        document.add_heading("Diagnoses", level=1)
        document.add_paragraph(f"Primary: {payload.get('consult_primary_diagnosis') or 'Pending'}")
        document.add_paragraph(f"Secondary: {payload.get('consult_secondary_diagnosis') or 'Pending'}")

        document.add_heading("Clinical Summary", level=1)
        document.add_paragraph(
            "The Veteran presents with persistent, function-limiting low back pain refractory to conservative management. Symptoms include:"
        )
        symptom_items = [
            (payload.get("consult_symptom_axial_pain"), "Chronic axial lumbar pain"),
            (payload.get("consult_symptom_activity_exacerbation"), "Activity-related exacerbation (sitting, standing, bending)"),
            (payload.get("consult_symptom_reduced_tolerance"), "Reduced tolerance for prolonged positioning"),
            (payload.get("consult_symptom_functional_impairment"), "Functional impairment affecting occupational and daily activities"),
        ]
        for enabled, label in symptom_items:
            if enabled:
                document.add_paragraph(label, style="List Bullet")

        document.add_paragraph(f"MRI Date: {payload.get('consult_mri_date') or 'Pending'}")
        document.add_paragraph(f"MRI Findings: {payload.get('consult_mri_findings') or 'Pending'}")
        document.add_paragraph("These findings are reviewed alongside the clinical presentation and prior treatment history.")
        document.add_paragraph("Conservative management has included:")
        conservative_items = [
            (payload.get("consult_include_physical_therapy"), "Physical therapy"),
            (payload.get("consult_include_nsaids"), "NSAIDs and non-opioid analgesics"),
            (payload.get("consult_include_activity_modification"), "Activity modification"),
            (payload.get("consult_include_home_exercise"), "Home exercise program"),
            (payload.get("consult_include_interventional_history"), "Epidural steroid injections and/or other interventional procedures (if applicable)"),
        ]
        for enabled, label in conservative_items:
            if enabled:
                document.add_paragraph(label, style="List Bullet")
        document.add_paragraph(
            "Despite appropriate treatment for greater than "
            + (payload.get("consult_conservative_duration") or "Pending")
            + ", the Veteran continues to experience persistent pain and impaired function."
        )

        document.add_heading("Requested Services", level=1)
        document.add_paragraph("Authorization is requested for:")
        service_items = [
            (payload.get("consult_include_pain_management_consultation"), "Comprehensive pain management consultation"),
            (payload.get("consult_include_procedural_planning"), "Diagnostic confirmation and procedural planning"),
            (payload.get("consult_include_annulargram"), "Diagnostic Annulargram"),
            (
                payload.get("consult_include_fibrin_injection"),
                "Intraannular Fibrin injection"
                + (f" at {payload.get('consult_fibrin_levels')}" if payload.get("consult_fibrin_levels") else "")
                + " if indicated",
            ),
            (payload.get("consult_include_follow_up"), "Routine post-procedure follow-up visit(s)"),
        ]
        for enabled, label in service_items:
            if enabled:
                document.add_paragraph(label, style="List Bullet")
        document.add_paragraph(payload.get("consult_scope_exclusion_text") or "Pending")

        document.add_heading("Medical Rationale", level=1)
        document.add_paragraph(payload.get("consult_medical_rationale_text") or "Pending")
        document.add_paragraph("Clinical goals include:")
        goal_items = [
            (payload.get("consult_goal_pain_reduction"), "Pain reduction"),
            (payload.get("consult_goal_functional_improvement"), "Functional improvement"),
            (payload.get("consult_goal_reduce_analgesics"), "Decreased reliance on chronic analgesic therapy"),
            (payload.get("consult_goal_prevent_surgery"), "Prevention of progression requiring more invasive surgical intervention"),
        ]
        for enabled, label in goal_items:
            if enabled:
                document.add_paragraph(label, style="List Bullet")
        document.add_paragraph(payload.get("consult_risk_without_treatment") or "Pending")

        document.add_heading("Duration and Scope of Care", level=1)
        document.add_paragraph(payload.get("consult_duration_scope_text") or "Pending")
        document.add_paragraph(payload.get("consult_contact_statement") or "Pending")

        document.add_paragraph("Sincerely,")
        document.add_paragraph(
            "\n".join(
                [
                    f"{payload.get('provider') or payload.get('ordering_doctor') or 'Pending'}"
                    + (f", {payload.get('provider_credentials')}" if payload.get("provider_credentials") else ""),
                    payload.get("provider_specialty") or "Pending",
                    f"NPI: {payload.get('provider_npi') or 'Pending'}",
                    payload.get("practice_name") or "Pending",
                    payload.get("provider_address") or "Pending",
                    payload.get("provider_phone") or "Pending",
                    payload.get("provider_fax") or "Pending",
                    payload.get("provider_email") or "Pending",
                ]
            )
        )
        document.save(path)

    def _write_clinical_documentation_doc(self, path, payload):
        document = Document()

        title = document.add_paragraph()
        title.alignment = 1
        title_run = title.add_run(payload.get("packet_title") or payload.get("clinical_doc_title") or "Clinical Documentation Template")
        title_run.bold = True
        title_run.font.size = Pt(18)

        document.add_heading("I. Chief Complaint", level=1)
        document.add_paragraph(payload.get("clinical_doc_chief_complaint") or "Pending")

        document.add_heading("II. History of Present Illness", level=1)
        document.add_paragraph("Duration of Symptoms:")
        duration_items = [
            (payload.get("clinical_doc_duration_gt_3m"), "> 3 months"),
            (payload.get("clinical_doc_duration_gt_6m"), "> 6 months"),
            (payload.get("clinical_doc_duration_gt_12m"), "> 12 months"),
        ]
        for enabled, label in duration_items:
            if enabled:
                document.add_paragraph(label, style="List Bullet")
        document.add_paragraph(f"Exact duration: {payload.get('clinical_doc_exact_duration') or 'Pending'}")

        document.add_paragraph("Pain Characteristics:")
        pain_items = [
            (payload.get("clinical_doc_pain_axial"), "Axial lumbar pain"),
            (payload.get("clinical_doc_pain_discogenic"), "Discogenic pattern"),
            (payload.get("clinical_doc_pain_activity_exacerbation"), "Activity-related exacerbation"),
            (payload.get("clinical_doc_pain_sitting_intolerance"), "Sitting intolerance"),
            (payload.get("clinical_doc_pain_standing_intolerance"), "Standing intolerance"),
            (payload.get("clinical_doc_pain_bending_lifting"), "Bending/lifting provocation"),
        ]
        for enabled, label in pain_items:
            if enabled:
                document.add_paragraph(label, style="List Bullet")
        document.add_paragraph(f"Pain severity (0-10): {payload.get('clinical_doc_pain_severity') or 'Pending'}")

        document.add_heading("III. Functional Impairment", level=1)
        document.add_paragraph("Patient reports limitation in:")
        limitation_items = [
            (payload.get("clinical_doc_limit_occupational"), "Occupational duties"),
            (payload.get("clinical_doc_limit_prolonged_sitting"), "Prolonged sitting"),
            (payload.get("clinical_doc_limit_prolonged_standing"), "Prolonged standing"),
            (payload.get("clinical_doc_limit_ambulation"), "Ambulation tolerance"),
            (payload.get("clinical_doc_limit_household"), "Household activities"),
            (payload.get("clinical_doc_limit_sleep"), "Sleep disturbance"),
        ]
        for enabled, label in limitation_items:
            if enabled:
                document.add_paragraph(label, style="List Bullet")
        document.add_paragraph(
            "Describe specific functional impact: "
            + (payload.get("clinical_doc_functional_impact") or "Pending")
        )

        document.add_heading("IV. Conservative Therapy History", level=1)
        document.add_paragraph("The patient has completed and/or attempted:")
        conservative_items = [
            (payload.get("clinical_doc_conservative_pt"), "Physical therapy"),
            (payload.get("clinical_doc_conservative_home_exercise"), "Home exercise program"),
            (payload.get("clinical_doc_conservative_nsaids"), "NSAIDs"),
            (payload.get("clinical_doc_conservative_non_opioid"), "Non-opioid analgesics"),
            (payload.get("clinical_doc_conservative_activity_modification"), "Activity modification"),
            (payload.get("clinical_doc_conservative_esi"), "Epidural steroid injections"),
            (payload.get("clinical_doc_conservative_other_interventional"), "Other interventional procedures"),
        ]
        for enabled, label in conservative_items:
            if enabled:
                document.add_paragraph(label, style="List Bullet")
        document.add_paragraph(
            "Duration of conservative treatment: "
            + (payload.get("clinical_doc_conservative_duration") or "Pending")
        )
        document.add_paragraph(
            "Response to prior ESI (if applicable): "
            + (payload.get("clinical_doc_esi_response") or "Pending")
        )

        document.add_heading("V. Imaging Findings", level=1)
        document.add_paragraph(f"MRI Date: {payload.get('clinical_doc_mri_date') or 'Pending'}")
        imaging_items = [
            (payload.get("clinical_doc_imaging_annular_tear"), "Annular tear"),
            (payload.get("clinical_doc_imaging_disc_degeneration"), "Disc degeneration"),
            (payload.get("clinical_doc_imaging_disc_protrusion"), "Disc protrusion"),
            (payload.get("clinical_doc_imaging_disc_displacement"), "Disc displacement"),
        ]
        document.add_paragraph("Findings:")
        for enabled, label in imaging_items:
            if enabled:
                document.add_paragraph(label, style="List Bullet")
        document.add_paragraph(f"Affected Levels: {payload.get('clinical_doc_affected_levels') or 'Pending'}")
        document.add_paragraph("Imaging correlates with clinical presentation.")

        document.add_heading("VI. Assessment", level=1)
        document.add_paragraph(f"Primary: {payload.get('clinical_doc_primary_diagnosis') or 'Pending'}")
        document.add_paragraph(f"Secondary: {payload.get('clinical_doc_secondary_diagnosis') or 'Pending'}")
        document.add_paragraph(payload.get("clinical_doc_assessment_summary") or "Pending")

        document.add_heading("VII. Treatment Plan", level=1)
        document.add_paragraph(payload.get("clinical_doc_treatment_plan_intro") or "Pending")
        document.add_paragraph("Plan includes:")
        plan_items = [
            (payload.get("clinical_doc_plan_diagnostic_confirmation"), "Diagnostic confirmation if necessary"),
            (payload.get("clinical_doc_plan_intradiscal_intervention"), "Intradiscal annular intervention if indicated"),
            (payload.get("clinical_doc_plan_follow_up"), "Standard post-procedure follow-up"),
        ]
        for enabled, label in plan_items:
            if enabled:
                document.add_paragraph(label, style="List Bullet")
        document.add_paragraph(payload.get("clinical_doc_plan_exclusion") or "Pending")

        document.add_heading("Physician Narrative Paragraph", level=1)
        document.add_paragraph(payload.get("clinical_doc_physician_narrative") or "Pending")

        document.save(path)

    def _write_seoc_request_doc_v2(self, path, payload):
        document = Document()
        section = document.sections[0]
        section.top_margin = Pt(46)
        section.bottom_margin = Pt(46)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)

        normal_style = document.styles["Normal"]
        normal_style.font.name = "Calibri"
        normal_style.font.size = Pt(10.5)
        body_indent = Pt(18)
        bullet_indent = Pt(30)

        def fill_text(value, blank="______________"):
            text = str(value or "").strip()
            return text or blank

        def add_line(text, bold=False, space_after=Pt(4), alignment=None, left_indent=None):
            paragraph = document.add_paragraph()
            if alignment is not None:
                paragraph.alignment = alignment
            if left_indent is not None:
                paragraph.paragraph_format.left_indent = left_indent
            paragraph.paragraph_format.space_after = space_after
            run = paragraph.add_run(text)
            run.bold = bold
            return paragraph

        def add_bullets(items, fallback_text):
            selected = [label for enabled, label in items if enabled]
            if not selected:
                selected = [fallback_text]
            for label in selected:
                add_line("- " + label, space_after=Pt(1), left_indent=bullet_indent)

        diagnosis_line = fill_text(payload.get("episode_diagnosis"), "________________________________________")

        objective_lines = [line.strip() for line in str(payload.get("clinical_objectives") or "").splitlines() if line.strip()]
        if not objective_lines:
            objective_lines = ["________________________________________"]
        scope_summary = str(payload.get("seoc_scope_text") or "").strip()
        continuity_text = str(payload.get("seoc_continuity_text") or "").strip()

        add_line(fill_text(payload.get("seoc_request_date"), "[Date]"))
        add_line(
            "Department of Veterans Affairs\nCommunity Care Office\n"
            + fill_text(payload.get("va_medical_center_name"), "[VA Medical Center Name]"),
            space_after=Pt(10),
        )
        add_line("RE: Single Episode of Care (SEOC) Request", bold=True, space_after=Pt(2))
        add_line(f"Veteran Name: {fill_text(payload.get('patient_name'), '[Full Name]')}", space_after=Pt(2))
        add_line(f"DOB: {fill_text(payload.get('date_of_birth'), '[MM/DD/YYYY]')}", space_after=Pt(2))
        add_line(f"Last Four SSN: {fill_text(payload.get('last_four_ssn'), '[XXXX]')}", space_after=Pt(8))
        add_line(
            "This request is for authorization of a defined, time-limited Single Episode of Care (SEOC) related to the diagnosis or condition listed below.",
            space_after=Pt(6),
            left_indent=body_indent,
        )
        add_line("Episode Diagnosis / Condition:", bold=True, space_after=Pt(2))
        add_line(f"Diagnosis: {diagnosis_line}", space_after=Pt(2), left_indent=body_indent)
        add_line(f"ICD-10 Code: {fill_text(payload.get('episode_icd_code'), '____________')}", space_after=Pt(6), left_indent=body_indent)
        add_line("Scope of Requested Episode:", bold=True, space_after=Pt(2))
        add_line(
            "This SEOC includes only the following services directly related to treatment of the above diagnosis:",
            space_after=Pt(3),
            left_indent=body_indent,
        )
        add_bullets(
            [
                (payload.get("seoc_include_preprocedure_eval"), "Pre-procedure evaluation and procedural planning"),
                (payload.get("seoc_include_annulargram"), "Diagnostic Annulargram"),
                (payload.get("seoc_include_fibrin_injection"), "Inter Annular Fibrin Injections if indicated"),
                (payload.get("seoc_include_follow_up"), "Standard post-procedure follow-up visit(s) within the global postoperative period"),
            ],
            "Pending scope items",
        )
        add_line(
            scope_summary or "No additional pain management services, long-term medication management, or unrelated spine care are requested under this episode.",
            space_after=Pt(6),
            left_indent=body_indent,
        )
        add_line("Estimated Duration of Episode:", bold=True, space_after=Pt(2))
        add_line(
            "Expected episode duration: " + fill_text(payload.get("estimated_duration_text"), "__________________"),
            space_after=Pt(6),
            left_indent=body_indent,
        )
        add_line("Clinical Objective:", bold=True, space_after=Pt(2))
        for line in objective_lines:
            add_line("- " + line, space_after=Pt(1), left_indent=bullet_indent)
        add_line("Continuity of Care:", bold=True, space_after=Pt(2))
        add_line(
            continuity_text or "Upon completion of the episode, a summary of treatment rendered and clinical outcome will be forwarded to the referring VA provider. Any need for further care beyond this defined episode will require separate evaluation and authorization.",
            space_after=Pt(3),
            left_indent=body_indent,
        )
        add_line("Sincerely,", space_after=Pt(8))
        provider_header = fill_text(payload.get("provider") or payload.get("ordering_doctor"), "[Provider Name]")
        credentials = str(payload.get("provider_credentials") or "").strip()
        if credentials:
            provider_header += f", {credentials}"
        add_line(provider_header, space_after=Pt(2))
        add_line(
            f"{fill_text(payload.get('provider_specialty'), '[Specialty]')} | NPI Number: {fill_text(payload.get('provider_npi'), '[NPI Number]')}",
            space_after=Pt(2),
        )
        add_line(fill_text(payload.get("practice_name") or payload.get("community_facility"), "[Practice Name]"), space_after=Pt(2))
        add_line(
            "Phone: "
            + fill_text(payload.get("provider_phone"), "[Phone]")
            + "    Fax: "
            + fill_text(payload.get("provider_fax"), "[Fax]"),
            space_after=Pt(0),
        )
        document.save(path)

    def _write_lomn_doc_v2(self, path, payload):
        document = Document()
        section = document.sections[0]
        section.top_margin = Pt(46)
        section.bottom_margin = Pt(46)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)

        normal_style = document.styles["Normal"]
        normal_style.font.name = "Calibri"
        normal_style.font.size = Pt(10.5)
        body_indent = Pt(18)
        bullet_indent = Pt(30)

        def fill_text(value, blank="______________"):
            text = str(value or "").strip()
            return text or blank

        def add_line(text, bold=False, space_after=Pt(4), alignment=None, left_indent=None):
            paragraph = document.add_paragraph()
            if alignment is not None:
                paragraph.alignment = alignment
            if left_indent is not None:
                paragraph.paragraph_format.left_indent = left_indent
            paragraph.paragraph_format.space_after = space_after
            run = paragraph.add_run(text)
            run.bold = bold
            return paragraph

        def add_bullets(items, fallback_text):
            selected = [label for enabled, label in items if enabled]
            if not selected:
                selected = [fallback_text]
            for label in selected:
                add_line("- " + label, space_after=Pt(1), left_indent=bullet_indent)

        provider_header = fill_text(payload.get("provider") or payload.get("ordering_doctor"), "[Provider Name]")
        credentials = str(payload.get("provider_credentials") or "").strip()
        if credentials:
            provider_header += f", {credentials}"
        practice_name = fill_text(payload.get("practice_name") or payload.get("community_facility"), "[Practice Name]")
        provider_address = str(payload.get("provider_address") or "").strip()
        provider_phone = fill_text(payload.get("provider_phone"), "[Phone]")
        provider_fax = fill_text(payload.get("provider_fax"), "[Fax]")

        title = payload.get("packet_title") or "LETTER OF MEDICAL NECESSITY"
        add_line(title, bold=True, alignment=1, space_after=Pt(2))
        add_line("(Pain Management - Interventional Spine)", bold=True, alignment=1, space_after=Pt(6))
        add_line(practice_name, alignment=1, space_after=Pt(1))
        add_line(provider_header, alignment=1, space_after=Pt(1))
        if provider_address:
            add_line(provider_address, alignment=1, space_after=Pt(1))
        add_line(f"Phone: {provider_phone} | Fax: {provider_fax}", alignment=1, space_after=Pt(8))

        add_line(fill_text(payload.get("lmn_request_date"), "[Date]"))
        va_block = [
            "Department of Veterans Affairs",
            "Community Care Office",
            fill_text(payload.get("va_medical_center_name"), "[VA Medical Center Name]"),
        ]
        facility = str(payload.get("facility") or "").strip()
        if facility and facility != va_block[-1]:
            va_block.append(facility)
        add_line("\n".join(va_block), space_after=Pt(10))

        add_line("RE: Letter of Medical Necessity", bold=True, space_after=Pt(2))
        add_line(f"Veteran Name: {fill_text(payload.get('patient_name'), '[Full Name]')}", space_after=Pt(2))
        add_line(f"DOB: {fill_text(payload.get('date_of_birth'), '[MM/DD/YYYY]')}", space_after=Pt(2))
        add_line(f"Last Four SSN: {fill_text(payload.get('last_four_ssn'), '[XXXX]')}", space_after=Pt(2))
        add_line(f"VA Claim Number: {fill_text(payload.get('lmn_va_claim_number'), '[If Applicable]')}", space_after=Pt(8))

        add_line("To Whom It May Concern:", bold=True, space_after=Pt(4))
        add_line(
            "I am writing to formally document the medical necessity of interventional spine treatment for the above-named Veteran.",
            space_after=Pt(6),
            left_indent=body_indent,
        )
        add_line("Clinical Diagnoses:", bold=True, space_after=Pt(2))
        add_line(f"Primary: {fill_text(payload.get('lmn_primary_diagnosis'), '______________________________')}", space_after=Pt(2), left_indent=body_indent)
        add_line(f"Secondary: {fill_text(payload.get('lmn_secondary_diagnosis'), '______________________________')}", space_after=Pt(6), left_indent=body_indent)

        add_line("Clinical Basis:", bold=True, space_after=Pt(2))
        add_line(
            str(payload.get("lmn_clinical_summary") or "").strip()
            or "____________________________________________________________",
            space_after=Pt(3),
            left_indent=body_indent,
        )
        add_line("Imaging Support:", bold=True, space_after=Pt(2))
        add_line(f"MRI Date: {fill_text(payload.get('lmn_mri_date'), '________________')}", space_after=Pt(2), left_indent=body_indent)
        add_line(f"MRI Findings: {fill_text(payload.get('lmn_mri_findings'), '________________________________________')}", space_after=Pt(3), left_indent=body_indent)
        add_line("The Veteran has completed extensive conservative management, including:", space_after=Pt(3), left_indent=body_indent)
        add_bullets(
            [
                (payload.get("lmn_include_physical_therapy"), "Structured physical therapy"),
                (payload.get("lmn_include_nsaids"), "NSAIDs and non-opioid analgesics"),
                (payload.get("lmn_include_activity_modification"), "Activity modification"),
                (payload.get("lmn_include_home_exercise"), "Home exercise program"),
                (payload.get("lmn_include_epidural_steroid_injections"), "Epidural Steroid injections (if applicable)"),
            ],
            "________________________________________",
        )
        add_line(
            "Despite appropriate and guideline-based conservative treatment for greater than "
            + fill_text(payload.get("lmn_conservative_duration"), "________________")
            + ", the Veteran continues to experience persistent pain and functional limitation.",
            space_after=Pt(6),
            left_indent=body_indent,
        )

        add_line("Basis for Medical Necessity:", bold=True, space_after=Pt(2))
        add_line(
            str(payload.get("lmn_medical_necessity_statement") or "").strip()
            or "____________________________________________________________",
            space_after=Pt(3),
            left_indent=body_indent,
        )
        add_line("Requested Treatment Objectives:", bold=True, space_after=Pt(2))
        add_line("This intervention is indicated to:", space_after=Pt(3), left_indent=body_indent)
        add_bullets(
            [
                (payload.get("lmn_indication_reduce_pain"), "Reduce pain severity"),
                (payload.get("lmn_indication_improve_function"), "Improve functional capacity"),
                (payload.get("lmn_indication_prevent_degeneration"), "Prevent further disc degeneration"),
                (payload.get("lmn_indication_reduce_opioid_reliance"), "Decrease reliance on long-term opioid therapy"),
                (payload.get("lmn_indication_prevent_surgery"), "Potentially prevent need for more invasive surgical intervention"),
            ],
            "________________________________________",
        )
        add_line("Risk if Treatment Is Delayed or Denied:", bold=True, space_after=Pt(2))
        add_line(
            str(payload.get("lmn_risk_statement") or "").strip()
            or "____________________________________________________________",
            space_after=Pt(3),
            left_indent=body_indent,
        )
        add_line("Reasonable and Necessary Determination:", bold=True, space_after=Pt(2))
        add_line(
            str(payload.get("lmn_reasonable_necessary_statement") or "").strip()
            or "____________________________________________________________",
            space_after=Pt(3),
            left_indent=body_indent,
        )
        add_line("Provider Contact Statement:", bold=True, space_after=Pt(2))
        add_line(
            str(payload.get("lmn_contact_statement") or "").strip()
            or "____________________________________________________________",
            space_after=Pt(10),
            left_indent=body_indent,
        )
        add_line("Sincerely,", space_after=Pt(8))
        add_line(provider_header, space_after=Pt(2))
        add_line(
            f"{fill_text(payload.get('provider_specialty'), '[Specialty - Pain Management / Interventional Spine]')} | NPI Number: {fill_text(payload.get('provider_npi'), '[NPI Number]')}",
            space_after=Pt(2),
        )
        add_line(practice_name, space_after=Pt(2))
        add_line(f"Phone: {provider_phone} | Fax: {provider_fax}", space_after=Pt(0))
        document.save(path)

    def _write_consult_request_doc_v2(self, path, payload):
        document = Document()
        section = document.sections[0]
        section.top_margin = Pt(46)
        section.bottom_margin = Pt(46)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)

        normal_style = document.styles["Normal"]
        normal_style.font.name = "Calibri"
        normal_style.font.size = Pt(10.5)
        body_indent = Pt(18)
        bullet_indent = Pt(30)

        def fill_text(value, blank="______________"):
            text = str(value or "").strip()
            return text or blank

        def add_line(text, bold=False, space_after=Pt(4), alignment=None, left_indent=None):
            paragraph = document.add_paragraph()
            if alignment is not None:
                paragraph.alignment = alignment
            if left_indent is not None:
                paragraph.paragraph_format.left_indent = left_indent
            paragraph.paragraph_format.space_after = space_after
            run = paragraph.add_run(text)
            run.bold = bold
            return paragraph

        def add_bullets(items, fallback_text):
            selected = [label for enabled, label in items if enabled]
            if not selected:
                selected = [fallback_text]
            for label in selected:
                add_line("- " + label, space_after=Pt(1), left_indent=bullet_indent)

        title = payload.get("packet_title") or "CONSULTATION AND TREATMENT REQUEST"
        provider_header = fill_text(payload.get("provider") or payload.get("ordering_doctor"), "[Provider Name]")
        credentials = str(payload.get("provider_credentials") or "").strip()
        if credentials:
            provider_header += f", {credentials}"
        practice_name = fill_text(payload.get("practice_name") or payload.get("community_facility"), "[Practice Name]")
        provider_address = fill_text(payload.get("provider_address"), "[Address]")
        provider_phone = fill_text(payload.get("provider_phone"), "[Phone]")
        provider_fax = fill_text(payload.get("provider_fax"), "[Fax]")
        provider_email = fill_text(payload.get("provider_email"), "[Email]")

        add_line(title, bold=True, alignment=1, space_after=Pt(2))
        add_line("(Pain Management - Interventional Spine)", bold=True, alignment=1, space_after=Pt(6))
        add_line(practice_name, alignment=1, space_after=Pt(1))
        add_line(provider_header, alignment=1, space_after=Pt(1))
        add_line(provider_address, alignment=1, space_after=Pt(1))
        add_line(f"Phone: {provider_phone} | Fax: {provider_fax}", alignment=1, space_after=Pt(1))
        add_line(provider_email, alignment=1, space_after=Pt(6))

        add_line(fill_text(payload.get("consult_request_date"), "[Date]"))
        va_block = [
            "Department of Veterans Affairs",
            "Community Care Office",
            fill_text(payload.get("va_medical_center_name"), "[VA Medical Center Name]"),
        ]
        facility = str(payload.get("facility") or "").strip()
        if facility and facility != va_block[-1]:
            va_block.append(facility)
        add_line("\n".join(va_block), space_after=Pt(10))

        add_line("RE: Consultation and Treatment Request", bold=True, space_after=Pt(2))
        add_line(f"Veteran Name: {fill_text(payload.get('patient_name'), '[Full Name]')}", space_after=Pt(2))
        add_line(f"DOB: {fill_text(payload.get('date_of_birth'), '[MM/DD/YYYY]')}", space_after=Pt(2))
        add_line(f"Last Four SSN: {fill_text(payload.get('last_four_ssn'), '[XXXX]')}", space_after=Pt(2))
        add_line(f"VA Claim Number: {fill_text(payload.get('consult_va_claim_number'), '[If Applicable]')}", space_after=Pt(2))
        add_line(f"Referring VA Provider: {fill_text(payload.get('consult_referring_va_provider'), '[Name, if applicable]')}", space_after=Pt(8))

        add_line("Reason for Consultation", bold=True, space_after=Pt(2))
        add_line(
            str(payload.get("consult_reason_text") or "").strip()
            or "____________________________________________________________",
            space_after=Pt(6),
            left_indent=body_indent,
        )
        add_line("Diagnoses", bold=True, space_after=Pt(2))
        add_line(f"Primary: {fill_text(payload.get('consult_primary_diagnosis'), '______________________________')}", space_after=Pt(2), left_indent=body_indent)
        add_line(f"Secondary: {fill_text(payload.get('consult_secondary_diagnosis'), '______________________________')}", space_after=Pt(6), left_indent=body_indent)

        add_line("Clinical Summary", bold=True, space_after=Pt(2))
        add_line(
            "Symptoms and functional concerns documented for this consultation include:",
            space_after=Pt(3),
            left_indent=body_indent,
        )
        add_bullets(
            [
                (payload.get("consult_symptom_axial_pain"), "Chronic axial lumbar pain"),
                (payload.get("consult_symptom_activity_exacerbation"), "Activity-related exacerbation (sitting, standing, bending)"),
                (payload.get("consult_symptom_reduced_tolerance"), "Reduced tolerance for prolonged positioning"),
                (payload.get("consult_symptom_functional_impairment"), "Functional impairment affecting occupational and daily activities"),
            ],
            "________________________________________",
        )
        add_line("Imaging Review", bold=True, space_after=Pt(2))
        add_line(f"MRI Date: {fill_text(payload.get('consult_mri_date'), '________________')}", space_after=Pt(2), left_indent=body_indent)
        add_line(f"MRI Findings: {fill_text(payload.get('consult_mri_findings'), '________________________________________')}", space_after=Pt(3), left_indent=body_indent)
        add_line("Conservative management has included:", space_after=Pt(3), left_indent=body_indent)
        add_bullets(
            [
                (payload.get("consult_include_physical_therapy"), "Physical therapy"),
                (payload.get("consult_include_nsaids"), "NSAIDs and non-opioid analgesics"),
                (payload.get("consult_include_activity_modification"), "Activity modification"),
                (payload.get("consult_include_home_exercise"), "Home exercise program"),
                (payload.get("consult_include_interventional_history"), "Epidural steroid injections and/or other interventional procedures (if applicable)"),
            ],
            "________________________________________",
        )
        add_line(
            "Despite appropriate treatment for greater than "
            + fill_text(payload.get("consult_conservative_duration"), "________________")
            + ", the Veteran continues to experience persistent pain and impaired function.",
            space_after=Pt(6),
            left_indent=body_indent,
        )

        add_line("Requested Services", bold=True, space_after=Pt(2))
        add_line("Authorization is requested for:", space_after=Pt(3), left_indent=body_indent)
        fibrin_line = "Intraannular Fibrin injection"
        if str(payload.get("consult_fibrin_levels") or "").strip():
            fibrin_line += f" at {payload.get('consult_fibrin_levels').strip()}"
        fibrin_line += " if indicated"
        add_bullets(
            [
                (payload.get("consult_include_pain_management_consultation"), "Comprehensive pain management consultation"),
                (payload.get("consult_include_procedural_planning"), "Diagnostic confirmation and procedural planning"),
                (payload.get("consult_include_annulargram"), "Diagnostic Annulargram"),
                (payload.get("consult_include_fibrin_injection"), fibrin_line),
                (payload.get("consult_include_follow_up"), "Routine post-procedure follow-up visit(s)"),
            ],
            "________________________________________",
        )
        add_line("Scope of Request", bold=True, space_after=Pt(2))
        add_line(
            str(payload.get("consult_scope_exclusion_text") or "").strip()
            or "____________________________________________________________",
            space_after=Pt(6),
            left_indent=body_indent,
        )

        add_line("Medical Rationale", bold=True, space_after=Pt(2))
        add_line(
            str(payload.get("consult_medical_rationale_text") or "").strip()
            or "____________________________________________________________",
            space_after=Pt(3),
            left_indent=body_indent,
        )
        add_line("Clinical goals include:", space_after=Pt(3), left_indent=body_indent)
        add_bullets(
            [
                (payload.get("consult_goal_pain_reduction"), "Pain reduction"),
                (payload.get("consult_goal_functional_improvement"), "Functional improvement"),
                (payload.get("consult_goal_reduce_analgesics"), "Decreased reliance on chronic analgesic therapy"),
                (payload.get("consult_goal_prevent_surgery"), "Prevention of progression requiring more invasive surgical intervention"),
            ],
            "________________________________________",
        )
        add_line("Risk Without Requested Treatment", bold=True, space_after=Pt(2))
        add_line(
            str(payload.get("consult_risk_without_treatment") or "").strip()
            or "____________________________________________________________",
            space_after=Pt(6),
            left_indent=body_indent,
        )

        add_line("Duration and Scope of Care", bold=True, space_after=Pt(2))
        add_line(
            str(payload.get("consult_duration_scope_text") or "").strip()
            or "____________________________________________________________",
            space_after=Pt(3),
            left_indent=body_indent,
        )
        add_line("Provider Contact Statement", bold=True, space_after=Pt(2))
        add_line(
            str(payload.get("consult_contact_statement") or "").strip()
            or "____________________________________________________________",
            space_after=Pt(10),
            left_indent=body_indent,
        )
        add_line("Sincerely,", space_after=Pt(8))
        add_line(provider_header, space_after=Pt(2))
        add_line(
            f"{fill_text(payload.get('provider_specialty'), '[Specialty - Pain Management / Interventional Spine]')} | NPI Number: {fill_text(payload.get('provider_npi'), '[NPI Number]')}",
            space_after=Pt(2),
        )
        add_line(practice_name, space_after=Pt(2))
        add_line(provider_address, space_after=Pt(2))
        add_line(f"Phone: {provider_phone} | Fax: {provider_fax} | Email: {provider_email}", space_after=Pt(0))
        document.save(path)

    def _write_clinical_documentation_doc_v2(self, path, payload):
        document = Document()
        section = document.sections[0]
        section.top_margin = Pt(46)
        section.bottom_margin = Pt(46)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)

        normal_style = document.styles["Normal"]
        normal_style.font.name = "Calibri"
        normal_style.font.size = Pt(10.5)
        body_indent = Pt(18)
        checkbox_indent = Pt(30)

        title = document.add_paragraph()
        title.alignment = 1
        title_run = title.add_run(payload.get("packet_title") or payload.get("clinical_doc_title") or "Clinical Documentation Template")
        title_run.bold = True
        title_run.font.size = Pt(16)
        title.paragraph_format.space_after = Pt(10)

        def fill_text(value, blank="_____________"):
            text = str(value or "").strip()
            return text or blank

        def add_line(text, bold=False, space_after=Pt(4), left_indent=None):
            paragraph = document.add_paragraph()
            if left_indent is not None:
                paragraph.paragraph_format.left_indent = left_indent
            paragraph.paragraph_format.space_after = space_after
            run = paragraph.add_run(text)
            run.bold = bold
            return paragraph

        def marker(enabled):
            return "[X]" if enabled else "[ ]"

        def add_checkbox_lines(items):
            for enabled, label in items:
                add_line(f"{marker(enabled)} {label}", space_after=Pt(1), left_indent=checkbox_indent)

        add_line("I. Chief Complaint", bold=True)
        add_line(fill_text(payload.get("clinical_doc_chief_complaint"), "____________________________"), left_indent=body_indent)

        add_line("II. History of Present Illness", bold=True)
        add_line("Duration of Symptoms:", space_after=Pt(2), left_indent=body_indent)
        add_checkbox_lines(
            [
                (payload.get("clinical_doc_duration_gt_3m"), "> 3 months"),
                (payload.get("clinical_doc_duration_gt_6m"), "> 6 months"),
                (payload.get("clinical_doc_duration_gt_12m"), "> 12 months"),
            ]
        )
        add_line(f"Exact duration: {fill_text(payload.get('clinical_doc_exact_duration'), '______________________')}", left_indent=body_indent)
        add_line("Pain Characteristics:", space_after=Pt(2), left_indent=body_indent)
        add_checkbox_lines(
            [
                (payload.get("clinical_doc_pain_axial"), "Axial lumbar pain"),
                (payload.get("clinical_doc_pain_discogenic"), "Discogenic pattern"),
                (payload.get("clinical_doc_pain_activity_exacerbation"), "Activity-related exacerbation"),
                (payload.get("clinical_doc_pain_sitting_intolerance"), "Sitting intolerance"),
                (payload.get("clinical_doc_pain_standing_intolerance"), "Standing intolerance"),
                (payload.get("clinical_doc_pain_bending_lifting"), "Bending/lifting provocation"),
            ]
        )
        add_line(f"Pain severity (0-10): {fill_text(payload.get('clinical_doc_pain_severity'), '______')}", left_indent=body_indent)

        add_line("III. Functional Impairment", bold=True)
        add_line("Patient reports limitation in:", left_indent=body_indent)
        add_checkbox_lines(
            [
                (payload.get("clinical_doc_limit_occupational"), "Occupational duties"),
                (payload.get("clinical_doc_limit_prolonged_sitting"), "Prolonged sitting"),
                (payload.get("clinical_doc_limit_prolonged_standing"), "Prolonged standing"),
                (payload.get("clinical_doc_limit_ambulation"), "Ambulation tolerance"),
                (payload.get("clinical_doc_limit_household"), "Household activities"),
                (payload.get("clinical_doc_limit_sleep"), "Sleep disturbance"),
            ]
        )
        add_line("Describe specific functional impact:", left_indent=body_indent)
        add_line(fill_text(payload.get("clinical_doc_functional_impact"), "____________________________"), left_indent=body_indent)

        add_line("IV. Conservative Therapy History", bold=True)
        add_line("The patient has completed and/or attempted:", left_indent=body_indent)
        add_checkbox_lines(
            [
                (payload.get("clinical_doc_conservative_pt"), "Physical therapy"),
                (payload.get("clinical_doc_conservative_home_exercise"), "Home exercise program"),
                (payload.get("clinical_doc_conservative_nsaids"), "NSAIDs"),
                (payload.get("clinical_doc_conservative_non_opioid"), "Non-opioid analgesics"),
                (payload.get("clinical_doc_conservative_activity_modification"), "Activity modification"),
                (payload.get("clinical_doc_conservative_esi"), "Epidural steroid injections"),
                (payload.get("clinical_doc_conservative_other_interventional"), "Other interventional procedures"),
            ]
        )
        add_line(f"Duration of conservative treatment: {fill_text(payload.get('clinical_doc_conservative_duration'), '____________________')}", left_indent=body_indent)
        add_line("Response to prior ESI (if applicable):", space_after=Pt(2), left_indent=body_indent)
        esi_response = str(payload.get("clinical_doc_esi_response") or "").strip().lower()
        add_checkbox_lines(
            [
                (esi_response == "temporary relief", "Temporary relief"),
                (esi_response == "partial relief", "Partial relief"),
                (esi_response == "no relief", "No relief"),
            ]
        )

        add_line("V. Imaging Findings", bold=True)
        add_line(f"MRI Date: {fill_text(payload.get('clinical_doc_mri_date'), '____________________')}", left_indent=body_indent)
        add_line("Findings:", left_indent=body_indent)
        add_checkbox_lines(
            [
                (payload.get("clinical_doc_imaging_annular_tear"), "Annular tear"),
                (payload.get("clinical_doc_imaging_disc_degeneration"), "Disc degeneration"),
                (payload.get("clinical_doc_imaging_disc_protrusion"), "Disc protrusion"),
                (payload.get("clinical_doc_imaging_disc_displacement"), "Disc displacement"),
            ]
        )
        add_line(f"Affected Levels: {fill_text(payload.get('clinical_doc_affected_levels'), '____________________')}", left_indent=body_indent)
        add_line("Imaging correlates with clinical presentation.", left_indent=body_indent)

        add_line("VI. Assessment", bold=True)
        add_line(f"Primary Diagnosis / ICD-10: {fill_text(payload.get('clinical_doc_primary_diagnosis'), '__________________________')}", left_indent=body_indent)
        add_line(f"Secondary Diagnosis / ICD-10: {fill_text(payload.get('clinical_doc_secondary_diagnosis'), '__________________________')}", left_indent=body_indent)
        add_line(
            fill_text(
                payload.get("clinical_doc_assessment_summary"),
                "____________________________________________________________",
            ),
            left_indent=body_indent,
        )

        add_line("VII. Treatment Plan", bold=True)
        add_line(
            fill_text(
                payload.get("clinical_doc_treatment_plan_intro"),
                "____________________________________________________________",
            ),
            left_indent=body_indent,
        )
        selected_plan_items = []
        if payload.get("clinical_doc_plan_diagnostic_confirmation"):
            selected_plan_items.append("Diagnostic confirmation if necessary")
        if payload.get("clinical_doc_plan_intradiscal_intervention"):
            selected_plan_items.append("Intradiscal annular intervention if indicated")
        if payload.get("clinical_doc_plan_follow_up"):
            selected_plan_items.append("Standard post-procedure follow-up")
        if not selected_plan_items:
            selected_plan_items = ["________________________________________"]
        add_line("Plan includes:", left_indent=body_indent)
        for item in selected_plan_items:
            add_line("- " + item, space_after=Pt(1), left_indent=checkbox_indent)
        add_line(
            fill_text(
                payload.get("clinical_doc_plan_exclusion"),
                "____________________________________________________________",
            ),
            left_indent=body_indent,
        )

        add_line("Physician Narrative Paragraph", bold=True, space_after=Pt(4))
        add_line(
            fill_text(
                payload.get("clinical_doc_physician_narrative"),
                "____________________________________________________________",
            ),
            left_indent=body_indent,
        )
        document.save(path)
