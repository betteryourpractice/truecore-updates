from __future__ import annotations

import base64
import html
import hashlib
import io
import json
import os
import re
import subprocess
import tempfile
import unicodedata
from datetime import datetime
from typing import Dict, List

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_TAB_ALIGNMENT
from PyPDF2 import PdfReader, PdfWriter
from PySide6.QtCore import QByteArray, QBuffer, QMarginsF, QObject, QRectF, QSizeF, Qt, QThread, QTimer, QSignalBlocker, Signal
from PySide6.QtGui import QColor, QFont, QFontMetrics, QGuiApplication, QImage, QPageLayout, QPainter, QPageSize, QPdfWriter, QPen, QPixmap, QTextDocument, QTextOption
from PySide6.QtPdf import QPdfDocument
from PySide6.QtPdfWidgets import QPdfView
from PySide6.QtPrintSupport import QPrinter
from PySide6.QtWidgets import (
    QCheckBox,
    QComboBox,
    QDialog,
    QFileDialog,
    QFormLayout,
    QFrame,
    QGridLayout,
    QGroupBox,
    QHBoxLayout,
    QInputDialog,
    QLabel,
    QLineEdit,
    QListWidget,
    QListWidgetItem,
    QMessageBox,
    QPlainTextEdit,
    QPushButton,
    QScrollArea,
    QSizePolicy,
    QSplitter,
    QStackedWidget,
    QTabWidget,
    QTextEdit,
    QVBoxLayout,
    QWidget,
)

from TrueCoreIntel.data.packet_model import Packet
from TrueCoreIntel.intelligence.packet_rubric import build_packet_rubric
from TrueCoreIntel.review.review_engine import ReviewEngine
from TrueCore.utils.runtime_info import resource_path, runtime_data_path


DEV_TOOLS_CONFIG_PATH = runtime_data_path("dev_system", "dev_tools_config.json")
PACKET_LIBRARY_DIR = runtime_data_path("dev_system", "packet_library")
REFERRAL_REQUEST_PROFILE = "Community Care Referral Request"
PATIENT_PACKET_PROFILE_ORDER = [
    "Submission Cover Sheet",
    "Virtual Consent Form",
    "VA Form 10-10172",
    "Single Episode of Care Request Template",
    "Consultation & Treatment Request Template",
    "Letter of Medical Necessity Template",
    "Clinical Notes Template",
]

PACKET_BUILDER_PROFILES = [
    REFERRAL_REQUEST_PROFILE,
    *PATIENT_PACKET_PROFILE_ORDER,
]
BUNDLE_EXPORT_EXTENSIONS = {
    "word": ["docx"],
    "pdf": ["pdf"],
    "both": ["docx", "pdf"],
}
LEGACY_IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".bmp", ".webp", ".ico"}
WORDING_REVIEW_DECISIONS = {"accepted", "edited", "keep_original"}
SIGNATURE_IMAGE_FIELDS = {
    "patient_signature_name": "patient_signature_image",
    "office_staff_signature": "office_staff_signature_image",
    "va10172_signature_text": "va10172_signature_image",
}


SIGNATURE_ROLE_LABELS = {
    "patient_signature_name": "Veteran / patient signature",
    "office_staff_signature": "program user / office staff signature",
    "va10172_signature_text": "CCN ordering provider signature",
}
WORDING_STATUS_COLORS = {
    "approved": ("#173B29", "#2C8B57", "#EAFBF1"),
    "needs_review": ("#4A3410", "#C7922E", "#FFF6DB"),
    "needs_facts": ("#441A1A", "#B55050", "#FFECEC"),
}
WORDING_FACT_LABELS = {
    "diagnosis": "Primary diagnosis",
    "secondary_diagnosis": "Secondary diagnosis",
    "icd_codes": "ICD-10 code",
    "requested_service": "Requested service",
    "mri_date": "MRI date",
    "mri_findings": "MRI findings",
    "affected_levels": "Affected spinal level(s)",
    "symptom_summary": "Clinical summary / symptoms",
    "functional_impact": "Functional impairment",
    "conservative_history": "Conservative treatment history",
    "consult_requested_services": "Requested consult services",
    "seoc_scope_items": "SEOC scope items",
    "clinical_plan_items": "Treatment plan items",
    "clinical_limitations": "Functional limitation detail",
    "estimated_duration_text": "Estimated duration",
    "ordering_doctor": "VA referring / ordering provider",
    "facility": "VA facility / medical center",
}


def _first_nonempty_text(*values):
    for value in values:
        text = str(value or "").strip()
        if text:
            return text
    return ""


def _ensure_sentence(text):
    normalized = sanitize_packet_builder_text(text)
    if not normalized:
        return ""
    if normalized[-1] not in ".!?":
        normalized += "."
    return normalized


def _human_join(items):
    cleaned = []
    seen = set()
    for item in items:
        text = str(item or "").strip()
        if not text:
            continue
        normalized = re.sub(r"\s+", " ", text).strip().lower()
        if normalized in seen:
            continue
        seen.add(normalized)
        cleaned.append(text)
    if not cleaned:
        return ""
    if len(cleaned) == 1:
        return cleaned[0]
    if len(cleaned) == 2:
        return f"{cleaned[0]} and {cleaned[1]}"
    return ", ".join(cleaned[:-1]) + f", and {cleaned[-1]}"


def _selected_labels(packet, option_pairs):
    labels = []
    for field_name, label in option_pairs:
        if packet.get(field_name):
            labels.append(label)
    return labels


def _consult_requested_service_labels(packet):
    labels = _selected_labels(packet, [
        ("consult_include_pain_management_consultation", "specialty pain management consultation"),
        ("consult_include_procedural_planning", "diagnostic confirmation and procedural planning"),
        ("consult_include_annulargram", "diagnostic annulargram if clinically indicated"),
        ("consult_include_fibrin_injection", "inter annular fibrin injection if clinically indicated"),
        ("consult_include_follow_up", "standard post-procedure follow-up"),
    ])
    if not labels:
        requested_service = str(packet.get("requested_service") or "").strip()
        if requested_service:
            labels.append(requested_service)
    return labels


def _seoc_scope_labels(packet):
    labels = _selected_labels(packet, [
        ("seoc_include_preprocedure_eval", "pre-procedure evaluation and procedural planning"),
        ("seoc_include_annulargram", "diagnostic annulargram"),
        ("seoc_include_fibrin_injection", "inter annular fibrin injection if clinically indicated"),
        ("seoc_include_follow_up", "standard post-procedure follow-up"),
    ])
    if not labels:
        requested_service = str(packet.get("requested_service") or "").strip()
        if requested_service:
            labels.append(requested_service)
    return labels


def _clinical_plan_labels(packet):
    labels = _selected_labels(packet, [
        ("clinical_doc_plan_diagnostic_confirmation", "diagnostic confirmation if clinically indicated"),
        ("clinical_doc_plan_intradiscal_intervention", "the requested interventional treatment if clinically indicated"),
        ("clinical_doc_plan_follow_up", "standard post-procedure follow-up"),
    ])
    if not labels:
        requested_service = str(packet.get("requested_service") or "").strip()
        if requested_service:
            labels.append(requested_service)
    return labels


def _clinical_limitation_labels(packet):
    labels = _selected_labels(packet, [
        ("clinical_doc_limit_occupational", "occupational duties"),
        ("clinical_doc_limit_prolonged_sitting", "prolonged sitting"),
        ("clinical_doc_limit_prolonged_standing", "prolonged standing"),
        ("clinical_doc_limit_ambulation", "ambulation tolerance"),
        ("clinical_doc_limit_household", "household activities"),
        ("clinical_doc_limit_sleep", "sleep quality"),
    ])
    direct_text = _clean_functional_impact_text(packet.get("clinical_doc_functional_impact"))
    if direct_text:
        labels.insert(0, direct_text)
    return labels


def _conservative_history_labels(packet):
    labels = []
    labels.extend(_selected_labels(packet, [
        ("consult_include_physical_therapy", "physical therapy"),
        ("consult_include_nsaids", "NSAIDs"),
        ("consult_include_activity_modification", "activity modification"),
        ("consult_include_home_exercise", "home exercise program"),
        ("consult_include_interventional_history", "prior interventional treatment"),
        ("lmn_include_physical_therapy", "physical therapy"),
        ("lmn_include_nsaids", "NSAIDs"),
        ("lmn_include_activity_modification", "activity modification"),
        ("lmn_include_home_exercise", "home exercise program"),
        ("lmn_include_epidural_steroid_injections", "epidural steroid injection history"),
        ("clinical_doc_conservative_pt", "physical therapy"),
        ("clinical_doc_conservative_home_exercise", "home exercise program"),
        ("clinical_doc_conservative_nsaids", "NSAIDs"),
        ("clinical_doc_conservative_non_opioid", "non-opioid analgesics"),
        ("clinical_doc_conservative_activity_modification", "activity modification"),
        ("clinical_doc_conservative_esi", "epidural steroid injections"),
        ("clinical_doc_conservative_other_interventional", "other interventional procedures"),
    ]))
    duration = _first_nonempty_text(
        packet.get("consult_conservative_duration"),
        packet.get("lmn_conservative_duration"),
        packet.get("clinical_doc_conservative_duration"),
    )
    if duration:
        labels.append(f"documented conservative management over {duration}")
    deduped = []
    for label in labels:
        if label not in deduped:
            deduped.append(label)
    return deduped


def _clean_functional_impact_text(value):
    text = sanitize_packet_builder_text(value)
    if not text:
        return ""
    patterns = [
        r"^(the condition continues to interfere with)\s+",
        r"^(current symptoms continue to limit)\s+",
        r"^(the veteran'?s symptoms interfere with)\s+",
        r"^(the veteran'?s ability to)\s+",
        r"^(symptoms continue to impair)\s+",
        r"^(functional impact is evident in)\s+",
    ]
    for pattern in patterns:
        text = re.sub(pattern, "", text, flags=re.IGNORECASE).strip(" :;-")
    return text


def _wording_fact_text(packet, fact_key):
    packet = normalize_packet_builder_payload(packet)
    if fact_key == "diagnosis":
        return _first_nonempty_text(
            packet.get("diagnosis"),
            packet.get("consult_primary_diagnosis"),
            packet.get("lmn_primary_diagnosis"),
            packet.get("clinical_doc_primary_diagnosis"),
            packet.get("episode_diagnosis"),
            packet.get("va10172_diagnosis_description"),
        )
    if fact_key == "secondary_diagnosis":
        return _first_nonempty_text(
            packet.get("secondary_diagnosis"),
            packet.get("consult_secondary_diagnosis"),
            packet.get("lmn_secondary_diagnosis"),
            packet.get("clinical_doc_secondary_diagnosis"),
        )
    if fact_key == "icd_codes":
        return _first_nonempty_text(
            packet.get("icd_codes"),
            packet.get("episode_icd_code"),
            packet.get("primary_diagnosis_code"),
        )
    if fact_key == "requested_service":
        return _first_nonempty_text(
            packet.get("requested_service"),
            packet.get("va10172_description_cpt_hcpcs_code"),
        )
    if fact_key == "mri_date":
        return _first_nonempty_text(
            packet.get("master_mri_date"),
            packet.get("consult_mri_date"),
            packet.get("lmn_mri_date"),
            packet.get("clinical_doc_mri_date"),
        )
    if fact_key == "mri_findings":
        return _first_nonempty_text(
            packet.get("master_mri_findings"),
            packet.get("consult_mri_findings"),
            packet.get("lmn_mri_findings"),
        )
    if fact_key == "affected_levels":
        return _first_nonempty_text(
            packet.get("master_affected_levels"),
            packet.get("clinical_doc_affected_levels"),
        )
    if fact_key == "symptom_summary":
        return _first_nonempty_text(
            packet.get("clinical_summary"),
            packet.get("clinical_doc_chief_complaint"),
            packet.get("lmn_clinical_summary"),
            packet.get("consult_reason_text"),
        )
    if fact_key == "functional_impact":
        return _first_nonempty_text(
            packet.get("clinical_doc_functional_impact"),
            _human_join(_clinical_limitation_labels(packet)),
        )
    if fact_key == "conservative_history":
        return _human_join(_conservative_history_labels(packet))
    if fact_key == "consult_requested_services":
        return _human_join(_consult_requested_service_labels(packet))
    if fact_key == "seoc_scope_items":
        return _human_join(_seoc_scope_labels(packet))
    if fact_key == "clinical_plan_items":
        return _human_join(_clinical_plan_labels(packet))
    if fact_key == "clinical_limitations":
        return _human_join(_clinical_limitation_labels(packet))
    if fact_key == "estimated_duration_text":
        return str(packet.get("estimated_duration_text") or "").strip()
    if fact_key in {"ordering_doctor", "facility"}:
        return str(packet.get(fact_key) or "").strip()
    return str(packet.get(fact_key) or "").strip()


def _diagnosis_phrase(packet):
    diagnosis = _wording_fact_text(packet, "diagnosis")
    icd_code = _first_icd_code(_wording_fact_text(packet, "icd_codes"))
    if not diagnosis:
        return ""
    if icd_code and icd_code not in diagnosis:
        return f"{diagnosis} ({icd_code})"
    return diagnosis


def _diagnosis_with_levels_phrase(packet):
    diagnosis_phrase = _diagnosis_phrase(packet)
    levels = _wording_fact_text(packet, "affected_levels")
    if diagnosis_phrase and levels and levels.lower() not in diagnosis_phrase.lower():
        return f"{diagnosis_phrase} involving {levels}"
    return diagnosis_phrase


def _wording_cycle_index(packet, context=""):
    current_key = str(packet.get("__wording_cycle_key") or "").strip().lower()
    if current_key:
        try:
            return max(0, int(packet.get("__wording_cycle_index") or 0))
        except Exception:
            return 0
    return 0


def _wording_variation_seed(packet, context=""):
    seed_parts = [
        context,
        _wording_fact_text(packet, "diagnosis"),
        _wording_fact_text(packet, "requested_service"),
        _wording_fact_text(packet, "mri_findings"),
        _wording_fact_text(packet, "functional_impact"),
        _wording_fact_text(packet, "conservative_history"),
    ]
    return "|".join(part for part in seed_parts if part)


def _pick_wording_variant(packet, context, *options):
    choices = [option for option in options if option]
    if not choices:
        return ""
    seed = _wording_variation_seed(packet, context)
    digest = hashlib.sha1(seed.encode("utf-8")).hexdigest()
    index = (int(digest, 16) + _wording_cycle_index(packet, context)) % len(choices)
    return choices[index]


def _pick_wording_index(packet, context, count):
    if count <= 0:
        return 0
    seed = _wording_variation_seed(packet, context)
    digest = hashlib.sha1(seed.encode("utf-8")).hexdigest()
    return (int(digest, 16) + _wording_cycle_index(packet, context)) % count


def _select_wording_blueprint(packet, context, scenarios):
    valid = [scenario for scenario in scenarios if scenario.get("text")]
    if not valid:
        return {
            "suggestion_text": "",
            "scenario_label": "",
            "scenario_use_when": "",
        }
    selected = valid[_pick_wording_index(packet, context, len(valid))]
    return {
        "suggestion_text": sanitize_packet_builder_text(selected.get("text") or ""),
        "scenario_label": str(selected.get("label") or "").strip(),
        "scenario_use_when": str(selected.get("use_when") or "").strip(),
    }


def _make_wording_scenario(label, use_when, text, enabled=True):
    return {
        "label": label,
        "use_when": use_when,
        "text": sanitize_packet_builder_text(text) if enabled and text else "",
    }


def _clean_lower_text(value):
    return sanitize_packet_builder_text(value).lower()


def _text_contains_any(value, *needles):
    haystack = _clean_lower_text(value)
    return any(str(needle or "").strip().lower() in haystack for needle in needles if str(needle or "").strip())


def _combined_wording_source_text(packet):
    parts = [
        _wording_fact_text(packet, "diagnosis"),
        _wording_fact_text(packet, "mri_findings"),
        _wording_fact_text(packet, "symptom_summary"),
        _wording_fact_text(packet, "requested_service"),
        _wording_fact_text(packet, "functional_impact"),
        _wording_fact_text(packet, "conservative_history"),
        _wording_fact_text(packet, "affected_levels"),
    ]
    return " | ".join(part for part in parts if part)


def _duration_source_text(packet):
    return _first_nonempty_text(
        packet.get("consult_conservative_duration"),
        packet.get("lmn_conservative_duration"),
        packet.get("clinical_doc_conservative_duration"),
        packet.get("clinical_doc_exact_duration"),
        packet.get("estimated_duration_text"),
    )


def _scenario_choices(packet, field_name):
    value = str(packet.get(field_name) or "").strip()
    if not value or value.lower() == "auto":
        return set()
    parts = [part.strip() for part in value.replace("\n", "|").split("|") if part.strip()]
    return {
        re.sub(r"[^a-z0-9]+", "_", part.lower()).strip("_")
        for part in parts
        if part.lower() != "auto"
    }


def _scenario_has(packet, field_name, *values):
    selected = _scenario_choices(packet, field_name)
    if not selected:
        return False
    wanted = {
        re.sub(r"[^a-z0-9]+", "_", str(value or "").lower()).strip("_")
        for value in values
        if str(value or "").strip()
    }
    return any(item in selected for item in wanted)


def _duration_bucket(packet):
    explicit = _scenario_choices(packet, "scenario_conservative_duration")
    for choice in ("over_12_months", "over_6_months", "over_90_days"):
        if choice in explicit:
            return choice
    source = _clean_lower_text(_duration_source_text(packet))
    if packet.get("clinical_doc_duration_gt_12m"):
        return "over_12_months"
    if packet.get("clinical_doc_duration_gt_6m"):
        return "over_6_months"
    if packet.get("clinical_doc_duration_gt_3m"):
        return "over_90_days"
    if not source:
        return ""
    if re.search(r"\b(12|twelve)\s*(months?|mos?)\b", source) or re.search(r"\b(1|one)\s*(years?|yrs?)\b", source):
        return "over_12_months"
    if re.search(r"\b(6|six)\s*(months?|mos?)\b", source):
        return "over_6_months"
    if re.search(r"\b(90|ninety)\b", source) and re.search(r"\b(days?)\b", source):
        return "over_90_days"
    if re.search(r"\b(3|three)\s*(months?|mos?)\b", source):
        return "over_90_days"
    return ""


def _has_annular_pattern(packet):
    if _scenario_has(packet, "scenario_pathology_pattern", "discogenic_annular"):
        return True
    if _scenario_has(packet, "scenario_symptom_pattern", "discogenic_pattern"):
        return True
    return (
        bool(packet.get("clinical_doc_imaging_annular_tear"))
        or _text_contains_any(_combined_wording_source_text(packet), "annular", "discogenic")
    )


def _has_radicular_pattern(packet):
    if _scenario_has(packet, "scenario_pathology_pattern", "radicular"):
        return True
    if _scenario_has(packet, "scenario_symptom_pattern", "radicular_symptoms"):
        return True
    return _text_contains_any(
        _combined_wording_source_text(packet),
        "radicular",
        "radiculopathy",
        "radiating pain",
        "numbness",
        "tingling",
        "sciatica",
    )


def _has_disc_displacement_pattern(packet):
    if _scenario_has(packet, "scenario_pathology_pattern", "disc_displacement_herniation"):
        return True
    return _text_contains_any(
        _combined_wording_source_text(packet),
        "herniat",
        "displacement",
        "protrusion",
        "extrusion",
    )


def _has_multilevel_pattern(packet):
    if _scenario_has(packet, "scenario_pathology_pattern", "multilevel"):
        return True
    levels = _clean_lower_text(_wording_fact_text(packet, "affected_levels"))
    if not levels:
        return False
    if "," in levels or "/" in levels or " and " in levels:
        return True
    level_matches = re.findall(r"[clt]\d+\s*-\s*[clt]\d+", levels)
    return len(level_matches) > 1


def _has_failed_esi_history(packet):
    if _scenario_has(packet, "scenario_prior_esi_response", "temporary_relief", "partial_relief", "no_relief"):
        return True
    if _scenario_has(packet, "scenario_prior_esi_response", "none_not_documented"):
        return False
    if _scenario_has(packet, "scenario_conservative_modalities", "prior_esi_interventional"):
        return True
    return any(
        bool(packet.get(flag))
        for flag in [
            "clinical_doc_conservative_esi",
            "lmn_include_epidural_steroid_injections",
            "consult_include_interventional_history",
        ]
    ) or _text_contains_any(_combined_wording_source_text(packet), "epidural steroid", "esi")


def _has_failed_pt_history(packet):
    if _scenario_has(packet, "scenario_conservative_modalities", "physical_therapy"):
        return True
    return any(
        bool(packet.get(flag))
        for flag in [
            "clinical_doc_conservative_pt",
            "lmn_include_physical_therapy",
            "consult_include_physical_therapy",
        ]
    )


def _has_medication_history(packet):
    if _scenario_has(packet, "scenario_conservative_modalities", "nsaids_analgesics"):
        return True
    return any(
        bool(packet.get(flag))
        for flag in [
            "clinical_doc_conservative_nsaids",
            "clinical_doc_conservative_non_opioid",
            "lmn_include_nsaids",
            "consult_include_nsaids",
        ]
    ) or _text_contains_any(_combined_wording_source_text(packet), "nsaid", "analgesic")


def _has_sleep_limitation(packet):
    if _scenario_has(packet, "scenario_functional_emphasis", "sleep_disruption"):
        return True
    return bool(packet.get("clinical_doc_limit_sleep")) or _text_contains_any(
        _wording_fact_text(packet, "functional_impact"),
        "sleep",
    )


def _has_occupational_limitation(packet):
    if _scenario_has(packet, "scenario_functional_emphasis", "work_capacity"):
        return True
    return bool(packet.get("clinical_doc_limit_occupational")) or _text_contains_any(
        _wording_fact_text(packet, "functional_impact"),
        "work",
        "occupational",
        "job",
    )


def _has_position_tolerance_limitation(packet):
    if _scenario_has(packet, "scenario_functional_emphasis", "sitting_standing_tolerance"):
        return True
    if _scenario_has(packet, "scenario_symptom_pattern", "position_intolerance", "activity_related_exacerbation"):
        return True
    return any(
        bool(packet.get(flag))
        for flag in [
            "clinical_doc_limit_prolonged_sitting",
            "clinical_doc_limit_prolonged_standing",
            "clinical_doc_limit_ambulation",
            "clinical_doc_limit_household",
        ]
    ) or _text_contains_any(
        _wording_fact_text(packet, "functional_impact"),
        "sitting",
        "standing",
        "walking",
        "bending",
        "lifting",
        "household",
    )


def _has_specific_requested_procedure(packet):
    if _scenario_has(packet, "scenario_request_framing", "specific_requested_procedure", "defined_interventional_pathway"):
        return True
    explicit = _scenario_choices(packet, "scenario_request_framing")
    if explicit == {"procedure_neutral_evaluation"}:
        return False
    return _text_contains_any(
        _wording_fact_text(packet, "requested_service"),
        "annulargram",
        "fibrin",
        "injection",
        "procedure",
        "intervention",
    )


def _has_imaging_correlation(packet):
    if _scenario_has(packet, "scenario_review_concern", "imaging_correlation_emphasis"):
        return True
    return bool(_wording_fact_text(packet, "mri_findings"))


def _symptom_summary_sentence(packet):
    symptom_summary = _wording_fact_text(packet, "symptom_summary")
    if not symptom_summary:
        return ""
    cleaned = sanitize_packet_builder_text(symptom_summary)
    weak_patterns = [
        r"\bhaving pain\b",
        r"\binability to function normally\b",
        r"\blower lumbar injuries\b",
        r"\bback pain\b",
    ]
    if any(re.search(pattern, cleaned, flags=re.IGNORECASE) for pattern in weak_patterns):
        diagnosis_phrase = _diagnosis_with_levels_phrase(packet) or "the documented lumbar condition"
        return _pick_wording_variant(
            packet,
            "symptom_summary_rewrite",
            _ensure_sentence(f"The Veteran has chronic, function-limiting pain related to {diagnosis_phrase}"),
            _ensure_sentence(f"Symptoms remain consistent with chronic, function-limiting pain associated with {diagnosis_phrase}"),
            _ensure_sentence(f"The current presentation reflects persistent pain and functional limitation related to {diagnosis_phrase}"),
            _ensure_sentence(f"The Veteran continues to report persistent pain and activity-related limitation related to {diagnosis_phrase}"),
        )
    return _ensure_sentence(cleaned)


def _imaging_correlation_sentence(packet):
    mri_date = _wording_fact_text(packet, "mri_date")
    mri_findings = _wording_fact_text(packet, "mri_findings")
    if mri_date and mri_findings:
        return _pick_wording_variant(
            packet,
            "imaging_correlation_with_date",
            _ensure_sentence(
                f"MRI dated {mri_date} demonstrates {mri_findings}, and the imaging findings correlate with the clinical presentation"
            ),
            _ensure_sentence(
                f"MRI dated {mri_date} shows {mri_findings}, and those findings correlate with the reported symptom pattern and examination context"
            ),
            _ensure_sentence(
                f"Imaging obtained on {mri_date} demonstrates {mri_findings}, with radiographic findings that support the clinical presentation"
            ),
        )
    if mri_findings:
        return _pick_wording_variant(
            packet,
            "imaging_correlation_no_date",
            _ensure_sentence(
                f"Imaging findings demonstrate {mri_findings}, and the imaging findings correlate with the clinical presentation"
            ),
            _ensure_sentence(
                f"Available imaging shows {mri_findings}, and those findings are consistent with the clinical presentation"
            ),
            _ensure_sentence(
                f"Imaging review demonstrates {mri_findings}, with findings that correlate with the Veteran's reported symptoms"
            ),
        )
    return ""


def _conservative_history_sentence(packet):
    conservative_history = _wording_fact_text(packet, "conservative_history")
    if not conservative_history:
        return ""
    return _pick_wording_variant(
        packet,
        "conservative_history",
        _ensure_sentence(f"Symptoms have persisted despite {conservative_history}"),
        _ensure_sentence(f"Persistent symptoms remain despite conservative management that has included {conservative_history}"),
        _ensure_sentence(f"The Veteran has not achieved sustained relief despite conservative treatment, including {conservative_history}"),
    )


def _functional_impact_sentence(packet):
    functional_impact = _clean_functional_impact_text(_wording_fact_text(packet, "functional_impact"))
    if not functional_impact:
        return ""
    return _pick_wording_variant(
        packet,
        "functional_impact",
        _ensure_sentence(f"The condition continues to interfere with {functional_impact}"),
        _ensure_sentence(f"Functional limitation remains evident in {functional_impact}"),
        _ensure_sentence(f"Symptoms continue to impair {functional_impact}"),
    )


def _requested_service_sentence(packet):
    requested_service = _wording_fact_text(packet, "requested_service")
    requested_cpt = _wording_fact_text(packet, "requested_cpt_code")
    if requested_service:
        if requested_cpt and requested_cpt not in requested_service:
            return _pick_wording_variant(
                packet,
                "requested_service_with_cpt",
                _ensure_sentence(f"Authorization is requested for {requested_service} associated with CPT/HCPCS code {requested_cpt}"),
                _ensure_sentence(f"The requested service is {requested_service}, submitted in association with CPT/HCPCS code {requested_cpt}"),
                _ensure_sentence(f"Review is requested for {requested_service} under CPT/HCPCS code {requested_cpt} if clinically indicated"),
            )
        return _pick_wording_variant(
            packet,
            "requested_service_no_cpt",
            _ensure_sentence(f"Authorization is requested for {requested_service}"),
            _ensure_sentence(f"The requested care includes {requested_service}"),
            _ensure_sentence(f"Review is requested for {requested_service} if clinically indicated"),
        )
    return _pick_wording_variant(
        packet,
        "requested_service_fallback",
        "Authorization is requested for specialty evaluation, diagnostic confirmation, and indicated interventional treatment planning.",
        "Authorization is requested for specialty evaluation, diagnostic confirmation, and clinically indicated interventional treatment planning.",
        "Review is requested for specialty evaluation, diagnostic confirmation, and an indicated interventional treatment pathway.",
    )


def _build_consult_reason_suggestion(packet, raw_text):
    requested_service = _wording_fact_text(packet, "requested_service") or "specialty interventional spine consultation and treatment planning"
    diagnosis_phrase = _diagnosis_phrase(packet) or "the documented lumbar condition"
    diagnosis_with_levels = _diagnosis_with_levels_phrase(packet) or diagnosis_phrase
    mri_findings = _wording_fact_text(packet, "mri_findings") or "documented imaging findings"
    levels = _wording_fact_text(packet, "affected_levels") or "the affected lumbar level(s)"
    conservative = _wording_fact_text(packet, "conservative_history") or "appropriate conservative management"
    duration = _duration_source_text(packet) or "the documented treatment period"
    impact = _clean_functional_impact_text(_wording_fact_text(packet, "functional_impact")) or "daily function and activity tolerance"
    scenarios = [
        _make_wording_scenario(
            "Discogenic / Annular pathology consult",
            "Use when annular or discogenic pathology is documented and specialty correlation is needed.",
            (
                f"Consultation is requested for symptoms consistent with discogenic lumbar pain, with imaging findings demonstrating {mri_findings} at {levels}. "
                f"Specialty review is needed to correlate symptoms, imaging, and treatment history and to determine whether {requested_service} is clinically indicated."
            ),
            enabled=_has_annular_pattern(packet),
        ),
        _make_wording_scenario(
            "Radicular / herniation consult",
            "Use when disc displacement, herniation, or radicular symptoms are part of the presentation.",
            (
                f"Authorization is requested for consultation related to documented lumbar disc pathology involving {levels}, with persistent symptoms concerning for radicular involvement. "
                f"Diagnostic confirmation and treatment planning are needed for {diagnosis_phrase} after inadequate response to {conservative}."
            ),
            enabled=_has_radicular_pattern(packet) or _has_disc_displacement_pattern(packet),
        ),
        _make_wording_scenario(
            "Imaging correlation consult",
            "Use when MRI correlation should be stated clearly.",
            (
                f"MRI dated {_wording_fact_text(packet, 'mri_date') or 'the documented study date'} demonstrates {mri_findings}, which correlates with the Veteran's reported pain pattern and functional limitations. "
                f"Consultation and treatment planning are requested for {diagnosis_with_levels}."
            ),
            enabled=_has_imaging_correlation(packet),
        ),
        _make_wording_scenario(
            "Failed epidural / interventional history consult",
            "Use when prior epidural steroid injection or other interventional care did not produce sustained relief.",
            (
                f"Consultation is requested because the Veteran continues to report persistent, function-limiting symptoms despite prior epidural steroid injection history and other conservative measures, including {conservative}. "
                f"Specialty review is needed to reevaluate {diagnosis_with_levels} and determine the clinically appropriate next step within a defined scope of care."
            ),
            enabled=_has_failed_esi_history(packet),
        ),
        _make_wording_scenario(
            "Extended conservative care consult",
            "Use when symptoms persisted despite prolonged conservative care.",
            (
                f"Authorization is requested because symptoms have persisted despite {duration} of conservative treatment, including {conservative}. "
                f"Specialty consultation is needed to evaluate {diagnosis_with_levels} and determine indicated treatment options within the requested pathway."
            ),
            enabled=_duration_bucket(packet) in {"over_90_days", "over_6_months", "over_12_months"},
        ),
        _make_wording_scenario(
            "Functional impairment consult",
            "Use when functional limitation is the clearest reason for escalation.",
            (
                f"Consultation is requested because the Veteran reports ongoing functional limitation affecting {impact} in association with {diagnosis_with_levels}. "
                f"Symptoms remain persistent despite {conservative}, and further specialty review is needed before proceeding with {requested_service}."
            ),
            enabled=bool(_wording_fact_text(packet, "functional_impact")),
        ),
        _make_wording_scenario(
            "High-quality default consult",
            "Use when a broad but denial-resistant default is needed.",
            (
                f"Authorization is requested for {requested_service} related to documented {diagnosis_with_levels}. "
                f"The Veteran remains function-limited despite {conservative}, and the available imaging findings support the current clinical presentation."
            ),
            enabled=True,
        ),
    ]
    return _select_wording_blueprint(packet, "consult_reason", scenarios)


def _build_consult_rationale_suggestion(packet, raw_text):
    diagnosis_with_levels = _diagnosis_with_levels_phrase(packet) or "the documented lumbar condition"
    conservative = _wording_fact_text(packet, "conservative_history") or "appropriate conservative management"
    impact = _clean_functional_impact_text(_wording_fact_text(packet, "functional_impact")) or "daily function"
    mri_findings = _wording_fact_text(packet, "mri_findings") or "documented imaging findings"
    scenarios = [
        _make_wording_scenario(
            "Imaging-supported rationale",
            "Use when MRI correlation is central to medical justification.",
            (
                f"The requested consultation and treatment pathway is supported by documented {diagnosis_with_levels}, with imaging demonstrating {mri_findings} that correlate with the clinical presentation. "
                f"Persistent symptoms and reduced function despite {conservative} support specialist review and procedural planning."
            ),
            enabled=_has_imaging_correlation(packet),
        ),
        _make_wording_scenario(
            "Failed conservative care rationale",
            "Use when conservative-care failure is the primary support.",
            (
                f"The requested care pathway is supported by documented {diagnosis_with_levels}, persistent symptoms, and lack of sustained relief despite conservative treatment that has included {conservative}. "
                f"Further specialty evaluation is appropriate to determine whether the requested intervention remains clinically indicated."
            ),
            enabled=bool(_wording_fact_text(packet, "conservative_history")),
        ),
        _make_wording_scenario(
            "Occupational / daily function rationale",
            "Use when occupational or activity tolerance impairment is prominent.",
            (
                f"The requested specialty evaluation is supported by ongoing functional impairment affecting {impact}, together with documented {diagnosis_with_levels}. "
                f"Because symptoms persist despite {conservative}, additional specialty review and treatment planning are medically appropriate."
            ),
            enabled=bool(_wording_fact_text(packet, "functional_impact")),
        ),
        _make_wording_scenario(
            "Records-supported rationale",
            "Use when the packet clearly includes the diagnosis, imaging, and treatment history.",
            (
                f"This request is supported by the attached diagnosis, imaging findings, treatment history, and functional limitation details related to {diagnosis_with_levels}. "
                f"Those records support specialty consultation to confirm the pain generator and define the appropriate treatment pathway."
            ),
            enabled=True,
        ),
    ]
    return _select_wording_blueprint(packet, "consult_rationale", scenarios)


def _build_consult_scope_suggestion(packet, raw_text):
    services = _wording_fact_text(packet, "consult_requested_services") or _wording_fact_text(packet, "requested_service") or "specialty evaluation, diagnostic confirmation, indicated intervention if clinically appropriate, and standard follow-up"
    diagnosis_phrase = _diagnosis_phrase(packet) or "the documented lumbar condition"
    scenarios = [
        _make_wording_scenario(
            "Diagnostic-first consult scope",
            "Use when evaluation and confirmation should occur before any intervention is assumed.",
            (
                f"This consultation request is limited to specialty evaluation, diagnostic confirmation, and procedural planning related to {diagnosis_phrase}. "
                f"Any intervention should occur only if clinically indicated after review of the documented condition."
            ),
            enabled=not _has_specific_requested_procedure(packet),
        ),
        _make_wording_scenario(
            "Procedure-cautious consult scope",
            "Use when a procedure is named but should not sound preapproved.",
            (
                f"This request is limited to {services} related only to {diagnosis_phrase}. "
                f"The requested intervention should proceed only if clinically appropriate after specialist evaluation and diagnostic confirmation."
            ),
            enabled=_has_specific_requested_procedure(packet),
        ),
        _make_wording_scenario(
            "Non-open-ended consult scope",
            "Use when the reviewer needs reassurance that care is contained.",
            (
                f"This request is limited to {services} for {diagnosis_phrase}. "
                "It is not a request for open-ended pain management, unrelated spine care, or long-term medication management; additional or unrelated services require separate authorization."
            ),
            enabled=True,
        ),
    ]
    return _select_wording_blueprint(packet, "consult_scope", scenarios)


def _build_lomn_clinical_basis_suggestion(packet, raw_text):
    diagnosis_with_levels = _diagnosis_with_levels_phrase(packet) or "the documented lumbar condition"
    conservative = _wording_fact_text(packet, "conservative_history") or "appropriate conservative management"
    mri_date = _wording_fact_text(packet, "mri_date") or "the documented study date"
    mri_findings = _wording_fact_text(packet, "mri_findings") or "documented lumbar pathology"
    impact = _clean_functional_impact_text(_wording_fact_text(packet, "functional_impact")) or "daily function"
    scenarios = [
        _make_wording_scenario(
            "Annular / discogenic clinical basis",
            "Use when annular pathology or discogenic pain should anchor the clinical basis.",
            (
                f"The Veteran presents with symptoms consistent with discogenic lumbar pain associated with {diagnosis_with_levels}. "
                f"MRI dated {mri_date} demonstrates {mri_findings}, which correlate with the Veteran's functional limitations and persistent symptoms despite {conservative}."
            ),
            enabled=_has_annular_pattern(packet),
        ),
        _make_wording_scenario(
            "Radicular / displacement clinical basis",
            "Use when disc displacement or radicular features are part of the record.",
            (
                f"The Veteran has documented lumbar disc pathology involving {diagnosis_with_levels}, with persistent symptoms concerning for radicular involvement and reduced function in {impact}. "
                f"Those symptoms continue despite {conservative}, and imaging findings support the current clinical presentation."
            ),
            enabled=_has_radicular_pattern(packet) or _has_disc_displacement_pattern(packet),
        ),
        _make_wording_scenario(
            "Extended conservative care clinical basis",
            "Use when long-standing symptoms and conservative-care failure are the clearest support.",
            (
                f"The Veteran continues to report chronic, function-limiting symptoms related to {diagnosis_with_levels} despite {conservative} over {_duration_source_text(packet) or 'the documented treatment period'}. "
                f"MRI findings showing {mri_findings} support the need for the requested specialty evaluation and treatment pathway."
            ),
            enabled=_duration_bucket(packet) in {"over_90_days", "over_6_months", "over_12_months"},
        ),
        _make_wording_scenario(
            "Default clinical basis",
            "Use when a broad medical-necessity foundation is needed without overstating the case.",
            (
                f"The Veteran has documented {diagnosis_with_levels} with persistent, function-limiting symptoms affecting {impact}. "
                f"Imaging demonstrates {mri_findings}, and symptoms remain refractory to {conservative}."
            ),
            enabled=True,
        ),
    ]
    return _select_wording_blueprint(packet, "lomn_clinical_basis", scenarios)


def _build_lomn_medical_necessity_suggestion(packet, raw_text):
    diagnosis_phrase = _diagnosis_with_levels_phrase(packet) or "the documented lumbar condition"
    requested_service = _wording_fact_text(packet, "requested_service") or "the requested interventional pathway"
    conservative = _wording_fact_text(packet, "conservative_history") or "appropriate conservative management"
    mri_findings = _wording_fact_text(packet, "mri_findings") or "documented imaging findings"
    scenarios = [
        _make_wording_scenario(
            "Imaging-correlation necessity",
            "Use when imaging correlation is central to necessity.",
            (
                f"Based on documented {diagnosis_phrase}, imaging findings demonstrating {mri_findings}, persistent symptoms despite {conservative}, and associated functional limitation, {requested_service} is medically reasonable and necessary if clinically indicated."
            ),
            enabled=_has_imaging_correlation(packet),
        ),
        _make_wording_scenario(
            "Failed conservative care necessity",
            "Use when conservative-care failure is the strongest support.",
            (
                f"Given persistent symptoms related to {diagnosis_phrase}, inadequate sustained response to {conservative}, and ongoing functional impairment, {requested_service} is medically reasonable and necessary if clinically indicated."
            ),
            enabled=bool(_wording_fact_text(packet, "conservative_history")),
        ),
        _make_wording_scenario(
            "Procedure-specific but restrained necessity",
            "Use when the requested procedure should be named without sounding guaranteed or promotional.",
            (
                f"The requested {requested_service} is medically reasonable and necessary if clinically indicated after specialty evaluation and diagnostic confirmation. "
                f"This request is supported by documented {diagnosis_phrase}, persistent symptoms, and lack of sustained relief with conservative care."
            ),
            enabled=_has_specific_requested_procedure(packet),
        ),
        _make_wording_scenario(
            "Default medical necessity",
            "Use when a broad, denial-resistant necessity statement is needed.",
            (
                f"Based on documented {diagnosis_phrase}, persistent symptoms despite appropriate conservative management, functional limitation, and imaging findings that correlate with the clinical presentation, {requested_service} is medically reasonable and necessary if clinically indicated."
            ),
            enabled=True,
        ),
    ]
    return _select_wording_blueprint(packet, "lomn_medical_necessity", scenarios)


def _build_lomn_risk_suggestion(packet, raw_text):
    diagnosis_phrase = _diagnosis_phrase(packet) or "the documented condition"
    impact = _clean_functional_impact_text(_wording_fact_text(packet, "functional_impact")) or "daily activities"
    scenarios = [
        _make_wording_scenario(
            "Sleep / daily function risk",
            "Use when pain is disrupting sleep and routine daily function.",
            (
                f"Without timely specialty evaluation and indicated treatment for {diagnosis_phrase}, the Veteran may continue to experience persistent pain, sleep disruption, and worsening limitation in {impact}."
            ),
            enabled=_has_sleep_limitation(packet),
        ),
        _make_wording_scenario(
            "Occupational risk",
            "Use when the condition is interfering with work or role function.",
            (
                f"If specialty evaluation and indicated treatment are delayed or denied for {diagnosis_phrase}, the Veteran may continue to experience persistent pain, reduced occupational capacity, and ongoing limitation in {impact}."
            ),
            enabled=_has_occupational_limitation(packet),
        ),
        _make_wording_scenario(
            "Default risk if delayed",
            "Use when a restrained risk statement is needed.",
            (
                f"Without timely specialty evaluation and indicated treatment for {diagnosis_phrase}, the Veteran remains at risk for persistent pain, worsening functional limitation, and continued impairment of daily activities."
            ),
            enabled=True,
        ),
    ]
    return _select_wording_blueprint(packet, "lomn_risk", scenarios)


def _build_seoc_diagnosis_suggestion(packet, raw_text):
    diagnosis_with_levels = _diagnosis_with_levels_phrase(packet)
    if not diagnosis_with_levels:
        return {
            "suggestion_text": "",
            "scenario_label": "",
            "scenario_use_when": "",
        }
    scenarios = [
        _make_wording_scenario(
            "Multi-level SEOC diagnosis",
            "Use when more than one lumbar level is involved.",
            f"{diagnosis_with_levels}.",
            enabled=_has_multilevel_pattern(packet),
        ),
        _make_wording_scenario(
            "Focused SEOC diagnosis",
            "Use when the diagnosis should remain tightly aligned to the documented episode.",
            f"{diagnosis_with_levels}.",
            enabled=True,
        ),
    ]
    return _select_wording_blueprint(packet, "seoc_diagnosis", scenarios)


def _build_seoc_scope_suggestion(packet, raw_text):
    services = _wording_fact_text(packet, "seoc_scope_items") or _wording_fact_text(packet, "requested_service") or "specialty evaluation, diagnostic confirmation, indicated intervention if clinically appropriate, and standard follow-up"
    diagnosis_phrase = _diagnosis_phrase(packet) or "the documented lumbar condition"
    scenarios = [
        _make_wording_scenario(
            "Diagnostic-first SEOC",
            "Use when evaluation and procedural planning must come before intervention.",
            (
                f"This SEOC is limited to specialty evaluation, diagnostic confirmation, and procedural planning related only to {diagnosis_phrase}. "
                "Any intervention should occur only if clinically indicated after the authorized evaluation."
            ),
            enabled=not _has_specific_requested_procedure(packet),
        ),
        _make_wording_scenario(
            "Procedure-included SEOC",
            "Use when the episode should expressly include the indicated intervention and follow-up.",
            (
                f"This SEOC is limited to {services} related only to {diagnosis_phrase}. "
                "No open-ended pain management, unrelated spine care, or long-term medication management is requested under this episode."
            ),
            enabled=_has_specific_requested_procedure(packet),
        ),
        _make_wording_scenario(
            "Contained-scope SEOC",
            "Use when the reviewer needs a tight scope statement.",
            (
                f"This SEOC is limited to {services} for {diagnosis_phrase}. "
                "Any additional or unrelated services will require separate evaluation and authorization."
            ),
            enabled=True,
        ),
    ]
    return _select_wording_blueprint(packet, "seoc_scope", scenarios)


def _build_seoc_duration_suggestion(packet, raw_text):
    duration = _wording_fact_text(packet, "estimated_duration_text")
    if not duration:
        return {
            "suggestion_text": "",
            "scenario_label": "",
            "scenario_use_when": "",
        }
    scenarios = [
        _make_wording_scenario(
            "Thirty-to-ninety-day duration",
            "Use when the episode is a common 30- to 90-day authorization window.",
            _ensure_sentence(f"The authorized episode is expected to conclude within {duration}"),
            enabled=_text_contains_any(duration, "30", "ninety", "90"),
        ),
        _make_wording_scenario(
            "Brief procedure-focused duration",
            "Use when the episode should read as short and procedure-focused.",
            _ensure_sentence(f"This defined episode of care is expected to be completed within {duration}"),
            enabled=True,
        ),
    ]
    return _select_wording_blueprint(packet, "seoc_duration", scenarios)


def _build_seoc_continuity_suggestion(packet, raw_text):
    scenarios = [
        _make_wording_scenario(
            "VA handoff continuity",
            "Use when continuity language should emphasize communication back to the VA referrer.",
            "Upon completion of the authorized episode, a treatment summary and clinical outcome update will be provided to the referring VA provider. Any additional or unrelated care will require separate evaluation and authorization.",
            enabled=True,
        ),
        _make_wording_scenario(
            "Contained continuity",
            "Use when the emphasis is that any further care needs new authorization.",
            "Following completion of the authorized episode, the referring VA provider will receive a treatment summary and outcome update. Additional or unrelated services will require separate review and authorization.",
            enabled=True,
        ),
    ]
    return _select_wording_blueprint(packet, "seoc_continuity", scenarios)


def _build_clinical_functional_impact_suggestion(packet, raw_text):
    impact = _wording_fact_text(packet, "functional_impact")
    if not impact:
        return {
            "suggestion_text": "",
            "scenario_label": "",
            "scenario_use_when": "",
        }
    impact = _clean_functional_impact_text(impact)
    scenarios = [
        _make_wording_scenario(
            "Sleep-focused functional impact",
            "Use when sleep disruption is prominent.",
            _ensure_sentence(f"The Veteran reports sleep disruption and reduced daily function related to limitation in {impact}"),
            enabled=_has_sleep_limitation(packet),
        ),
        _make_wording_scenario(
            "Occupational functional impact",
            "Use when work capacity is being affected.",
            _ensure_sentence(f"The Veteran reports functional limitation affecting {impact}, including reduced occupational tolerance and day-to-day activity capacity"),
            enabled=_has_occupational_limitation(packet),
        ),
        _make_wording_scenario(
            "Default functional impact",
            "Use when a general provider-style function statement is needed.",
            _ensure_sentence(f"The Veteran's symptoms interfere with {impact}"),
            enabled=True,
        ),
    ]
    return _select_wording_blueprint(packet, "clinical_functional_impact", scenarios)


def _build_clinical_assessment_suggestion(packet, raw_text):
    diagnosis_with_levels = _diagnosis_with_levels_phrase(packet) or "the documented lumbar condition"
    mri_findings = _wording_fact_text(packet, "mri_findings") or "documented imaging findings"
    conservative = _wording_fact_text(packet, "conservative_history") or "appropriate conservative management"
    scenarios = [
        _make_wording_scenario(
            "Discogenic assessment",
            "Use when the presentation is discogenic or annular in pattern.",
            (
                f"The clinical presentation is consistent with discogenic lumbar pain associated with {diagnosis_with_levels}. "
                f"Imaging demonstrates {mri_findings}, and symptoms remain persistent despite {conservative}."
            ),
            enabled=_has_annular_pattern(packet),
        ),
        _make_wording_scenario(
            "Radicular / disc displacement assessment",
            "Use when radicular features or displacement findings are present.",
            (
                f"Current findings support {diagnosis_with_levels} with persistent symptoms concerning for radicular involvement. "
                f"Imaging demonstrates {mri_findings}, and symptoms remain refractory to {conservative}."
            ),
            enabled=_has_radicular_pattern(packet) or _has_disc_displacement_pattern(packet),
        ),
        _make_wording_scenario(
            "Imaging-correlation assessment",
            "Use when tying the clinical picture directly to the MRI is most important.",
            (
                f"The overall assessment is consistent with {diagnosis_with_levels}. "
                f"Imaging findings demonstrating {mri_findings} correlate with the Veteran's pain pattern and persistent symptoms despite {conservative}."
            ),
            enabled=_has_imaging_correlation(packet),
        ),
        _make_wording_scenario(
            "Default assessment",
            "Use when a concise provider-style assessment is needed.",
            (
                f"The overall assessment is consistent with documented {diagnosis_with_levels}. "
                f"Symptoms remain persistent despite {conservative} and continue to warrant specialty evaluation."
            ),
            enabled=True,
        ),
    ]
    return _select_wording_blueprint(packet, "clinical_assessment", scenarios)


def _build_clinical_treatment_plan_suggestion(packet, raw_text):
    plan_items = _wording_fact_text(packet, "clinical_plan_items") or _wording_fact_text(packet, "requested_service")
    if not plan_items:
        return {
            "suggestion_text": "",
            "scenario_label": "",
            "scenario_use_when": "",
        }
    scenarios = [
        _make_wording_scenario(
            "Diagnostic-first treatment plan",
            "Use when the provider should confirm indication before intervention.",
            (
                f"The treatment plan is limited to diagnostic confirmation and procedural planning, followed by {plan_items} only if clinically indicated. "
                "No open-ended pain management or unrelated care is requested."
            ),
            enabled=not _has_specific_requested_procedure(packet),
        ),
        _make_wording_scenario(
            "Contained intervention plan",
            "Use when the plan should remain tightly limited to the documented condition.",
            (
                f"The requested plan of care is limited to {plan_items} related only to the documented lumbar condition. "
                "This plan does not request unrelated care, long-term medication management, or broad ongoing pain management."
            ),
            enabled=True,
        ),
    ]
    return _select_wording_blueprint(packet, "clinical_treatment_plan", scenarios)


def _build_clinical_narrative_suggestion(packet, raw_text):
    diagnosis_with_levels = _diagnosis_with_levels_phrase(packet) or "the documented lumbar condition"
    requested_service = _wording_fact_text(packet, "requested_service") or "specialty evaluation and indicated treatment planning"
    conservative = _wording_fact_text(packet, "conservative_history") or "appropriate conservative management"
    mri_date = _wording_fact_text(packet, "mri_date") or "the documented study date"
    mri_findings = _wording_fact_text(packet, "mri_findings") or "documented imaging findings"
    symptom_summary = _wording_fact_text(packet, "symptom_summary") or "persistent function-limiting lumbar symptoms"
    impact = _clean_functional_impact_text(_wording_fact_text(packet, "functional_impact")) or "daily function"
    scenarios = [
        _make_wording_scenario(
            "Discogenic narrative",
            "Use when the packet points to discogenic or annular pathology.",
            (
                f"The Veteran presents with chronic, function-limiting lumbar symptoms associated with {diagnosis_with_levels}. "
                f"Symptoms include {symptom_summary.lower()} and interfere with {impact}. MRI dated {mri_date} demonstrates {mri_findings}, which correlate with the clinical presentation. "
                f"Symptoms have persisted despite {conservative}. Based on the documented condition, imaging correlation, and failed conservative management, {requested_service} is medically appropriate if clinically indicated."
            ),
            enabled=_has_annular_pattern(packet),
        ),
        _make_wording_scenario(
            "Radicular / displacement narrative",
            "Use when displacement findings or radicular features are present.",
            (
                f"The Veteran presents with {diagnosis_with_levels} and persistent symptoms concerning for radicular involvement. "
                f"These symptoms continue to impair {impact} despite {conservative}. MRI dated {mri_date} demonstrates {mri_findings}, and the imaging findings correlate with the current clinical picture. "
                f"Specialty evaluation and {requested_service} are medically appropriate for consideration if clinically indicated."
            ),
            enabled=_has_radicular_pattern(packet) or _has_disc_displacement_pattern(packet),
        ),
        _make_wording_scenario(
            "Long-duration narrative",
            "Use when chronicity and failed conservative care are the strongest support.",
            (
                f"The Veteran has chronic, function-limiting symptoms related to {diagnosis_with_levels} that have persisted despite {conservative} over {_duration_source_text(packet) or 'the documented treatment period'}. "
                f"MRI dated {mri_date} demonstrates {mri_findings}, and those findings correlate with the clinical presentation. "
                f"Given the persistent symptoms, documented functional limitation, and failure of conservative management, {requested_service} is medically reasonable for consideration if clinically indicated."
            ),
            enabled=_duration_bucket(packet) in {"over_90_days", "over_6_months", "over_12_months"},
        ),
        _make_wording_scenario(
            "Default physician narrative",
            "Use when a broad provider-authored narrative is needed.",
            (
                f"The Veteran presents with chronic, function-limiting lumbar pain associated with {diagnosis_with_levels}. "
                f"Symptoms interfere with {impact}. MRI dated {mri_date} demonstrates {mri_findings}, which correlate with the clinical presentation. "
                f"Symptoms have persisted despite {conservative}. Based on the documented condition, imaging correlation, and failed conservative management, {requested_service} is medically appropriate if clinically indicated."
            ),
            enabled=True,
        ),
    ]
    return _select_wording_blueprint(packet, "clinical_narrative", scenarios)


def _build_va_reason_for_request_suggestion(packet, raw_text):
    diagnosis_with_levels = _diagnosis_with_levels_phrase(packet) or "the documented lumbar condition"
    requested_service = _wording_fact_text(packet, "requested_service") or "specialty interventional spine evaluation and treatment planning"
    conservative = _wording_fact_text(packet, "conservative_history") or "appropriate conservative management"
    mri_date = _wording_fact_text(packet, "mri_date") or "the documented study date"
    mri_findings = _wording_fact_text(packet, "mri_findings") or "documented imaging findings"
    duration = _duration_source_text(packet) or "the documented treatment period"
    impact = _clean_functional_impact_text(_wording_fact_text(packet, "functional_impact")) or "daily function and activity tolerance"
    scenarios = [
        _make_wording_scenario(
            "Annular / discogenic 10-10172 request",
            "Use when MRI documents annular pathology or the case reads as discogenic pain.",
            (
                f"Authorization is requested for evaluation and treatment planning for discogenic lumbar pain associated with {mri_findings} at {_wording_fact_text(packet, 'affected_levels') or 'the documented lumbar level(s)'}. "
                f"Symptoms persist despite conservative management including {conservative}. Requested services are limited to specialty consultation, diagnostic confirmation, indicated intervention if clinically appropriate, and routine follow-up."
            ),
            enabled=_has_annular_pattern(packet),
        ),
        _make_wording_scenario(
            "Radicular / herniation 10-10172 request",
            "Use when disc displacement or radicular symptoms are documented.",
            (
                f"The Veteran has documented lumbar disc pathology involving {diagnosis_with_levels}, with persistent symptoms concerning for radicular involvement despite conservative care. "
                f"Authorization is requested for specialty evaluation, diagnostic confirmation, and indicated interventional treatment planning related to the documented condition."
            ),
            enabled=_has_radicular_pattern(packet) or _has_disc_displacement_pattern(packet),
        ),
        _make_wording_scenario(
            "Imaging correlation 10-10172 request",
            "Use when MRI correlation should anchor the reason-for-request field.",
            (
                f"Authorization is requested because MRI dated {mri_date} demonstrates {mri_findings}, which correlates with the Veteran's clinical presentation and functional limitations. "
                f"Symptoms have persisted despite {conservative}. Requested care is limited to specialty evaluation, procedural planning, indicated intervention if clinically appropriate, and standard follow-up."
            ),
            enabled=_has_imaging_correlation(packet),
        ),
        _make_wording_scenario(
            "Extended conservative care 10-10172 request",
            "Use when conservative care has been prolonged without sustained relief.",
            (
                f"Authorization is requested for a defined specialty interventional spine evaluation and treatment pathway for {diagnosis_with_levels}. "
                f"The Veteran has chronic, function-limiting symptoms that persist despite {conservative} over {duration}. Requested services are limited to diagnostic confirmation, indicated intervention if clinically appropriate, and routine post-procedure follow-up."
            ),
            enabled=_duration_bucket(packet) in {"over_90_days", "over_6_months", "over_12_months"},
        ),
        _make_wording_scenario(
            "SEOC-limited 10-10172 request",
            "Use when the reason-for-request field should clearly reject open-ended care.",
            (
                f"This request is for a defined specialty evaluation and treatment pathway related only to {diagnosis_with_levels}. "
                f"It is not a request for open-ended pain management, unrelated spine care, or long-term medication management. The Veteran remains function-limited in {impact} despite {conservative}."
            ),
            enabled=True,
        ),
        _make_wording_scenario(
            "High-quality default 10-10172 request",
            "Use when a general best-default reason-for-request is needed.",
            (
                f"Authorization is requested for {requested_service} related to documented {diagnosis_with_levels}. "
                f"The Veteran has chronic, function-limiting symptoms that correlate with imaging findings and persist despite {conservative}. Requested care is limited to diagnostic confirmation, indicated intervention if clinically appropriate, and routine follow-up."
            ),
            enabled=True,
        ),
    ]
    return _select_wording_blueprint(packet, "va_reason_for_request", scenarios)


WORDING_ASSIST_SPECS = {
    "va_form_10_10172": [
        {
            "key": "va_reason_for_request",
            "label": "Reason for Request",
            "field_name": "va10172_reason_for_request",
            "required_facts": ["diagnosis", "icd_codes", "requested_service", "mri_findings", "symptom_summary", "conservative_history"],
            "builder": _build_va_reason_for_request_suggestion,
        },
    ],
    "lomn": [
        {
            "key": "lomn_clinical_basis",
            "label": "Clinical Basis",
            "field_name": "lmn_clinical_summary",
            "required_facts": ["diagnosis", "icd_codes", "mri_date", "mri_findings", "symptom_summary", "conservative_history"],
            "builder": _build_lomn_clinical_basis_suggestion,
        },
        {
            "key": "lomn_medical_necessity",
            "label": "Medical Necessity",
            "field_name": "lmn_medical_necessity_statement",
            "required_facts": ["diagnosis", "requested_service", "functional_impact", "conservative_history", "mri_findings"],
            "builder": _build_lomn_medical_necessity_suggestion,
        },
        {
            "key": "lomn_risk_if_delayed",
            "label": "Risk if Delayed / Denied",
            "field_name": "lmn_risk_statement",
            "required_facts": ["diagnosis", "functional_impact"],
            "builder": _build_lomn_risk_suggestion,
        },
    ],
    "consult": [
        {
            "key": "consult_reason",
            "label": "Reason for Consultation",
            "field_name": "consult_reason_text",
            "required_facts": ["diagnosis", "requested_service", "symptom_summary"],
            "builder": _build_consult_reason_suggestion,
        },
        {
            "key": "consult_rationale",
            "label": "Medical Rationale",
            "field_name": "consult_medical_rationale_text",
            "required_facts": ["diagnosis", "functional_impact", "conservative_history", "mri_findings"],
            "builder": _build_consult_rationale_suggestion,
        },
        {
            "key": "consult_scope",
            "label": "Scope of Request",
            "field_name": "consult_scope_exclusion_text",
            "required_facts": ["diagnosis", "consult_requested_services"],
            "builder": _build_consult_scope_suggestion,
        },
    ],
    "seoc": [
        {
            "key": "seoc_diagnosis_language",
            "label": "Diagnosis Language",
            "field_name": "episode_diagnosis",
            "required_facts": ["diagnosis", "icd_codes"],
            "builder": _build_seoc_diagnosis_suggestion,
        },
        {
            "key": "seoc_scope_language",
            "label": "Scope of Request",
            "field_name": "seoc_scope_text",
            "required_facts": ["diagnosis", "seoc_scope_items"],
            "builder": _build_seoc_scope_suggestion,
        },
        {
            "key": "seoc_duration_language",
            "label": "Duration Language",
            "field_name": "estimated_duration_text",
            "required_facts": ["estimated_duration_text"],
            "builder": _build_seoc_duration_suggestion,
        },
        {
            "key": "seoc_continuity_language",
            "label": "Continuity of Care Language",
            "field_name": "seoc_continuity_text",
            "required_facts": ["ordering_doctor", "facility"],
            "builder": _build_seoc_continuity_suggestion,
        },
    ],
    "clinical": [
        {
            "key": "clinical_functional_impact",
            "label": "Functional Impact",
            "field_name": "clinical_doc_functional_impact",
            "required_facts": ["clinical_limitations"],
            "builder": _build_clinical_functional_impact_suggestion,
        },
        {
            "key": "clinical_assessment",
            "label": "Assessment",
            "field_name": "clinical_doc_assessment_summary",
            "required_facts": ["symptom_summary", "diagnosis", "mri_findings", "conservative_history"],
            "builder": _build_clinical_assessment_suggestion,
        },
        {
            "key": "clinical_treatment_plan",
            "label": "Treatment Plan",
            "field_name": "clinical_doc_treatment_plan_intro",
            "required_facts": ["requested_service", "clinical_plan_items"],
            "builder": _build_clinical_treatment_plan_suggestion,
        },
        {
            "key": "clinical_physician_narrative",
            "label": "Physician Narrative Paragraph",
            "field_name": "clinical_doc_physician_narrative",
            "required_facts": ["symptom_summary", "diagnosis", "functional_impact", "conservative_history", "mri_findings"],
            "builder": _build_clinical_narrative_suggestion,
        },
    ],
}


def wording_assist_specs_for_profile(profile_name):
    if is_va_10172_profile(profile_name):
        return WORDING_ASSIST_SPECS["va_form_10_10172"]
    if is_lomn_profile(profile_name):
        return WORDING_ASSIST_SPECS["lomn"]
    if is_consult_request_profile(profile_name):
        return WORDING_ASSIST_SPECS["consult"]
    if is_seoc_request_profile(profile_name):
        return WORDING_ASSIST_SPECS["seoc"]
    if is_clinical_documentation_profile(profile_name):
        return WORDING_ASSIST_SPECS["clinical"]
    return []


def build_wording_risk_messages(text):
    value = sanitize_packet_builder_text(text)
    if not value:
        return ["No wording entered yet for this field."]
    risks = []
    lowered = value.lower()
    if re.search(r"\b(heal|heals|cure|cures|guarantee|guarantees|fix|fixes|eliminate|eliminates)\b", lowered):
        risks.append("Avoid overpromising words such as heal, cure, fix, eliminate, or guarantee.")
    if "will prevent" in lowered or "will reduce" in lowered:
        risks.append("Avoid absolute outcome claims such as will prevent or will reduce. Prefer intended to or if clinically indicated.")
    if re.search(r"\b(best|superior|revolutionary|cutting-edge|innovative)\b", lowered):
        risks.append("Avoid marketing or promotional language. Keep the tone clinical and documentation-based.")
    if re.search(r"\b(stuff|thing|things|really|just|bad back|fix pain)\b", lowered):
        risks.append("Tighten casual or vague wording so the request reads like provider documentation.")
    if "pain management" in lowered and "not open-ended pain management" not in lowered:
        risks.append("Clarify scope so the request does not sound like open-ended pain management.")
    return risks


def _wording_entry_fingerprint(spec, packet, raw_text, suggestion_text):
    payload = {
        "field_name": spec.get("field_name"),
        "raw_text": sanitize_packet_builder_text(raw_text),
        "suggestion_text": sanitize_packet_builder_text(suggestion_text),
        "facts": {
            fact_key: _wording_fact_text(packet, fact_key)
            for fact_key in spec.get("required_facts") or []
        },
    }
    return hashlib.sha1(json.dumps(payload, sort_keys=True, ensure_ascii=False).encode("utf-8")).hexdigest()


def build_wording_assist_entries(payload):
    packet = normalize_packet_builder_payload(payload)
    profile_name = packet.get("packet_profile") or ""
    review_state = dict(packet.get("wording_assist_state") or {})
    entries = []
    for spec in wording_assist_specs_for_profile(profile_name):
        raw_text = str(packet.get(spec.get("field_name")) or "").strip()
        missing_facts = [
            WORDING_FACT_LABELS.get(fact_key, fact_key.replace("_", " ").title())
            for fact_key in spec.get("required_facts") or []
            if not _wording_fact_text(packet, fact_key)
        ]
        suggestion_text = ""
        state = dict(review_state.get(spec["key"]) or {})
        cycle_index = 0
        try:
            cycle_index = max(0, int(state.get("cycle_index") or 0))
        except Exception:
            cycle_index = 0
        if not missing_facts:
            packet_for_builder = dict(packet)
            packet_for_builder["__wording_cycle_key"] = spec["key"]
            packet_for_builder["__wording_cycle_index"] = cycle_index
            builder_result = spec["builder"](packet_for_builder, raw_text)
            if isinstance(builder_result, dict):
                suggestion_text = sanitize_packet_builder_text(builder_result.get("suggestion_text") or "")
                scenario_label = sanitize_packet_builder_text(builder_result.get("scenario_label") or "")
                scenario_use_when = sanitize_packet_builder_text(builder_result.get("scenario_use_when") or "")
            else:
                suggestion_text = sanitize_packet_builder_text(builder_result)
                scenario_label = ""
                scenario_use_when = ""
        else:
            scenario_label = ""
            scenario_use_when = ""
        source_fingerprint = _wording_entry_fingerprint(spec, packet, raw_text, suggestion_text + "|" + scenario_label)
        decision = str(state.get("decision") or "").strip().lower()
        approved_text = sanitize_packet_builder_text(state.get("approved_text") or "")
        stale = bool(state) and str(state.get("source_fingerprint") or "") != source_fingerprint
        if missing_facts:
            status_key = "needs_facts"
            status_label = "Needs Facts"
        elif decision in WORDING_REVIEW_DECISIONS and approved_text and not stale:
            status_key = "approved"
            status_label = {
                "accepted": "Accepted Suggestion",
                "edited": "Edited Suggestion",
                "keep_original": "Kept Original",
            }.get(decision, "Approved")
        else:
            status_key = "needs_review"
            status_label = "Needs Review"
        entries.append(
            {
                "key": spec["key"],
                "label": spec["label"],
                "field_name": spec["field_name"],
                "raw_text": raw_text,
                "suggestion_text": suggestion_text,
                "missing_facts": missing_facts,
                "risk_messages": build_wording_risk_messages(raw_text),
                "status_key": status_key,
                "status_label": status_label,
                "decision": decision,
                "approved_text": approved_text,
                "source_fingerprint": source_fingerprint,
                "stale": stale,
                "cycle_index": cycle_index,
                "scenario_label": scenario_label,
                "scenario_use_when": scenario_use_when,
            }
        )
    return entries


def build_wording_assist_banner(payload):
    entries = build_wording_assist_entries(payload)
    if not entries:
        return ""
    pending = [entry for entry in entries if entry["status_key"] != "approved"]
    if not pending:
        return (
            "<div style='margin-bottom:14px; padding:10px 12px; background:#ECFDF3; border:1px solid #B8E3C9; color:#1F6F43; font-size:11pt; line-height:1.35;'>"
            "<div style='font-weight:700;'>Wording Assist Ready</div>"
            "<div style='margin-top:4px;'>High-risk wording fields on this form have been reviewed and approved.</div>"
            "</div>"
        )
    fact_blockers = [entry for entry in pending if entry["status_key"] == "needs_facts"]
    review_blockers = [entry for entry in pending if entry["status_key"] == "needs_review"]
    detail_lines = []
    if fact_blockers:
        detail_lines.append(
            f"{len(fact_blockers)} field(s) still need supporting facts before professional wording can be approved."
        )
    if review_blockers:
        detail_lines.append(
            f"{len(review_blockers)} field(s) still need wording review or approval."
        )
    return (
        "<div style='margin-bottom:14px; padding:10px 12px; background:#FFF7ED; border:1px solid #F1C996; color:#7C3400; font-size:11pt; line-height:1.35;'>"
        "<div style='font-weight:700;'>Wording Assist Pending</div>"
        "<div style='margin-top:4px;'>"
        + html.escape(" ".join(detail_lines) or "Review the Wording Assist tab before final export.")
        + "</div>"
        "</div>"
    )


def _critical_shared_field_count(packet):
    critical_shared_values = [
        str(packet.get("patient_name") or "").strip(),
        str(packet.get("date_of_birth") or "").strip(),
        str(packet.get("facility") or "").strip(),
        str(packet.get("ordering_doctor") or "").strip(),
        str(packet.get("diagnosis") or "").strip(),
        str(packet.get("authorization_number") or "").strip(),
    ]
    return sum(1 for value in critical_shared_values if value)


def should_enforce_wording_assist(payload, group_name=""):
    packet = normalize_packet_builder_payload(payload)
    return _critical_shared_field_count(packet) >= 3


def build_wording_export_blockers(payload, group_name=""):
    packet = normalize_packet_builder_payload(payload)
    if not should_enforce_wording_assist(packet, group_name=group_name):
        return []
    normalized_group = str(group_name or "").strip().lower()
    profiles = bundle_profiles_for_group(normalized_group) if normalized_group else [packet.get("packet_profile")]
    blockers = []
    for profile_name in profiles:
        if not profile_name:
            continue
        profile_payload = build_profile_export_payload(packet, profile_name) if normalized_group else normalize_packet_builder_payload(packet)
        for entry in build_wording_assist_entries(profile_payload):
            if entry["missing_facts"]:
                blockers.append(
                    f"{default_title_for_profile(profile_name) or profile_name} - {entry['label']}: missing "
                    + ", ".join(entry["missing_facts"])
                )
            elif entry["status_key"] != "approved":
                blockers.append(
                    f"{default_title_for_profile(profile_name) or profile_name} - {entry['label']}: wording review still needs approval"
                )
    return blockers


def _default_export_dir():
    home = os.path.expanduser("~")
    for candidate in ("Desktop", "Documents", "Downloads"):
        path = os.path.join(home, candidate)
        if os.path.isdir(path):
            return path
    return home


def _default_va_form_10172_template_path():
    candidates = [
        resource_path("ui/pyside_gui/assets/VA_Form_10-10172.pdf"),
        r"C:\Users\aaron\OneDrive\Desktop\TrueDisc_VA_Complete_Submission_Program_v1.0\03_Patient_Submission_Templates\VA_Form_10-10172.pdf",
    ]
    for path in candidates:
        normalized = str(path or "").strip()
        if normalized and os.path.exists(normalized):
            return normalized
    return ""


DEFAULT_DEV_TOOLS_CONFIG = {
    "legacy_gallery_paths": [],
    "packet_builder_export_dir": _default_export_dir(),
    "va_form_10172_template_path": _default_va_form_10172_template_path(),
}

PACKET_WIDGET_BINDINGS = {
    "packet_title": ("packet_title_input", "text"),
    "packet_profile": ("profile_combo", "currentText"),
    "patient_name": ("patient_name_input", "text"),
    "date_of_birth": ("dob_input", "text"),
    "authorization_number": ("auth_input", "text"),
    "va_icn": ("icn_input", "text"),
    "ordering_doctor": ("ordering_doctor_input", "text"),
    "provider": ("provider_input", "text"),
    "facility": ("facility_input", "text"),
    "community_facility": ("community_facility_input", "text"),
    "requested_service": ("requested_service_input", "text"),
    "diagnosis": ("diagnosis_input", "text"),
    "secondary_diagnosis": ("secondary_diagnosis_input", "text"),
    "icd_codes": ("icd_codes_input", "text"),
    "master_requested_cpt_code": ("master_requested_cpt_code_input", "text"),
    "master_provider_credentials": ("master_provider_credentials_input", "text"),
    "master_provider_specialty": ("master_provider_specialty_input", "text"),
    "master_provider_npi": ("master_provider_npi_input", "text"),
    "master_practice_name": ("master_practice_name_input", "text"),
    "master_provider_phone": ("master_provider_phone_input", "text"),
    "master_provider_fax": ("master_provider_fax_input", "text"),
    "master_provider_email": ("master_provider_email_input", "text"),
    "master_provider_address": ("master_provider_address_input", "toPlainText"),
    "master_mri_date": ("master_mri_date_input", "text"),
    "master_mri_findings": ("master_mri_findings_input", "text"),
    "master_affected_levels": ("master_affected_levels_input", "text"),
    "clinical_summary": ("clinical_summary_input", "toPlainText"),
    "packet_notes": ("packet_notes_input", "toPlainText"),
    "scenario_pathology_pattern": ("scenario_pathology_pattern_input", "text"),
    "scenario_conservative_duration": ("scenario_conservative_duration_input", "text"),
    "scenario_prior_esi_response": ("scenario_prior_esi_response_input", "text"),
    "scenario_functional_emphasis": ("scenario_functional_emphasis_input", "text"),
    "scenario_request_framing": ("scenario_request_framing_input", "text"),
    "scenario_symptom_pattern": ("scenario_symptom_pattern_input", "text"),
    "scenario_conservative_modalities": ("scenario_conservative_modalities_input", "text"),
    "scenario_review_concern": ("scenario_review_concern_input", "text"),
    "scenario_treatment_goals": ("scenario_treatment_goals_input", "text"),
    "referral_subtitle": ("referral_subtitle_input", "text"),
    "pcp_request_text": ("pcp_request_input", "text"),
    "referral_entry_text": ("referral_entry_input", "text"),
    "areas_of_concern": ("areas_of_concern_input", "text"),
    "group_npi": ("group_npi_input", "text"),
    "fax_number": ("fax_number_input", "text"),
    "liaison_contact_info": ("liaison_contact_input", "toPlainText"),
    "consent_form_title": ("consent_form_title_input", "text"),
    "street_address": ("street_address_input", "text"),
    "city": ("city_input", "text"),
    "state": ("state_input", "text"),
    "zip_code": ("zip_code_input", "text"),
    "home_phone": ("home_phone_input", "text"),
    "mobile_phone": ("mobile_phone_input", "text"),
    "work_phone": ("work_phone_input", "text"),
    "email_address": ("email_address_input", "text"),
    "ssn": ("ssn_input", "text"),
    "drivers_license": ("drivers_license_input", "text"),
    "drivers_license_state": ("drivers_license_state_input", "text"),
    "appointment_confirmation_method": ("appointment_confirmation_combo", "currentText"),
    "filed_for_disability": ("filed_for_disability_combo", "currentText"),
    "condition_work_related": ("condition_work_related_combo", "currentText"),
    "condition_due_to_accident": ("condition_due_to_accident_combo", "currentText"),
    "yes_response_explanation": ("yes_response_explanation_input", "toPlainText"),
    "has_attorney": ("has_attorney_combo", "currentText"),
    "attorney_name": ("attorney_name_input", "text"),
    "attorney_phone": ("attorney_phone_input", "text"),
    "emergency_contact_name": ("emergency_contact_name_input", "text"),
    "emergency_contact_relationship": ("emergency_contact_relationship_input", "text"),
    "emergency_contact_phone": ("emergency_contact_phone_input", "text"),
    "primary_insurance_carrier": ("primary_insurance_carrier_input", "text"),
    "primary_insurance_id": ("primary_insurance_id_input", "text"),
    "primary_insurance_phone": ("primary_insurance_phone_input", "text"),
    "secondary_insurance_carrier": ("secondary_insurance_carrier_input", "text"),
    "secondary_insurance_id": ("secondary_insurance_id_input", "text"),
    "secondary_insurance_phone": ("secondary_insurance_phone_input", "text"),
    "pcp_pcm_name": ("pcp_pcm_name_input", "text"),
    "pcp_pcm_phone": ("pcp_pcm_phone_input", "text"),
    "pcp_pcm_fax": ("pcp_pcm_fax_input", "text"),
    "consent_provider_name": ("consent_provider_name_input", "text"),
    "consent_initials": ("consent_initials_input", "text"),
    "minor_doctor_name": ("minor_doctor_name_input", "text"),
    "minor_consent_initials": ("minor_consent_initials_input", "text"),
    "service_authorization_name": ("service_authorization_name_input", "text"),
    "patient_signature_name": ("patient_signature_name_input", "text"),
    "patient_signature_date": ("patient_signature_date_input", "text"),
    "submission_cover_title": ("submission_cover_title_input", "text"),
    "submission_date": ("submission_date_input", "text"),
    "primary_diagnosis_code": ("primary_diagnosis_code_input", "text"),
    "included_virtual_consent_form": ("included_virtual_consent_checkbox", "isChecked"),
    "included_va_form_10_10172": ("included_va_form_checkbox", "isChecked"),
    "included_seoc_request": ("included_seoc_checkbox", "isChecked"),
    "included_consult_request": ("included_consult_checkbox", "isChecked"),
    "included_lomn": ("included_lomn_checkbox", "isChecked"),
    "included_clinical_notes": ("included_clinical_notes_checkbox", "isChecked"),
    "included_mri_report": ("included_mri_checkbox", "isChecked"),
    "submitting_office": ("submitting_office_input", "text"),
    "office_staff_name": ("office_staff_name_input", "text"),
    "office_staff_signature": ("office_staff_signature_input", "text"),
    "date_reviewed": ("date_reviewed_input", "text"),
    "seoc_request_date": ("seoc_request_date_input", "text"),
    "va_medical_center_name": ("va_medical_center_input", "text"),
    "last_four_ssn": ("last_four_ssn_input", "text"),
    "episode_diagnosis": ("episode_diagnosis_input", "text"),
    "episode_icd_code": ("episode_icd_code_input", "text"),
    "seoc_scope_text": ("seoc_scope_text_input", "toPlainText"),
    "estimated_duration_text": ("estimated_duration_input", "text"),
    "clinical_objectives": ("clinical_objectives_input", "toPlainText"),
    "seoc_continuity_text": ("seoc_continuity_text_input", "toPlainText"),
    "provider_credentials": ("provider_credentials_input", "text"),
    "provider_specialty": ("provider_specialty_input", "text"),
    "provider_npi": ("provider_npi_input", "text"),
    "practice_name": ("practice_name_input", "text"),
    "provider_phone": ("provider_phone_input", "text"),
    "provider_fax": ("provider_fax_input", "text"),
    "seoc_include_preprocedure_eval": ("seoc_preprocedure_checkbox", "isChecked"),
    "seoc_include_annulargram": ("seoc_annulargram_checkbox", "isChecked"),
    "seoc_include_fibrin_injection": ("seoc_fibrin_checkbox", "isChecked"),
    "seoc_include_follow_up": ("seoc_follow_up_checkbox", "isChecked"),
    "lmn_request_date": ("lomn_request_date_input", "text"),
    "lmn_va_claim_number": ("lomn_va_claim_number_input", "text"),
    "lmn_primary_diagnosis": ("lomn_primary_diagnosis_input", "text"),
    "lmn_secondary_diagnosis": ("lomn_secondary_diagnosis_input", "text"),
    "lmn_clinical_summary": ("lomn_clinical_summary_input", "toPlainText"),
    "lmn_mri_date": ("lomn_mri_date_input", "text"),
    "lmn_mri_findings": ("lomn_mri_findings_input", "text"),
    "lmn_conservative_duration": ("lomn_conservative_duration_input", "text"),
    "lmn_include_physical_therapy": ("lomn_physical_therapy_checkbox", "isChecked"),
    "lmn_include_nsaids": ("lomn_nsaids_checkbox", "isChecked"),
    "lmn_include_activity_modification": ("lomn_activity_modification_checkbox", "isChecked"),
    "lmn_include_home_exercise": ("lomn_home_exercise_checkbox", "isChecked"),
    "lmn_include_epidural_steroid_injections": ("lomn_esi_checkbox", "isChecked"),
    "lmn_medical_necessity_statement": ("lomn_medical_necessity_input", "toPlainText"),
    "lmn_indication_reduce_pain": ("lomn_reduce_pain_checkbox", "isChecked"),
    "lmn_indication_improve_function": ("lomn_improve_function_checkbox", "isChecked"),
    "lmn_indication_prevent_degeneration": ("lomn_prevent_degeneration_checkbox", "isChecked"),
    "lmn_indication_reduce_opioid_reliance": ("lomn_reduce_opioids_checkbox", "isChecked"),
    "lmn_indication_prevent_surgery": ("lomn_prevent_surgery_checkbox", "isChecked"),
    "lmn_risk_statement": ("lomn_risk_statement_input", "toPlainText"),
    "lmn_reasonable_necessary_statement": ("lomn_reasonable_statement_input", "toPlainText"),
    "lmn_contact_statement": ("lomn_contact_statement_input", "toPlainText"),
    "consult_request_date": ("consult_request_date_input", "text"),
    "consult_va_claim_number": ("consult_va_claim_number_input", "text"),
    "consult_referring_va_provider": ("consult_referring_va_provider_input", "text"),
    "consult_reason_text": ("consult_reason_input", "toPlainText"),
    "consult_primary_diagnosis": ("consult_primary_diagnosis_input", "text"),
    "consult_secondary_diagnosis": ("consult_secondary_diagnosis_input", "text"),
    "consult_symptom_axial_pain": ("consult_symptom_axial_pain_checkbox", "isChecked"),
    "consult_symptom_activity_exacerbation": ("consult_symptom_activity_checkbox", "isChecked"),
    "consult_symptom_reduced_tolerance": ("consult_symptom_tolerance_checkbox", "isChecked"),
    "consult_symptom_functional_impairment": ("consult_symptom_function_checkbox", "isChecked"),
    "consult_mri_date": ("consult_mri_date_input", "text"),
    "consult_mri_findings": ("consult_mri_findings_input", "text"),
    "consult_conservative_duration": ("consult_conservative_duration_input", "text"),
    "consult_include_physical_therapy": ("consult_physical_therapy_checkbox", "isChecked"),
    "consult_include_nsaids": ("consult_nsaids_checkbox", "isChecked"),
    "consult_include_activity_modification": ("consult_activity_mod_checkbox", "isChecked"),
    "consult_include_home_exercise": ("consult_home_exercise_checkbox", "isChecked"),
    "consult_include_interventional_history": ("consult_interventional_history_checkbox", "isChecked"),
    "consult_include_pain_management_consultation": ("consult_include_pm_consult_checkbox", "isChecked"),
    "consult_include_procedural_planning": ("consult_include_planning_checkbox", "isChecked"),
    "consult_include_annulargram": ("consult_include_annulargram_checkbox", "isChecked"),
    "consult_include_fibrin_injection": ("consult_include_fibrin_checkbox", "isChecked"),
    "consult_include_follow_up": ("consult_include_follow_up_checkbox", "isChecked"),
    "consult_fibrin_levels": ("consult_fibrin_levels_input", "text"),
    "consult_scope_exclusion_text": ("consult_scope_exclusion_input", "toPlainText"),
    "consult_medical_rationale_text": ("consult_medical_rationale_input", "toPlainText"),
    "consult_goal_pain_reduction": ("consult_goal_pain_checkbox", "isChecked"),
    "consult_goal_functional_improvement": ("consult_goal_function_checkbox", "isChecked"),
    "consult_goal_reduce_analgesics": ("consult_goal_analgesics_checkbox", "isChecked"),
    "consult_goal_prevent_surgery": ("consult_goal_surgery_checkbox", "isChecked"),
    "consult_risk_without_treatment": ("consult_risk_without_treatment_input", "toPlainText"),
    "consult_duration_scope_text": ("consult_duration_scope_input", "toPlainText"),
    "consult_contact_statement": ("consult_contact_statement_input", "toPlainText"),
    "provider_address": ("provider_address_input", "text"),
    "provider_email": ("provider_email_input", "text"),
    "clinical_doc_title": ("clinical_doc_title_input", "text"),
    "clinical_doc_chief_complaint": ("clinical_doc_chief_complaint_input", "text"),
    "clinical_doc_duration_gt_3m": ("clinical_doc_duration_3m_checkbox", "isChecked"),
    "clinical_doc_duration_gt_6m": ("clinical_doc_duration_6m_checkbox", "isChecked"),
    "clinical_doc_duration_gt_12m": ("clinical_doc_duration_12m_checkbox", "isChecked"),
    "clinical_doc_exact_duration": ("clinical_doc_exact_duration_input", "text"),
    "clinical_doc_pain_axial": ("clinical_doc_pain_axial_checkbox", "isChecked"),
    "clinical_doc_pain_discogenic": ("clinical_doc_pain_discogenic_checkbox", "isChecked"),
    "clinical_doc_pain_activity_exacerbation": ("clinical_doc_pain_activity_checkbox", "isChecked"),
    "clinical_doc_pain_sitting_intolerance": ("clinical_doc_pain_sitting_checkbox", "isChecked"),
    "clinical_doc_pain_standing_intolerance": ("clinical_doc_pain_standing_checkbox", "isChecked"),
    "clinical_doc_pain_bending_lifting": ("clinical_doc_pain_bending_checkbox", "isChecked"),
    "clinical_doc_pain_severity": ("clinical_doc_pain_severity_input", "text"),
    "clinical_doc_limit_occupational": ("clinical_doc_limit_occupational_checkbox", "isChecked"),
    "clinical_doc_limit_prolonged_sitting": ("clinical_doc_limit_sitting_checkbox", "isChecked"),
    "clinical_doc_limit_prolonged_standing": ("clinical_doc_limit_standing_checkbox", "isChecked"),
    "clinical_doc_limit_ambulation": ("clinical_doc_limit_ambulation_checkbox", "isChecked"),
    "clinical_doc_limit_household": ("clinical_doc_limit_household_checkbox", "isChecked"),
    "clinical_doc_limit_sleep": ("clinical_doc_limit_sleep_checkbox", "isChecked"),
    "clinical_doc_functional_impact": ("clinical_doc_functional_impact_input", "toPlainText"),
    "clinical_doc_conservative_pt": ("clinical_doc_conservative_pt_checkbox", "isChecked"),
    "clinical_doc_conservative_home_exercise": ("clinical_doc_conservative_home_exercise_checkbox", "isChecked"),
    "clinical_doc_conservative_nsaids": ("clinical_doc_conservative_nsaids_checkbox", "isChecked"),
    "clinical_doc_conservative_non_opioid": ("clinical_doc_conservative_non_opioid_checkbox", "isChecked"),
    "clinical_doc_conservative_activity_modification": ("clinical_doc_conservative_activity_checkbox", "isChecked"),
    "clinical_doc_conservative_esi": ("clinical_doc_conservative_esi_checkbox", "isChecked"),
    "clinical_doc_conservative_other_interventional": ("clinical_doc_conservative_other_checkbox", "isChecked"),
    "clinical_doc_conservative_duration": ("clinical_doc_conservative_duration_input", "text"),
    "clinical_doc_esi_response": ("clinical_doc_esi_response_combo", "currentText"),
    "clinical_doc_mri_date": ("clinical_doc_mri_date_input", "text"),
    "clinical_doc_imaging_annular_tear": ("clinical_doc_imaging_annular_checkbox", "isChecked"),
    "clinical_doc_imaging_disc_degeneration": ("clinical_doc_imaging_degeneration_checkbox", "isChecked"),
    "clinical_doc_imaging_disc_protrusion": ("clinical_doc_imaging_protrusion_checkbox", "isChecked"),
    "clinical_doc_imaging_disc_displacement": ("clinical_doc_imaging_displacement_checkbox", "isChecked"),
    "clinical_doc_affected_levels": ("clinical_doc_affected_levels_input", "text"),
    "clinical_doc_primary_diagnosis": ("clinical_doc_primary_diagnosis_input", "text"),
    "clinical_doc_secondary_diagnosis": ("clinical_doc_secondary_diagnosis_input", "text"),
    "clinical_doc_assessment_summary": ("clinical_doc_assessment_summary_input", "toPlainText"),
    "clinical_doc_treatment_plan_intro": ("clinical_doc_treatment_plan_intro_input", "toPlainText"),
    "clinical_doc_plan_diagnostic_confirmation": ("clinical_doc_plan_diagnostic_checkbox", "isChecked"),
    "clinical_doc_plan_intradiscal_intervention": ("clinical_doc_plan_intervention_checkbox", "isChecked"),
    "clinical_doc_plan_follow_up": ("clinical_doc_plan_follow_up_checkbox", "isChecked"),
    "clinical_doc_plan_exclusion": ("clinical_doc_plan_exclusion_input", "toPlainText"),
    "clinical_doc_physician_narrative": ("clinical_doc_physician_narrative_input", "toPlainText"),
    "va10172_va_facility_address": ("va10172_facility_address_input", "toPlainText"),
    "va10172_ordering_provider_office_address": ("va10172_provider_office_address_input", "toPlainText"),
    "va10172_is_ihs_provider": ("va10172_is_ihs_combo", "currentText"),
    "va10172_ordering_provider_phone": ("va10172_provider_phone_input", "text"),
    "va10172_ordering_provider_fax": ("va10172_provider_fax_input", "text"),
    "va10172_ordering_provider_secure_email": ("va10172_provider_email_input", "text"),
    "va10172_care_needed_within_48_hours": ("va10172_48h_combo", "currentText"),
    "va10172_is_continuation_of_care": ("va10172_continuation_combo", "currentText"),
    "va10172_referral_to_specialty": ("va10172_referral_specialty_combo", "currentText"),
    "va10172_referral_specialty_text": ("va10172_referral_specialty_text_input", "text"),
    "va10172_diagnosis_description": ("va10172_diagnosis_description_input", "toPlainText"),
    "va10172_requested_cpt_hcpcs_code": ("va10172_requested_cpt_input", "text"),
    "va10172_description_cpt_hcpcs_code": ("va10172_requested_cpt_description_input", "text"),
    "va10172_geriatric_care_option": ("va10172_geriatric_option_combo", "currentText"),
    "va10172_reason_for_request": ("va10172_reason_for_request_input", "toPlainText"),
    "va10172_ordering_provider_name_printed": ("va10172_provider_name_printed_input", "text"),
    "va10172_ordering_provider_npi": ("va10172_provider_npi_input", "text"),
    "va10172_signature_text": ("va10172_signature_text_input", "text"),
    "va10172_today_date": ("va10172_today_date_input", "text"),
}


def _deep_merge(base, updates):
    merged = dict(base or {})
    for key, value in dict(updates or {}).items():
        if isinstance(value, dict) and isinstance(merged.get(key), dict):
            merged[key] = _deep_merge(merged.get(key) or {}, value)
        else:
            merged[key] = value
    return merged


def normalize_dev_tools_config(data=None):
    payload = _deep_merge(DEFAULT_DEV_TOOLS_CONFIG, data or {})
    paths = []
    for path in payload.get("legacy_gallery_paths") or []:
        normalized = str(path or "").strip()
        if normalized and normalized not in paths:
            paths.append(normalized)
    payload["legacy_gallery_paths"] = paths
    export_dir = str(payload.get("packet_builder_export_dir") or "").strip() or _default_export_dir()
    payload["packet_builder_export_dir"] = export_dir
    template_path = str(payload.get("va_form_10172_template_path") or "").strip()
    if not template_path:
        template_path = _default_va_form_10172_template_path()
    payload["va_form_10172_template_path"] = template_path
    return payload


def load_dev_tools_config():
    os.makedirs(os.path.dirname(DEV_TOOLS_CONFIG_PATH), exist_ok=True)
    if not os.path.exists(DEV_TOOLS_CONFIG_PATH):
        save_dev_tools_config({})
    try:
        with open(DEV_TOOLS_CONFIG_PATH, "r", encoding="utf-8") as handle:
            return normalize_dev_tools_config(json.load(handle))
    except Exception:
        return normalize_dev_tools_config({})


def save_dev_tools_config(data):
    payload = normalize_dev_tools_config(data)
    os.makedirs(os.path.dirname(DEV_TOOLS_CONFIG_PATH), exist_ok=True)
    with open(DEV_TOOLS_CONFIG_PATH, "w", encoding="utf-8") as handle:
        json.dump(payload, handle, indent=4)
    return payload


def update_dev_tools_config(changes):
    return save_dev_tools_config(_deep_merge(load_dev_tools_config(), changes or {}))


def sanitize_builder_filename(value):
    text = str(value or "").strip()
    if not text:
        return "truecore_packet"
    safe = []
    for char in text:
        if char.isalnum() or char in {"-", "_", " "}:
            safe.append(char)
    compact = "".join(safe).strip().replace(" ", "_")
    return compact or "truecore_packet"


def get_packet_export_group(profile_name):
    normalized = str(profile_name or "").strip()
    if normalized == REFERRAL_REQUEST_PROFILE:
        return "referral_request"
    if normalized in PATIENT_PACKET_PROFILE_ORDER:
        return "patient_packet"
    return "standalone"


def get_patient_packet_position(profile_name):
    normalized = str(profile_name or "").strip()
    if normalized not in PATIENT_PACKET_PROFILE_ORDER:
        return None
    return PATIENT_PACKET_PROFILE_ORDER.index(normalized) + 1


def describe_packet_export_context(profile_name):
    group = get_packet_export_group(profile_name)
    if group == "referral_request":
        return "Standalone export: Community Care Referral Request (pre-packet document)."
    if group == "patient_packet":
        position = get_patient_packet_position(profile_name)
        total = len(PATIENT_PACKET_PROFILE_ORDER)
        return f"Patient packet component: item {position} of {total}. Exports as part of the ordered patient packet bundle."
    return "Standalone builder profile."


def bundle_profiles_for_group(group_name):
    normalized = str(group_name or "").strip().lower()
    if normalized == "referral_request":
        return [REFERRAL_REQUEST_PROFILE]
    if normalized == "patient_packet":
        return list(PATIENT_PACKET_PROFILE_ORDER)
    return []


def bundle_folder_name(base_filename, group_name):
    safe_base = sanitize_builder_filename(base_filename)
    normalized = str(group_name or "").strip().lower()
    suffix = "packet_bundle"
    if normalized == "referral_request":
        suffix = "referral_request"
    elif normalized == "patient_packet":
        suffix = "patient_packet"
    return f"{safe_base}_{suffix}"


def compiled_packet_filename(base_filename, group_name):
    safe_base = sanitize_builder_filename(base_filename)
    normalized = str(group_name or "").strip().lower()
    if normalized == "patient_packet":
        return f"{safe_base}_patient_packet.pdf"
    if normalized == "referral_request":
        return f"{safe_base}_referral_request.pdf"
    return f"{safe_base}_packet.pdf"


def _is_legacy_image_path(path):
    suffix = os.path.splitext(str(path or "").strip())[1].lower()
    return suffix in LEGACY_IMAGE_SUFFIXES


def _find_reference_asset(root_path, file_names):
    normalized_root = str(root_path or "").strip()
    if not normalized_root or not os.path.isdir(normalized_root):
        return ""
    wanted = {str(name or "").lower() for name in file_names or []}
    for current_root, _, files in os.walk(normalized_root):
        lower_map = {name.lower(): os.path.join(current_root, name) for name in files}
        for candidate in wanted:
            if candidate in lower_map:
                return lower_map[candidate]
    return ""


def discover_legacy_reference_entries(reference_path):
    normalized = str(reference_path or "").strip()
    if not normalized or not os.path.exists(normalized):
        return []

    if os.path.isfile(normalized) and _is_legacy_image_path(normalized):
        return [
            {
                "entry_id": f"image::{normalized}",
                "label": os.path.basename(normalized),
                "kind": "image",
                "path": normalized,
                "source_path": normalized,
                "builtin": False,
            }
        ]

    root_path = normalized if os.path.isdir(normalized) else os.path.dirname(normalized)
    display_name = os.path.basename(normalized.rstrip("\\/")) or os.path.basename(root_path.rstrip("\\/")) or "Legacy Reference"
    launcher_code = _find_reference_asset(root_path, ["launcher_window.py", "launcher.py"])
    gui_code = _find_reference_asset(root_path, ["main_window.py"])
    launcher_background = _find_reference_asset(
        root_path,
        [
            "launcher_background.png",
            "launcher_background.jpg",
            "launcher_background.jpeg",
            "background.png",
            "background.jpg",
        ],
    )
    gui_background = _find_reference_asset(
        root_path,
        [
            "launcher_background.png",
            "gui_background.png",
            "background.png",
            "background.jpg",
        ],
    )
    logo_path = _find_reference_asset(
        root_path,
        ["truecore_logo.png", "logo.png", "truecore_logo.jpg", "truecore_logo.ico", "logo.ico"],
    )

    entries = []
    if launcher_code or "launcher" in root_path.lower():
        entries.append(
            {
                "entry_id": f"launcher::{normalized}",
                "label": f"Launcher Preview - {display_name}",
                "kind": "launcher_preview",
                "path": root_path,
                "source_path": normalized,
                "builtin": False,
                "background_path": launcher_background,
                "logo_path": logo_path,
                "code_path": launcher_code or normalized,
            }
        )
    if gui_code or "gui" in root_path.lower():
        entries.append(
            {
                "entry_id": f"gui::{normalized}",
                "label": f"GUI Preview - {display_name}",
                "kind": "gui_preview",
                "path": root_path,
                "source_path": normalized,
                "builtin": False,
                "background_path": gui_background or launcher_background,
                "logo_path": logo_path,
                "code_path": gui_code or normalized,
            }
        )

    if entries:
        return entries

    if os.path.isfile(normalized):
        return [
            {
                "entry_id": f"file::{normalized}",
                "label": f"Legacy Source - {display_name}",
                "kind": "source_reference",
                "path": normalized,
                "source_path": normalized,
                "builtin": False,
            }
        ]
    return []


def bundle_member_filename(profile_name, group_name):
    normalized_group = str(group_name or "").strip().lower()
    title = default_title_for_profile(profile_name) or str(profile_name or "packet_document").strip() or "packet_document"
    safe_title = sanitize_builder_filename(title)
    if normalized_group == "patient_packet":
        position = get_patient_packet_position(profile_name)
        if position is not None:
            return f"{position:02d}_{safe_title}"
    return safe_title


def normalize_bundle_export_format(format_name):
    normalized = str(format_name or "").strip().lower()
    if normalized in BUNDLE_EXPORT_EXTENSIONS:
        return normalized
    return "both"


def bundle_extensions_for_format(format_name, include_real_va_pdf=False):
    normalized = normalize_bundle_export_format(format_name)
    extensions = list(BUNDLE_EXPORT_EXTENSIONS.get(normalized) or BUNDLE_EXPORT_EXTENSIONS["both"])
    if include_real_va_pdf and "pdf" not in extensions:
        extensions.append("pdf")
    return extensions


def build_profile_export_payload(base_payload, profile_name):
    payload = normalize_packet_builder_payload(base_payload)
    normalized_profile = str(profile_name or "").strip()
    payload["packet_profile"] = normalized_profile
    payload["packet_title"] = default_title_for_profile(normalized_profile) or normalized_profile or payload.get("packet_title") or ""
    if is_virtual_consent_profile(normalized_profile):
        payload["consent_form_title"] = payload.get("consent_form_title") or default_title_for_profile(normalized_profile)
    elif is_submission_cover_profile(normalized_profile):
        payload["submission_cover_title"] = payload.get("submission_cover_title") or default_title_for_profile(normalized_profile)
    elif is_clinical_documentation_profile(normalized_profile):
        payload["clinical_doc_title"] = payload.get("clinical_doc_title") or default_title_for_profile(normalized_profile)
    return apply_packet_builder_shared_field_sync(payload)


SHARED_PACKET_HEADER_FIELDS = [
    ("patient_name", "Veteran Name"),
    ("date_of_birth", "Date of Birth"),
    ("authorization_number", "VA Authorization"),
    ("va_icn", "VA ICN"),
    ("ordering_doctor", "VA Referring / Ordering Provider"),
    ("facility", "VA Referring Facility"),
    ("diagnosis", "Primary Diagnosis"),
    ("icd_codes", "ICD-10 Codes"),
    ("master_provider_npi", "Community Care Treating Provider NPI"),
    ("master_practice_name", "Community Care Practice Name"),
]


def _any_present(packet, field_names):
    return any(_has_meaningful_value(packet.get(field_name)) for field_name in field_names or [])


def _any_checked(packet, field_names):
    return any(bool(packet.get(field_name)) for field_name in field_names or [])


def _normalized_compare_value(value):
    return re.sub(r"[^a-z0-9]+", "", str(value or "").strip().lower())


def _last_four_digits(value):
    digits = "".join(ch for ch in str(value or "") if ch.isdigit())
    return digits[-4:] if len(digits) >= 4 else digits


def _values_conflict(left_value, right_value, mode="text", allow_contains=False):
    left_raw = str(left_value or "").strip()
    right_raw = str(right_value or "").strip()
    if not left_raw or not right_raw:
        return False
    if mode == "icd_first":
        left_raw = _first_icd_code(left_raw)
        right_raw = _first_icd_code(right_raw)
    elif mode == "last_four":
        left_raw = _last_four_digits(left_raw)
        right_raw = _last_four_digits(right_raw)
    left_norm = _normalized_compare_value(left_raw)
    right_norm = _normalized_compare_value(right_raw)
    if not left_norm or not right_norm:
        return False
    if allow_contains and (left_norm in right_norm or right_norm in left_norm):
        return False
    return left_norm != right_norm


def build_profile_current_form_checks(profile_name, packet):
    if is_referral_request_profile(profile_name):
        return [
            ("Packet Title", _has_meaningful_value(packet.get("packet_title"))),
            ("Subtitle", _has_meaningful_value(packet.get("referral_subtitle"))),
            ("PCP Request Wording", _has_meaningful_value(packet.get("pcp_request_text"))),
            ("Referral Entry Wording", _has_meaningful_value(packet.get("referral_entry_text"))),
            ("Areas of Concern", _has_meaningful_value(packet.get("areas_of_concern"))),
            ("Group NPI", _has_meaningful_value(packet.get("group_npi"))),
            ("Fax Number", _has_meaningful_value(packet.get("fax_number"))),
            ("Liaison Contact", _has_meaningful_value(packet.get("liaison_contact_info"))),
        ]
    if is_submission_cover_profile(profile_name):
        return [
            ("Packet Title", _has_meaningful_value(packet.get("packet_title"))),
            ("Submission Date", _has_meaningful_value(packet.get("submission_date"))),
            ("Primary Diagnosis Code", _has_meaningful_value(packet.get("primary_diagnosis_code"))),
            ("Submitting Office", _has_meaningful_value(packet.get("submitting_office"))),
            ("Office Staff Name", _has_meaningful_value(packet.get("office_staff_name"))),
            ("Program User / Office Staff Signature", _has_meaningful_value(packet.get("office_staff_signature")) or _has_meaningful_value(packet.get("office_staff_signature_image"))),
            ("Date Reviewed", _has_meaningful_value(packet.get("date_reviewed"))),
            ("Document Checklist", _any_checked(packet, [
                "included_virtual_consent_form",
                "included_va_form_10_10172",
                "included_seoc_request",
                "included_consult_request",
                "included_lomn",
                "included_clinical_notes",
                "included_mri_report",
            ])),
        ]
    if is_virtual_consent_profile(profile_name):
        return [
            ("Packet Title", _has_meaningful_value(packet.get("packet_title"))),
            ("Street Address", _has_meaningful_value(packet.get("street_address"))),
            ("City", _has_meaningful_value(packet.get("city"))),
            ("State", _has_meaningful_value(packet.get("state"))),
            ("ZIP Code", _has_meaningful_value(packet.get("zip_code"))),
            ("Phone", _any_present(packet, ["mobile_phone", "home_phone", "work_phone"])),
            ("Email Address", _has_meaningful_value(packet.get("email_address"))),
            ("SSN", _has_meaningful_value(packet.get("ssn"))),
            ("Driver License", _has_meaningful_value(packet.get("drivers_license"))),
            ("Driver License State", _has_meaningful_value(packet.get("drivers_license_state"))),
            ("Emergency Contact Name", _has_meaningful_value(packet.get("emergency_contact_name"))),
            ("Emergency Contact Relationship", _has_meaningful_value(packet.get("emergency_contact_relationship"))),
            ("Emergency Contact Phone", _has_meaningful_value(packet.get("emergency_contact_phone"))),
            ("PCP / PCM Name", _has_meaningful_value(packet.get("pcp_pcm_name"))),
            ("PCP / PCM Phone", _has_meaningful_value(packet.get("pcp_pcm_phone"))),
            ("PCP / PCM Fax", _has_meaningful_value(packet.get("pcp_pcm_fax"))),
            ("Consent Provider Name", _has_meaningful_value(packet.get("consent_provider_name"))),
            ("Consent Initials", _has_meaningful_value(packet.get("consent_initials"))),
            ("Service Authorization Name", _has_meaningful_value(packet.get("service_authorization_name"))),
            ("Veteran / Patient Signature", _has_meaningful_value(packet.get("patient_signature_name")) or _has_meaningful_value(packet.get("patient_signature_image"))),
            ("Patient Signature Date", _has_meaningful_value(packet.get("patient_signature_date"))),
        ]
    if is_va_10172_profile(profile_name):
        return [
            ("VA Facility Address", _has_meaningful_value(packet.get("va10172_va_facility_address"))),
            ("Provider Office Address", _has_meaningful_value(packet.get("va10172_ordering_provider_office_address"))),
            ("Provider Phone", _has_meaningful_value(packet.get("va10172_ordering_provider_phone"))),
            ("Provider Fax", _has_meaningful_value(packet.get("va10172_ordering_provider_fax"))),
            ("Provider Secure Email", _has_meaningful_value(packet.get("va10172_ordering_provider_secure_email"))),
            ("Specialty Text", _has_meaningful_value(packet.get("va10172_referral_specialty_text"))),
            ("Diagnosis Description", _has_meaningful_value(packet.get("va10172_diagnosis_description"))),
            ("Requested CPT / HCPCS", _has_meaningful_value(packet.get("va10172_requested_cpt_hcpcs_code"))),
            ("Requested Service Description", _has_meaningful_value(packet.get("va10172_description_cpt_hcpcs_code"))),
            ("Reason For Request", _has_meaningful_value(packet.get("va10172_reason_for_request"))),
            ("Printed CCN Ordering Provider Name", _has_meaningful_value(packet.get("va10172_ordering_provider_name_printed"))),
            ("CCN Ordering Provider NPI", _has_meaningful_value(packet.get("va10172_ordering_provider_npi"))),
            ("CCN Ordering Provider Signature", _has_meaningful_value(packet.get("va10172_signature_text")) or _has_meaningful_value(packet.get("va10172_signature_image"))),
            ("Today's Date", _has_meaningful_value(packet.get("va10172_today_date"))),
        ]
    if is_seoc_request_profile(profile_name):
        return [
            ("Request Date", _has_meaningful_value(packet.get("seoc_request_date"))),
            ("VA Medical Center", _has_meaningful_value(packet.get("va_medical_center_name"))),
            ("Last Four SSN", _has_meaningful_value(packet.get("last_four_ssn"))),
            ("Episode Diagnosis", _has_meaningful_value(packet.get("episode_diagnosis"))),
            ("Episode ICD Code", _has_meaningful_value(packet.get("episode_icd_code"))),
            ("Scope Language", _has_meaningful_value(packet.get("seoc_scope_text"))),
            ("Estimated Duration", _has_meaningful_value(packet.get("estimated_duration_text"))),
            ("Clinical Objectives", _has_meaningful_value(packet.get("clinical_objectives"))),
            ("Continuity of Care", _has_meaningful_value(packet.get("seoc_continuity_text"))),
            ("Community Care Provider Credentials", _has_meaningful_value(packet.get("provider_credentials"))),
            ("Community Care Provider Specialty", _has_meaningful_value(packet.get("provider_specialty"))),
            ("Community Care Provider NPI", _has_meaningful_value(packet.get("provider_npi"))),
            ("Community Care Practice Name", _has_meaningful_value(packet.get("practice_name"))),
            ("Community Care Practice Phone", _has_meaningful_value(packet.get("provider_phone"))),
            ("Community Care Practice Fax", _has_meaningful_value(packet.get("provider_fax"))),
            ("Scope Items", _any_checked(packet, [
                "seoc_include_preprocedure_eval",
                "seoc_include_annulargram",
                "seoc_include_fibrin_injection",
                "seoc_include_follow_up",
            ])),
        ]
    if is_consult_request_profile(profile_name):
        return [
            ("Request Date", _has_meaningful_value(packet.get("consult_request_date"))),
            ("VA Claim Number", _has_meaningful_value(packet.get("consult_va_claim_number"))),
            ("Referring VA Provider", _has_meaningful_value(packet.get("consult_referring_va_provider"))),
            ("Reason For Consultation", _has_meaningful_value(packet.get("consult_reason_text"))),
            ("Primary Diagnosis", _has_meaningful_value(packet.get("consult_primary_diagnosis"))),
            ("Secondary Diagnosis", _has_meaningful_value(packet.get("consult_secondary_diagnosis"))),
            ("MRI Date", _has_meaningful_value(packet.get("consult_mri_date"))),
            ("MRI Findings", _has_meaningful_value(packet.get("consult_mri_findings"))),
            ("Conservative Duration", _has_meaningful_value(packet.get("consult_conservative_duration"))),
            ("Scope Statement", _has_meaningful_value(packet.get("consult_scope_exclusion_text"))),
            ("Medical Rationale", _has_meaningful_value(packet.get("consult_medical_rationale_text"))),
            ("Risk Without Treatment", _has_meaningful_value(packet.get("consult_risk_without_treatment"))),
            ("Duration / Scope", _has_meaningful_value(packet.get("consult_duration_scope_text"))),
            ("Community Care Provider Contact Statement", _has_meaningful_value(packet.get("consult_contact_statement"))),
            ("Community Care Provider Credentials", _has_meaningful_value(packet.get("provider_credentials"))),
            ("Community Care Provider Specialty", _has_meaningful_value(packet.get("provider_specialty"))),
            ("Community Care Provider NPI", _has_meaningful_value(packet.get("provider_npi"))),
            ("Community Care Practice Name", _has_meaningful_value(packet.get("practice_name"))),
            ("Community Care Practice Address", _has_meaningful_value(packet.get("provider_address"))),
            ("Community Care Practice Phone", _has_meaningful_value(packet.get("provider_phone"))),
            ("Community Care Practice Fax", _has_meaningful_value(packet.get("provider_fax"))),
            ("Community Care Secure Email", _has_meaningful_value(packet.get("provider_email"))),
            ("Symptom Set", _any_checked(packet, [
                "consult_symptom_axial_pain",
                "consult_symptom_activity_exacerbation",
                "consult_symptom_reduced_tolerance",
                "consult_symptom_functional_impairment",
            ])),
            ("Conservative History", _any_checked(packet, [
                "consult_include_physical_therapy",
                "consult_include_nsaids",
                "consult_include_activity_modification",
                "consult_include_home_exercise",
                "consult_include_interventional_history",
            ])),
            ("Requested Services", _any_checked(packet, [
                "consult_include_pain_management_consultation",
                "consult_include_procedural_planning",
                "consult_include_annulargram",
                "consult_include_fibrin_injection",
                "consult_include_follow_up",
            ])),
            ("Clinical Goals", _any_checked(packet, [
                "consult_goal_pain_reduction",
                "consult_goal_functional_improvement",
                "consult_goal_reduce_analgesics",
                "consult_goal_prevent_surgery",
            ])),
        ]
    if is_lomn_profile(profile_name):
        return [
            ("Request Date", _has_meaningful_value(packet.get("lmn_request_date"))),
            ("VA Claim Number", _has_meaningful_value(packet.get("lmn_va_claim_number"))),
            ("Primary Diagnosis", _has_meaningful_value(packet.get("lmn_primary_diagnosis"))),
            ("Secondary Diagnosis", _has_meaningful_value(packet.get("lmn_secondary_diagnosis"))),
            ("Clinical Summary", _has_meaningful_value(packet.get("lmn_clinical_summary"))),
            ("MRI Date", _has_meaningful_value(packet.get("lmn_mri_date"))),
            ("MRI Findings", _has_meaningful_value(packet.get("lmn_mri_findings"))),
            ("Conservative Duration", _has_meaningful_value(packet.get("lmn_conservative_duration"))),
            ("Medical Necessity Statement", _has_meaningful_value(packet.get("lmn_medical_necessity_statement"))),
            ("Risk Statement", _has_meaningful_value(packet.get("lmn_risk_statement"))),
            ("Reasonable / Necessary Statement", _has_meaningful_value(packet.get("lmn_reasonable_necessary_statement"))),
            ("Contact Statement", _has_meaningful_value(packet.get("lmn_contact_statement"))),
            ("Community Care Provider Credentials", _has_meaningful_value(packet.get("provider_credentials"))),
            ("Community Care Provider Specialty", _has_meaningful_value(packet.get("provider_specialty"))),
            ("Community Care Provider NPI", _has_meaningful_value(packet.get("provider_npi"))),
            ("Community Care Practice Name", _has_meaningful_value(packet.get("practice_name"))),
            ("Community Care Practice Phone", _has_meaningful_value(packet.get("provider_phone"))),
            ("Community Care Practice Fax", _has_meaningful_value(packet.get("provider_fax"))),
            ("Conservative History", _any_checked(packet, [
                "lmn_include_physical_therapy",
                "lmn_include_nsaids",
                "lmn_include_activity_modification",
                "lmn_include_home_exercise",
                "lmn_include_epidural_steroid_injections",
            ])),
            ("Treatment Objectives", _any_checked(packet, [
                "lmn_indication_reduce_pain",
                "lmn_indication_improve_function",
                "lmn_indication_prevent_degeneration",
                "lmn_indication_reduce_opioid_reliance",
                "lmn_indication_prevent_surgery",
            ])),
        ]
    if is_clinical_documentation_profile(profile_name):
        return [
            ("Chief Complaint", _has_meaningful_value(packet.get("clinical_doc_chief_complaint"))),
            ("Duration Detail", _has_meaningful_value(packet.get("clinical_doc_exact_duration"))),
            ("Pain Severity", _has_meaningful_value(packet.get("clinical_doc_pain_severity"))),
            ("Functional Impact", _has_meaningful_value(packet.get("clinical_doc_functional_impact"))),
            ("Conservative Duration", _has_meaningful_value(packet.get("clinical_doc_conservative_duration"))),
            ("MRI Date", _has_meaningful_value(packet.get("clinical_doc_mri_date"))),
            ("Affected Levels", _has_meaningful_value(packet.get("clinical_doc_affected_levels"))),
            ("Primary Diagnosis", _has_meaningful_value(packet.get("clinical_doc_primary_diagnosis"))),
            ("Secondary Diagnosis", _has_meaningful_value(packet.get("clinical_doc_secondary_diagnosis"))),
            ("Assessment Summary", _has_meaningful_value(packet.get("clinical_doc_assessment_summary"))),
            ("Treatment Plan Intro", _has_meaningful_value(packet.get("clinical_doc_treatment_plan_intro"))),
            ("Plan Exclusion", _has_meaningful_value(packet.get("clinical_doc_plan_exclusion"))),
            ("Physician Narrative", _has_meaningful_value(packet.get("clinical_doc_physician_narrative"))),
            ("Pain Pattern", _any_checked(packet, [
                "clinical_doc_pain_axial",
                "clinical_doc_pain_discogenic",
                "clinical_doc_pain_activity_exacerbation",
                "clinical_doc_pain_sitting_intolerance",
                "clinical_doc_pain_standing_intolerance",
                "clinical_doc_pain_bending_lifting",
            ])),
            ("Functional Limitation Set", _any_checked(packet, [
                "clinical_doc_limit_occupational",
                "clinical_doc_limit_prolonged_sitting",
                "clinical_doc_limit_prolonged_standing",
                "clinical_doc_limit_ambulation",
                "clinical_doc_limit_household",
                "clinical_doc_limit_sleep",
            ])),
            ("Conservative Care Set", _any_checked(packet, [
                "clinical_doc_conservative_pt",
                "clinical_doc_conservative_home_exercise",
                "clinical_doc_conservative_nsaids",
                "clinical_doc_conservative_non_opioid",
                "clinical_doc_conservative_activity_modification",
                "clinical_doc_conservative_esi",
                "clinical_doc_conservative_other_interventional",
            ])),
            ("Imaging Findings Set", _any_checked(packet, [
                "clinical_doc_imaging_annular_tear",
                "clinical_doc_imaging_disc_degeneration",
                "clinical_doc_imaging_disc_protrusion",
                "clinical_doc_imaging_disc_displacement",
            ])),
            ("Treatment Plan Set", _any_checked(packet, [
                "clinical_doc_plan_diagnostic_confirmation",
                "clinical_doc_plan_intradiscal_intervention",
                "clinical_doc_plan_follow_up",
            ])),
        ]
    return []


def build_packet_inconsistency_messages(packet):
    profile_name = str(packet.get("packet_profile") or "").strip()
    messages = []

    def add_conflict(label, left_value, right_value, mode="text", allow_contains=False):
        if _values_conflict(left_value, right_value, mode=mode, allow_contains=allow_contains):
            messages.append(
                f"{label} is inconsistent. Shared value '{str(left_value).strip()}' does not match this form's value '{str(right_value).strip()}'."
            )

    if is_submission_cover_profile(profile_name):
        add_conflict("Primary diagnosis code", packet.get("icd_codes"), packet.get("primary_diagnosis_code"), mode="icd_first")

    if is_seoc_request_profile(profile_name):
        add_conflict("VA Medical Center", packet.get("facility"), packet.get("va_medical_center_name"))
        add_conflict("Episode diagnosis", packet.get("diagnosis"), packet.get("episode_diagnosis"), allow_contains=True)
        add_conflict("Episode ICD-10", packet.get("icd_codes"), packet.get("episode_icd_code"), mode="icd_first")
        add_conflict("Last four SSN", packet.get("ssn"), packet.get("last_four_ssn"), mode="last_four")
        add_conflict("Community Care treating provider NPI", packet.get("master_provider_npi"), packet.get("provider_npi"))
        add_conflict("Community Care practice name", packet.get("master_practice_name"), packet.get("practice_name"), allow_contains=True)
        add_conflict("Community Care practice phone", packet.get("master_provider_phone"), packet.get("provider_phone"))
        add_conflict("Community Care practice fax", packet.get("master_provider_fax"), packet.get("provider_fax"))

    if is_consult_request_profile(profile_name):
        add_conflict("VA claim number", packet.get("authorization_number"), packet.get("consult_va_claim_number"))
        add_conflict("VA referring / ordering provider", packet.get("ordering_doctor"), packet.get("consult_referring_va_provider"), allow_contains=True)
        add_conflict("Primary diagnosis", packet.get("diagnosis"), packet.get("consult_primary_diagnosis"), allow_contains=True)
        add_conflict("Secondary diagnosis", packet.get("secondary_diagnosis"), packet.get("consult_secondary_diagnosis"), allow_contains=True)
        add_conflict("VA Medical Center", packet.get("facility"), packet.get("va_medical_center_name"))
        add_conflict("MRI date", packet.get("master_mri_date"), packet.get("consult_mri_date"))
        add_conflict("MRI findings", packet.get("master_mri_findings"), packet.get("consult_mri_findings"), allow_contains=True)
        add_conflict("Community Care treating provider NPI", packet.get("master_provider_npi"), packet.get("provider_npi"))
        add_conflict("Community Care practice name", packet.get("master_practice_name"), packet.get("practice_name"), allow_contains=True)
        add_conflict("Community Care practice address", packet.get("master_provider_address"), packet.get("provider_address"), allow_contains=True)
        add_conflict("Community Care practice phone", packet.get("master_provider_phone"), packet.get("provider_phone"))
        add_conflict("Community Care practice fax", packet.get("master_provider_fax"), packet.get("provider_fax"))
        add_conflict("Community Care secure email", packet.get("master_provider_email"), packet.get("provider_email"))

    if is_lomn_profile(profile_name):
        add_conflict("VA claim number", packet.get("authorization_number"), packet.get("lmn_va_claim_number"))
        add_conflict("Primary diagnosis", packet.get("diagnosis"), packet.get("lmn_primary_diagnosis"), allow_contains=True)
        add_conflict("Secondary diagnosis", packet.get("secondary_diagnosis"), packet.get("lmn_secondary_diagnosis"), allow_contains=True)
        add_conflict("MRI date", packet.get("master_mri_date"), packet.get("lmn_mri_date"))
        add_conflict("MRI findings", packet.get("master_mri_findings"), packet.get("lmn_mri_findings"), allow_contains=True)
        add_conflict("Community Care treating provider NPI", packet.get("master_provider_npi"), packet.get("provider_npi"))
        add_conflict("Community Care practice name", packet.get("master_practice_name"), packet.get("practice_name"), allow_contains=True)
        add_conflict("Community Care practice phone", packet.get("master_provider_phone"), packet.get("provider_phone"))
        add_conflict("Community Care practice fax", packet.get("master_provider_fax"), packet.get("provider_fax"))

    if is_clinical_documentation_profile(profile_name):
        add_conflict("Primary diagnosis", packet.get("diagnosis"), packet.get("clinical_doc_primary_diagnosis"), allow_contains=True)
        add_conflict("Secondary diagnosis", packet.get("secondary_diagnosis"), packet.get("clinical_doc_secondary_diagnosis"), allow_contains=True)
        add_conflict("MRI date", packet.get("master_mri_date"), packet.get("clinical_doc_mri_date"))
        add_conflict("Affected spinal levels", packet.get("master_affected_levels"), packet.get("clinical_doc_affected_levels"), allow_contains=True)

    if is_va_10172_profile(profile_name):
        add_conflict(
            "10-10172 CCN ordering provider name",
            packet.get("provider") or packet.get("ordering_doctor"),
            packet.get("va10172_ordering_provider_name_printed"),
            allow_contains=True,
        )
        add_conflict("10-10172 ordering provider NPI", packet.get("master_provider_npi"), packet.get("va10172_ordering_provider_npi"))
        add_conflict("10-10172 ordering provider phone", packet.get("master_provider_phone"), packet.get("va10172_ordering_provider_phone"))
        add_conflict("10-10172 ordering provider fax", packet.get("master_provider_fax"), packet.get("va10172_ordering_provider_fax"))
        add_conflict("10-10172 ordering provider secure email", packet.get("master_provider_email"), packet.get("va10172_ordering_provider_secure_email"))
        add_conflict("10-10172 ordering provider office address", packet.get("master_provider_address"), packet.get("va10172_ordering_provider_office_address"), allow_contains=True)
        add_conflict("Diagnosis description", packet.get("diagnosis"), packet.get("va10172_diagnosis_description"), allow_contains=True)
        add_conflict("Requested CPT / HCPCS", packet.get("master_requested_cpt_code"), packet.get("va10172_requested_cpt_hcpcs_code"))

    return messages


def _has_meaningful_value(value):
    if isinstance(value, bool):
        return bool(value)
    return bool(str(value or "").strip())


def _first_icd_code(value):
    raw = str(value or "").strip()
    if not raw:
        return ""
    return raw.split(",")[0].strip()


def _can_replace_with_shared_value(current_value, fallback_values=None):
    current = str(current_value or "").strip()
    if not current:
        return True
    for fallback in fallback_values or []:
        if current == str(fallback or "").strip():
            return True
    return False


def apply_packet_builder_shared_field_sync(payload):
    packet = dict(payload or {})
    defaults = default_packet_builder_payload()

    auth_number = str(packet.get("authorization_number") or "").strip()
    icd_codes = str(packet.get("icd_codes") or "").strip()
    ordering_doctor = str(packet.get("ordering_doctor") or "").strip()
    provider_name = str(packet.get("provider") or "").strip()
    facility = str(packet.get("facility") or "").strip()
    diagnosis = str(packet.get("diagnosis") or "").strip()
    ssn = str(packet.get("ssn") or "").strip()
    secondary_diagnosis = str(packet.get("secondary_diagnosis") or "").strip()
    provider_credentials = str(packet.get("master_provider_credentials") or "").strip()
    provider_specialty = str(packet.get("master_provider_specialty") or "").strip()
    provider_npi = str(packet.get("master_provider_npi") or "").strip()
    practice_name = str(packet.get("master_practice_name") or "").strip()
    provider_phone = str(packet.get("master_provider_phone") or "").strip()
    provider_fax = str(packet.get("master_provider_fax") or "").strip()
    provider_email = str(packet.get("master_provider_email") or "").strip()
    provider_address = str(packet.get("master_provider_address") or "").strip()
    requested_cpt_code = str(packet.get("master_requested_cpt_code") or "").strip()
    mri_date = str(packet.get("master_mri_date") or "").strip()
    mri_findings = str(packet.get("master_mri_findings") or "").strip()
    affected_levels = str(packet.get("master_affected_levels") or "").strip()
    requested_service = str(packet.get("requested_service") or "").strip()

    if auth_number:
        if _can_replace_with_shared_value(packet.get("lmn_va_claim_number"), [defaults.get("lmn_va_claim_number")]):
            packet["lmn_va_claim_number"] = auth_number
        if _can_replace_with_shared_value(packet.get("consult_va_claim_number"), [defaults.get("consult_va_claim_number")]):
            packet["consult_va_claim_number"] = auth_number

    if icd_codes:
        first_icd_code = _first_icd_code(icd_codes)
        if _can_replace_with_shared_value(packet.get("episode_icd_code"), [defaults.get("episode_icd_code")]):
            packet["episode_icd_code"] = first_icd_code or icd_codes
        if _can_replace_with_shared_value(packet.get("primary_diagnosis_code"), [defaults.get("primary_diagnosis_code")]):
            packet["primary_diagnosis_code"] = first_icd_code

    if ordering_doctor:
        if _can_replace_with_shared_value(packet.get("consult_referring_va_provider"), [defaults.get("consult_referring_va_provider")]):
            packet["consult_referring_va_provider"] = ordering_doctor

    if provider_name or ordering_doctor:
        va10172_provider_name = provider_name or ordering_doctor
        if _can_replace_with_shared_value(packet.get("va10172_ordering_provider_name_printed"), [defaults.get("va10172_ordering_provider_name_printed")]):
            packet["va10172_ordering_provider_name_printed"] = va10172_provider_name

    if facility:
        if _can_replace_with_shared_value(packet.get("va_medical_center_name"), [defaults.get("va_medical_center_name")]):
            packet["va_medical_center_name"] = facility

    if diagnosis:
        if _can_replace_with_shared_value(packet.get("episode_diagnosis"), [defaults.get("episode_diagnosis")]):
            packet["episode_diagnosis"] = diagnosis
        if _can_replace_with_shared_value(packet.get("lmn_primary_diagnosis"), [defaults.get("lmn_primary_diagnosis")]):
            packet["lmn_primary_diagnosis"] = diagnosis
        if _can_replace_with_shared_value(packet.get("consult_primary_diagnosis"), [defaults.get("consult_primary_diagnosis")]):
            packet["consult_primary_diagnosis"] = diagnosis
        if _can_replace_with_shared_value(packet.get("clinical_doc_primary_diagnosis"), [defaults.get("clinical_doc_primary_diagnosis")]):
            packet["clinical_doc_primary_diagnosis"] = diagnosis
        if _can_replace_with_shared_value(packet.get("va10172_diagnosis_description"), [defaults.get("va10172_diagnosis_description")]):
            packet["va10172_diagnosis_description"] = diagnosis

    if secondary_diagnosis:
        if _can_replace_with_shared_value(packet.get("lmn_secondary_diagnosis"), [defaults.get("lmn_secondary_diagnosis")]):
            packet["lmn_secondary_diagnosis"] = secondary_diagnosis
        if _can_replace_with_shared_value(packet.get("consult_secondary_diagnosis"), [defaults.get("consult_secondary_diagnosis")]):
            packet["consult_secondary_diagnosis"] = secondary_diagnosis
        if _can_replace_with_shared_value(packet.get("clinical_doc_secondary_diagnosis"), [defaults.get("clinical_doc_secondary_diagnosis")]):
            packet["clinical_doc_secondary_diagnosis"] = secondary_diagnosis

    if provider_credentials:
        if _can_replace_with_shared_value(packet.get("provider_credentials"), [defaults.get("provider_credentials")]):
            packet["provider_credentials"] = provider_credentials

    if provider_specialty:
        if _can_replace_with_shared_value(packet.get("provider_specialty"), [defaults.get("provider_specialty")]):
            packet["provider_specialty"] = provider_specialty

    if provider_npi:
        if _can_replace_with_shared_value(packet.get("provider_npi"), [defaults.get("provider_npi")]):
            packet["provider_npi"] = provider_npi
        if _can_replace_with_shared_value(packet.get("va10172_ordering_provider_npi"), [defaults.get("va10172_ordering_provider_npi")]):
            packet["va10172_ordering_provider_npi"] = provider_npi

    if practice_name:
        if _can_replace_with_shared_value(packet.get("practice_name"), [defaults.get("practice_name")]):
            packet["practice_name"] = practice_name

    if provider_phone:
        if _can_replace_with_shared_value(packet.get("provider_phone"), [defaults.get("provider_phone")]):
            packet["provider_phone"] = provider_phone
        if _can_replace_with_shared_value(packet.get("va10172_ordering_provider_phone"), [defaults.get("va10172_ordering_provider_phone")]):
            packet["va10172_ordering_provider_phone"] = provider_phone

    if provider_fax:
        if _can_replace_with_shared_value(packet.get("provider_fax"), [defaults.get("provider_fax")]):
            packet["provider_fax"] = provider_fax
        if _can_replace_with_shared_value(packet.get("va10172_ordering_provider_fax"), [defaults.get("va10172_ordering_provider_fax")]):
            packet["va10172_ordering_provider_fax"] = provider_fax

    if provider_email:
        if _can_replace_with_shared_value(packet.get("provider_email"), [defaults.get("provider_email")]):
            packet["provider_email"] = provider_email
        if _can_replace_with_shared_value(packet.get("va10172_ordering_provider_secure_email"), [defaults.get("va10172_ordering_provider_secure_email")]):
            packet["va10172_ordering_provider_secure_email"] = provider_email

    if provider_address:
        if _can_replace_with_shared_value(packet.get("provider_address"), [defaults.get("provider_address")]):
            packet["provider_address"] = provider_address
        if _can_replace_with_shared_value(packet.get("va10172_ordering_provider_office_address"), [defaults.get("va10172_ordering_provider_office_address")]):
            packet["va10172_ordering_provider_office_address"] = provider_address

    if requested_cpt_code:
        if _can_replace_with_shared_value(packet.get("va10172_requested_cpt_hcpcs_code"), [defaults.get("va10172_requested_cpt_hcpcs_code")]):
            packet["va10172_requested_cpt_hcpcs_code"] = requested_cpt_code

    if requested_service:
        if _can_replace_with_shared_value(packet.get("va10172_description_cpt_hcpcs_code"), [defaults.get("va10172_description_cpt_hcpcs_code")]):
            packet["va10172_description_cpt_hcpcs_code"] = requested_service

    if mri_date:
        if _can_replace_with_shared_value(packet.get("lmn_mri_date"), [defaults.get("lmn_mri_date")]):
            packet["lmn_mri_date"] = mri_date
        if _can_replace_with_shared_value(packet.get("consult_mri_date"), [defaults.get("consult_mri_date")]):
            packet["consult_mri_date"] = mri_date
        if _can_replace_with_shared_value(packet.get("clinical_doc_mri_date"), [defaults.get("clinical_doc_mri_date")]):
            packet["clinical_doc_mri_date"] = mri_date

    if mri_findings:
        if _can_replace_with_shared_value(packet.get("lmn_mri_findings"), [defaults.get("lmn_mri_findings")]):
            packet["lmn_mri_findings"] = mri_findings
        if _can_replace_with_shared_value(packet.get("consult_mri_findings"), [defaults.get("consult_mri_findings")]):
            packet["consult_mri_findings"] = mri_findings

    if affected_levels:
        if _can_replace_with_shared_value(packet.get("clinical_doc_affected_levels"), [defaults.get("clinical_doc_affected_levels")]):
            packet["clinical_doc_affected_levels"] = affected_levels

    if ssn and not str(packet.get("last_four_ssn") or "").strip():
        digits = "".join(ch for ch in ssn if ch.isdigit())
        if len(digits) >= 4:
            packet["last_four_ssn"] = digits[-4:]

    return packet


def build_packet_lab_report(payload):
    packet = apply_packet_builder_shared_field_sync(default_packet_builder_payload() | dict(payload or {}))
    profile_name = packet.get("packet_profile") or ""

    shared_checks = [
        (label, _has_meaningful_value(packet.get(field_name)))
        for field_name, label in SHARED_PACKET_HEADER_FIELDS
    ]
    current_form_checks = list(build_profile_current_form_checks(profile_name, packet))
    inconsistency_messages = build_packet_inconsistency_messages(packet)
    wording_entries = build_wording_assist_entries(packet)
    wording_checks = []
    for entry in wording_entries:
        wording_checks.append(
            (
                f"Wording Review: {entry['label']}",
                entry["status_key"] == "approved",
            )
        )
    current_form_checks.extend(wording_checks)

    patient_packet_checks = []
    if get_packet_export_group(profile_name) == "patient_packet":
        patient_packet_checks = [
            ("Virtual Consent Included", bool(packet.get("included_virtual_consent_form"))),
            ("VA Form 10-10172 Included", bool(packet.get("included_va_form_10_10172"))),
            ("SEOC Request Included", bool(packet.get("included_seoc_request"))),
            ("Consult Request Included", bool(packet.get("included_consult_request"))),
            ("Letter Of Medical Necessity Included", bool(packet.get("included_lomn"))),
            ("Clinical Notes Included", bool(packet.get("included_clinical_notes"))),
            ("MRI Report Included", bool(packet.get("included_mri_report"))),
        ]

    shared_missing = [label for label, present in shared_checks if not present]
    current_missing = [label for label, present in current_form_checks if not present]
    packet_missing = [label for label, present in patient_packet_checks if not present]

    shared_complete = sum(1 for _, present in shared_checks if present)
    current_complete = sum(1 for _, present in current_form_checks if present)
    packet_complete = sum(1 for _, present in patient_packet_checks if present)

    score = 100
    score -= len(shared_missing) * 8
    score -= len(current_missing) * 7
    score -= len(packet_missing) * 3
    score -= len(inconsistency_messages) * 10
    score = max(0, min(100, score))

    if inconsistency_messages:
        status = "Needs Consistency Fixes"
    elif shared_missing:
        status = "Needs Shared Packet Header Details"
    elif current_missing or packet_missing:
        status = "Needs Current Form Work"
    else:
        status = "Ready To Export"

    return {
        "status": status,
        "score": score,
        "profile_name": profile_name,
        "shared_checks": shared_checks,
        "current_form_checks": current_form_checks,
        "patient_packet_checks": patient_packet_checks,
        "shared_missing": shared_missing,
        "current_missing": current_missing,
        "packet_missing": packet_missing,
        "shared_complete": shared_complete,
        "current_complete": current_complete,
        "packet_complete": packet_complete,
        "inconsistency_messages": inconsistency_messages,
        "wording_entries": wording_entries,
        "wording_checks": wording_checks,
    }


def classify_packet_lab_completion(report):
    current_checks = list(report.get("current_form_checks") or [])
    current_complete = int(report.get("current_complete") or 0)
    current_missing = list(report.get("current_missing") or [])
    shared_missing = list(report.get("shared_missing") or [])
    inconsistency_messages = list(report.get("inconsistency_messages") or [])

    if current_checks and current_complete <= 0:
        return "not_started", "Not Started"
    if shared_missing or current_missing or inconsistency_messages:
        return "in_progress", "In Progress"
    return "complete", "Complete"


def packet_profile_status_palette(status_key):
    palettes = {
        "complete": {
            "item_background": QColor("#103522"),
            "item_foreground": QColor("#CFF8DE"),
            "combo_background": "#173B29",
            "combo_border": "#2C8B57",
            "combo_text": "#EAFBF1",
        },
        "in_progress": {
            "item_background": QColor("#4A3410"),
            "item_foreground": QColor("#FFF2C4"),
            "combo_background": "#4A3410",
            "combo_border": "#C7922E",
            "combo_text": "#FFF6DB",
        },
        "not_started": {
            "item_background": QColor("#441A1A"),
            "item_foreground": QColor("#FFD8D8"),
            "combo_background": "#441A1A",
            "combo_border": "#B55050",
            "combo_text": "#FFECEC",
        },
    }
    return palettes.get(status_key, palettes["not_started"])


def build_packet_lab_html(payload):
    report = build_packet_lab_report(payload)

    def checklist_html(rows):
        if not rows:
            return "<div style='color:#64748B; font-style:italic;'>No checks for this scope.</div>"
        parts = []
        for label, present in rows:
            color = "#1F6F43" if present else "#A63A3A"
            marker = "&#10003;" if present else "&#10007;"
            parts.append(
                f"<div style='margin-top:6px; color:{color};'><span style='font-weight:700;'>{marker}</span> {html.escape(label)}</div>"
            )
        return "".join(parts)

    missing_lines = []
    if report["shared_missing"]:
        missing_lines.append(
            "<div style='margin-top:8px;'><strong>Shared header still missing:</strong> "
            + ", ".join(html.escape(item) for item in report["shared_missing"])
            + "</div>"
        )
    if report["current_missing"]:
        missing_lines.append(
            "<div style='margin-top:8px;'><strong>Current form still missing:</strong> "
            + ", ".join(html.escape(item) for item in report["current_missing"])
            + "</div>"
        )
    if report["packet_missing"]:
        missing_lines.append(
            "<div style='margin-top:8px;'><strong>Patient packet items not marked included:</strong> "
            + ", ".join(html.escape(item) for item in report["packet_missing"])
            + "</div>"
        )
    if report["inconsistency_messages"]:
        missing_lines.append(
            "<div style='margin-top:8px;'><strong>Consistency issues:</strong></div>"
            + "".join(
                f"<div style='margin-top:6px;'>{html.escape(item)}</div>"
                for item in report["inconsistency_messages"]
            )
        )
    pending_wording = [
        entry for entry in report.get("wording_entries") or []
        if entry.get("status_key") != "approved"
    ]
    if pending_wording:
        missing_lines.append(
            "<div style='margin-top:8px;'><strong>Wording assist still needs attention:</strong></div>"
            + "".join(
                (
                    f"<div style='margin-top:6px;'>{html.escape(entry['label'])}: "
                    + (
                        "missing " + ", ".join(html.escape(item) for item in entry.get("missing_facts") or [])
                        if entry.get("missing_facts")
                        else "review or approval still pending"
                    )
                    + "</div>"
                )
                for entry in pending_wording
            )
        )

    markup = (
        "<div style='font-family:Calibri, Arial, sans-serif; color:#111827; line-height:1.55;'>"
        "<div style='font-size:18pt; font-weight:700;'>Packet Studio Lab</div>"
        f"<div style='margin-top:8px; font-size:11pt;'><strong>Current form:</strong> {html.escape(report['profile_name'] or 'Unknown')}</div>"
        f"<div style='margin-top:6px; font-size:11pt;'><strong>Draft status:</strong> {html.escape(report['status'])}</div>"
        f"<div style='margin-top:6px; font-size:11pt;'><strong>Lab score:</strong> {report['score']}%</div>"
        "<div style='margin-top:18px; font-size:12pt; font-weight:700; color:#1F3A5F;'>Shared Packet Header</div>"
        f"<div style='margin-top:6px; color:#4B5563;'>{report['shared_complete']} of {len(report['shared_checks'])} complete. These values carry across forms where expected.</div>"
        + checklist_html(report["shared_checks"])
        + "<div style='margin-top:18px; font-size:12pt; font-weight:700; color:#1F3A5F;'>Current Form Essentials</div>"
        f"<div style='margin-top:6px; color:#4B5563;'>{report['current_complete']} of {len(report['current_form_checks'])} complete.</div>"
        + checklist_html(report["current_form_checks"])
    )
    if report["patient_packet_checks"]:
        markup += (
            "<div style='margin-top:18px; font-size:12pt; font-weight:700; color:#1F3A5F;'>Patient Packet Bundle Marks</div>"
            f"<div style='margin-top:6px; color:#4B5563;'>{report['packet_complete']} of {len(report['patient_packet_checks'])} marked included.</div>"
            + checklist_html(report["patient_packet_checks"])
        )
    if report["inconsistency_messages"]:
        markup += (
            "<div style='margin-top:18px; padding:12px 14px; background:#FEF3C7; border:1px solid #EAB308;'>"
            "<div style='font-weight:700; color:#92400E;'>Consistency Flags</div>"
            + "".join(
                f"<div style='margin-top:6px; color:#92400E;'>{html.escape(item)}</div>"
                for item in report["inconsistency_messages"]
            )
            + "</div>"
        )
    if missing_lines:
        markup += (
            "<div style='margin-top:18px; padding:12px 14px; background:#FFF7ED; border:1px solid #F3D6A4;'>"
            "<div style='font-weight:700; color:#9A4D00;'>What To Fix Before Export</div>"
            + "".join(missing_lines)
            + "</div>"
        )
    else:
        markup += (
            "<div style='margin-top:18px; padding:12px 14px; background:#ECFDF3; border:1px solid #B8E3C9;'>"
            "<div style='font-weight:700; color:#1F6F43;'>This draft is in a strong place for export.</div>"
            "<div style='margin-top:6px;'>Shared packet facts are complete, and the current form has its key page-specific details filled in.</div>"
            "</div>"
        )

    markup += "</div>"
    return render_packet_builder_document_preview(markup)


def ensure_packet_library_dir():
    os.makedirs(PACKET_LIBRARY_DIR, exist_ok=True)
    return PACKET_LIBRARY_DIR


def packet_library_record_path(draft_id):
    safe_id = sanitize_builder_filename(draft_id or "packet_draft")
    return os.path.join(ensure_packet_library_dir(), f"{safe_id}.json")


def generate_packet_library_draft_id():
    return datetime.now().strftime("draft_%Y%m%d_%H%M%S_%f")


def packet_library_display_name(payload, base_filename=""):
    packet = normalize_packet_builder_payload(payload)
    patient_name = str(packet.get("patient_name") or "").strip()
    profile_name = str(packet.get("packet_profile") or "").strip()
    packet_title = str(packet.get("packet_title") or "").strip()
    normalized_base = sanitize_builder_filename(base_filename)

    if patient_name:
        return patient_name
    if normalized_base and normalized_base != "truecore_packet":
        return normalized_base.replace("_", " ")
    if packet_title and packet_title != default_title_for_profile(profile_name):
        return packet_title
    fallback_title = default_title_for_profile(profile_name) or "TrueCore Packet"
    return f"{fallback_title} Draft"


def build_packet_library_completion(payload):
    packet = normalize_packet_builder_payload(payload)
    group_name = get_packet_export_group(packet.get("packet_profile"))
    profile_names = bundle_profiles_for_group(group_name) or [packet.get("packet_profile")]
    profile_reports = [
        build_packet_lab_report(build_profile_export_payload(packet, profile_name))
        for profile_name in profile_names
        if profile_name
    ]
    if not profile_reports:
        profile_reports = [build_packet_lab_report(packet)]

    shared_checks = list(profile_reports[0].get("shared_checks") or [])
    packet_checks = list(profile_reports[0].get("patient_packet_checks") or [])
    current_total = sum(len(report.get("current_form_checks") or []) for report in profile_reports)
    current_complete = sum(int(report.get("current_complete") or 0) for report in profile_reports)
    packet_total = len(packet_checks)
    packet_complete = int(profile_reports[0].get("packet_complete") or 0)
    shared_total = len(shared_checks)
    shared_complete = int(profile_reports[0].get("shared_complete") or 0)

    shared_missing = list(dict.fromkeys(profile_reports[0].get("shared_missing") or []))
    packet_missing = list(dict.fromkeys(profile_reports[0].get("packet_missing") or []))
    current_missing = []
    for report in profile_reports:
        for label in report.get("current_missing") or []:
            if label not in current_missing:
                current_missing.append(label)

    total_checks = shared_total + current_total + packet_total
    total_complete = shared_complete + current_complete + packet_complete
    completion_score = round((100.0 * total_complete / total_checks), 2) if total_checks else 0.0

    if total_complete <= 0:
        status_key, status_label = "not_started", "Not Started"
    elif not shared_missing and not current_missing and not packet_missing:
        status_key, status_label = "complete", "Complete"
    else:
        status_key, status_label = "in_progress", "In Progress"

    return {
        "group_name": group_name,
        "profile_names": profile_names,
        "shared_checks": shared_checks,
        "packet_checks": packet_checks,
        "shared_total": shared_total,
        "shared_complete": shared_complete,
        "current_total": current_total,
        "current_complete": current_complete,
        "packet_total": packet_total,
        "packet_complete": packet_complete,
        "shared_missing": shared_missing,
        "current_missing": current_missing,
        "packet_missing": packet_missing,
        "status_key": status_key,
        "status_label": status_label,
        "completion_score": completion_score,
    }


def build_export_warning_context(payload, group_name=""):
    packet = normalize_packet_builder_payload(payload)
    normalized_group = str(group_name or get_packet_export_group(packet.get("packet_profile")) or "").strip().lower()
    critical_shared_values = [
        str(packet.get("patient_name") or "").strip(),
        str(packet.get("date_of_birth") or "").strip(),
        str(packet.get("facility") or "").strip(),
        str(packet.get("ordering_doctor") or "").strip(),
        str(packet.get("diagnosis") or "").strip(),
        str(packet.get("authorization_number") or "").strip(),
    ]
    critical_shared_complete = sum(1 for value in critical_shared_values if value)

    if normalized_group == "patient_packet":
        completion = build_packet_library_completion(packet)
        completion_score = float(completion.get("completion_score") or 0.0)
        if completion.get("status_key") == "not_started" or completion_score <= 10.0 or critical_shared_complete <= 1:
            return {
                "title": "Packet Mostly Empty",
                "message": (
                    "This patient packet is still mostly empty.\n\n"
                    f"Completion score: {completion_score:.0f}%\n"
                    f"Core packet facts: {critical_shared_complete} of 6\n"
                    f"Shared header: {completion.get('shared_complete', 0)} of {completion.get('shared_total', 0)}\n"
                    f"Forms complete: {completion.get('current_complete', 0)} of {completion.get('current_total', 0)}\n\n"
                    "Export anyway?"
                ),
            }
        return None

    report = build_packet_lab_report(packet)
    status_key, _ = classify_packet_lab_completion(report)
    if status_key == "not_started" or critical_shared_complete <= 1:
        return {
            "title": "Form Not Started",
            "message": (
                "This form has not really been started yet.\n\n"
                f"Core packet facts: {critical_shared_complete} of 6\n"
                f"Shared header: {report.get('shared_complete', 0)} of {len(report.get('shared_checks') or [])}\n"
                f"Current form: {report.get('current_complete', 0)} of {len(report.get('current_form_checks') or [])}\n\n"
                "Export anyway?"
            ),
        }
    return None


def _add_packet_field_observation(packet, field_name, document_type, value):
    normalized_value = value if isinstance(value, bool) else str(value or "").strip()
    if normalized_value in ("", None, []):
        return
    packet.fields[field_name] = normalized_value
    packet.field_observations.setdefault(field_name, []).append(
        {
            "document_type": str(document_type or "").strip().lower(),
            "value": normalized_value,
        }
    )


def build_packet_library_production_metrics(payload):
    packet_payload = normalize_packet_builder_payload(payload)
    completion = build_packet_library_completion(packet_payload)
    packet = Packet()
    group_name = completion.get("group_name")
    packet.packet_profile = "authorization_request" if group_name == "referral_request" else "full_submission"
    packet.packet_profile_label = "Authorization Request" if group_name == "referral_request" else "Full Submission"
    packet.packet_assembly_score = float(completion.get("completion_score") or 0.0)
    packet.packet_confidence = round(max(0.32, min(0.96, packet.packet_assembly_score / 100.0)), 2)

    patient_name = str(packet_payload.get("patient_name") or "").strip()
    dob = str(packet_payload.get("date_of_birth") or "").strip()
    auth_number = str(packet_payload.get("authorization_number") or "").strip()
    diagnosis = str(packet_payload.get("diagnosis") or packet_payload.get("consult_primary_diagnosis") or "").strip()
    icd_codes = str(packet_payload.get("icd_codes") or "").strip()
    ordering_provider = str(packet_payload.get("ordering_doctor") or "").strip()
    facility = str(packet_payload.get("facility") or "").strip()
    requested_service = str(packet_payload.get("requested_service") or "").strip()
    clinical_summary = str(
        packet_payload.get("clinical_summary")
        or packet_payload.get("clinical_doc_chief_complaint")
        or packet_payload.get("consult_reason_text")
        or ""
    ).strip()

    packet.fields.update(
        {
            "name": patient_name,
            "dob": dob,
            "authorization_number": auth_number,
            "claim_number": auth_number,
            "va_icn": str(packet_payload.get("va_icn") or "").strip(),
            "ordering_provider": ordering_provider,
            "provider": str(packet_payload.get("provider") or packet_payload.get("master_practice_name") or "").strip(),
            "facility": facility,
            "diagnosis": diagnosis,
            "icd_codes": icd_codes,
            "reason_for_request": requested_service or str(packet_payload.get("consult_reason_text") or "").strip(),
            "procedure": requested_service,
            "symptom": clinical_summary,
        }
    )

    detected_documents = set()
    if group_name == "referral_request":
        detected_documents.update({"rfs", "approved_referral"})
    else:
        if packet_payload.get("included_va_form_10_10172"):
            detected_documents.update({"rfs", "approved_referral"})
        if packet_payload.get("included_virtual_consent_form"):
            detected_documents.add("consent")
        if packet_payload.get("included_seoc_request"):
            detected_documents.add("seoc")
        if packet_payload.get("included_consult_request"):
            detected_documents.add("consult_request")
        if packet_payload.get("included_lomn"):
            detected_documents.add("lomn")
        if packet_payload.get("included_clinical_notes"):
            detected_documents.add("clinical_notes")
        if packet_payload.get("included_mri_report"):
            detected_documents.add("imaging_report")
    packet.detected_documents = detected_documents

    if "rfs" in detected_documents:
        for field_name, value in {
            "name": patient_name,
            "dob": dob,
            "authorization_number": auth_number,
            "diagnosis": diagnosis or str(packet_payload.get("va10172_diagnosis_description") or "").strip(),
            "icd_codes": icd_codes or str(packet_payload.get("primary_diagnosis_code") or "").strip(),
        }.items():
            _add_packet_field_observation(packet, field_name, "rfs", value)

    if "seoc" in detected_documents:
        for field_name, value in {
            "name": patient_name,
            "dob": dob,
            "reason_for_request": str(packet_payload.get("clinical_objectives") or requested_service).strip(),
            "diagnosis": str(packet_payload.get("episode_diagnosis") or diagnosis).strip(),
            "icd_codes": str(packet_payload.get("episode_icd_code") or _first_icd_code(icd_codes)).strip(),
        }.items():
            _add_packet_field_observation(packet, field_name, "seoc", value)

    if "consult_request" in detected_documents:
        for field_name, value in {
            "ordering_provider": str(packet_payload.get("consult_referring_va_provider") or ordering_provider).strip(),
            "reason_for_request": str(packet_payload.get("consult_reason_text") or requested_service).strip(),
            "diagnosis": str(packet_payload.get("consult_primary_diagnosis") or diagnosis).strip(),
            "icd_codes": icd_codes,
        }.items():
            _add_packet_field_observation(packet, field_name, "consult_request", value)

    if "lomn" in detected_documents:
        for field_name, value in {
            "reason_for_request": str(packet_payload.get("lmn_medical_necessity_statement") or requested_service).strip(),
            "diagnosis": str(packet_payload.get("lmn_primary_diagnosis") or diagnosis).strip(),
            "icd_codes": icd_codes,
        }.items():
            _add_packet_field_observation(packet, field_name, "lomn", value)

    if "consent" in detected_documents:
        _add_packet_field_observation(packet, "name", "consent", patient_name)
        _add_packet_field_observation(packet, "dob", "consent", dob)
        if str(packet_payload.get("patient_signature_name") or "").strip() and str(packet_payload.get("patient_signature_date") or "").strip():
            _add_packet_field_observation(packet, "signature_present", "consent", True)
        else:
            packet.unfilled_documents.add("consent")

    if "clinical_notes" in detected_documents:
        for field_name, value in {
            "diagnosis": str(packet_payload.get("clinical_doc_primary_diagnosis") or diagnosis).strip(),
            "icd_codes": icd_codes,
            "symptom": str(packet_payload.get("clinical_doc_chief_complaint") or clinical_summary).strip(),
        }.items():
            _add_packet_field_observation(packet, field_name, "clinical_notes", value)

    packet.packet_rubric = build_packet_rubric(packet)
    packet.packet_score = packet.packet_rubric.get("score")
    packet.packet_strength = packet.packet_rubric.get("score_band")
    packet.packet_main_blocker = packet.packet_rubric.get("main_blocker")

    decision = ReviewEngine().build_submission_decision(packet)
    return {
        "score": float(packet.packet_rubric.get("score") or 0.0),
        "strength": str(packet.packet_rubric.get("score_band") or "weak"),
        "main_blocker": packet.packet_rubric.get("main_blocker"),
        "readiness": str(decision.get("readiness") or "hold"),
        "next_action": str(decision.get("next_action") or "").strip(),
        "decision": decision,
        "completion": completion,
    }


def list_packet_library_records():
    ensure_packet_library_dir()
    records = []
    for name in os.listdir(PACKET_LIBRARY_DIR):
        if not str(name).lower().endswith(".json"):
            continue
        path = os.path.join(PACKET_LIBRARY_DIR, name)
        try:
            with open(path, "r", encoding="utf-8") as handle:
                record = dict(json.load(handle) or {})
        except Exception:
            continue
        payload = normalize_packet_builder_payload(record.get("payload") or {})
        record["draft_id"] = str(record.get("draft_id") or os.path.splitext(name)[0]).strip()
        record["base_filename"] = sanitize_builder_filename(record.get("base_filename") or "truecore_packet")
        record["display_name"] = str(
            record.get("display_name")
            or packet_library_display_name(payload, record.get("base_filename"))
        ).strip()
        record["path"] = path
        record["payload"] = payload
        record["artifacts"] = dict(record.get("artifacts") or {})
        record["updated_at"] = str(record.get("updated_at") or record.get("saved_at") or "").strip()
        records.append(record)
    records.sort(key=lambda item: item.get("updated_at") or "", reverse=True)
    return records


def save_packet_library_record(record):
    draft_id = str(record.get("draft_id") or generate_packet_library_draft_id()).strip()
    record = dict(record or {})
    record["draft_id"] = draft_id
    path = packet_library_record_path(draft_id)
    with open(path, "w", encoding="utf-8") as handle:
        json.dump(record, handle, indent=4)
    record["path"] = path
    return record


def delete_packet_library_record(draft_id):
    path = packet_library_record_path(draft_id)
    if os.path.exists(path):
        os.remove(path)


def find_word_executable():
    candidates = [
        r"C:\Program Files\Microsoft Office\root\Office16\WINWORD.EXE",
        r"C:\Program Files (x86)\Microsoft Office\root\Office16\WINWORD.EXE",
        os.path.join(os.environ.get("LOCALAPPDATA", ""), "Microsoft", "WindowsApps", "WINWORD.EXE"),
    ]
    for candidate in candidates:
        normalized = str(candidate or "").strip()
        if normalized and os.path.exists(normalized):
            return normalized
    return ""


def convert_docx_to_pdf_via_word(docx_path, pdf_path):
    docx_path = os.path.abspath(str(docx_path or "").strip())
    pdf_path = os.path.abspath(str(pdf_path or "").strip())
    if not docx_path or not os.path.exists(docx_path):
        raise RuntimeError("The DOCX preview source could not be found.")
    os.makedirs(os.path.dirname(pdf_path), exist_ok=True)
    docx_literal = docx_path.replace("'", "''")
    pdf_literal = pdf_path.replace("'", "''")
    script = (
        "$ErrorActionPreference = 'Stop'\n"
        "$word = $null\n"
        "$doc = $null\n"
        f"$docxPath = '{docx_literal}'\n"
        f"$pdfPath = '{pdf_literal}'\n"
        "try {\n"
        "  $word = New-Object -ComObject Word.Application\n"
        "  $word.Visible = $false\n"
        "  $word.DisplayAlerts = 0\n"
        "  $doc = $word.Documents.Open($docxPath)\n"
        "  $doc.SaveAs([ref]$pdfPath, [ref]17)\n"
        "} finally {\n"
        "  if ($doc -ne $null) { $doc.Close([ref]$false) }\n"
        "  if ($word -ne $null) { $word.Quit() }\n"
        "}\n"
    )
    startupinfo = None
    if hasattr(subprocess, "STARTUPINFO"):
        startupinfo = subprocess.STARTUPINFO()
        startupinfo.dwFlags |= getattr(subprocess, "STARTF_USESHOWWINDOW", 0)
        startupinfo.wShowWindow = 0
    completed = subprocess.run(
        ["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-Command", script],
        capture_output=True,
        text=True,
        timeout=120,
        startupinfo=startupinfo,
        creationflags=getattr(subprocess, "CREATE_NO_WINDOW", 0),
    )
    if completed.returncode != 0 or not os.path.exists(pdf_path):
        details = "\n".join(part for part in [completed.stdout.strip(), completed.stderr.strip()] if part)
        raise RuntimeError(details or "Word could not render the DOCX preview.")
    return pdf_path


def _signature_image_bytes(payload, field_name):
    image_field = SIGNATURE_IMAGE_FIELDS.get(field_name, field_name)
    raw = str((payload or {}).get(image_field) or "").strip()
    if not raw:
        typed_text = sanitize_packet_builder_text((payload or {}).get(field_name) or "")
        if typed_text:
            return _typed_signature_image_bytes(typed_text)
        return b""
    try:
        return base64.b64decode(raw)
    except Exception:
        return b""


def _signature_image_html(payload, field_name, width_px=220, height_px=54):
    image_bytes = _signature_image_bytes(payload, field_name)
    if not image_bytes:
        return ""
    raw = base64.b64encode(image_bytes).decode("ascii")
    return (
        f"<img src='data:image/png;base64,{raw}' "
        f"style='max-width:{int(width_px)}px; max-height:{int(height_px)}px; display:block;'/>"
    )


def _typed_signature_image_bytes(text):
    cleaned = sanitize_packet_builder_text(text)
    if not cleaned:
        return b""
    image = QImage(420, 110, QImage.Format_ARGB32_Premultiplied)
    image.fill(Qt.transparent)
    painter = QPainter(image)
    painter.setRenderHint(QPainter.Antialiasing, True)
    painter.setRenderHint(QPainter.TextAntialiasing, True)
    font_candidates = ["Segoe Script", "Lucida Handwriting", "Brush Script MT", "Segoe Print", "Comic Sans MS"]
    font = QFont(font_candidates[0], 28)
    for family in font_candidates:
        candidate = QFont(family, 28)
        metrics = QFontMetrics(candidate)
        if metrics.horizontalAdvance(cleaned) <= 380:
            font = candidate
            break
    painter.setFont(font)
    painter.setPen(QColor("#111827"))
    painter.drawText(QRectF(12, 10, 396, 86), Qt.AlignLeft | Qt.AlignVCenter, cleaned)
    painter.end()
    array = QByteArray()
    buffer = QBuffer(array)
    buffer.open(QBuffer.WriteOnly)
    image.save(buffer, "PNG")
    return bytes(array)


class SignaturePadWidget(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setMinimumSize(420, 150)
        self.setStyleSheet("background:#FFFFFF; border:1px solid #C9D4E2; border-radius:8px;")
        self._pixmap = QPixmap(self.size())
        self._pixmap.fill(Qt.white)
        self._last_point = None
        self._drawing = False

    def resizeEvent(self, event):
        if self.width() <= 0 or self.height() <= 0:
            return
        new_pixmap = QPixmap(self.size())
        new_pixmap.fill(Qt.white)
        if not self._pixmap.isNull():
            painter = QPainter(new_pixmap)
            painter.drawPixmap(0, 0, self._pixmap)
            painter.end()
        self._pixmap = new_pixmap
        super().resizeEvent(event)

    def paintEvent(self, event):
        painter = QPainter(self)
        painter.drawPixmap(0, 0, self._pixmap)
        painter.end()
        super().paintEvent(event)

    def mousePressEvent(self, event):
        if event.button() == Qt.LeftButton:
            self._drawing = True
            self._last_point = event.position().toPoint()
        super().mousePressEvent(event)

    def mouseMoveEvent(self, event):
        if self._drawing and self._last_point is not None:
            painter = QPainter(self._pixmap)
            painter.setRenderHint(QPainter.Antialiasing, True)
            painter.setPen(QPen(QColor("#111827"), 2.2, Qt.SolidLine, Qt.RoundCap, Qt.RoundJoin))
            current_point = event.position().toPoint()
            painter.drawLine(self._last_point, current_point)
            painter.end()
            self._last_point = current_point
            self.update()
        super().mouseMoveEvent(event)

    def mouseReleaseEvent(self, event):
        if event.button() == Qt.LeftButton:
            self._drawing = False
            self._last_point = None
        super().mouseReleaseEvent(event)

    def clear_signature(self):
        self._pixmap.fill(Qt.white)
        self.update()

    def load_signature_bytes(self, data):
        self.clear_signature()
        if not data:
            return
        image = QImage()
        if image.loadFromData(data, "PNG"):
            scaled = QPixmap.fromImage(image).scaled(
                self.size(),
                Qt.KeepAspectRatio,
                Qt.SmoothTransformation,
            )
            self._pixmap.fill(Qt.white)
            painter = QPainter(self._pixmap)
            x = max(0, (self.width() - scaled.width()) // 2)
            y = max(0, (self.height() - scaled.height()) // 2)
            painter.drawPixmap(x, y, scaled)
            painter.end()
            self.update()

    def export_signature_base64(self):
        image = self._pixmap.toImage()
        if image.isNull():
            return ""
        # Detect whether anything other than white has been drawn.
        nonwhite = False
        step_x = max(1, image.width() // 40)
        step_y = max(1, image.height() // 20)
        for x in range(0, image.width(), step_x):
            for y in range(0, image.height(), step_y):
                if image.pixelColor(x, y) != QColor(Qt.white):
                    nonwhite = True
                    break
            if nonwhite:
                break
        if not nonwhite:
            return ""
        array = QByteArray()
        buffer = QBuffer(array)
        buffer.open(QBuffer.WriteOnly)
        self._pixmap.save(buffer, "PNG")
        return bytes(array.toBase64()).decode("ascii")


class SignatureCaptureDialog(QDialog):
    def __init__(self, label_text="", existing_base64="", parent=None):
        super().__init__(parent)
        self.setWindowTitle("Draw Signature")
        self.setModal(True)
        self.resize(520, 260)
        self.setStyleSheet("QDialog { background:#0F1823; color:#E5E7EB; }")
        layout = QVBoxLayout(self)
        layout.setContentsMargins(14, 14, 14, 14)
        layout.setSpacing(10)

        help_label = QLabel(label_text or "Draw the signature below, then accept or reset it.")
        help_label.setWordWrap(True)
        help_label.setStyleSheet("color:#E5E7EB;")
        layout.addWidget(help_label)

        self.pad = SignaturePadWidget()
        layout.addWidget(self.pad, stretch=1)
        if existing_base64:
            self.pad.load_signature_bytes(_signature_image_bytes({ "image": existing_base64 }, "image"))

        button_row = QHBoxLayout()
        button_row.addStretch(1)
        self.cancel_button = QPushButton("Cancel")
        self.reset_button = QPushButton("Reset")
        self.accept_button = QPushButton("Accept")
        for button in [self.cancel_button, self.reset_button, self.accept_button]:
            button.setMinimumWidth(100)
        self.cancel_button.clicked.connect(self.reject)
        self.reset_button.clicked.connect(self.pad.clear_signature)
        self.accept_button.clicked.connect(self.accept)
        button_row.addWidget(self.cancel_button)
        button_row.addWidget(self.reset_button)
        button_row.addWidget(self.accept_button)
        layout.addLayout(button_row)

    def signature_base64(self):
        return self.pad.export_signature_base64()


class MultiSelectPromptDialog(QDialog):
    def __init__(self, title, options, selected_values=None, parent=None):
        super().__init__(parent)
        self.setWindowTitle(title or "Choose Options")
        self.setModal(True)
        self.resize(430, 330)
        self.setStyleSheet("QDialog { background:#0F1823; color:#E5E7EB; }")
        layout = QVBoxLayout(self)
        layout.setContentsMargins(14, 14, 14, 14)
        layout.setSpacing(10)

        help_label = QLabel("Choose one or more prompts that fit this packet. Leave everything clear to keep the field on Auto.")
        help_label.setWordWrap(True)
        help_label.setStyleSheet("color:#BFD0E3;")
        layout.addWidget(help_label)

        self._checks = []
        selected = {str(value or "").strip() for value in (selected_values or [])}
        options_scroll = QScrollArea()
        options_scroll.setWidgetResizable(True)
        options_scroll.setFrameShape(QFrame.NoFrame)
        options_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        options_container = QWidget()
        options_layout = QVBoxLayout(options_container)
        options_layout.setContentsMargins(0, 0, 0, 0)
        options_layout.setSpacing(8)
        for option in options or []:
            checkbox = QCheckBox(str(option))
            checkbox.setChecked(str(option) in selected)
            checkbox.setStyleSheet("QCheckBox { color:#EAF2FF; font-size:14px; padding:2px 0; }")
            options_layout.addWidget(checkbox)
            self._checks.append(checkbox)
        options_layout.addStretch(1)
        options_scroll.setWidget(options_container)
        layout.addWidget(options_scroll, stretch=1)

        button_row = QHBoxLayout()
        button_row.addStretch(1)
        clear_button = QPushButton("Reset to Auto")
        clear_button.clicked.connect(self._clear_all)
        cancel_button = QPushButton("Cancel")
        cancel_button.clicked.connect(self.reject)
        accept_button = QPushButton("Accept")
        accept_button.clicked.connect(self.accept)
        for button in (clear_button, cancel_button, accept_button):
            button_metrics = QFontMetrics(button.font())
            button_width = max(84, min(132, button_metrics.horizontalAdvance(button.text()) + 34))
            button.setFixedHeight(32)
            button.setFixedWidth(button_width)
            button.setStyleSheet(
                "QPushButton {"
                "background-color:#1D2A3A; color:#FFFFFF; border:1px solid #34506B; border-radius:8px; "
                "padding:4px 10px; font-size:12px; font-weight:600; }"
                "QPushButton:hover { background-color:#223246; }"
            )
            button_row.addWidget(button)
        layout.addLayout(button_row)

    def _clear_all(self):
        for checkbox in self._checks:
            checkbox.setChecked(False)

    def selected_items(self):
        return [checkbox.text() for checkbox in self._checks if checkbox.isChecked()]


class MultiSelectPromptField(QWidget):
    textChanged = Signal(str)

    def __init__(self, title, options, parent=None):
        super().__init__(parent)
        self._title = str(title or "Choose Prompts")
        self._options = [str(option) for option in (options or [])]
        self._selected = []
        layout = QHBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(8)

        self.summary_input = QLineEdit()
        self.summary_input.setReadOnly(True)
        self.summary_input.setPlaceholderText("Auto")
        self.summary_input.setMinimumWidth(0)
        layout.addWidget(self.summary_input, stretch=1)

        self.choose_button = QPushButton("Choose")
        self.choose_button.setFixedHeight(32)
        self.choose_button.setMaximumWidth(88)
        self.choose_button.setStyleSheet(
            "QPushButton {"
            "background-color:#1D2A3A; color:#FFFFFF; border:1px solid #34506B; border-radius:8px; "
            "padding:4px 10px; font-size:12px; font-weight:600; }"
            "QPushButton:hover { background-color:#223246; }"
        )
        self.choose_button.clicked.connect(self._open_picker)
        layout.addWidget(self.choose_button, 0)
        self._refresh_summary(emit_signal=False)

    def _open_picker(self):
        dialog = MultiSelectPromptDialog(self._title, self._options, self._selected, parent=self)
        if dialog.exec() != QDialog.Accepted:
            return
        self._selected = dialog.selected_items()
        self._refresh_summary()

    def _refresh_summary(self, emit_signal=True):
        summary = "Auto" if not self._selected else "; ".join(self._selected)
        with QSignalBlocker(self.summary_input):
            self.summary_input.setText(summary)
        self.summary_input.setToolTip(summary)
        if emit_signal:
            self.textChanged.emit(self.text())

    def text(self):
        return "Auto" if not self._selected else " | ".join(self._selected)

    def setText(self, value):
        raw = str(value or "").strip()
        if not raw or raw.lower() == "auto":
            self._selected = []
        else:
            pieces = [part.strip() for part in raw.replace("\n", "|").split("|") if part.strip()]
            option_map = {re.sub(r"\s+", " ", option).strip().lower(): option for option in self._options}
            selected = []
            for piece in pieces:
                key = re.sub(r"\s+", " ", piece).strip().lower()
                selected.append(option_map.get(key, piece))
            deduped = []
            for item in selected:
                if item not in deduped:
                    deduped.append(item)
            self._selected = deduped
        self._refresh_summary()


class AutoSizingPlainTextEdit(QPlainTextEdit):
    def __init__(self, minimum_height=140, parent=None):
        super().__init__(parent)
        self._minimum_editor_height = int(minimum_height or 140)
        self.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        self.textChanged.connect(self.update_content_height)
        QTimer.singleShot(0, self.update_content_height)

    def resizeEvent(self, event):
        super().resizeEvent(event)
        QTimer.singleShot(0, self.update_content_height)

    def update_content_height(self):
        document = self.document()
        layout = document.documentLayout()
        if layout is None:
            return
        content_height = layout.documentSize().height()
        margin = document.documentMargin()
        frame = self.frameWidth() * 2
        target_height = int(max(self._minimum_editor_height, content_height + (margin * 2) + frame + 8))
        self.setFixedHeight(target_height)


class ExactPreviewRenderWorker(QObject):
    finished = Signal(str, str)
    failed = Signal(str, str)

    def __init__(self, writer_owner, render_key, payload, output_dir):
        super().__init__()
        self.writer_owner = writer_owner
        self.render_key = str(render_key or "").strip()
        self.payload = dict(payload or {})
        self.output_dir = str(output_dir or "").strip()

    def run(self):
        try:
            os.makedirs(self.output_dir, exist_ok=True)
            pdf_path = os.path.join(self.output_dir, f"{self.render_key}.pdf")
            if is_va_10172_profile(self.payload.get("packet_profile")):
                PacketBuilderTab._write_pdf_doc_for_payload(self.writer_owner, pdf_path, self.payload)
            else:
                docx_path = os.path.join(self.output_dir, f"{self.render_key}.docx")
                PacketBuilderTab._write_word_doc_for_payload(self.writer_owner, docx_path, self.payload)
                convert_docx_to_pdf_via_word(docx_path, pdf_path)
            self.finished.emit(self.render_key, pdf_path)
        except Exception as exc:
            self.failed.emit(self.render_key, str(exc))


class PacketExportWorker(QObject):
    finished = Signal(dict)
    failed = Signal(str)

    def __init__(self, writer_owner, request):
        super().__init__()
        self.writer_owner = writer_owner
        self.request = dict(request or {})

    def run(self):
        try:
            result = PacketBuilderTab._run_export_request(self.writer_owner, self.request)
            self.finished.emit(result)
        except Exception as exc:
            self.failed.emit(str(exc))


def build_bundle_export_plan(base_filename, group_name, format_name):
    normalized_group = str(group_name or "").strip().lower()
    extensions = bundle_extensions_for_format(format_name)
    documents = []
    for profile_name in bundle_profiles_for_group(normalized_group):
        member_stem = bundle_member_filename(profile_name, normalized_group)
        member_extensions = list(extensions)
        if is_va_10172_profile(profile_name) and "pdf" not in member_extensions:
            member_extensions.append("pdf")
        for extension in member_extensions:
            documents.append(
                {
                    "profile_name": profile_name,
                    "filename": f"{member_stem}.{extension}",
                    "extension": extension,
                }
            )
    return {
        "group_name": normalized_group,
        "folder_name": bundle_folder_name(base_filename, normalized_group),
        "documents": documents,
    }


def default_packet_builder_payload():
    return {
        "packet_title": "",
        "packet_profile": PACKET_BUILDER_PROFILES[0],
        "patient_name": "",
        "date_of_birth": "",
        "authorization_number": "",
        "va_icn": "",
        "ordering_doctor": "",
        "provider": "",
        "facility": "",
        "community_facility": "",
        "requested_service": "",
        "diagnosis": "",
        "icd_codes": "",
        "clinical_summary": "",
        "packet_notes": "",
        "scenario_pathology_pattern": "Auto",
        "scenario_conservative_duration": "Auto",
        "scenario_prior_esi_response": "Auto",
        "scenario_functional_emphasis": "Auto",
        "scenario_request_framing": "Auto",
        "scenario_symptom_pattern": "Auto",
        "scenario_conservative_modalities": "Auto",
        "scenario_review_concern": "Auto",
        "scenario_treatment_goals": "Auto",
        "referral_subtitle": "Spine and Pain Evaluation",
        "pcp_request_text": "I am requesting a Community Care referral for pain management and spine evaluation to determine the most appropriate treatment plan.",
        "referral_entry_text": "Specialty spine and pain evaluation and consultation",
        "areas_of_concern": "Cervical, Thoracic, Lumbar, Joints (if applicable)",
        "group_npi": "",
        "fax_number": "",
        "liaison_contact_info": "",
        "consent_form_title": "TELEHEALTH VIRTUAL CONSENT FORM",
        "consent_provider_name": "",
        "consent_initials": "",
        "minor_doctor_name": "",
        "minor_consent_initials": "",
        "service_authorization_name": "",
        "street_address": "",
        "city": "",
        "state": "",
        "zip_code": "",
        "home_phone": "",
        "mobile_phone": "",
        "work_phone": "",
        "email_address": "",
        "ssn": "",
        "drivers_license": "",
        "drivers_license_state": "",
        "appointment_confirmation_method": "Phone",
        "filed_for_disability": "No",
        "condition_work_related": "No",
        "condition_due_to_accident": "No",
        "yes_response_explanation": "",
        "has_attorney": "No",
        "attorney_name": "",
        "attorney_phone": "",
        "emergency_contact_name": "",
        "emergency_contact_relationship": "",
        "emergency_contact_phone": "",
        "primary_insurance_carrier": "",
        "primary_insurance_id": "",
        "primary_insurance_phone": "",
        "secondary_insurance_carrier": "",
        "secondary_insurance_id": "",
        "secondary_insurance_phone": "",
        "pcp_pcm_name": "",
        "pcp_pcm_phone": "",
        "pcp_pcm_fax": "",
        "patient_signature_name": "",
        "patient_signature_image": "",
        "patient_signature_date": "",
        "submission_cover_title": "VA Submission Cover Sheet",
        "submission_date": "",
        "primary_diagnosis_code": "",
        "included_virtual_consent_form": True,
        "included_va_form_10_10172": True,
        "included_seoc_request": True,
        "included_consult_request": True,
        "included_lomn": True,
        "included_clinical_notes": True,
        "included_mri_report": True,
        "submitting_office": "",
        "office_staff_name": "",
        "office_staff_signature": "",
        "office_staff_signature_image": "",
        "date_reviewed": "",
        "seoc_request_date": "",
        "va_medical_center_name": "",
        "last_four_ssn": "",
        "episode_diagnosis": "",
        "episode_icd_code": "",
        "seoc_scope_text": "",
        "estimated_duration_text": "Thirty to ninety days, including routine post-procedure follow-up related only to the authorized episode of care.",
        "clinical_objectives": "Reduce pain severity\nImprove functional capacity\nSupport activity tolerance related to the documented condition",
        "seoc_continuity_text": (
            "Upon completion of the authorized episode, a treatment summary and clinical outcome update will be provided "
            "to the referring VA provider. Any additional or unrelated care will require separate evaluation and authorization."
        ),
        "provider_credentials": "",
        "provider_specialty": "",
        "provider_npi": "",
        "practice_name": "",
        "provider_phone": "",
        "provider_fax": "",
        "seoc_include_preprocedure_eval": True,
        "seoc_include_annulargram": True,
        "seoc_include_fibrin_injection": True,
        "seoc_include_follow_up": True,
        "lmn_request_date": "",
        "lmn_va_claim_number": "",
        "lmn_primary_diagnosis": "",
        "lmn_secondary_diagnosis": "",
        "lmn_clinical_summary": "",
        "lmn_mri_date": "",
        "lmn_mri_findings": "",
        "lmn_conservative_duration": "",
        "lmn_include_physical_therapy": True,
        "lmn_include_nsaids": True,
        "lmn_include_activity_modification": True,
        "lmn_include_home_exercise": True,
        "lmn_include_epidural_steroid_injections": True,
        "lmn_medical_necessity_statement": "Based on chronic symptoms, failed conservative management, functional limitation, and imaging findings that correlate with the clinical presentation, the requested intervention is medically reasonable and necessary.",
        "lmn_indication_reduce_pain": True,
        "lmn_indication_improve_function": True,
        "lmn_indication_prevent_degeneration": True,
        "lmn_indication_reduce_opioid_reliance": True,
        "lmn_indication_prevent_surgery": True,
        "lmn_risk_statement": "Without appropriate treatment, the patient remains at risk for persistent pain, worsening functional limitation, and continued impairment of daily activities.",
        "lmn_reasonable_necessary_statement": "The requested care is consistent with the documented diagnosis, the failure of conservative treatment, and the current clinical findings.",
        "lmn_contact_statement": "Additional supporting documentation can be provided upon request.",
        "consult_request_date": "",
        "consult_va_claim_number": "",
        "consult_referring_va_provider": "",
        "consult_reason_text": "Authorization is requested for specialty consultation and treatment planning related to documented lumbar disc pathology and persistent pain despite conservative management.",
        "consult_primary_diagnosis": "",
        "consult_secondary_diagnosis": "",
        "consult_symptom_axial_pain": True,
        "consult_symptom_activity_exacerbation": True,
        "consult_symptom_reduced_tolerance": True,
        "consult_symptom_functional_impairment": True,
        "consult_mri_date": "",
        "consult_mri_findings": "",
        "consult_conservative_duration": "",
        "consult_include_physical_therapy": True,
        "consult_include_nsaids": True,
        "consult_include_activity_modification": True,
        "consult_include_home_exercise": True,
        "consult_include_interventional_history": True,
        "consult_include_pain_management_consultation": True,
        "consult_include_procedural_planning": True,
        "consult_include_annulargram": True,
        "consult_include_fibrin_injection": True,
        "consult_include_follow_up": True,
        "consult_fibrin_levels": "",
        "consult_scope_exclusion_text": "This request is limited to evaluation, procedural planning, the indicated intervention if clinically appropriate, and standard post-procedure follow-up related only to the documented lumbar condition.",
        "consult_medical_rationale_text": "The requested evaluation and treatment pathway is supported by the patient’s ongoing symptoms, functional impairment, failure of conservative care, and imaging findings that correlate with the clinical presentation.",
        "consult_goal_pain_reduction": True,
        "consult_goal_functional_improvement": True,
        "consult_goal_reduce_analgesics": True,
        "consult_goal_prevent_surgery": True,
        "consult_risk_without_treatment": "Without appropriate specialty evaluation and treatment, the patient may continue to experience persistent pain, reduced function, and progression of disability.",
        "consult_duration_scope_text": "This consultation and treatment request is limited to a defined evaluation and treatment pathway rather than open-ended pain management.",
        "consult_contact_statement": "Please contact the treating office if additional records or clarification are needed for review.",
        "provider_address": "",
        "provider_email": "",
        "clinical_doc_title": "Clinical Documentation Template",
        "clinical_doc_chief_complaint": "",
        "clinical_doc_duration_gt_3m": True,
        "clinical_doc_duration_gt_6m": False,
        "clinical_doc_duration_gt_12m": False,
        "clinical_doc_exact_duration": "",
        "clinical_doc_pain_axial": True,
        "clinical_doc_pain_discogenic": True,
        "clinical_doc_pain_activity_exacerbation": True,
        "clinical_doc_pain_sitting_intolerance": True,
        "clinical_doc_pain_standing_intolerance": True,
        "clinical_doc_pain_bending_lifting": True,
        "clinical_doc_pain_severity": "",
        "clinical_doc_limit_occupational": True,
        "clinical_doc_limit_prolonged_sitting": True,
        "clinical_doc_limit_prolonged_standing": True,
        "clinical_doc_limit_ambulation": False,
        "clinical_doc_limit_household": True,
        "clinical_doc_limit_sleep": True,
        "clinical_doc_functional_impact": "",
        "clinical_doc_conservative_pt": True,
        "clinical_doc_conservative_home_exercise": True,
        "clinical_doc_conservative_nsaids": True,
        "clinical_doc_conservative_non_opioid": True,
        "clinical_doc_conservative_activity_modification": True,
        "clinical_doc_conservative_esi": True,
        "clinical_doc_conservative_other_interventional": False,
        "clinical_doc_conservative_duration": "",
        "clinical_doc_esi_response": "",
        "clinical_doc_mri_date": "",
        "clinical_doc_imaging_annular_tear": True,
        "clinical_doc_imaging_disc_degeneration": True,
        "clinical_doc_imaging_disc_protrusion": False,
        "clinical_doc_imaging_disc_displacement": False,
        "clinical_doc_affected_levels": "",
        "clinical_doc_primary_diagnosis": "",
        "clinical_doc_secondary_diagnosis": "",
        "clinical_doc_assessment_summary": "",
        "clinical_doc_treatment_plan_intro": "",
        "clinical_doc_plan_diagnostic_confirmation": True,
        "clinical_doc_plan_intradiscal_intervention": True,
        "clinical_doc_plan_follow_up": True,
        "clinical_doc_plan_exclusion": "This plan does not request open-ended medication management or unrelated pain treatment outside the documented lumbar condition.",
        "clinical_doc_physician_narrative": "",
        "va10172_va_facility_address": "",
        "va10172_ordering_provider_office_address": "",
        "va10172_is_ihs_provider": "No",
        "va10172_ordering_provider_phone": "",
        "va10172_ordering_provider_fax": "",
        "va10172_ordering_provider_secure_email": "",
        "va10172_care_needed_within_48_hours": "No",
        "va10172_is_continuation_of_care": "No",
        "va10172_referral_to_specialty": "No",
        "va10172_referral_specialty_text": "",
        "va10172_diagnosis_description": "",
        "va10172_requested_cpt_hcpcs_code": "",
        "va10172_description_cpt_hcpcs_code": "",
        "va10172_geriatric_care_option": "None",
        "va10172_reason_for_request": "",
        "va10172_ordering_provider_name_printed": "",
        "va10172_ordering_provider_npi": "",
        "va10172_signature_text": "",
        "va10172_signature_image": "",
        "va10172_today_date": "",
        "wording_assist_state": {},
    }


def is_referral_request_profile(profile_name):
    return str(profile_name or "").strip().lower() == "community care referral request"


def is_virtual_consent_profile(profile_name):
    return str(profile_name or "").strip().lower() == "virtual consent form"


def is_submission_cover_profile(profile_name):
    return str(profile_name or "").strip().lower() == "submission cover sheet"


def is_seoc_request_profile(profile_name):
    normalized = str(profile_name or "").strip().lower()
    return normalized in {
        "seoc request template",
        "single episode of care request template",
    }


def is_lomn_profile(profile_name):
    normalized = str(profile_name or "").strip().lower()
    return normalized in {
        "letter of medical necessity template",
        "letter of medical necessity",
    }


def is_consult_request_profile(profile_name):
    normalized = str(profile_name or "").strip().lower()
    return normalized in {
        "consultation request template",
        "consultation & treatment request template",
        "consultation and treatment request",
    }


def is_clinical_documentation_profile(profile_name):
    normalized = str(profile_name or "").strip().lower()
    return normalized in {
        "clinical documentation template",
        "clinical notes template",
        "clinical documentation",
        "clinical notes",
    }


def is_va_10172_profile(profile_name):
    normalized = str(profile_name or "").strip().lower()
    return normalized in {
        "va form 10-10172",
        "va form 10172",
        "10-10172",
    }


def default_title_for_profile(profile_name):
    if is_referral_request_profile(profile_name):
        return "Community Care Referral Request"
    if is_virtual_consent_profile(profile_name):
        return "TELEHEALTH VIRTUAL CONSENT FORM"
    if is_submission_cover_profile(profile_name):
        return "VA Submission Cover Sheet"
    if is_seoc_request_profile(profile_name):
        return "Single Episode of Care (SEOC) Request"
    if is_lomn_profile(profile_name):
        return "LETTER OF MEDICAL NECESSITY"
    if is_consult_request_profile(profile_name):
        return "CONSULTATION AND TREATMENT REQUEST"
    if is_clinical_documentation_profile(profile_name):
        return "Clinical Documentation Template"
    if is_va_10172_profile(profile_name):
        return "VA Form 10-10172"
    return ""


LEGACY_PACKET_TEXT_REPLACEMENTS = {
    "clinical_doc_treatment_plan_intro": {
        "The patient is an appropriate candidate for TrueDisc intradiscal biologic repair under a Single Episode of Care (SEOC) model.": (
            "The patient is an appropriate candidate for focused spine evaluation and consideration of the requested intervention if confirmed clinically appropriate."
        ),
    },
}


TEXT_ARTIFACT_REPLACEMENTS = {
    "â€™": "'",
    "â€œ": '"',
    "â€\x9d": '"',
    "â€“": "-",
    "â€”": "-",
    "â€¦": "...",
    "\u2018": "'",
    "\u2019": "'",
    "\u201c": '"',
    "\u201d": '"',
    "\u2013": "-",
    "\u2014": "-",
    "\u2022": "-",
    "\u2611": "[X]",
    "\u2610": "[ ]",
    "\ufb00": "ff",
    "\ufb01": "fi",
    "\ufb02": "fl",
    "\ufb03": "ffi",
    "\ufb04": "ffl",
}


def sanitize_packet_builder_text(value):
    text = str(value or "")
    if not text:
        return ""
    for source, target in TEXT_ARTIFACT_REPLACEMENTS.items():
        text = text.replace(source, target)
    text = unicodedata.normalize("NFKC", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)
    return text.strip()


def export_checkbox_marker(enabled):
    return "[X]" if enabled else "[ ]"


def normalize_packet_builder_payload(payload):
    packet = dict(default_packet_builder_payload())
    packet.update(dict(payload or {}))
    for field_name, value in list(packet.items()):
        if isinstance(value, str):
            packet[field_name] = sanitize_packet_builder_text(value)
    for field_name, replacements in LEGACY_PACKET_TEXT_REPLACEMENTS.items():
        current_value = str(packet.get(field_name) or "").strip()
        if current_value in replacements:
            packet[field_name] = replacements[current_value]
    return apply_packet_builder_shared_field_sync(packet)


def _normalize_packet_builder_document_markup(markup):
    replacements = {
        "font-family:Segoe UI; color:#E8F1FC; line-height:1.55;": "font-family:Calibri, Arial, sans-serif; color:#111827; line-height:1.6; font-size:11pt;",
        "font-family:Segoe UI; color:#E8F1FC; line-height:1.5;": "font-family:Calibri, Arial, sans-serif; color:#111827; line-height:1.6; font-size:11pt;",
        "#F2F5F9": "#111827",
        "#E8F1FC": "#111827",
        "#C8D8E8": "#374151",
        "#9BB3CC": "#6B7280",
        "#8FA6C1": "#4B5563",
        "#69BCFF": "#1B4F72",
        "#0F1823": "#FFFFFF",
        "#13202E": "#FBF7E8",
        "#243446": "#D6DCE5",
        "#28415A": "#D6DCE5",
        "#F7E3AA": "#8A6A0A",
    }
    rendered = str(markup or "")
    for source, target in replacements.items():
        rendered = rendered.replace(source, target)
    return re.sub(
        r"(?<![A-Za-z])Pending(?![A-Za-z])",
        "<span style='color:#64748B; font-style:italic; font-weight:600;'>Pending</span>",
        rendered,
    )


def render_packet_builder_document_preview(markup):
    rendered = _normalize_packet_builder_document_markup(markup)
    return (
        "<div style='background:#0D1520; padding:18px 16px;'>"
        "<div style='width:100%; max-width:920px; margin:0 auto; box-sizing:border-box; background:#FFFFFF; border:1px solid #D6DCE5; "
        "box-shadow:0 18px 42px rgba(0,0,0,0.18); padding:32px 34px 38px 34px; overflow-wrap:anywhere; word-break:normal;'>"
        + rendered
        + "</div></div>"
    )


def render_packet_builder_document_export(markup):
    rendered = _normalize_packet_builder_document_markup(markup)
    return (
        "<html><body style='margin:0; background:#FFFFFF;'>"
        "<div style='background:#FFFFFF; padding:18px 22px 24px 22px; color:#111827; overflow-wrap:anywhere; word-break:normal;'>"
        + rendered
        + "</div></body></html>"
    )


def build_packet_builder_sections(payload):
    packet = dict(default_packet_builder_payload())
    packet.update(payload or {})
    return [
        (
            "Packet Identity",
            [
                ("Packet Title", packet.get("packet_title")),
                ("Packet Profile", packet.get("packet_profile")),
            ],
        ),
        (
            "Patient And VA Information",
            [
                ("Patient Name", packet.get("patient_name")),
                ("Date of Birth", packet.get("date_of_birth")),
                ("Authorization Number", packet.get("authorization_number")),
                ("VA ICN", packet.get("va_icn")),
            ],
        ),
        (
            "Clinical Routing",
            [
                ("Ordering Doctor", packet.get("ordering_doctor")),
                ("Provider", packet.get("provider")),
                ("Facility", packet.get("facility")),
                ("Community Facility", packet.get("community_facility")),
                ("Requested Service", packet.get("requested_service")),
            ],
        ),
        (
            "Clinical Support",
            [
                ("Diagnosis", packet.get("diagnosis")),
                ("ICD Codes", packet.get("icd_codes")),
                ("Clinical Summary", packet.get("clinical_summary")),
                ("Packet Notes", packet.get("packet_notes")),
            ],
        ),
    ]


def build_packet_builder_document_markup(payload):
    packet = dict(default_packet_builder_payload())
    packet.update(payload or {})

    if is_referral_request_profile(packet.get("packet_profile")):
        return build_referral_request_preview_html(packet)
    if is_virtual_consent_profile(packet.get("packet_profile")):
        return build_virtual_consent_preview_html(packet)
    if is_submission_cover_profile(packet.get("packet_profile")):
        return build_submission_cover_preview_html(packet)
    if is_seoc_request_profile(packet.get("packet_profile")):
        return build_seoc_request_preview_markup(packet)
    if is_lomn_profile(packet.get("packet_profile")):
        return build_lomn_preview_html(packet)
    if is_consult_request_profile(packet.get("packet_profile")):
        return build_consult_request_preview_html(packet)
    if is_clinical_documentation_profile(packet.get("packet_profile")):
        return build_clinical_documentation_preview_html(packet)
    if is_va_10172_profile(packet.get("packet_profile")):
        return build_va_10172_preview_html(packet)

    title = html.escape(packet.get("packet_title") or "Untitled Packet Draft")
    summary_rows = []
    for section_title, rows in build_packet_builder_sections(packet):
        row_html = []
        for label, value in rows:
            rendered = html.escape(str(value or "Pending").strip() or "Pending")
            row_html.append(
                f"<tr><td style='padding:8px 10px; color:#8FA6C1; width:32%; vertical-align:top;'>{html.escape(label)}</td>"
                f"<td style='padding:8px 10px; color:#E8F1FC;'>{rendered}</td></tr>"
            )
        summary_rows.append(
            f"<div style='margin-top:16px;'>"
            f"<div style='font-size:15px; font-weight:700; color:#69BCFF; margin-bottom:6px;'>{html.escape(section_title)}</div>"
            f"<table style='width:100%; border-collapse:collapse; background:#0F1823; border:1px solid #243446; border-radius:10px; table-layout:fixed;'>"
            + "".join(row_html)
            + "</table></div>"
        )

    return (
        "<div style='font-family:Segoe UI; color:#E8F1FC; line-height:1.5;'>"
        f"<div style='font-size:24px; font-weight:800; color:#F2F5F9;'>{title}</div>"
        f"<div style='font-size:13px; color:#9BB3CC; margin-top:4px;'>Profile: {packet.get('packet_profile') or 'Pending'}</div>"
        "<div style='margin-top:12px; font-size:13px; color:#C8D8E8;'>"
        "This is the developer-side packet-builder preview. It is a clean drafting shell for packet creation and export."
        "</div>"
        + "".join(summary_rows)
        + "</div>"
    )


def build_packet_builder_inconsistency_banner(payload):
    packet = apply_packet_builder_shared_field_sync(default_packet_builder_payload() | dict(payload or {}))
    messages = build_packet_inconsistency_messages(packet)
    if not messages:
        return ""
    items = "".join(f"<li style='margin-bottom:6px;'>{html.escape(message)}</li>" for message in messages)
    return (
        "<div style='margin-bottom:14px; padding:10px 12px; background:#FFF7ED; border:1px solid #F1C996; color:#7C3400; font-size:11pt; line-height:1.35;'>"
        "<div style='font-weight:700;'>Consistency Warning</div>"
        "<div style='margin-top:4px;'>Some repeated values on this form do not match the shared packet header.</div>"
        f"<ul style='margin:6px 0 0 18px; padding:0;'>{items}</ul>"
        "</div>"
    )


def build_packet_builder_preview_html(payload):
    banner = build_packet_builder_inconsistency_banner(payload)
    wording_banner = build_wording_assist_banner(payload)
    return render_packet_builder_document_preview(banner + wording_banner + build_packet_builder_document_markup(payload))


def build_packet_builder_export_html(payload):
    return render_packet_builder_document_export(build_packet_builder_document_markup(payload))


def build_referral_request_preview_html(packet):
    title = html.escape(packet.get("packet_title") or "Community Care Referral Request")
    subtitle = html.escape(packet.get("referral_subtitle") or "Spine & Pain Evaluation")
    pcp_request = html.escape(packet.get("pcp_request_text") or "Pending")
    referral_entry = html.escape(packet.get("referral_entry_text") or "Pending")
    areas_of_concern = html.escape(packet.get("areas_of_concern") or "Pending")
    group_npi = html.escape(packet.get("group_npi") or "Pending")
    fax_number = html.escape(packet.get("fax_number") or "Pending")
    contact_info = html.escape(packet.get("liaison_contact_info") or "Pending").replace("\n", "<br/>")

    return (
        "<div style='font-family:Calibri, Arial, sans-serif; color:#111827; line-height:1.55;'>"
        f"<div style='text-align:center; font-size:19pt; font-weight:700; letter-spacing:0.2px;'>{title}</div>"
        f"<div style='text-align:center; font-size:12.5pt; font-weight:700; color:#1F3A5F; margin-top:6px;'>{subtitle}</div>"
        "<p style='margin-top:18px;'>"
        "Please provide this request to your VA Primary Care Provider to begin the Community Care referral process. "
        "This document supports specialty spine and pain evaluation."
        "</p>"
        "<div style='margin-top:16px; padding:12px 14px; background:#F7F9FC; border:1px solid #D5DCE5;'>"
        "<div style='font-weight:700;'>Important:</div>"
        "<div style='margin-top:6px;'>This request is for specialty evaluation and consultation. It is not a request for a specific procedure.</div>"
        "</div>"
        "<div style='margin-top:18px; font-size:11pt; font-weight:700; text-transform:uppercase; letter-spacing:0.4px; color:#1F3A5F;'>Instructions for PCP Referral Entry</div>"
        "<ol style='margin-top:10px; padding-left:22px;'>"
        "<li style='margin-bottom:8px;'>Schedule a visit with your VA Primary Care Provider.</li>"
        f"<li style='margin-bottom:8px;'>Ask the provider to enter the referral as: <strong>{referral_entry}</strong>.</li>"
        f"<li style='margin-bottom:8px;'>Use the following request language: &ldquo;{pcp_request}&rdquo;</li>"
        f"<li style='margin-bottom:8px;'>Include all relevant diagnoses and areas of concern: <strong>{areas_of_concern}</strong>.</li>"
        "<li style='margin-bottom:8px;'>Send any relevant imaging and clinical records with the referral.</li>"
        "</ol>"
        "<div style='margin-top:18px; font-size:11pt; font-weight:700; text-transform:uppercase; letter-spacing:0.4px; color:#1F3A5F;'>Routing Information</div>"
        "<table style='width:100%; border-collapse:collapse; table-layout:fixed; margin-top:10px;'>"
        f"<tr><td style='width:28%; padding:8px 10px; border:1px solid #D5DCE5; font-weight:700; vertical-align:top;'>Group NPI</td><td style='padding:8px 10px; border:1px solid #D5DCE5; vertical-align:top;'>{group_npi}</td></tr>"
        f"<tr><td style='padding:8px 10px; border:1px solid #D5DCE5; font-weight:700; vertical-align:top;'>Fax Number</td><td style='padding:8px 10px; border:1px solid #D5DCE5; vertical-align:top;'>{fax_number}</td></tr>"
        "</table>"
        "<div style='margin-top:18px; font-size:11pt; font-weight:700; text-transform:uppercase; letter-spacing:0.4px; color:#1F3A5F;'>Veteran Liaison Contact</div>"
        f"<div style='margin-top:8px;'>{contact_info}</div>"
        "</div>"
    )


def build_seoc_request_preview_markup(packet):
    def filled(value, blank_length):
        raw = str(value or "").strip()
        return html.escape(raw) if raw else "_" * blank_length

    def bullet_lines(items):
        return "".join(f"<div style='margin-top:6px;'>- {html.escape(item)}</div>" for item in items)

    request_date = filled(packet.get("seoc_request_date"), 12)
    va_medical_center_name = filled(packet.get("va_medical_center_name"), 22)
    patient_name = filled(packet.get("patient_name"), 18)
    dob = filled(packet.get("date_of_birth"), 12)
    last_four_ssn = filled(packet.get("last_four_ssn"), 8)
    episode_diagnosis = filled(packet.get("episode_diagnosis"), 44)
    episode_icd_code = str(packet.get("episode_icd_code") or "").strip()
    estimated_duration = filled(packet.get("estimated_duration_text"), 18)
    provider_name = filled(packet.get("provider") or packet.get("ordering_doctor"), 18)
    provider_credentials = filled(packet.get("provider_credentials"), 8)
    provider_specialty = filled(packet.get("provider_specialty"), 18)
    provider_npi = filled(packet.get("provider_npi"), 14)
    practice_name = filled(packet.get("practice_name") or packet.get("community_facility"), 20)
    provider_phone = filled(packet.get("provider_phone"), 14)
    provider_fax = filled(packet.get("provider_fax"), 14)
    provider_header = provider_name
    raw_credentials = str(packet.get("provider_credentials") or "").strip()
    if raw_credentials:
        provider_header += f", {html.escape(raw_credentials)}"
    provider_header = provider_name
    raw_credentials = str(packet.get("provider_credentials") or "").strip()
    if raw_credentials:
        provider_header += f", {html.escape(raw_credentials)}"
    provider_header = provider_name
    raw_credentials = str(packet.get("provider_credentials") or "").strip()
    if raw_credentials:
        provider_header += f", {html.escape(raw_credentials)}"

    scope_items = []
    if packet.get("seoc_include_preprocedure_eval"):
        scope_items.append("Pre-procedure evaluation and procedural planning")
    if packet.get("seoc_include_annulargram"):
        scope_items.append("Diagnostic Annulargram")
    if packet.get("seoc_include_fibrin_injection"):
        scope_items.append("Inter Annular Fibrin Injections if indicated")
    if packet.get("seoc_include_follow_up"):
        scope_items.append("Standard post-procedure follow-up visit(s) within the global postoperative period")
    if not scope_items:
        scope_items.append("________________________________________")

    objective_lines = [line.strip() for line in str(packet.get("clinical_objectives") or "").splitlines() if line.strip()]
    if not objective_lines:
        objective_lines = ["________________________________________"]
    scope_summary = str(packet.get("seoc_scope_text") or "").strip()
    continuity_text = str(packet.get("seoc_continuity_text") or "").strip()

    diagnosis_line = episode_diagnosis
    provider_header = provider_name
    raw_credentials = str(packet.get("provider_credentials") or "").strip()
    if raw_credentials:
        provider_header += f", {html.escape(raw_credentials)}"

    return (
        "<div style='font-family:Calibri, Arial, sans-serif; color:#111827; line-height:1.5;'>"
        f"<div>{request_date}</div>"
        "<div style='margin-top:10px;'>Department of Veterans Affairs<br/>Community Care Office<br/>"
        f"{va_medical_center_name}</div>"
        "<div style='margin-top:16px; font-weight:700;'>RE: Single Episode of Care (SEOC) Request</div>"
        f"<div style='margin-top:4px;'>Veteran Name: {patient_name}</div>"
        f"<div>DOB: {dob}</div>"
        f"<div>Last Four SSN: {last_four_ssn}</div>"
        "<div style='margin-top:12px;'>This request is for authorization of a defined, time-limited Single Episode of Care (SEOC) for the condition identified below.</div>"
        "<div style='margin-top:12px; font-weight:700;'>Episode Diagnosis / Condition</div>"
        f"<div style='margin-top:4px;'>Diagnosis: {diagnosis_line}</div>"
        f"<div>ICD-10 Code: {html.escape(episode_icd_code) if episode_icd_code else '____________'}</div>"
        "<div style='margin-top:12px; font-weight:700;'>Scope of Requested Episode:</div>"
        "<div style='margin-top:4px;'>This SEOC includes only the following services directly related to treatment of the above diagnosis:</div>"
        + bullet_lines(scope_items)
        + f"<div style='margin-top:8px;'>{html.escape(scope_summary) if scope_summary else 'No additional pain management services, long-term medication management, or unrelated spine care are requested under this episode.'}</div>"
        + "<div style='margin-top:12px; font-weight:700;'>Estimated Duration of Episode:</div>"
        + f"<div style='margin-top:4px;'>Expected episode duration: {estimated_duration}</div>"
        + "<div style='margin-top:12px; font-weight:700;'>Clinical Objective:</div>"
        + bullet_lines(objective_lines)
        + "<div style='margin-top:12px; font-weight:700;'>Continuity of Care:</div>"
        + f"<div style='margin-top:4px;'>{html.escape(continuity_text) if continuity_text else 'Upon completion of the episode, a summary of treatment rendered and clinical outcome will be forwarded to the referring VA provider. Any need for further care beyond this defined episode will require separate evaluation and authorization.'}</div>"
        + "<div style='margin-top:16px;'>Sincerely,</div>"
        + f"<div style='margin-top:10px;'>{provider_header}<br/>{provider_specialty}<br/>NPI Number: {provider_npi}<br/>{practice_name}<br/>{provider_phone}<br/>{provider_fax}</div>"
        + "</div>"
    )


def build_virtual_consent_preview_html(packet):
    title = html.escape(packet.get("packet_title") or packet.get("consent_form_title") or "TELEHEALTH VIRTUAL CONSENT FORM")
    def filled(value, blank_length):
        raw = str(value or "").strip()
        return html.escape(raw) if raw else "_" * blank_length

    def checkbox_row(label, value):
        selected = str(value or "").strip().lower()
        yes_box = "[X]" if selected == "yes" else "[ ]"
        no_box = "[X]" if selected == "no" else "[ ]"
        return (
            f"<div style='margin-top:8px;'><strong>{html.escape(label)}</strong> "
            f"{yes_box} Yes&nbsp;&nbsp;&nbsp;{no_box} No</div>"
        )

    def options_row(label, selected_value, options):
        selected = str(selected_value or "").strip().lower()
        rendered = []
        for option in options:
            marker = "[X]" if selected == option.lower() else "[ ]"
            rendered.append(f"{marker} {html.escape(option)}")
        return (
            f"<div style='margin-top:8px;'><strong>{html.escape(label)}</strong> "
            + "&nbsp;&nbsp;&nbsp;".join(rendered)
            + "</div>"
        )

    def two_line_table(rows, widths):
        colgroup = "".join(f"<col style='width:{width}%'>" for width in widths)
        blocks = []
        for labels, values in rows:
            label_cells = "".join(
                f"<td style='padding:0 16px 3px 0; font-weight:700; vertical-align:bottom;'>{html.escape(label)}</td>"
                for label in labels
            )
            value_cells = "".join(
                f"<td style='padding:0 16px 10px 0; vertical-align:top;'>{value}</td>"
                for value in values
            )
            blocks.append(
                "<table style='width:100%; border-collapse:collapse; table-layout:fixed; margin-top:8px;'>"
                f"<colgroup>{colgroup}</colgroup>"
                f"<tr>{label_cells}</tr><tr>{value_cells}</tr></table>"
            )
        return "".join(blocks)

    telehealth_lines = [
        "I understand that through an interactive video connection, telehealth might be used to perform my consultation and that we may use telehealth to connect while working together.",
        "I understand the benefits and risks involved with telehealth technology:",
        "I do not need to travel to the consultation location.",
        "I have access to a specialist through this consultation.",
        "There may be interruptions, unauthorized access and technical difficulties and my health care provider(s) or myself can discontinue the telemedicine consult/visit if it is felt that the videoconferencing connections are not adequate for the situation.",
        "My healthcare information may be shared with other individuals for scheduling and billing purposes.",
        "I also understand that other individuals may need to use the telehealth platform and that reasonable steps will be taken to maintain confidentiality of the information obtained.",
        "I understand the initial virtual meet and greet discussion is primarily intended to determine any and all treatment I might pursue, in addition to the best approach for scheduling these treatments. The virtual discussion will be followed by individual cost estimates (if applicable).",
        "I have read this document in its entirety and understand the risks and benefits of telehealth consultation/post op visits and have had my questions explained. I hereby consent to participate in telehealth sessions under the conditions described in this document.",
    ]

    patient_name = filled(packet.get("patient_name"), 28)
    dob = filled(packet.get("date_of_birth"), 14)
    mobile_phone = filled(packet.get("mobile_phone") or packet.get("phone"), 14)
    street_address = filled(packet.get("street_address"), 24)
    city = filled(packet.get("city"), 12)
    state = filled(packet.get("state"), 6)
    zip_code = filled(packet.get("zip_code"), 10)
    home_phone = filled(packet.get("home_phone"), 14)
    work_phone = filled(packet.get("work_phone"), 14)
    email_address = filled(packet.get("email_address"), 22)
    ssn = filled(packet.get("ssn"), 14)
    drivers_license = filled(packet.get("drivers_license"), 14)
    drivers_license_state = filled(packet.get("drivers_license_state"), 10)
    yes_explanation = filled(packet.get("yes_response_explanation"), 54)
    attorney_name = filled(packet.get("attorney_name"), 22)
    attorney_phone = filled(packet.get("attorney_phone"), 20)
    emergency_contact_name = filled(packet.get("emergency_contact_name"), 18)
    emergency_contact_relationship = filled(packet.get("emergency_contact_relationship"), 14)
    emergency_contact_phone = filled(packet.get("emergency_contact_phone"), 16)
    primary_insurance_carrier = filled(packet.get("primary_insurance_carrier"), 22)
    primary_insurance_id = filled(packet.get("primary_insurance_id"), 14)
    primary_insurance_phone = filled(packet.get("primary_insurance_phone"), 16)
    secondary_insurance_carrier = filled(packet.get("secondary_insurance_carrier"), 22)
    secondary_insurance_id = filled(packet.get("secondary_insurance_id"), 14)
    secondary_insurance_phone = filled(packet.get("secondary_insurance_phone"), 16)
    pcp_pcm_name = filled(packet.get("pcp_pcm_name"), 22)
    pcp_pcm_phone = filled(packet.get("pcp_pcm_phone"), 14)
    pcp_pcm_fax = filled(packet.get("pcp_pcm_fax"), 14)
    consent_provider_name = filled(packet.get("consent_provider_name"), 18)
    consent_initials = filled(packet.get("consent_initials"), 6)
    minor_doctor_name = filled(packet.get("minor_doctor_name"), 18)
    minor_consent_initials = filled(packet.get("minor_consent_initials"), 6)
    service_authorization_name = filled(packet.get("service_authorization_name"), 18)
    patient_signature_name = filled(packet.get("patient_signature_name") or packet.get("patient_name"), 30)
    patient_signature_image_html = _signature_image_html(packet, "patient_signature_name", width_px=240, height_px=56)
    patient_signature_date = filled(packet.get("patient_signature_date"), 12)

    return (
        "<div style='font-family:Calibri, Arial, sans-serif; color:#111827; line-height:1.45;'>"
        f"<div style='text-align:center; font-size:18pt; font-weight:700;'>{title}</div>"
        "<div style='margin-top:16px; font-weight:700;'>Demographic Information:</div>"
        + two_line_table(
            [
                (
                    ["Full Name (Shown on your Ins Card):", "Date of Birth:", "Phone:"],
                    [patient_name, dob, mobile_phone],
                ),
                (
                    ["Street Address:", "City:", "State:", "Zip:"],
                    [street_address, city, state, zip_code],
                ),
                (
                    ["Home Phone:", "Mobile Phone:", "Work Phone:"],
                    [home_phone, mobile_phone, work_phone],
                ),
                (
                    ["Email Address:", "SSN:", "DL:", "State:"],
                    [email_address, ssn, drivers_license, drivers_license_state],
                ),
            ],
            [42, 18, 20, 20],
        )
        + options_row("Appointment Confirmation Method:", packet.get("appointment_confirmation_method"), ["Phone", "Text", "Email"])
        + checkbox_row("Have you filed for Disability?", packet.get("filed_for_disability"))
        + checkbox_row("Is your condition work related?", packet.get("condition_work_related"))
        + checkbox_row("Is your condition due to an accident?", packet.get("condition_due_to_accident"))
        + "<div style='margin-top:10px; font-weight:700;'>Please explain any responses you answered 'Yes' to:</div>"
        + f"<div style='margin-top:4px;'>{yes_explanation}</div>"
        + checkbox_row("Do you have an Attorney?", packet.get("has_attorney"))
        + two_line_table(
            [
                (
                    ["Attorney Full Name:", "Attorney Phone Number:"],
                    [attorney_name, attorney_phone],
                ),
                (
                    ["Emergency Contact:", "Relationship:", "Phone:"],
                    [emergency_contact_name, emergency_contact_relationship, emergency_contact_phone],
                ),
            ],
            [42, 28, 30],
        )
        + "<div style='margin-top:10px; font-weight:700;'>Insurance Information - If your condition is due to a work injury or other Personal Injury/Accident, there is a separate form you will need to complete. See accident information form.</div>"
        + two_line_table(
            [
                (
                    ["Primary Insurance Carrier:", "ID Number:", "Insurance Phone:"],
                    [primary_insurance_carrier, primary_insurance_id, primary_insurance_phone],
                ),
                (
                    ["Secondary Insurance Carrier:", "ID Number:", "Insurance Phone:"],
                    [secondary_insurance_carrier, secondary_insurance_id, secondary_insurance_phone],
                ),
                (
                    ["PCP/PCM:", "Phone:", "Fax:"],
                    [pcp_pcm_name, pcp_pcm_phone, pcp_pcm_fax],
                ),
            ],
            [44, 22, 34],
        )
        + "<div style='page-break-before:always;'></div>"
        + "<div style='margin-top:8px; font-weight:700;'>CONSENT FOR MEDICAL CARE AND TREATMENT</div>"
        + (
            "<p style='margin-top:10px;'>"
            f"I, {patient_name} hereby agree and give my consent for {consent_provider_name} to furnish medical care and treatment considered necessary and proper in evaluating or treating my physical condition. {consent_initials} (initial)"
            "</p>"
        )
        + "<div style='margin-top:10px; font-weight:700;'>FOR MINORS ONLY CONSENT FOR CARE:</div>"
        + (
            "<p style='margin-top:6px;'>"
            f"As parent and/or legal guardian, I authorize {minor_doctor_name} (Doctors Name) to treat the minor patient named in the attached form while I am not present. {minor_consent_initials} (Parent/Guardian initial)"
            "</p>"
        )
        + "<div style='margin-top:14px; font-weight:700;'>TELEHEALTH CONSENT</div>"
        + "".join(f"<p style='margin-top:8px;'>{html.escape(line)}</p>" for line in telehealth_lines)
        + (
            "<p style='margin-top:12px;'>"
            f"By signing below, I agree that all the above information is correct, and that I authorize {service_authorization_name} to provide me with medical services and to furnish my physician, insurance company or attorney information concerning my injury and/or treatment."
            "</p>"
        )
        + two_line_table(
            [
                (
                    ["Veteran / Patient Signature (Parent/Guardian if applicable):", "Date:"],
                    [patient_signature_image_html or patient_signature_name, patient_signature_date],
                ),
            ],
            [72, 28],
        )
        + "</div>"
    )


def build_submission_cover_preview_html(packet):
    title = html.escape(packet.get("packet_title") or packet.get("submission_cover_title") or "VA Submission Cover Sheet")
    patient_name = html.escape(packet.get("patient_name") or "")
    dob = html.escape(packet.get("date_of_birth") or "")
    facility = html.escape(packet.get("facility") or "")
    ordering_provider = html.escape(packet.get("ordering_doctor") or packet.get("provider") or "")
    submission_date = html.escape(packet.get("submission_date") or "")
    diagnosis_code = html.escape(packet.get("primary_diagnosis_code") or packet.get("icd_codes") or "")
    submitting_office = html.escape(packet.get("submitting_office") or "")
    office_staff_name = html.escape(packet.get("office_staff_name") or "")
    office_staff_signature = html.escape(packet.get("office_staff_signature") or "")
    office_staff_signature_image_html = _signature_image_html(packet, "office_staff_signature", width_px=220, height_px=48)
    date_reviewed = html.escape(packet.get("date_reviewed") or "")

    included_docs = [
        ("Virtual Consent Form completed and signed", bool(packet.get("included_virtual_consent_form"))),
        ("VA Form 10-10172 completed and signed", bool(packet.get("included_va_form_10_10172"))),
        ("SEOC request signed by CCN ordering provider", bool(packet.get("included_seoc_request"))),
        ("Consultation & Treatment Request completed", bool(packet.get("included_consult_request"))),
        ("Letter of Medical Necessity completed", bool(packet.get("included_lomn"))),
        ("Clinical Notes included", bool(packet.get("included_clinical_notes"))),
        ("MRI Report included", bool(packet.get("included_mri_report"))),
    ]

    checklist_html = "".join(
        f"<div style='margin-bottom:6px;'>{'[X]' if checked else '[ ]'} {html.escape(label)}</div>"
        for label, checked in included_docs
    )

    def field_line(label, value, blank_length, raw_html=False):
        display = value if value else "_" * blank_length
        rendered = display if raw_html else html.escape(str(display))
        return (
            "<div style='margin-bottom:10px;'>"
            f"<span style='font-weight:700;'>{html.escape(label)}:</span> "
            f"<span>{rendered}</span>"
            "</div>"
        )

    return (
        "<div style='font-family:Calibri, Arial, sans-serif; color:#111827; line-height:1.5;'>"
        f"<div style='text-align:center; font-size:18pt; font-weight:700;'>{title}</div>"
        + "<div style='margin-top:20px;'>"
        + field_line("Patient Name", patient_name, 24)
        + field_line("DOB", dob, 14)
        + field_line("VA Facility", facility, 28)
        + field_line("Ordering Provider", ordering_provider, 24)
        + field_line("Date of Submission", submission_date, 18)
        + field_line("Primary Diagnosis Code", diagnosis_code, 18)
        + "</div>"
        + "<div style='margin-top:18px; font-weight:700;'>Documents Included (check all):</div>"
        + "<div style='margin-top:10px; padding-left:4px;'>"
        + checklist_html
        + "</div>"
        + "<div style='margin-top:22px;'>"
        + field_line("Submitting Office", submitting_office, 24)
        + field_line("Office Staff Name", office_staff_name, 20)
        + field_line("Program User / Office Staff Signature", office_staff_signature_image_html or office_staff_signature, 20, raw_html=bool(office_staff_signature_image_html))
        + field_line("Date Reviewed", date_reviewed, 18)
        + "</div>"
        + "</div>"
    )


def build_seoc_request_preview_html(packet):
    title = "Single Episode of Care (SEOC) Request"
    request_date = html.escape(packet.get("seoc_request_date") or "Pending")
    va_medical_center_name = html.escape(packet.get("va_medical_center_name") or "Pending")
    patient_name = html.escape(packet.get("patient_name") or "Pending")
    dob = html.escape(packet.get("date_of_birth") or "Pending")
    last_four_ssn = html.escape(packet.get("last_four_ssn") or "Pending")
    episode_diagnosis = html.escape(packet.get("episode_diagnosis") or "Pending")
    episode_icd_code = html.escape(packet.get("episode_icd_code") or "Pending")
    estimated_duration = html.escape(packet.get("estimated_duration_text") or "Pending")
    provider_name = html.escape(packet.get("provider") or packet.get("ordering_doctor") or "Pending")
    provider_credentials = html.escape(packet.get("provider_credentials") or "Pending")
    provider_specialty = html.escape(packet.get("provider_specialty") or "Pending")
    provider_npi = html.escape(packet.get("provider_npi") or "Pending")
    practice_name = html.escape(packet.get("practice_name") or packet.get("community_facility") or "Pending")
    provider_phone = html.escape(packet.get("provider_phone") or "Pending")
    provider_fax = html.escape(packet.get("provider_fax") or "Pending")

    scope_items = []
    if packet.get("seoc_include_preprocedure_eval"):
        scope_items.append("Pre-procedure evaluation and procedural planning")
    if packet.get("seoc_include_annulargram"):
        scope_items.append("Diagnostic Annulargram")
    if packet.get("seoc_include_fibrin_injection"):
        scope_items.append("Inter Annular Fibrin Injections if indicated")
    if packet.get("seoc_include_follow_up"):
        scope_items.append("Standard post-procedure follow-up visit(s) within the global postoperative period")
    if not scope_items:
        scope_items.append("Pending scope items")

    objective_lines = [line.strip() for line in str(packet.get("clinical_objectives") or "").splitlines() if line.strip()]
    if not objective_lines:
        objective_lines = ["Pending clinical objective"]
    scope_summary = str(packet.get("seoc_scope_text") or "").strip()
    continuity_text = str(packet.get("seoc_continuity_text") or "").strip()

    def bullets(items):
        return "".join(
            f"<div style='margin-top:6px; color:#E8F1FC;'>- {html.escape(item)}</div>" for item in items
        )

    return (
        "<div style='font-family:Segoe UI; color:#E8F1FC; line-height:1.55;'>"
        f"<div style='color:#8FA6C1;'>{request_date}</div>"
        "<div style='margin-top:12px; color:#E8F1FC;'>Department of Veterans Affairs<br/>Community Care Office<br/>"
        f"{va_medical_center_name}</div>"
        f"<div style='margin-top:16px; font-size:24px; font-weight:800; color:#F2F5F9;'>{title}</div>"
        "<div style='margin-top:10px; padding:12px 14px; background:#0F1823; border:1px solid #243446; border-radius:10px;'>"
        f"<div><span style='color:#8FA6C1;'>Veteran Name:</span> <span style='color:#E8F1FC;'>{patient_name}</span></div>"
        f"<div style='margin-top:6px;'><span style='color:#8FA6C1;'>DOB:</span> <span style='color:#E8F1FC;'>{dob}</span></div>"
        f"<div style='margin-top:6px;'><span style='color:#8FA6C1;'>Last Four SSN:</span> <span style='color:#E8F1FC;'>{last_four_ssn}</span></div>"
        "</div>"
        "<div style='margin-top:14px; color:#E8F1FC;'>"
        "This request is for authorization of a defined, time-limited Single Episode of Care (SEOC) for treatment of lumbar disc pathology."
        "</div>"
        "<div style='margin-top:14px; font-size:16px; font-weight:800; color:#69BCFF;'>Episode Diagnosis</div>"
        f"<div style='margin-top:6px; color:#E8F1FC;'>{episode_diagnosis}"
        + (f" - {episode_icd_code}" if episode_icd_code else "")
        + "</div>"
        "<div style='margin-top:14px; font-size:16px; font-weight:800; color:#69BCFF;'>Scope of Requested Episode</div>"
        "<div style='margin-top:6px; color:#C8D8E8;'>This SEOC includes only the following services directly related to treatment of the above diagnosis:</div>"
        + bullets(scope_items)
        + f"<div style='color:#C8D8E8;'>{html.escape(scope_summary) if scope_summary else 'No additional pain management services, long-term medication management, or unrelated spine care are requested under this episode.'}</div>"
        + "<div style='margin-top:14px; font-size:16px; font-weight:800; color:#69BCFF;'>Estimated Duration of Episode</div>"
        + f"<div style='margin-top:6px; color:#E8F1FC;'>The episode is expected to begin upon authorization and conclude within approximately {estimated_duration}</div>"
        + "<div style='margin-top:14px; font-size:16px; font-weight:800; color:#69BCFF;'>Clinical Objective</div>"
        + bullets(objective_lines)
        + "<div style='margin-top:14px; font-size:16px; font-weight:800; color:#69BCFF;'>Continuity of Care</div>"
        + f"<div style='margin-top:6px; color:#E8F1FC;'>{html.escape(continuity_text) if continuity_text else 'Upon completion of the episode, a summary of treatment rendered and clinical outcome will be forwarded to the referring VA provider. Any need for further care beyond this defined episode will require separate evaluation and authorization.'}</div>"
        + "<div style='margin-top:10px; color:#E8F1FC;'>This request represents a discrete, procedure-based intervention and meets criteria for authorization under the Single Episode of Care (SEOC) framework.</div>"
        + "<div style='margin-top:18px; color:#E8F1FC;'>Sincerely,</div>"
        + f"<div style='margin-top:12px; color:#E8F1FC;'><strong>{provider_name}</strong>{', ' + provider_credentials if provider_credentials and provider_credentials != 'Pending' else ''}<br/>"
        + f"{provider_specialty}<br/>NPI: {provider_npi}<br/>{practice_name}<br/>{provider_phone}<br/>{provider_fax}</div>"
        + "</div>"
    )


def build_lomn_preview_html(packet):
    def filled(value, blank_length):
        raw = str(value or "").strip()
        return html.escape(raw) if raw else "_" * blank_length

    def block_text(value, blank_length=48):
        raw = str(value or "").strip()
        return html.escape(raw).replace("\n", "<br/>") if raw else "_" * blank_length

    def bullet_lines(items):
        return "".join(f"<div style='margin-top:6px;'>- {html.escape(item)}</div>" for item in items)

    title = html.escape(packet.get("packet_title") or "LETTER OF MEDICAL NECESSITY")
    request_date = filled(packet.get("lmn_request_date"), 12)
    va_medical_center_name = filled(packet.get("va_medical_center_name"), 22)
    facility = filled(packet.get("facility"), 20)
    patient_name = filled(packet.get("patient_name"), 18)
    dob = filled(packet.get("date_of_birth"), 12)
    last_four_ssn = filled(packet.get("last_four_ssn"), 8)
    claim_number = filled(packet.get("lmn_va_claim_number"), 14)
    primary_diagnosis = filled(packet.get("lmn_primary_diagnosis"), 30)
    secondary_diagnosis = filled(packet.get("lmn_secondary_diagnosis"), 30)
    clinical_summary = block_text(packet.get("lmn_clinical_summary"))
    mri_date = filled(packet.get("lmn_mri_date"), 12)
    mri_findings = filled(packet.get("lmn_mri_findings"), 28)
    conservative_duration = filled(packet.get("lmn_conservative_duration"), 14)
    medical_necessity_statement = block_text(packet.get("lmn_medical_necessity_statement"))
    risk_statement = block_text(packet.get("lmn_risk_statement"))
    reasonable_statement = block_text(packet.get("lmn_reasonable_necessary_statement"))
    contact_statement = block_text(packet.get("lmn_contact_statement"))
    provider_name = filled(packet.get("provider") or packet.get("ordering_doctor"), 18)
    provider_credentials = filled(packet.get("provider_credentials"), 8)
    provider_specialty = filled(packet.get("provider_specialty"), 18)
    provider_npi = filled(packet.get("provider_npi"), 14)
    practice_name = filled(packet.get("practice_name") or packet.get("community_facility"), 20)
    provider_phone = filled(packet.get("provider_phone"), 14)
    provider_fax = filled(packet.get("provider_fax"), 14)
    provider_header = provider_name
    raw_credentials = str(packet.get("provider_credentials") or "").strip()
    if raw_credentials:
        provider_header += f", {html.escape(raw_credentials)}"

    conservative_items = []
    if packet.get("lmn_include_physical_therapy"):
        conservative_items.append("Structured physical therapy")
    if packet.get("lmn_include_nsaids"):
        conservative_items.append("NSAIDs and non-opioid analgesics")
    if packet.get("lmn_include_activity_modification"):
        conservative_items.append("Activity modification")
    if packet.get("lmn_include_home_exercise"):
        conservative_items.append("Home exercise program")
    if packet.get("lmn_include_epidural_steroid_injections"):
        conservative_items.append("Epidural Steroid injections (if applicable)")
    if not conservative_items:
        conservative_items.append("________________________________________")

    indication_items = []
    if packet.get("lmn_indication_reduce_pain"):
        indication_items.append("Reduce pain severity")
    if packet.get("lmn_indication_improve_function"):
        indication_items.append("Improve functional capacity")
    if packet.get("lmn_indication_prevent_degeneration"):
        indication_items.append("Prevent further disc degeneration")
    if packet.get("lmn_indication_reduce_opioid_reliance"):
        indication_items.append("Decrease reliance on long-term opioid therapy")
    if packet.get("lmn_indication_prevent_surgery"):
        indication_items.append("Potentially prevent need for more invasive surgical intervention")
    if not indication_items:
        indication_items.append("________________________________________")

    return (
        "<div style='font-family:Calibri, Arial, sans-serif; color:#111827; line-height:1.5;'>"
        f"<div style='text-align:center; font-size:18pt; font-weight:700;'>{title}</div>"
        "<div style='text-align:center; margin-top:2px;'>(Pain Management - Interventional Spine)</div>"
        f"<div style='margin-top:12px;'>{request_date}</div>"
        "<div style='margin-top:8px;'>Department of Veterans Affairs<br/>Community Care Office<br/>"
        f"{va_medical_center_name}<br/>{facility}</div>"
        "<div style='margin-top:12px; font-weight:700;'>RE: Letter of Medical Necessity</div>"
        f"<div>Veteran Name: {patient_name}</div>"
        f"<div>DOB: {dob}</div>"
        f"<div>Last Four SSN: {last_four_ssn}</div>"
        f"<div>VA Claim Number: {claim_number}</div>"
        + "<div style='margin-top:10px;'>To Whom It May Concern:</div>"
        + "<div style='margin-top:6px;'>I am writing to formally document the medical necessity of interventional spine treatment for the above-named Veteran.</div>"
        + "<div style='margin-top:12px; font-weight:700;'>Clinical Diagnoses</div>"
        + f"<div style='margin-top:4px;'>Primary: {primary_diagnosis}</div>"
        + f"<div>Secondary: {secondary_diagnosis}</div>"
        + "<div style='margin-top:12px; font-weight:700;'>Clinical Basis</div>"
        + f"<div style='margin-top:4px;'>{clinical_summary}</div>"
        + "<div style='margin-top:12px; font-weight:700;'>Imaging Support</div>"
        + f"<div style='margin-top:4px;'>MRI Date: {mri_date}</div>"
        + f"<div>MRI Findings: {mri_findings}</div>"
        + "<div style='margin-top:8px;'>The Veteran has completed extensive conservative management, including:</div>"
        + bullet_lines(conservative_items)
        + f"<div style='margin-top:8px;'>Despite appropriate and guideline-based conservative treatment for greater than {conservative_duration}, the Veteran continues to experience persistent pain and functional limitation.</div>"
        + "<div style='margin-top:12px; font-weight:700;'>Basis for Medical Necessity</div>"
        + f"<div style='margin-top:4px;'>{medical_necessity_statement}</div>"
        + "<div style='margin-top:12px; font-weight:700;'>Requested Treatment Objectives</div>"
        + "<div style='margin-top:4px;'>This intervention is indicated to:</div>"
        + bullet_lines(indication_items)
        + "<div style='margin-top:12px; font-weight:700;'>Risk if Treatment Is Delayed or Denied</div>"
        + f"<div style='margin-top:4px;'>{risk_statement}</div>"
        + "<div style='margin-top:12px; font-weight:700;'>Reasonable and Necessary Determination</div>"
        + f"<div style='margin-top:4px;'>{reasonable_statement}</div>"
        + "<div style='margin-top:12px; font-weight:700;'>Provider Contact Statement</div>"
        + f"<div style='margin-top:4px;'>{contact_statement}</div>"
        + "<div style='margin-top:16px;'>Sincerely,</div>"
        + f"<div style='margin-top:10px;'>{provider_header}<br/>{provider_specialty}<br/>NPI Number: {provider_npi}<br/>{practice_name}<br/>{provider_phone}<br/>{provider_fax}</div>"
        + "</div>"
    )


def build_consult_request_preview_html(packet):
    def filled(value, blank_length):
        raw = str(value or "").strip()
        return html.escape(raw) if raw else "_" * blank_length

    def block_text(value, blank_length=48):
        raw = str(value or "").strip()
        return html.escape(raw).replace("\n", "<br/>") if raw else "_" * blank_length

    def bullet_lines(items):
        return "".join(f"<div style='margin-top:6px;'>- {html.escape(item)}</div>" for item in items)

    title = html.escape(packet.get("packet_title") or "CONSULTATION AND TREATMENT REQUEST")
    request_date = filled(packet.get("consult_request_date"), 12)
    va_medical_center_name = filled(packet.get("va_medical_center_name"), 22)
    facility = filled(packet.get("facility"), 20)
    patient_name = filled(packet.get("patient_name"), 18)
    dob = filled(packet.get("date_of_birth"), 12)
    last_four_ssn = filled(packet.get("last_four_ssn"), 8)
    claim_number = filled(packet.get("consult_va_claim_number"), 14)
    referring_provider = filled(packet.get("consult_referring_va_provider"), 18)
    reason_text = block_text(packet.get("consult_reason_text"))
    primary_diagnosis = filled(packet.get("consult_primary_diagnosis"), 30)
    secondary_diagnosis = filled(packet.get("consult_secondary_diagnosis"), 30)
    mri_date = filled(packet.get("consult_mri_date"), 12)
    mri_findings = filled(packet.get("consult_mri_findings"), 28)
    conservative_duration = filled(packet.get("consult_conservative_duration"), 14)
    scope_exclusion_text = block_text(packet.get("consult_scope_exclusion_text"))
    medical_rationale_text = block_text(packet.get("consult_medical_rationale_text"))
    risk_without_treatment = block_text(packet.get("consult_risk_without_treatment"))
    duration_scope_text = block_text(packet.get("consult_duration_scope_text"))
    consult_contact_statement = block_text(packet.get("consult_contact_statement"))
    provider_name = filled(packet.get("provider") or packet.get("ordering_doctor"), 18)
    provider_credentials = filled(packet.get("provider_credentials"), 8)
    provider_specialty = filled(packet.get("provider_specialty"), 18)
    provider_npi = filled(packet.get("provider_npi"), 14)
    practice_name = filled(packet.get("practice_name") or packet.get("community_facility"), 20)
    provider_address = filled(packet.get("provider_address"), 22)
    provider_phone = filled(packet.get("provider_phone"), 14)
    provider_fax = filled(packet.get("provider_fax"), 14)
    provider_email = filled(packet.get("provider_email"), 20)
    provider_header = provider_name
    raw_credentials = str(packet.get("provider_credentials") or "").strip()
    if raw_credentials:
        provider_header += f", {html.escape(raw_credentials)}"

    symptom_items = []
    if packet.get("consult_symptom_axial_pain"):
        symptom_items.append("Chronic axial lumbar pain")
    if packet.get("consult_symptom_activity_exacerbation"):
        symptom_items.append("Activity-related exacerbation (sitting, standing, bending)")
    if packet.get("consult_symptom_reduced_tolerance"):
        symptom_items.append("Reduced tolerance for prolonged positioning")
    if packet.get("consult_symptom_functional_impairment"):
        symptom_items.append("Functional impairment affecting occupational and daily activities")
    if not symptom_items:
        symptom_items.append("________________________________________")

    conservative_items = []
    if packet.get("consult_include_physical_therapy"):
        conservative_items.append("Physical therapy")
    if packet.get("consult_include_nsaids"):
        conservative_items.append("NSAIDs and non-opioid analgesics")
    if packet.get("consult_include_activity_modification"):
        conservative_items.append("Activity modification")
    if packet.get("consult_include_home_exercise"):
        conservative_items.append("Home exercise program")
    if packet.get("consult_include_interventional_history"):
        conservative_items.append("Epidural steroid injections and/or other interventional procedures (if applicable)")
    if not conservative_items:
        conservative_items.append("________________________________________")

    service_items = []
    if packet.get("consult_include_pain_management_consultation"):
        service_items.append("Comprehensive pain management consultation")
    if packet.get("consult_include_procedural_planning"):
        service_items.append("Diagnostic confirmation and procedural planning")
    if packet.get("consult_include_annulargram"):
        service_items.append("Diagnostic Annulargram")
    if packet.get("consult_include_fibrin_injection"):
        fibrin_line = "Intraannular Fibrin injection"
        if str(packet.get("consult_fibrin_levels") or "").strip():
            fibrin_line += f" at {packet.get('consult_fibrin_levels').strip()}"
        fibrin_line += " if indicated"
        service_items.append(fibrin_line)
    if packet.get("consult_include_follow_up"):
        service_items.append("Routine post-procedure follow-up visit(s)")
    if not service_items:
        service_items.append("________________________________________")

    goal_items = []
    if packet.get("consult_goal_pain_reduction"):
        goal_items.append("Pain reduction")
    if packet.get("consult_goal_functional_improvement"):
        goal_items.append("Functional improvement")
    if packet.get("consult_goal_reduce_analgesics"):
        goal_items.append("Decreased reliance on chronic analgesic therapy")
    if packet.get("consult_goal_prevent_surgery"):
        goal_items.append("Prevention of progression requiring more invasive surgical intervention")
    if not goal_items:
        goal_items.append("________________________________________")

    return (
        "<div style='font-family:Calibri, Arial, sans-serif; color:#111827; line-height:1.5;'>"
        f"<div style='text-align:center; font-size:18pt; font-weight:700;'>{title}</div>"
        "<div style='text-align:center; margin-top:2px;'>(Pain Management - Interventional Spine)</div>"
        f"<div style='margin-top:12px;'>{request_date}</div>"
        "<div style='margin-top:8px;'>Department of Veterans Affairs<br/>Community Care Office<br/>"
        f"{va_medical_center_name}<br/>{facility}</div>"
        + "<div style='margin-top:12px; font-weight:700;'>RE: Consultation and Treatment Request</div>"
        + f"<div>Veteran Name: {patient_name}</div>"
        + f"<div>DOB: {dob}</div>"
        + f"<div>Last Four SSN: {last_four_ssn}</div>"
        + f"<div>VA Claim Number: {claim_number}</div>"
        + f"<div>Referring VA Provider: {referring_provider}</div>"
        + "<div style='margin-top:12px; font-weight:700;'>Reason for Consultation</div>"
        + f"<div style='margin-top:4px;'>{reason_text}</div>"
        + "<div style='margin-top:12px; font-weight:700;'>Diagnoses</div>"
        + f"<div style='margin-top:4px;'>Primary: {primary_diagnosis}</div>"
        + f"<div>Secondary: {secondary_diagnosis}</div>"
        + "<div style='margin-top:12px; font-weight:700;'>Clinical Summary</div>"
        + "<div style='margin-top:4px;'>Symptoms and functional concerns documented for this consultation include:</div>"
        + bullet_lines(symptom_items)
        + "<div style='margin-top:12px; font-weight:700;'>Imaging Review</div>"
        + f"<div style='margin-top:4px;'>MRI Date: {mri_date}</div>"
        + f"<div>MRI Findings: {mri_findings}</div>"
        + "<div style='margin-top:4px;'>These findings are reviewed alongside the reported symptoms and prior treatment history.</div>"
        + "<div style='margin-top:8px;'>Conservative management has included:</div>"
        + bullet_lines(conservative_items)
        + f"<div style='margin-top:8px;'>Despite appropriate treatment for greater than {conservative_duration}, the Veteran continues to experience persistent pain and impaired function.</div>"
        + "<div style='margin-top:12px; font-weight:700;'>Requested Services</div>"
        + "<div style='margin-top:4px;'>Authorization is requested for:</div>"
        + bullet_lines(service_items)
        + "<div style='margin-top:12px; font-weight:700;'>Scope of Request</div>"
        + f"<div style='margin-top:4px;'>{scope_exclusion_text}</div>"
        + "<div style='margin-top:12px; font-weight:700;'>Medical Rationale</div>"
        + f"<div style='margin-top:4px;'>{medical_rationale_text}</div>"
        + "<div style='margin-top:8px;'>Clinical goals include:</div>"
        + bullet_lines(goal_items)
        + "<div style='margin-top:12px; font-weight:700;'>Risk Without Requested Treatment</div>"
        + f"<div style='margin-top:4px;'>{risk_without_treatment}</div>"
        + "<div style='margin-top:12px; font-weight:700;'>Duration and Scope of Care</div>"
        + f"<div style='margin-top:4px;'>{duration_scope_text}</div>"
        + "<div style='margin-top:12px; font-weight:700;'>Provider Contact Statement</div>"
        + f"<div style='margin-top:4px;'>{consult_contact_statement}</div>"
        + "<div style='margin-top:16px;'>Sincerely,</div>"
        + f"<div style='margin-top:10px;'>{provider_header}<br/>{provider_specialty}<br/>NPI Number: {provider_npi}<br/>{practice_name}<br/>{provider_address}<br/>{provider_phone}<br/>{provider_fax}<br/>{provider_email}</div>"
        + "</div>"
    )


def build_clinical_documentation_preview_html(packet):
    def filled(value, blank_length):
        raw = str(value or "").strip()
        return html.escape(raw) if raw else "_" * blank_length

    def block_text(value, blank_length=40):
        raw = str(value or "").strip()
        return html.escape(raw).replace("\n", "<br/>") if raw else "_" * blank_length

    def checkbox_items(items):
        rendered = []
        for enabled, label in items:
            marker = "[X]" if enabled else "[ ]"
            rendered.append(f"<div style='margin-top:4px; padding-left:12px;'>{marker} {html.escape(label)}</div>")
        return "".join(rendered)

    title = html.escape(packet.get("packet_title") or packet.get("clinical_doc_title") or "Clinical Documentation Template")
    chief_complaint = filled(packet.get("clinical_doc_chief_complaint"), 24)
    exact_duration = filled(packet.get("clinical_doc_exact_duration"), 20)
    pain_severity = filled(packet.get("clinical_doc_pain_severity"), 6)
    functional_impact = block_text(packet.get("clinical_doc_functional_impact"), 28)
    conservative_duration = filled(packet.get("clinical_doc_conservative_duration"), 18)
    imaging_date = filled(packet.get("clinical_doc_mri_date"), 18)
    affected_levels = filled(packet.get("clinical_doc_affected_levels"), 20)
    primary_diagnosis = filled(packet.get("clinical_doc_primary_diagnosis"), 26)
    secondary_diagnosis = filled(packet.get("clinical_doc_secondary_diagnosis"), 26)
    assessment_summary = block_text(packet.get("clinical_doc_assessment_summary"))
    treatment_plan_intro = block_text(packet.get("clinical_doc_treatment_plan_intro"))
    plan_exclusion = block_text(packet.get("clinical_doc_plan_exclusion"))
    physician_narrative = block_text(packet.get("clinical_doc_physician_narrative"))

    duration_items = [
        (packet.get("clinical_doc_duration_gt_3m"), "> 3 months"),
        (packet.get("clinical_doc_duration_gt_6m"), "> 6 months"),
        (packet.get("clinical_doc_duration_gt_12m"), "> 12 months"),
    ]
    pain_items = [
        (packet.get("clinical_doc_pain_axial"), "Axial lumbar pain"),
        (packet.get("clinical_doc_pain_discogenic"), "Discogenic pattern"),
        (packet.get("clinical_doc_pain_activity_exacerbation"), "Activity-related exacerbation"),
        (packet.get("clinical_doc_pain_sitting_intolerance"), "Sitting intolerance"),
        (packet.get("clinical_doc_pain_standing_intolerance"), "Standing intolerance"),
        (packet.get("clinical_doc_pain_bending_lifting"), "Bending/lifting provocation"),
    ]
    limitation_items = [
        (packet.get("clinical_doc_limit_occupational"), "Occupational duties"),
        (packet.get("clinical_doc_limit_prolonged_sitting"), "Prolonged sitting"),
        (packet.get("clinical_doc_limit_prolonged_standing"), "Prolonged standing"),
        (packet.get("clinical_doc_limit_ambulation"), "Ambulation tolerance"),
        (packet.get("clinical_doc_limit_household"), "Household activities"),
        (packet.get("clinical_doc_limit_sleep"), "Sleep disturbance"),
    ]
    conservative_items = [
        (packet.get("clinical_doc_conservative_pt"), "Physical therapy"),
        (packet.get("clinical_doc_conservative_home_exercise"), "Home exercise program"),
        (packet.get("clinical_doc_conservative_nsaids"), "NSAIDs"),
        (packet.get("clinical_doc_conservative_non_opioid"), "Non-opioid analgesics"),
        (packet.get("clinical_doc_conservative_activity_modification"), "Activity modification"),
        (packet.get("clinical_doc_conservative_esi"), "Epidural steroid injections"),
        (packet.get("clinical_doc_conservative_other_interventional"), "Other interventional procedures"),
    ]
    imaging_items = [
        (packet.get("clinical_doc_imaging_annular_tear"), "Annular tear"),
        (packet.get("clinical_doc_imaging_disc_degeneration"), "Disc degeneration"),
        (packet.get("clinical_doc_imaging_disc_protrusion"), "Disc protrusion"),
        (packet.get("clinical_doc_imaging_disc_displacement"), "Disc displacement"),
    ]
    plan_items = [
        (packet.get("clinical_doc_plan_diagnostic_confirmation"), "Diagnostic confirmation if necessary"),
        (packet.get("clinical_doc_plan_intradiscal_intervention"), "Intradiscal annular intervention if indicated"),
        (packet.get("clinical_doc_plan_follow_up"), "Standard post-procedure follow-up"),
    ]

    return (
        "<div style='font-family:Calibri, Arial, sans-serif; color:#111827; line-height:1.45;'>"
        f"<div style='text-align:center; font-size:18pt; font-weight:700;'>{title}</div>"
        + "<div style='margin-top:14px; font-weight:700;'>I. Chief Complaint</div>"
        + f"<div style='margin-top:4px;'>{chief_complaint}</div>"
        + "<div style='margin-top:12px; font-weight:700;'>II. History of Present Illness</div>"
        + "<div style='margin-top:4px;'>Duration of Symptoms:</div>"
        + checkbox_items(duration_items)
        + f"<div style='margin-top:6px;'>Exact duration: {exact_duration}</div>"
        + "<div style='margin-top:8px;'>Pain Characteristics:</div>"
        + checkbox_items(pain_items)
        + f"<div style='margin-top:8px;'>Pain severity (0-10): {pain_severity}</div>"
        + "<div style='margin-top:12px; font-weight:700;'>III. Functional Impairment</div>"
        + "<div style='margin-top:4px;'>Patient reports limitation in:</div>"
        + checkbox_items(limitation_items)
        + f"<div style='margin-top:8px;'>Describe specific functional impact: {functional_impact}</div>"
        + "<div style='margin-top:12px; font-weight:700;'>IV. Conservative Therapy History</div>"
        + "<div style='margin-top:4px;'>The patient has completed and/or attempted:</div>"
        + checkbox_items(conservative_items)
        + f"<div style='margin-top:8px;'>Duration of conservative treatment: {conservative_duration}</div>"
        + f"<div style='margin-top:8px;'>Response to prior ESI (if applicable): {filled(packet.get('clinical_doc_esi_response'), 18)}</div>"
        + "<div style='margin-top:12px; font-weight:700;'>V. Imaging Findings</div>"
        + f"<div style='margin-top:4px;'>MRI Date: {imaging_date}</div>"
        + "<div style='margin-top:4px;'>Findings:</div>"
        + checkbox_items(imaging_items)
        + f"<div style='margin-top:8px;'>Affected Levels: {affected_levels}</div>"
        + "<div style='margin-top:4px;'>Imaging correlates with clinical presentation.</div>"
        + "<div style='margin-top:12px; font-weight:700;'>VI. Assessment</div>"
        + f"<div style='margin-top:4px;'>Primary Diagnosis / ICD-10: {primary_diagnosis}</div>"
        + f"<div>Secondary Diagnosis / ICD-10: {secondary_diagnosis}</div>"
        + f"<div style='margin-top:6px;'>{assessment_summary}</div>"
        + "<div style='margin-top:12px; font-weight:700;'>VII. Treatment Plan</div>"
        + f"<div style='margin-top:4px;'>{treatment_plan_intro}</div>"
        + "<div style='margin-top:6px;'>Plan includes:</div>"
        + checkbox_items([(enabled, label) for enabled, label in plan_items if enabled] or [(False, "________________________________________")])
        + f"<div style='margin-top:6px;'>{plan_exclusion}</div>"
        + "<div style='margin-top:12px; font-weight:700;'>Physician Narrative Paragraph</div>"
        + f"<div style='margin-top:6px;'>{physician_narrative}</div>"
        + "</div>"
    )


def build_va_10172_preview_html(packet, template_path=""):
    if not template_path:
        template_path = packet.get("va10172_template_path") or ""
    title = html.escape(packet.get("packet_title") or "VA Form 10-10172")
    patient_name = html.escape(packet.get("patient_name") or "Pending")
    dob = html.escape(packet.get("date_of_birth") or "Pending")
    auth = html.escape(packet.get("authorization_number") or "Pending")
    va_facility_address = html.escape(packet.get("va10172_va_facility_address") or "Pending").replace("\n", "<br/>")
    provider_office_address = html.escape(packet.get("va10172_ordering_provider_office_address") or "Pending").replace("\n", "<br/>")
    provider_phone = html.escape(packet.get("va10172_ordering_provider_phone") or "Pending")
    provider_fax = html.escape(packet.get("va10172_ordering_provider_fax") or "Pending")
    provider_email = html.escape(packet.get("va10172_ordering_provider_secure_email") or "Pending")
    specialty = html.escape(packet.get("va10172_referral_specialty_text") or "Pending")
    icd_codes = html.escape(packet.get("icd_codes") or "Pending")
    diagnosis_description = html.escape(packet.get("va10172_diagnosis_description") or packet.get("diagnosis") or "Pending")
    requested_cpt = html.escape(packet.get("va10172_requested_cpt_hcpcs_code") or "Pending")
    cpt_description = html.escape(packet.get("va10172_description_cpt_hcpcs_code") or "Pending")
    reason_for_request = html.escape(packet.get("va10172_reason_for_request") or "Pending").replace("\n", "<br/>")
    provider_name_printed = html.escape(packet.get("va10172_ordering_provider_name_printed") or packet.get("provider") or packet.get("ordering_doctor") or "Pending")
    provider_npi = html.escape(packet.get("va10172_ordering_provider_npi") or packet.get("provider_npi") or "Pending")
    today_date = html.escape(packet.get("va10172_today_date") or "Pending")
    template_label = html.escape(template_path or "No template selected")

    def kv(label, value):
        return f"<div style='margin-top:6px;'><span style='color:#8FA6C1;'>{html.escape(label)}:</span> <span style='color:#E8F1FC;'>{value}</span></div>"

    return (
        "<div style='font-family:Segoe UI; color:#E8F1FC; line-height:1.55;'>"
        f"<div style='font-size:24px; font-weight:800; color:#F2F5F9;'>{title}</div>"
        "<div style='margin-top:6px; color:#69BCFF; font-weight:700;'>Actual VA PDF Export Profile</div>"
        "<div style='margin-top:10px; color:#C8D8E8;'>"
        "This profile exports onto the real VA Form 10-10172 template. PDF export fills the actual form page rather than producing a visual clone."
        "</div>"
        + f"<div style='margin-top:10px; color:#8FA6C1;'>Template PDF: {template_label}</div>"
        + "<div style='margin-top:16px; padding:12px 14px; background:#0F1823; border:1px solid #243446; border-radius:10px;'>"
        + kv("Veteran Legal Full Name", patient_name)
        + kv("Date of Birth", dob)
        + kv("VA Authorization Number", auth)
        + kv("VA Facility & Address", va_facility_address)
        + kv("Ordering Provider Office Name & Address", provider_office_address)
        + kv("Ordering Provider Phone", provider_phone)
        + kv("Ordering Provider Fax", provider_fax)
        + kv("Ordering Provider Secure Email", provider_email)
        + kv("Care Needed Within 48 Hours", html.escape(packet.get("va10172_care_needed_within_48_hours") or "Pending"))
        + kv("Continuation of Care", html.escape(packet.get("va10172_is_continuation_of_care") or "Pending"))
        + kv("Referral to Another Specialty", html.escape(packet.get("va10172_referral_to_specialty") or "Pending"))
        + kv("Specialty", specialty)
        + kv("Diagnosis Codes", icd_codes)
        + kv("Diagnosis Description", diagnosis_description)
        + kv("Requested CPT/HCPCS Code", requested_cpt)
        + kv("Description CPT/HCPCS Code", cpt_description)
        + kv("Reason for Request", reason_for_request)
        + kv("CCN Ordering Provider Name (Printed)", provider_name_printed)
        + kv("CCN Ordering Provider NPI", provider_npi)
        + kv("Today's Date", today_date)
        + "</div>"
        + "<div style='margin-top:12px; color:#C8D8E8;'>"
        "Note: the CCN ordering provider signature field can be staged with a typed placeholder or drawn signature for drafting, but formal signing requirements may still need office review."
        "</div>"
        + "</div>"
    )


class LegacyGalleryTab(QWidget):
    BUILTIN_REFERENCES = [
        {
            "entry_id": "builtin_launcher_preview",
            "label": "Current Launcher Preview",
            "kind": "launcher_preview",
            "path": resource_path("launcher"),
            "source_path": resource_path("launcher"),
            "builtin": True,
            "background_path": resource_path("launcher/assets/launcher_background.png"),
            "logo_path": resource_path("launcher/assets/truecore_logo.png"),
            "code_path": resource_path("launcher/launcher_window.py"),
        },
        {
            "entry_id": "builtin_gui_preview",
            "label": "Current GUI Preview",
            "kind": "gui_preview",
            "path": resource_path("ui/pyside_gui"),
            "source_path": resource_path("ui/pyside_gui"),
            "builtin": True,
            "background_path": resource_path("ui/pyside_gui/assets/launcher_background.png"),
            "logo_path": resource_path("ui/pyside_gui/assets/truecore_logo.png"),
            "code_path": resource_path("ui/pyside_gui/main_window.py"),
        },
    ]

    def __init__(self, config, save_config_callback, parent=None):
        super().__init__(parent)
        self.config = dict(config or {})
        self.save_config_callback = save_config_callback
        self.entries: List[Dict[str, str]] = []
        self.current_entry = None

        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(12)

        intro = QLabel(
            "Legacy UI Gallery is for visual reference only. Add old launcher or GUI folders/files here and the studio will show preview-only references without running old logic."
        )
        intro.setWordWrap(True)
        intro.setStyleSheet("color:#BFD0E3;")
        layout.addWidget(intro)

        splitter = QSplitter(Qt.Horizontal)
        splitter.setChildrenCollapsible(False)
        layout.addWidget(splitter, stretch=1)

        left_panel = QWidget()
        left_layout = QVBoxLayout(left_panel)
        left_layout.setContentsMargins(0, 0, 0, 0)
        left_layout.setSpacing(10)

        self.reference_list = QListWidget()
        self.reference_list.itemSelectionChanged.connect(self.on_selection_changed)
        left_layout.addWidget(self.reference_list, stretch=1)

        button_row = QHBoxLayout()
        add_files_button = QPushButton("Add Files")
        add_files_button.clicked.connect(self.add_reference_files)
        add_folder_button = QPushButton("Add Folder")
        add_folder_button.clicked.connect(self.add_reference_folder)
        remove_button = QPushButton("Remove Selected")
        remove_button.clicked.connect(self.remove_selected_reference)
        button_row.addWidget(add_files_button)
        button_row.addWidget(add_folder_button)
        button_row.addWidget(remove_button)
        left_layout.addLayout(button_row)
        splitter.addWidget(left_panel)

        right_panel = QWidget()
        right_layout = QVBoxLayout(right_panel)
        right_layout.setContentsMargins(0, 0, 0, 0)
        right_layout.setSpacing(10)

        self.preview_title = QLabel("Select a legacy reference")
        self.preview_title.setStyleSheet("font-size:18px; font-weight:700; color:#69BCFF;")
        right_layout.addWidget(self.preview_title)

        self.preview_path = QLabel("No legacy reference selected yet.")
        self.preview_path.setWordWrap(True)
        self.preview_path.setStyleSheet("color:#8FA6C1;")
        right_layout.addWidget(self.preview_path)

        self.preview_frame = QFrame()
        self.preview_frame.setStyleSheet("background:#0F1823; border:1px solid #243446; border-radius:10px;")
        frame_layout = QVBoxLayout(self.preview_frame)
        frame_layout.setContentsMargins(12, 12, 12, 12)
        self.preview_label = QLabel("No preview selected.")
        self.preview_label.setAlignment(Qt.AlignCenter)
        self.preview_label.setMinimumSize(480, 300)
        self.preview_label.setStyleSheet("color:#9BB3CC;")
        frame_layout.addWidget(self.preview_label, stretch=1)
        right_layout.addWidget(self.preview_frame, stretch=1)
        splitter.addWidget(right_panel)
        splitter.setStretchFactor(0, 2)
        splitter.setStretchFactor(1, 5)
        splitter.setSizes([320, 780])

        self.reload_entries()

    def reload_entries(self):
        selected_entry_id = self.current_entry.get("entry_id") if self.current_entry else None
        self.reference_list.clear()
        self.entries = []

        for entry in self.BUILTIN_REFERENCES:
            self.entries.append(dict(entry))

        for path in self.config.get("legacy_gallery_paths") or []:
            self.entries.extend(discover_legacy_reference_entries(path))

        for entry in self.entries:
            item = QListWidgetItem(entry["label"])
            item.setData(Qt.UserRole, entry["entry_id"])
            if entry.get("builtin"):
                item.setToolTip("Built-in visual reference")
            else:
                item.setToolTip(entry.get("source_path") or entry.get("path") or entry["label"])
            self.reference_list.addItem(item)

        if self.entries:
            target_index = 0
            if selected_entry_id:
                for index, entry in enumerate(self.entries):
                    if entry["entry_id"] == selected_entry_id:
                        target_index = index
                        break
            self.reference_list.setCurrentRow(target_index)
        else:
            self.current_entry = None
            self.preview_title.setText("Select a legacy reference")
            self.preview_path.setText("No legacy reference selected yet.")
            self.preview_label.setPixmap(QPixmap())
            self.preview_label.setText("No preview selected.")

    def save_config(self):
        payload = self.save_config_callback({"legacy_gallery_paths": self.config.get("legacy_gallery_paths") or []})
        self.config = dict(payload or {})

    def add_reference_files(self):
        file_paths, _ = QFileDialog.getOpenFileNames(
            self,
            "Add Legacy Reference Files",
            self.config.get("packet_builder_export_dir") or _default_export_dir(),
            "Reference Files (*.py *.qss *.png *.jpg *.jpeg *.bmp *.webp *.ico);;All Files (*.*)",
        )
        if not file_paths:
            return

        existing = list(self.config.get("legacy_gallery_paths") or [])
        for path in file_paths:
            if path not in existing:
                existing.append(path)
        self.config["legacy_gallery_paths"] = existing
        self.save_config()
        self.reload_entries()

    def add_reference_folder(self):
        folder_path = QFileDialog.getExistingDirectory(
            self,
            "Add Legacy Reference Folder",
            self.config.get("packet_builder_export_dir") or _default_export_dir(),
        )
        if not folder_path:
            return
        existing = list(self.config.get("legacy_gallery_paths") or [])
        if folder_path not in existing:
            existing.append(folder_path)
        self.config["legacy_gallery_paths"] = existing
        self.save_config()
        self.reload_entries()

    def remove_selected_reference(self):
        item = self.reference_list.currentItem()
        if not item:
            return
        entry_id = item.data(Qt.UserRole)
        selected_entry = next((entry for entry in self.entries if entry.get("entry_id") == entry_id), None)
        if not selected_entry:
            return
        source_path = selected_entry.get("source_path")
        removable = [p for p in (self.config.get("legacy_gallery_paths") or []) if p != source_path]
        if selected_entry.get("builtin") or len(removable) == len(self.config.get("legacy_gallery_paths") or []):
            QMessageBox.information(
                self,
                "Built-In Reference",
                "Built-in references cannot be removed. Custom launcher and GUI references can.",
            )
            return
        self.config["legacy_gallery_paths"] = removable
        self.save_config()
        self.reload_entries()

    def on_selection_changed(self):
        item = self.reference_list.currentItem()
        if not item:
            return
        entry_id = item.data(Qt.UserRole)
        for entry in self.entries:
            if entry["entry_id"] == entry_id:
                self.current_entry = entry
                break
        self.render_preview()

    def render_preview(self):
        if not self.current_entry:
            return
        self.preview_title.setText(self.current_entry["label"])
        source_label = self.current_entry.get("source_path") or self.current_entry.get("path") or ""
        self.preview_path.setText(source_label or "Legacy preview reference")
        pixmap = self._build_preview_pixmap(self.current_entry)
        if pixmap.isNull():
            self.preview_label.setText("Preview could not be generated for this reference.")
            self.preview_label.setPixmap(QPixmap())
            return
        target_size = self.preview_label.size()
        if target_size.width() > 0 and target_size.height() > 0:
            pixmap = pixmap.scaled(target_size, Qt.KeepAspectRatio, Qt.SmoothTransformation)
        self.preview_label.setPixmap(pixmap)
        self.preview_label.setText("")

    def resizeEvent(self, event):
        super().resizeEvent(event)
        if self.current_entry:
            self.render_preview()

    def _build_preview_pixmap(self, entry):
        size = self.preview_label.size()
        width = max(720, size.width())
        height = max(420, size.height())
        kind = str(entry.get("kind") or "").strip().lower()
        if kind == "image":
            pixmap = QPixmap(entry.get("path") or "")
            if pixmap.isNull():
                return QPixmap()
            return pixmap.scaled(width, height, Qt.KeepAspectRatio, Qt.SmoothTransformation)
        if kind == "launcher_preview":
            return self._render_composed_preview(entry, width, height, mode="launcher")
        if kind == "gui_preview":
            return self._render_composed_preview(entry, width, height, mode="gui")
        return self._render_reference_card(entry, width, height)

    def _render_reference_card(self, entry, width, height):
        canvas = QPixmap(width, height)
        canvas.fill(QColor("#0F1823"))
        painter = QPainter(canvas)
        painter.setRenderHint(QPainter.Antialiasing)
        painter.fillRect(0, 0, width, height, QColor("#101925"))
        painter.setPen(QColor("#243446"))
        painter.drawRoundedRect(18, 18, width - 36, height - 36, 14, 14)
        painter.setPen(QColor("#69BCFF"))
        painter.setFont(QFont("Segoe UI", 18, QFont.Bold))
        painter.drawText(QRectF(40, 40, width - 80, 34), Qt.AlignLeft | Qt.AlignVCenter, entry.get("label") or "Legacy Reference")
        painter.setPen(QColor("#C8D8E8"))
        painter.setFont(QFont("Segoe UI", 11))
        painter.drawText(
            QRectF(40, 92, width - 80, height - 132),
            Qt.TextWordWrap,
            "This source reference was added for preview-only context. No runnable launcher or GUI preview could be composed from the selected files yet.",
        )
        painter.end()
        return canvas

    def _render_composed_preview(self, entry, width, height, mode="launcher"):
        canvas = QPixmap(width, height)
        canvas.fill(QColor("#0F1823"))
        painter = QPainter(canvas)
        painter.setRenderHint(QPainter.Antialiasing)
        self._paint_preview_background(painter, entry.get("background_path"), width, height)
        if mode == "launcher":
            self._paint_launcher_mock(painter, entry, width, height)
        else:
            self._paint_gui_mock(painter, entry, width, height)
        painter.end()
        return canvas

    def _paint_preview_background(self, painter, background_path, width, height):
        background = QPixmap(background_path or "")
        if not background.isNull():
            scaled = background.scaled(width, height, Qt.KeepAspectRatioByExpanding, Qt.SmoothTransformation)
            x = max(0, (scaled.width() - width) // 2)
            y = max(0, (scaled.height() - height) // 2)
            painter.drawPixmap(0, 0, scaled, x, y, width, height)
        else:
            painter.fillRect(0, 0, width, height, QColor("#0F1823"))
        painter.fillRect(0, 0, width, height, QColor(9, 18, 29, 210))

    def _paint_logo(self, painter, logo_path, rect):
        logo = QPixmap(logo_path or "")
        if logo.isNull():
            return
        scaled = logo.scaled(int(rect.width()), int(rect.height()), Qt.KeepAspectRatio, Qt.SmoothTransformation)
        x = rect.x() + (rect.width() - scaled.width()) / 2
        y = rect.y() + (rect.height() - scaled.height()) / 2
        painter.drawPixmap(int(x), int(y), scaled)

    def _paint_launcher_mock(self, painter, entry, width, height):
        card_width = min(860, width - 70)
        card_height = min(500, height - 70)
        card_x = (width - card_width) / 2
        card_y = (height - card_height) / 2
        painter.fillRect(0, 0, width, height, QColor(5, 12, 20, 75))
        painter.setPen(QColor("#2B4258"))
        painter.setBrush(QColor(14, 24, 36, 238))
        painter.drawRoundedRect(QRectF(card_x, card_y, card_width, card_height), 18, 18)

        self._paint_logo(painter, entry.get("logo_path"), QRectF(card_x + card_width / 2 - 70, card_y + 18, 140, 84))

        left_x = card_x + 24
        left_y = card_y + 118
        left_w = card_width * 0.6 - 30
        left_h = card_height - 150
        right_x = card_x + card_width * 0.63
        right_y = left_y
        right_w = card_width * 0.31
        right_h = left_h

        painter.setPen(QColor("#34506B"))
        painter.setBrush(QColor(16, 26, 38, 246))
        painter.drawRoundedRect(QRectF(left_x, left_y, left_w, left_h), 14, 14)
        painter.drawRoundedRect(QRectF(right_x, right_y, right_w, right_h), 14, 14)

        painter.setPen(QColor("#69BCFF"))
        painter.setFont(QFont("Segoe UI", 18, QFont.Bold))
        painter.drawText(QRectF(left_x + 18, left_y + 16, left_w - 36, 26), Qt.AlignLeft | Qt.AlignVCenter, "Legacy Launcher Preview")
        painter.setPen(QColor("#C8D8E8"))
        painter.setFont(QFont("Segoe UI", 10))
        lines = [
            "Launcher version: preview-only",
            "Update notes hidden",
            "This mock is generated from the added legacy reference",
        ]
        for index, line in enumerate(lines):
            painter.drawText(QRectF(left_x + 20, left_y + 66 + index * 28, left_w - 40, 24), Qt.AlignLeft | Qt.AlignVCenter, line)

        painter.setPen(QColor("#34506B"))
        for index in range(4):
            box_y = left_y + 170 + index * 48
            painter.drawRoundedRect(QRectF(left_x + 20, box_y, left_w - 40, 34), 8, 8)

        painter.setPen(QColor("#69BCFF"))
        painter.setFont(QFont("Segoe UI", 18, QFont.Bold))
        painter.drawText(QRectF(right_x + 16, right_y + 18, right_w - 32, 28), Qt.AlignLeft | Qt.AlignVCenter, "Sign In")
        painter.setPen(QColor("#34506B"))
        for index in range(5):
            field_y = right_y + 64 + index * 44
            painter.drawRoundedRect(QRectF(right_x + 16, field_y, right_w - 32, 30), 8, 8)
        painter.drawRoundedRect(QRectF(right_x + 16, right_y + right_h - 60, right_w - 32, 38), 10, 10)

    def _paint_gui_mock(self, painter, entry, width, height):
        frame_width = min(980, width - 70)
        frame_height = min(560, height - 70)
        frame_x = (width - frame_width) / 2
        frame_y = (height - frame_height) / 2

        painter.setPen(QColor("#2B4258"))
        painter.setBrush(QColor(14, 24, 36, 238))
        painter.drawRoundedRect(QRectF(frame_x, frame_y, frame_width, frame_height), 16, 16)
        painter.fillRect(QRectF(frame_x, frame_y, frame_width, 54), QColor(18, 31, 46, 248))

        self._paint_logo(painter, entry.get("logo_path"), QRectF(frame_x + 18, frame_y + 6, 70, 42))

        painter.setPen(QColor("#69BCFF"))
        painter.setFont(QFont("Segoe UI", 16, QFont.Bold))
        painter.drawText(QRectF(frame_x + 102, frame_y + 11, 300, 30), Qt.AlignLeft | Qt.AlignVCenter, "Legacy GUI Preview")

        tab_titles = ["Overview", "Packet Details", "Admin"]
        for index, title in enumerate(tab_titles):
            tab_x = frame_x + 24 + index * 148
            tab_y = frame_y + 72
            painter.setPen(QColor("#34506B"))
            painter.setBrush(QColor(20, 33, 48, 248) if index else QColor(28, 48, 68, 250))
            painter.drawRoundedRect(QRectF(tab_x, tab_y, 130, 34), 10, 10)
            painter.setPen(QColor("#FFFFFF") if index == 0 else QColor("#C8D8E8"))
            painter.setFont(QFont("Segoe UI", 10, QFont.Bold if index == 0 else QFont.Normal))
            painter.drawText(QRectF(tab_x, tab_y, 130, 34), Qt.AlignCenter, title)

        content_x = frame_x + 24
        content_y = frame_y + 124
        left_w = frame_width * 0.58
        right_w = frame_width * 0.34
        content_h = frame_height - 150
        painter.setPen(QColor("#243446"))
        painter.setBrush(QColor(15, 24, 35, 240))
        painter.drawRoundedRect(QRectF(content_x, content_y, left_w, content_h), 12, 12)
        painter.drawRoundedRect(QRectF(frame_x + frame_width - right_w - 24, content_y, right_w, content_h), 12, 12)

        painter.setPen(QColor("#69BCFF"))
        painter.setFont(QFont("Segoe UI", 15, QFont.Bold))
        painter.drawText(QRectF(content_x + 16, content_y + 16, left_w - 32, 24), Qt.AlignLeft | Qt.AlignVCenter, "Packet Workspace")
        painter.drawText(QRectF(frame_x + frame_width - right_w - 8, content_y + 16, right_w - 32, 24), Qt.AlignLeft | Qt.AlignVCenter, "Details Panel")

        painter.setPen(QColor("#34506B"))
        for index in range(4):
            block_y = content_y + 58 + index * 68
            painter.drawRoundedRect(QRectF(content_x + 16, block_y, left_w - 32, 50), 8, 8)
        for index in range(5):
            block_y = content_y + 58 + index * 56
            painter.drawRoundedRect(QRectF(frame_x + frame_width - right_w - 8, block_y, right_w - 48, 40), 8, 8)


class PacketBuilderTab(QWidget):
    def __init__(self, config, save_config_callback, parent=None):
        super().__init__(parent)
        self.config = dict(config or {})
        self.save_config_callback = save_config_callback
        self.packet_payload = default_packet_builder_payload()
        self._builder_ready = False
        self._shared_sync_state = {}
        self._wording_assist_state = {}
        self._wording_entries = {}
        self._current_wording_entry_key = ""
        self.current_draft_id = None
        self.current_draft_record = {}
        self.word_executable_path = find_word_executable()
        self._preview_cache_dir = runtime_data_path("dev_system", "preview_cache")
        self._preview_render_thread = None
        self._preview_render_worker = None
        self._preview_latest_requested_key = ""
        self._preview_loaded_key = ""
        self._preview_current_key = ""
        self._export_thread = None
        self._export_worker = None
        self._export_controls = []
        self._signature_images = {}
        self._preview_refresh_timer = QTimer(self)
        self._preview_refresh_timer.setSingleShot(True)
        self._preview_refresh_timer.setInterval(120)
        self._preview_refresh_timer.timeout.connect(self._refresh_preview_now)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(12)

        intro = QLabel(
            "Packet Builder Studio is the dev-side drafting shell. It lets you stage packet data once, carry shared header details across forms, preview the document, and use Lab to catch gaps before export."
        )
        intro.setWordWrap(True)
        intro.setStyleSheet("color:#BFD0E3;")
        layout.addWidget(intro)

        self.current_packet_status_label = QLabel()
        self.current_packet_status_label.setWordWrap(True)
        self.current_packet_status_label.setStyleSheet("color:#8FA6C1; font-weight:600;")
        layout.addWidget(self.current_packet_status_label)

        splitter = QSplitter(Qt.Horizontal)
        splitter.setChildrenCollapsible(False)
        splitter.setHandleWidth(6)
        self.builder_splitter = splitter
        layout.addWidget(splitter, stretch=1)

        form_scroll = QScrollArea()
        form_scroll.setWidgetResizable(True)
        form_scroll.setFrameShape(QFrame.NoFrame)
        form_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        form_scroll.viewport().setStyleSheet("background:#0F1823;")
        form_scroll.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        self.form_scroll = form_scroll
        splitter.addWidget(form_scroll)

        form_panel = QWidget()
        form_panel.setStyleSheet("background:#0F1823;")
        form_panel.setMinimumWidth(0)
        form_panel.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
        form_layout = QVBoxLayout(form_panel)
        form_layout.setContentsMargins(0, 0, 6, 0)
        form_layout.setSpacing(12)
        self.form_panel = form_panel
        form_scroll.setWidget(form_panel)

        export_group = QGroupBox("Export")
        export_group.setStyleSheet("QGroupBox { font-weight:700; color:#E5E7EB; }")
        export_layout = QVBoxLayout(export_group)
        export_layout.setSpacing(10)

        self.export_dir_value = QLabel(self.config.get("packet_builder_export_dir") or _default_export_dir())
        self.export_dir_value.setWordWrap(True)
        self.export_dir_value.setStyleSheet("color:#8FA6C1;")
        choose_folder_button = QPushButton("Choose Folder")
        self._style_builder_button(choose_folder_button)
        choose_folder_button.setMaximumWidth(150)
        choose_folder_button.clicked.connect(self.choose_export_dir)
        export_layout.addWidget(QLabel("Export Folder"))
        export_path_row = QHBoxLayout()
        export_path_row.setContentsMargins(0, 0, 0, 0)
        export_path_row.setSpacing(10)
        export_path_row.addWidget(self.export_dir_value, stretch=1)
        export_path_row.addWidget(choose_folder_button, 0, Qt.AlignRight)
        export_layout.addLayout(export_path_row)

        self.base_filename_input = QLineEdit("truecore_packet")
        self.base_filename_input.textChanged.connect(self.refresh_preview)
        export_layout.addWidget(QLabel("Base Filename"))
        export_layout.addWidget(self.base_filename_input)
        form_layout.addWidget(export_group)

        editor_tabs = QTabWidget()
        editor_tabs.setDocumentMode(True)
        editor_tabs.setStyleSheet(
            "QTabWidget::pane { border:1px solid #243446; border-radius:12px; background:#0F1823; top:-1px; }"
            "QTabBar::tab { background:#13202E; color:#DDE8F5; padding:9px 16px; border-top-left-radius:10px; border-top-right-radius:10px; margin-right:4px; }"
            "QTabBar::tab:selected { background:#1D2B3A; color:#FFFFFF; font-weight:700; }"
            "QTabBar::tab:hover:!selected { background:#182433; }"
        )
        self.packet_builder_left_tabs = editor_tabs
        form_layout.addWidget(editor_tabs, stretch=1)

        intake_page = QWidget()
        intake_page.setStyleSheet("background:#0F1823;")
        intake_layout = QVBoxLayout(intake_page)
        intake_layout.setContentsMargins(12, 12, 12, 12)
        intake_layout.setSpacing(12)
        self.packet_intake_page = intake_page

        workspace_page = QWidget()
        workspace_page.setStyleSheet("background:#0F1823;")
        workspace_layout = QVBoxLayout(workspace_page)
        workspace_layout.setContentsMargins(12, 12, 12, 12)
        workspace_layout.setSpacing(12)
        self.packet_workspace_page = workspace_page

        editor_tabs.addTab(intake_page, "Packet Intake")
        editor_tabs.addTab(workspace_page, "Form Workspace")
        editor_tabs.setCurrentWidget(workspace_page)

        current_form_group = QGroupBox("Current Form")
        current_form_group.setStyleSheet("QGroupBox { font-weight:700; color:#E5E7EB; }")
        current_form_layout = QFormLayout(current_form_group)
        self._configure_form_layout(current_form_layout)

        current_form_help = QLabel(
            "Use Packet Intake for shared packet facts and scenario prompts. Only the active form below is export-facing."
        )
        current_form_help.setWordWrap(True)
        current_form_help.setStyleSheet("color:#8FA6C1;")
        current_form_layout.addRow(current_form_help)

        self.profile_combo = QComboBox()
        self.profile_combo.addItems(PACKET_BUILDER_PROFILES)
        self.profile_combo.currentTextChanged.connect(self.on_profile_changed)
        self.profile_combo.setMaximumWidth(360)
        self.profile_combo.setSizePolicy(QSizePolicy.Maximum, QSizePolicy.Fixed)
        current_form_layout.addRow("Packet Profile", self.profile_combo)

        self.export_context_label = QLabel("")
        self.export_context_label.setWordWrap(True)
        self.export_context_label.setStyleSheet("color:#8FA6C1;")
        current_form_layout.addRow("Export Context", self.export_context_label)

        self.packet_title_input = self._make_line_edit("Packet title")
        current_form_layout.addRow("Packet Title", self.packet_title_input)
        workspace_layout.addWidget(current_form_group)

        details_group = QGroupBox("Packet Intake")
        details_group.setStyleSheet("QGroupBox { font-weight:700; color:#E5E7EB; }")
        details_form = QFormLayout(details_group)
        self._configure_form_layout(details_form)

        shared_header_help = QLabel(
            "Enter repeated packet facts here once. This intake area is never exported. Packet Studio carries these values into matching form fields where expected, while page-specific reasoning stays manual.<br/><br/>"
            "<strong>VA Referring Source</strong> = VA facility and VA referring / ordering provider.<br/>"
            "<strong>Community Care Destination</strong> = the outside specialist / practice the Veteran is being referred to.<br/>"
            "<strong>PCP / PCM</strong> = the Veteran's primary care office, used only where the form specifically asks for it."
        )
        shared_header_help.setWordWrap(True)
        shared_header_help.setStyleSheet("color:#8FA6C1;")
        details_form.addRow(shared_header_help)

        def add_section_label(text):
            label = QLabel(text)
            label.setStyleSheet("color:#C8D8E8; font-weight:700; padding-top:6px;")
            details_form.addRow(label)

        add_section_label("Veteran / VA Identifiers")
        self.patient_name_input = self._make_line_edit("Patient name")
        details_form.addRow("Veteran Full Legal Name", self.patient_name_input)
        self.dob_input = self._make_line_edit("MM/DD/YYYY")
        details_form.addRow("Date of Birth", self.dob_input)
        self.auth_input = self._make_line_edit("VA authorization number")
        details_form.addRow("VA Authorization Number", self.auth_input)
        self.icn_input = self._make_line_edit("VA ICN")
        details_form.addRow("VA ICN", self.icn_input)
        add_section_label("VA Referring Source")
        self.ordering_doctor_input = self._make_line_edit("VA referring / ordering provider name")
        details_form.addRow("VA Referring / Ordering Provider", self.ordering_doctor_input)
        self.facility_input = self._make_line_edit("VA facility / VA medical center")
        details_form.addRow("VA Referring Facility / Medical Center", self.facility_input)
        add_section_label("Community Care Destination")
        self.provider_input = self._make_line_edit("Community Care treating provider name")
        details_form.addRow("Community Care Treating Provider", self.provider_input)
        self.community_facility_input = self._make_line_edit("Community Care referred-to practice / facility")
        details_form.addRow("Community Care Practice / Facility", self.community_facility_input)
        self.master_provider_credentials_input = self._make_line_edit("Community Care provider credentials")
        details_form.addRow("Community Care Provider Credentials", self.master_provider_credentials_input)
        self.master_provider_specialty_input = self._make_line_edit("Community Care provider specialty")
        details_form.addRow("Community Care Provider Specialty", self.master_provider_specialty_input)
        self.master_provider_npi_input = self._make_line_edit("Community Care provider NPI")
        details_form.addRow("Community Care Provider NPI", self.master_provider_npi_input)
        self.master_practice_name_input = self._make_line_edit("Community Care practice name")
        details_form.addRow("Community Care Practice Name", self.master_practice_name_input)
        self.master_provider_phone_input = self._make_line_edit("Community Care practice phone")
        details_form.addRow("Community Care Practice Phone", self.master_provider_phone_input)
        self.master_provider_fax_input = self._make_line_edit("Community Care practice fax")
        details_form.addRow("Community Care Practice Fax", self.master_provider_fax_input)
        self.master_provider_email_input = self._make_line_edit("Community Care secure email")
        details_form.addRow("Community Care Secure Email", self.master_provider_email_input)

        self.master_provider_address_input = QTextEdit()
        self.master_provider_address_input.setMinimumHeight(70)
        self.master_provider_address_input.textChanged.connect(self.refresh_preview)
        details_form.addRow("Community Care Practice Address", self.master_provider_address_input)

        add_section_label("Clinical Baseline Shared Across Forms")
        self.requested_service_input = self._make_line_edit("Requested service")
        details_form.addRow("Requested Service", self.requested_service_input)
        self.diagnosis_input = self._make_line_edit("Diagnosis")
        details_form.addRow("Primary Diagnosis", self.diagnosis_input)
        self.secondary_diagnosis_input = self._make_line_edit("Secondary diagnosis")
        details_form.addRow("Secondary Diagnosis", self.secondary_diagnosis_input)
        self.icd_codes_input = self._make_line_edit("ICD codes")
        details_form.addRow("ICD Codes", self.icd_codes_input)
        self.master_requested_cpt_code_input = self._make_line_edit("Requested CPT / HCPCS code")
        details_form.addRow("Requested CPT / HCPCS", self.master_requested_cpt_code_input)

        self.master_mri_date_input = self._make_line_edit("MM/DD/YYYY")
        details_form.addRow("MRI Date", self.master_mri_date_input)
        self.master_mri_findings_input = self._make_line_edit("MRI findings")
        details_form.addRow("MRI Findings", self.master_mri_findings_input)
        self.master_affected_levels_input = self._make_line_edit("Affected spinal levels")
        details_form.addRow("Affected Spinal Levels", self.master_affected_levels_input)

        self.clinical_summary_input = QTextEdit()
        self.clinical_summary_input.setMinimumHeight(90)
        self.clinical_summary_input.textChanged.connect(self.refresh_preview)
        details_form.addRow("Clinical Summary / Symptoms", self.clinical_summary_input)

        self.packet_notes_input = QTextEdit()
        self.packet_notes_input.setMinimumHeight(90)
        self.packet_notes_input.textChanged.connect(self.refresh_preview)
        details_form.addRow("Packet Notes", self.packet_notes_input)

        intake_layout.addWidget(details_group)

        scenario_group = QGroupBox("Scenario Prompts")
        scenario_group.setStyleSheet("QGroupBox { font-weight:700; color:#E5E7EB; }")
        scenario_form = QFormLayout(scenario_group)
        self._configure_form_layout(scenario_form)

        scenario_help = QLabel(
            "These prompts are not exported. Choose one or more where needed. They help Wording Assist choose cleaner blueprint families without forcing you to over-explain the same packet facts on every form."
        )
        scenario_help.setWordWrap(True)
        scenario_help.setStyleSheet("color:#8FA6C1;")
        scenario_form.addRow(scenario_help)

        self.scenario_pathology_pattern_input = MultiSelectPromptField(
            "Primary Pathology Pattern",
            [
                "Discogenic / Annular",
                "Radicular",
                "Disc Displacement / Herniation",
                "Multilevel",
            ],
        )
        self.scenario_pathology_pattern_input.textChanged.connect(self.refresh_preview)
        scenario_form.addRow("Primary Pathology Pattern", self.scenario_pathology_pattern_input)

        self.scenario_conservative_duration_input = MultiSelectPromptField(
            "Conservative Care Duration",
            [
                "Over 90 Days",
                "Over 6 Months",
                "Over 12 Months",
            ],
        )
        self.scenario_conservative_duration_input.textChanged.connect(self.refresh_preview)
        scenario_form.addRow("Conservative Care Duration", self.scenario_conservative_duration_input)

        self.scenario_prior_esi_response_input = MultiSelectPromptField(
            "Prior ESI Response",
            [
                "None / Not Documented",
                "Temporary Relief",
                "Partial Relief",
                "No Relief",
            ],
        )
        self.scenario_prior_esi_response_input.textChanged.connect(self.refresh_preview)
        scenario_form.addRow("Prior ESI Response", self.scenario_prior_esi_response_input)

        self.scenario_functional_emphasis_input = MultiSelectPromptField(
            "Functional Emphasis",
            [
                "Sleep Disruption",
                "Work Capacity",
                "Sitting / Standing Tolerance",
                "Daily Activities",
            ],
        )
        self.scenario_functional_emphasis_input.textChanged.connect(self.refresh_preview)
        scenario_form.addRow("Functional Emphasis", self.scenario_functional_emphasis_input)

        self.scenario_request_framing_input = MultiSelectPromptField(
            "Request Framing",
            [
                "Procedure-Neutral Evaluation",
                "Defined Interventional Pathway",
                "Specific Requested Procedure",
            ],
        )
        self.scenario_request_framing_input.textChanged.connect(self.refresh_preview)
        scenario_form.addRow("Request Framing", self.scenario_request_framing_input)

        self.scenario_symptom_pattern_input = MultiSelectPromptField(
            "Symptom Pattern",
            [
                "Axial Lumbar Pain",
                "Discogenic Pattern",
                "Radicular Symptoms",
                "Activity-Related Exacerbation",
                "Position Intolerance",
            ],
        )
        self.scenario_symptom_pattern_input.textChanged.connect(self.refresh_preview)
        scenario_form.addRow("Symptom Pattern", self.scenario_symptom_pattern_input)

        self.scenario_conservative_modalities_input = MultiSelectPromptField(
            "Conservative Modalities Already Tried",
            [
                "Physical Therapy",
                "Home Exercise Program",
                "NSAIDs / Analgesics",
                "Activity Modification",
                "Prior ESI / Interventional",
            ],
        )
        self.scenario_conservative_modalities_input.textChanged.connect(self.refresh_preview)
        scenario_form.addRow("Conservative Modalities Tried", self.scenario_conservative_modalities_input)

        self.scenario_review_concern_input = MultiSelectPromptField(
            "Primary Review Concern",
            [
                "Scope Must Stay Limited",
                "Medical Necessity Emphasis",
                "Imaging Correlation Emphasis",
                "Failed Conservative Care Emphasis",
            ],
        )
        self.scenario_review_concern_input.textChanged.connect(self.refresh_preview)
        scenario_form.addRow("Primary Review Concern", self.scenario_review_concern_input)

        self.scenario_treatment_goals_input = MultiSelectPromptField(
            "Treatment Goals to Emphasize",
            [
                "Pain Reduction",
                "Functional Improvement",
                "Reduce Analgesics",
                "Avoid Surgery",
                "Prevent Progression",
            ],
        )
        self.scenario_treatment_goals_input.textChanged.connect(self.refresh_preview)
        scenario_form.addRow("Treatment Goals", self.scenario_treatment_goals_input)
        intake_layout.addWidget(scenario_group)
        intake_layout.addStretch(1)

        self.referral_group = QGroupBox("Community Care Referral Request")
        self.referral_group.setStyleSheet("QGroupBox { font-weight:700; color:#E5E7EB; }")
        referral_form = QFormLayout(self.referral_group)
        self._configure_form_layout(referral_form)

        self.referral_subtitle_input = self._make_line_edit("Spine & Pain Evaluation")
        self.referral_subtitle_input.setText("Spine and Pain Evaluation")
        referral_form.addRow("Subtitle", self.referral_subtitle_input)

        self.pcp_request_input = self._make_line_edit("PCP request wording")
        self.pcp_request_input.setText(
            "I am requesting a Community Care referral for pain management and spine evaluation to determine the most appropriate treatment plan."
        )
        referral_form.addRow("PCP Request Wording", self.pcp_request_input)

        self.referral_entry_input = self._make_line_edit("Referral entry wording")
        self.referral_entry_input.setText("Specialty spine and pain evaluation and consultation")
        referral_form.addRow("Referral Entry Wording", self.referral_entry_input)

        self.areas_of_concern_input = self._make_line_edit("Cervical, Thoracic, Lumbar, Joints (if applicable)")
        self.areas_of_concern_input.setText("Cervical, Thoracic, Lumbar, Joints (if applicable)")
        referral_form.addRow("Areas of Concern", self.areas_of_concern_input)

        self.group_npi_input = self._make_line_edit("Group NPI")
        referral_form.addRow("Group NPI", self.group_npi_input)

        self.fax_number_input = self._make_line_edit("Fax number")
        referral_form.addRow("Fax Number", self.fax_number_input)

        self.liaison_contact_input = QTextEdit()
        self.liaison_contact_input.setMinimumHeight(80)
        self.liaison_contact_input.textChanged.connect(self.refresh_preview)
        self.liaison_contact_input.setPlaceholderText("Veteran liaison contact information")
        referral_form.addRow("Liaison Contact", self.liaison_contact_input)

        workspace_layout.addWidget(self.referral_group)

        self.virtual_consent_group = QGroupBox("Virtual Consent Form")
        self.virtual_consent_group.setStyleSheet("QGroupBox { font-weight:700; color:#E5E7EB; }")
        consent_form = QFormLayout(self.virtual_consent_group)
        self._configure_form_layout(consent_form)

        self.consent_form_title_input = self._make_line_edit("TELEHEALTH VIRTUAL CONSENT FORM")
        self.consent_form_title_input.setText("TELEHEALTH VIRTUAL CONSENT FORM")
        consent_form.addRow("Form Title", self.consent_form_title_input)

        self.street_address_input = self._make_line_edit("Street address")
        consent_form.addRow("Street Address", self.street_address_input)
        self.city_input = self._make_line_edit("City")
        consent_form.addRow("City", self.city_input)
        self.state_input = self._make_line_edit("State")
        consent_form.addRow("State", self.state_input)
        self.zip_code_input = self._make_line_edit("Zip code")
        consent_form.addRow("Zip Code", self.zip_code_input)

        self.home_phone_input = self._make_line_edit("Home phone")
        consent_form.addRow("Home Phone", self.home_phone_input)
        self.mobile_phone_input = self._make_line_edit("Mobile phone")
        consent_form.addRow("Mobile Phone", self.mobile_phone_input)
        self.work_phone_input = self._make_line_edit("Work phone")
        consent_form.addRow("Work Phone", self.work_phone_input)

        self.email_address_input = self._make_line_edit("Email address")
        consent_form.addRow("Email Address", self.email_address_input)
        self.ssn_input = self._make_line_edit("SSN")
        consent_form.addRow("SSN", self.ssn_input)
        self.drivers_license_input = self._make_line_edit("Driver license")
        consent_form.addRow("Driver License", self.drivers_license_input)
        self.drivers_license_state_input = self._make_line_edit("DL state")
        consent_form.addRow("DL State", self.drivers_license_state_input)

        self.appointment_confirmation_combo = self._make_choice_combo(["Phone", "Text", "Email"])
        consent_form.addRow("Appointment Confirmation", self.appointment_confirmation_combo)
        self.filed_for_disability_combo = self._make_choice_combo(["No", "Yes"])
        consent_form.addRow("Filed for Disability", self.filed_for_disability_combo)
        self.condition_work_related_combo = self._make_choice_combo(["No", "Yes"])
        consent_form.addRow("Condition Work Related", self.condition_work_related_combo)
        self.condition_due_to_accident_combo = self._make_choice_combo(["No", "Yes"])
        consent_form.addRow("Condition Due to Accident", self.condition_due_to_accident_combo)

        self.yes_response_explanation_input = QTextEdit()
        self.yes_response_explanation_input.setMinimumHeight(70)
        self.yes_response_explanation_input.textChanged.connect(self.refresh_preview)
        self.yes_response_explanation_input.setPlaceholderText("Explain any Yes responses")
        consent_form.addRow("Yes Response Explanation", self.yes_response_explanation_input)

        self.has_attorney_combo = self._make_choice_combo(["No", "Yes"])
        consent_form.addRow("Has Attorney", self.has_attorney_combo)
        self.attorney_name_input = self._make_line_edit("Attorney full name")
        consent_form.addRow("Attorney Name", self.attorney_name_input)
        self.attorney_phone_input = self._make_line_edit("Attorney phone")
        consent_form.addRow("Attorney Phone", self.attorney_phone_input)

        self.emergency_contact_name_input = self._make_line_edit("Emergency contact")
        consent_form.addRow("Emergency Contact", self.emergency_contact_name_input)
        self.emergency_contact_relationship_input = self._make_line_edit("Relationship")
        consent_form.addRow("Relationship", self.emergency_contact_relationship_input)
        self.emergency_contact_phone_input = self._make_line_edit("Emergency contact phone")
        consent_form.addRow("Emergency Contact Phone", self.emergency_contact_phone_input)

        self.primary_insurance_carrier_input = self._make_line_edit("Primary insurance carrier")
        consent_form.addRow("Primary Insurance Carrier", self.primary_insurance_carrier_input)
        self.primary_insurance_id_input = self._make_line_edit("Primary insurance ID")
        consent_form.addRow("Primary Insurance ID", self.primary_insurance_id_input)
        self.primary_insurance_phone_input = self._make_line_edit("Primary insurance phone")
        consent_form.addRow("Primary Insurance Phone", self.primary_insurance_phone_input)

        self.secondary_insurance_carrier_input = self._make_line_edit("Secondary insurance carrier")
        consent_form.addRow("Secondary Insurance Carrier", self.secondary_insurance_carrier_input)
        self.secondary_insurance_id_input = self._make_line_edit("Secondary insurance ID")
        consent_form.addRow("Secondary Insurance ID", self.secondary_insurance_id_input)
        self.secondary_insurance_phone_input = self._make_line_edit("Secondary insurance phone")
        consent_form.addRow("Secondary Insurance Phone", self.secondary_insurance_phone_input)

        self.pcp_pcm_name_input = self._make_line_edit("PCP/PCM")
        consent_form.addRow("PCP / PCM", self.pcp_pcm_name_input)
        self.pcp_pcm_phone_input = self._make_line_edit("PCP/PCM phone")
        consent_form.addRow("PCP / PCM Phone", self.pcp_pcm_phone_input)
        self.pcp_pcm_fax_input = self._make_line_edit("PCP/PCM fax")
        consent_form.addRow("PCP / PCM Fax", self.pcp_pcm_fax_input)

        self.consent_provider_name_input = self._make_line_edit("Care provider name")
        consent_form.addRow("Consent Provider Name", self.consent_provider_name_input)
        self.consent_initials_input = self._make_line_edit("Patient initials")
        consent_form.addRow("Consent Initials", self.consent_initials_input)
        self.minor_doctor_name_input = self._make_line_edit("Doctor name for minor consent")
        consent_form.addRow("Minor Doctor Name", self.minor_doctor_name_input)
        self.minor_consent_initials_input = self._make_line_edit("Parent / guardian initials")
        consent_form.addRow("Minor Consent Initials", self.minor_consent_initials_input)
        self.service_authorization_name_input = self._make_line_edit("Authorized service name")
        consent_form.addRow("Service Authorization Name", self.service_authorization_name_input)
        patient_signature_widget, self.patient_signature_name_input = self._make_signature_field("patient_signature_name", "Veteran / patient signature")
        consent_form.addRow("Veteran / Patient Signature", patient_signature_widget)
        self.patient_signature_date_input = self._make_line_edit("MM/DD/YYYY")
        consent_form.addRow("Veteran / Patient Signature Date", self.patient_signature_date_input)

        workspace_layout.addWidget(self.virtual_consent_group)

        self.submission_cover_group = QGroupBox("Submission Cover Sheet")
        self.submission_cover_group.setStyleSheet("QGroupBox { font-weight:700; color:#E5E7EB; }")
        cover_form = QFormLayout(self.submission_cover_group)
        self._configure_form_layout(cover_form)

        self.submission_cover_title_input = self._make_line_edit("VA Submission Cover Sheet")
        self.submission_cover_title_input.setText("VA Submission Cover Sheet")
        cover_form.addRow("Cover Title", self.submission_cover_title_input)

        self.submission_date_input = self._make_line_edit("MM/DD/YYYY")
        cover_form.addRow("Date of Submission", self.submission_date_input)
        self.primary_diagnosis_code_input = self._make_line_edit("Primary diagnosis code")
        cover_form.addRow("Primary Diagnosis Code", self.primary_diagnosis_code_input)

        self.included_virtual_consent_checkbox = self._make_checkbox("Virtual Consent Form completed and signed", True)
        cover_form.addRow("Included", self.included_virtual_consent_checkbox)
        self.included_va_form_checkbox = self._make_checkbox("VA Form 10-10172 completed and signed", True)
        cover_form.addRow("", self.included_va_form_checkbox)
        self.included_seoc_checkbox = self._make_checkbox("SEOC request signed by CCN ordering provider", True)
        cover_form.addRow("", self.included_seoc_checkbox)
        self.included_consult_checkbox = self._make_checkbox("Consultation & Treatment Request completed", True)
        cover_form.addRow("", self.included_consult_checkbox)
        self.included_lomn_checkbox = self._make_checkbox("Letter of Medical Necessity completed", True)
        cover_form.addRow("", self.included_lomn_checkbox)
        self.included_clinical_notes_checkbox = self._make_checkbox("Clinical Notes included", True)
        cover_form.addRow("", self.included_clinical_notes_checkbox)
        self.included_mri_checkbox = self._make_checkbox("MRI Report included", True)
        cover_form.addRow("", self.included_mri_checkbox)

        self.submitting_office_input = self._make_line_edit("Submitting office")
        cover_form.addRow("Submitting Office", self.submitting_office_input)
        self.office_staff_name_input = self._make_line_edit("Office staff name")
        cover_form.addRow("Office Staff Name", self.office_staff_name_input)
        office_signature_widget, self.office_staff_signature_input = self._make_signature_field("office_staff_signature", "Program user / office staff signature")
        cover_form.addRow("Program User / Office Staff Signature", office_signature_widget)
        self.date_reviewed_input = self._make_line_edit("MM/DD/YYYY")
        cover_form.addRow("Date Reviewed", self.date_reviewed_input)

        workspace_layout.addWidget(self.submission_cover_group)

        self.seoc_request_group = QGroupBox("Single Episode of Care Request Template")
        self.seoc_request_group.setStyleSheet("QGroupBox { font-weight:700; color:#E5E7EB; }")
        seoc_form = QFormLayout(self.seoc_request_group)
        self._configure_form_layout(seoc_form)

        self.seoc_request_date_input = self._make_line_edit("MM/DD/YYYY")
        seoc_form.addRow("Request Date", self.seoc_request_date_input)
        self.va_medical_center_input = self._make_line_edit("VA Medical Center Name")
        seoc_form.addRow("VA Medical Center", self.va_medical_center_input)
        self.last_four_ssn_input = self._make_line_edit("Last four SSN")
        seoc_form.addRow("Last Four SSN", self.last_four_ssn_input)
        self.episode_diagnosis_input = self._make_line_edit("Episode diagnosis / condition")
        seoc_form.addRow("Episode Diagnosis", self.episode_diagnosis_input)
        self.episode_icd_code_input = self._make_line_edit("ICD-10 code")
        seoc_form.addRow("Episode ICD Code", self.episode_icd_code_input)
        self.estimated_duration_input = self._make_line_edit("Thirty to ninety days, including routine post-procedure follow-up related only to the authorized episode of care.")
        self.estimated_duration_input.setText("Thirty to ninety days, including routine post-procedure follow-up related only to the authorized episode of care.")
        seoc_form.addRow("Estimated Duration", self.estimated_duration_input)

        self.seoc_preprocedure_checkbox = self._make_checkbox("Pre-procedure evaluation and procedural planning", True)
        seoc_form.addRow("Scope", self.seoc_preprocedure_checkbox)
        self.seoc_annulargram_checkbox = self._make_checkbox("Diagnostic Annulargram", True)
        seoc_form.addRow("", self.seoc_annulargram_checkbox)
        self.seoc_fibrin_checkbox = self._make_checkbox("Inter Annular Fibrin Injections if indicated", True)
        seoc_form.addRow("", self.seoc_fibrin_checkbox)
        self.seoc_follow_up_checkbox = self._make_checkbox("Standard post-procedure follow-up visit(s) within the global postoperative period", True)
        seoc_form.addRow("", self.seoc_follow_up_checkbox)
        self.seoc_scope_text_input = QTextEdit()
        self.seoc_scope_text_input.setMinimumHeight(78)
        self.seoc_scope_text_input.setPlaceholderText("Summarize the requested scope of this episode of care")
        self.seoc_scope_text_input.setPlainText(
            "This SEOC is limited to evaluation, procedural planning, the indicated intervention if clinically appropriate, and standard post-procedure follow-up related only to the documented condition. No unrelated care or open-ended pain management is requested under this episode."
        )
        self.seoc_scope_text_input.textChanged.connect(self.refresh_preview)
        seoc_form.addRow("Scope Language", self.seoc_scope_text_input)

        self.clinical_objectives_input = QTextEdit()
        self.clinical_objectives_input.setMinimumHeight(90)
        self.clinical_objectives_input.setPlaceholderText("One clinical objective per line")
        self.clinical_objectives_input.setPlainText(
            "Reduce pain severity\nImprove functional capacity\nSupport activity tolerance related to the documented condition"
        )
        self.clinical_objectives_input.textChanged.connect(self.refresh_preview)
        seoc_form.addRow("Clinical Objectives", self.clinical_objectives_input)
        self.seoc_continuity_text_input = QTextEdit()
        self.seoc_continuity_text_input.setMinimumHeight(78)
        self.seoc_continuity_text_input.setPlaceholderText("Describe continuity-of-care handling after the episode is complete")
        self.seoc_continuity_text_input.setPlainText(
            "Upon completion of the authorized episode, a treatment summary and clinical outcome update will be provided to the referring VA provider. Any additional or unrelated care will require separate evaluation and authorization."
        )
        self.seoc_continuity_text_input.textChanged.connect(self.refresh_preview)
        seoc_form.addRow("Continuity of Care", self.seoc_continuity_text_input)

        self.provider_credentials_input = self._make_line_edit("Community Care provider credentials")
        seoc_form.addRow("Community Care Provider Credentials", self.provider_credentials_input)
        self.provider_specialty_input = self._make_line_edit("Community Care provider specialty")
        seoc_form.addRow("Community Care Provider Specialty", self.provider_specialty_input)
        self.provider_npi_input = self._make_line_edit("Community Care provider NPI")
        seoc_form.addRow("Community Care Provider NPI", self.provider_npi_input)
        self.practice_name_input = self._make_line_edit("Community Care practice name")
        seoc_form.addRow("Community Care Practice Name", self.practice_name_input)
        self.provider_phone_input = self._make_line_edit("Community Care practice phone")
        seoc_form.addRow("Community Care Practice Phone", self.provider_phone_input)
        self.provider_fax_input = self._make_line_edit("Community Care practice fax")
        seoc_form.addRow("Community Care Practice Fax", self.provider_fax_input)

        workspace_layout.addWidget(self.seoc_request_group)

        self.lomn_group = QGroupBox("Letter of Medical Necessity Template")
        self.lomn_group.setStyleSheet("QGroupBox { font-weight:700; color:#E5E7EB; }")
        lomn_form = QFormLayout(self.lomn_group)
        self._configure_form_layout(lomn_form)

        self.lomn_request_date_input = self._make_line_edit("MM/DD/YYYY")
        lomn_form.addRow("Letter Date", self.lomn_request_date_input)
        self.lomn_va_claim_number_input = self._make_line_edit("VA claim number if applicable")
        lomn_form.addRow("VA Claim Number", self.lomn_va_claim_number_input)

        self.lomn_primary_diagnosis_input = self._make_line_edit("Primary diagnosis")
        lomn_form.addRow("Primary Diagnosis", self.lomn_primary_diagnosis_input)
        self.lomn_secondary_diagnosis_input = self._make_line_edit("Secondary diagnosis")
        lomn_form.addRow("Secondary Diagnosis", self.lomn_secondary_diagnosis_input)

        self.lomn_clinical_summary_input = QTextEdit()
        self.lomn_clinical_summary_input.setMinimumHeight(100)
        self.lomn_clinical_summary_input.setPlaceholderText("Summarize the clinical picture and functional impact")
        self.lomn_clinical_summary_input.setPlainText("")
        self.lomn_clinical_summary_input.textChanged.connect(self.refresh_preview)
        lomn_form.addRow("Clinical Basis", self.lomn_clinical_summary_input)

        self.lomn_mri_date_input = self._make_line_edit("MM/DD/YYYY")
        lomn_form.addRow("MRI Date", self.lomn_mri_date_input)
        self.lomn_mri_findings_input = self._make_line_edit("MRI findings")
        lomn_form.addRow("MRI Findings", self.lomn_mri_findings_input)
        self.lomn_conservative_duration_input = self._make_line_edit("Duration of conservative care")
        lomn_form.addRow("Conservative Care Duration", self.lomn_conservative_duration_input)

        self.lomn_physical_therapy_checkbox = self._make_checkbox("Structured physical therapy", True)
        lomn_form.addRow("Conservative Care", self.lomn_physical_therapy_checkbox)
        self.lomn_nsaids_checkbox = self._make_checkbox("NSAIDs and non-opioid analgesics", True)
        lomn_form.addRow("", self.lomn_nsaids_checkbox)
        self.lomn_activity_modification_checkbox = self._make_checkbox("Activity modification", True)
        lomn_form.addRow("", self.lomn_activity_modification_checkbox)
        self.lomn_home_exercise_checkbox = self._make_checkbox("Home exercise program", True)
        lomn_form.addRow("", self.lomn_home_exercise_checkbox)
        self.lomn_esi_checkbox = self._make_checkbox("Epidural Steroid injections (if applicable)", True)
        lomn_form.addRow("", self.lomn_esi_checkbox)

        self.lomn_medical_necessity_input = QTextEdit()
        self.lomn_medical_necessity_input.setMinimumHeight(80)
        self.lomn_medical_necessity_input.setPlaceholderText("Explain why the requested treatment is medically necessary")
        self.lomn_medical_necessity_input.setPlainText(
            "Based on chronic symptoms, failed conservative management, functional limitation, and imaging findings that correlate with the clinical presentation, the requested intervention is medically reasonable and necessary."
        )
        self.lomn_medical_necessity_input.textChanged.connect(self.refresh_preview)
        lomn_form.addRow("Medical Necessity", self.lomn_medical_necessity_input)

        self.lomn_reduce_pain_checkbox = self._make_checkbox("Reduce pain severity", True)
        lomn_form.addRow("Indications", self.lomn_reduce_pain_checkbox)
        self.lomn_improve_function_checkbox = self._make_checkbox("Improve functional capacity", True)
        lomn_form.addRow("", self.lomn_improve_function_checkbox)
        self.lomn_prevent_degeneration_checkbox = self._make_checkbox("Prevent further disc degeneration", True)
        lomn_form.addRow("", self.lomn_prevent_degeneration_checkbox)
        self.lomn_reduce_opioids_checkbox = self._make_checkbox("Decrease reliance on long-term opioid therapy", True)
        lomn_form.addRow("", self.lomn_reduce_opioids_checkbox)
        self.lomn_prevent_surgery_checkbox = self._make_checkbox("Potentially prevent need for more invasive surgical intervention", True)
        lomn_form.addRow("", self.lomn_prevent_surgery_checkbox)

        self.lomn_risk_statement_input = QTextEdit()
        self.lomn_risk_statement_input.setMinimumHeight(70)
        self.lomn_risk_statement_input.setPlaceholderText("Describe the risks of non-treatment")
        self.lomn_risk_statement_input.setPlainText(
            "Without appropriate treatment, the patient remains at risk for persistent pain, worsening functional limitation, and continued impairment of daily activities."
        )
        self.lomn_risk_statement_input.textChanged.connect(self.refresh_preview)
        lomn_form.addRow("Risk Statement", self.lomn_risk_statement_input)

        self.lomn_reasonable_statement_input = QTextEdit()
        self.lomn_reasonable_statement_input.setMinimumHeight(70)
        self.lomn_reasonable_statement_input.setPlaceholderText("State why the requested care is reasonable and necessary")
        self.lomn_reasonable_statement_input.setPlainText(
            "The requested care is consistent with the documented diagnosis, the failure of conservative treatment, and the current clinical findings."
        )
        self.lomn_reasonable_statement_input.textChanged.connect(self.refresh_preview)
        lomn_form.addRow("Reasonable / Necessary", self.lomn_reasonable_statement_input)

        self.lomn_contact_statement_input = QTextEdit()
        self.lomn_contact_statement_input.setMinimumHeight(60)
        self.lomn_contact_statement_input.setPlaceholderText("Optional contact statement")
        self.lomn_contact_statement_input.setPlainText(
            "Additional supporting documentation can be provided upon request."
        )
        self.lomn_contact_statement_input.textChanged.connect(self.refresh_preview)
        lomn_form.addRow("Contact Statement", self.lomn_contact_statement_input)

        workspace_layout.addWidget(self.lomn_group)

        self.consult_request_group = QGroupBox("Consultation & Treatment Request Template")
        self.consult_request_group.setStyleSheet("QGroupBox { font-weight:700; color:#E5E7EB; }")
        consult_form = QFormLayout(self.consult_request_group)
        self._configure_form_layout(consult_form)

        self.consult_request_date_input = self._make_line_edit("MM/DD/YYYY")
        consult_form.addRow("Request Date", self.consult_request_date_input)
        self.consult_va_claim_number_input = self._make_line_edit("VA claim number if applicable")
        consult_form.addRow("VA Claim Number", self.consult_va_claim_number_input)
        self.consult_referring_va_provider_input = self._make_line_edit("Referring VA provider")
        consult_form.addRow("Referring VA Provider", self.consult_referring_va_provider_input)

        self.consult_reason_input = QTextEdit()
        self.consult_reason_input.setMinimumHeight(70)
        self.consult_reason_input.setPlaceholderText("Summarize the reason for consultation and requested treatment")
        self.consult_reason_input.setPlainText(
            "Authorization is requested for specialty consultation and treatment planning related to documented lumbar disc pathology and persistent pain despite conservative management."
        )
        self.consult_reason_input.textChanged.connect(self.refresh_preview)
        consult_form.addRow("Reason for Consultation", self.consult_reason_input)

        self.consult_primary_diagnosis_input = self._make_line_edit("Primary diagnosis")
        consult_form.addRow("Primary Diagnosis", self.consult_primary_diagnosis_input)
        self.consult_secondary_diagnosis_input = self._make_line_edit("Secondary diagnosis")
        consult_form.addRow("Secondary Diagnosis", self.consult_secondary_diagnosis_input)

        self.consult_symptom_axial_pain_checkbox = self._make_checkbox("Chronic axial lumbar pain", True)
        consult_form.addRow("Symptoms", self.consult_symptom_axial_pain_checkbox)
        self.consult_symptom_activity_checkbox = self._make_checkbox("Activity-related exacerbation (sitting, standing, bending)", True)
        consult_form.addRow("", self.consult_symptom_activity_checkbox)
        self.consult_symptom_tolerance_checkbox = self._make_checkbox("Reduced tolerance for prolonged positioning", True)
        consult_form.addRow("", self.consult_symptom_tolerance_checkbox)
        self.consult_symptom_function_checkbox = self._make_checkbox("Functional impairment affecting occupational and daily activities", True)
        consult_form.addRow("", self.consult_symptom_function_checkbox)

        self.consult_mri_date_input = self._make_line_edit("MM/DD/YYYY")
        consult_form.addRow("MRI Date", self.consult_mri_date_input)
        self.consult_mri_findings_input = self._make_line_edit("MRI findings")
        consult_form.addRow("MRI Findings", self.consult_mri_findings_input)
        self.consult_conservative_duration_input = self._make_line_edit("Duration of prior treatment")
        consult_form.addRow("Conservative Care Duration", self.consult_conservative_duration_input)

        self.consult_physical_therapy_checkbox = self._make_checkbox("Physical therapy", True)
        consult_form.addRow("Conservative Care", self.consult_physical_therapy_checkbox)
        self.consult_nsaids_checkbox = self._make_checkbox("NSAIDs and non-opioid analgesics", True)
        consult_form.addRow("", self.consult_nsaids_checkbox)
        self.consult_activity_mod_checkbox = self._make_checkbox("Activity modification", True)
        consult_form.addRow("", self.consult_activity_mod_checkbox)
        self.consult_home_exercise_checkbox = self._make_checkbox("Home exercise program", True)
        consult_form.addRow("", self.consult_home_exercise_checkbox)
        self.consult_interventional_history_checkbox = self._make_checkbox("Epidural steroid injections and/or other interventional procedures (if applicable)", True)
        consult_form.addRow("", self.consult_interventional_history_checkbox)

        self.consult_include_pm_consult_checkbox = self._make_checkbox("Comprehensive pain management consultation", True)
        consult_form.addRow("Requested Services", self.consult_include_pm_consult_checkbox)
        self.consult_include_planning_checkbox = self._make_checkbox("Diagnostic confirmation and procedural planning", True)
        consult_form.addRow("", self.consult_include_planning_checkbox)
        self.consult_include_annulargram_checkbox = self._make_checkbox("Diagnostic Annulargram", True)
        consult_form.addRow("", self.consult_include_annulargram_checkbox)
        self.consult_include_fibrin_checkbox = self._make_checkbox("Intraannular Fibrin injection if indicated", True)
        consult_form.addRow("", self.consult_include_fibrin_checkbox)
        self.consult_fibrin_levels_input = self._make_line_edit("Specific lumbar level(s)")
        consult_form.addRow("Fibrin Levels", self.consult_fibrin_levels_input)
        self.consult_include_follow_up_checkbox = self._make_checkbox("Routine post-procedure follow-up visit(s)", True)
        consult_form.addRow("", self.consult_include_follow_up_checkbox)

        self.consult_scope_exclusion_input = QTextEdit()
        self.consult_scope_exclusion_input.setMinimumHeight(70)
        self.consult_scope_exclusion_input.setPlaceholderText("Optional scope or exclusion note")
        self.consult_scope_exclusion_input.setPlainText(
            "This request is limited to evaluation, procedural planning, the indicated intervention if clinically appropriate, and standard post-procedure follow-up related only to the documented lumbar condition."
        )
        self.consult_scope_exclusion_input.textChanged.connect(self.refresh_preview)
        consult_form.addRow("Scope of Request", self.consult_scope_exclusion_input)

        self.consult_medical_rationale_input = QTextEdit()
        self.consult_medical_rationale_input.setMinimumHeight(80)
        self.consult_medical_rationale_input.setPlaceholderText("Document the medical rationale")
        self.consult_medical_rationale_input.setPlainText(
            "The requested evaluation and treatment pathway is supported by the patient’s ongoing symptoms, functional impairment, failure of conservative care, and imaging findings that correlate with the clinical presentation."
        )
        self.consult_medical_rationale_input.textChanged.connect(self.refresh_preview)
        consult_form.addRow("Medical Rationale", self.consult_medical_rationale_input)

        self.consult_goal_pain_checkbox = self._make_checkbox("Pain reduction", True)
        consult_form.addRow("Clinical Goals", self.consult_goal_pain_checkbox)
        self.consult_goal_function_checkbox = self._make_checkbox("Functional improvement", True)
        consult_form.addRow("", self.consult_goal_function_checkbox)
        self.consult_goal_analgesics_checkbox = self._make_checkbox("Decreased reliance on chronic analgesic therapy", True)
        consult_form.addRow("", self.consult_goal_analgesics_checkbox)
        self.consult_goal_surgery_checkbox = self._make_checkbox("Prevention of progression requiring more invasive surgical intervention", True)
        consult_form.addRow("", self.consult_goal_surgery_checkbox)

        self.consult_risk_without_treatment_input = QTextEdit()
        self.consult_risk_without_treatment_input.setMinimumHeight(70)
        self.consult_risk_without_treatment_input.setPlaceholderText("Optional risk of non-treatment statement")
        self.consult_risk_without_treatment_input.setPlainText(
            "Without appropriate specialty evaluation and treatment, the patient may continue to experience persistent pain, reduced function, and progression of disability."
        )
        self.consult_risk_without_treatment_input.textChanged.connect(self.refresh_preview)
        consult_form.addRow("Risk Without Treatment", self.consult_risk_without_treatment_input)

        self.consult_duration_scope_input = QTextEdit()
        self.consult_duration_scope_input.setMinimumHeight(80)
        self.consult_duration_scope_input.setPlaceholderText("Describe the scope and expected duration of care")
        self.consult_duration_scope_input.setPlainText(
            "This consultation and treatment request is limited to a defined evaluation and treatment pathway rather than open-ended pain management."
        )
        self.consult_duration_scope_input.textChanged.connect(self.refresh_preview)
        consult_form.addRow("Duration and Scope", self.consult_duration_scope_input)

        self.consult_contact_statement_input = QTextEdit()
        self.consult_contact_statement_input.setMinimumHeight(60)
        self.consult_contact_statement_input.setPlaceholderText("Optional contact statement")
        self.consult_contact_statement_input.setPlainText(
            "Please contact the treating office if additional records or clarification are needed for review."
        )
        self.consult_contact_statement_input.textChanged.connect(self.refresh_preview)
        consult_form.addRow("Contact Statement", self.consult_contact_statement_input)

        self.provider_address_input = self._make_line_edit("Community Care practice address")
        consult_form.addRow("Community Care Practice Address", self.provider_address_input)
        self.provider_email_input = self._make_line_edit("Community Care secure email")
        consult_form.addRow("Community Care Secure Email", self.provider_email_input)

        workspace_layout.addWidget(self.consult_request_group)

        self.clinical_doc_group = QGroupBox("Clinical Notes Template")
        self.clinical_doc_group.setStyleSheet("QGroupBox { font-weight:700; color:#E5E7EB; }")
        clinical_doc_form = QFormLayout(self.clinical_doc_group)
        self._configure_form_layout(clinical_doc_form)

        self.clinical_doc_title_input = self._make_line_edit("Clinical Documentation Template")
        self.clinical_doc_title_input.setText("Clinical Documentation Template")
        clinical_doc_form.addRow("Template Title", self.clinical_doc_title_input)

        self.clinical_doc_chief_complaint_input = self._make_line_edit("Chief complaint")
        clinical_doc_form.addRow("Chief Complaint", self.clinical_doc_chief_complaint_input)

        self.clinical_doc_duration_3m_checkbox = self._make_checkbox("> 3 months", True)
        clinical_doc_form.addRow("Duration of Symptoms", self.clinical_doc_duration_3m_checkbox)
        self.clinical_doc_duration_6m_checkbox = self._make_checkbox("> 6 months", False)
        clinical_doc_form.addRow("", self.clinical_doc_duration_6m_checkbox)
        self.clinical_doc_duration_12m_checkbox = self._make_checkbox("> 12 months", False)
        clinical_doc_form.addRow("", self.clinical_doc_duration_12m_checkbox)
        self.clinical_doc_exact_duration_input = self._make_line_edit("Exact duration")
        clinical_doc_form.addRow("Exact Duration", self.clinical_doc_exact_duration_input)

        self.clinical_doc_pain_axial_checkbox = self._make_checkbox("Axial lumbar pain", True)
        clinical_doc_form.addRow("Pain Characteristics", self.clinical_doc_pain_axial_checkbox)
        self.clinical_doc_pain_discogenic_checkbox = self._make_checkbox("Discogenic pattern", True)
        clinical_doc_form.addRow("", self.clinical_doc_pain_discogenic_checkbox)
        self.clinical_doc_pain_activity_checkbox = self._make_checkbox("Activity-related exacerbation", True)
        clinical_doc_form.addRow("", self.clinical_doc_pain_activity_checkbox)
        self.clinical_doc_pain_sitting_checkbox = self._make_checkbox("Sitting intolerance", True)
        clinical_doc_form.addRow("", self.clinical_doc_pain_sitting_checkbox)
        self.clinical_doc_pain_standing_checkbox = self._make_checkbox("Standing intolerance", True)
        clinical_doc_form.addRow("", self.clinical_doc_pain_standing_checkbox)
        self.clinical_doc_pain_bending_checkbox = self._make_checkbox("Bending/lifting provocation", True)
        clinical_doc_form.addRow("", self.clinical_doc_pain_bending_checkbox)
        self.clinical_doc_pain_severity_input = self._make_line_edit("0-10")
        clinical_doc_form.addRow("Pain Severity", self.clinical_doc_pain_severity_input)

        self.clinical_doc_limit_occupational_checkbox = self._make_checkbox("Occupational duties", True)
        clinical_doc_form.addRow("Functional Limitation", self.clinical_doc_limit_occupational_checkbox)
        self.clinical_doc_limit_sitting_checkbox = self._make_checkbox("Prolonged sitting", True)
        clinical_doc_form.addRow("", self.clinical_doc_limit_sitting_checkbox)
        self.clinical_doc_limit_standing_checkbox = self._make_checkbox("Prolonged standing", True)
        clinical_doc_form.addRow("", self.clinical_doc_limit_standing_checkbox)
        self.clinical_doc_limit_ambulation_checkbox = self._make_checkbox("Ambulation tolerance", False)
        clinical_doc_form.addRow("", self.clinical_doc_limit_ambulation_checkbox)
        self.clinical_doc_limit_household_checkbox = self._make_checkbox("Household activities", True)
        clinical_doc_form.addRow("", self.clinical_doc_limit_household_checkbox)
        self.clinical_doc_limit_sleep_checkbox = self._make_checkbox("Sleep disturbance", True)
        clinical_doc_form.addRow("", self.clinical_doc_limit_sleep_checkbox)
        self.clinical_doc_functional_impact_input = QTextEdit()
        self.clinical_doc_functional_impact_input.setMinimumHeight(80)
        self.clinical_doc_functional_impact_input.textChanged.connect(self.refresh_preview)
        clinical_doc_form.addRow("Specific Functional Impact", self.clinical_doc_functional_impact_input)

        self.clinical_doc_conservative_pt_checkbox = self._make_checkbox("Physical therapy", True)
        clinical_doc_form.addRow("Conservative Therapy", self.clinical_doc_conservative_pt_checkbox)
        self.clinical_doc_conservative_home_exercise_checkbox = self._make_checkbox("Home exercise program", True)
        clinical_doc_form.addRow("", self.clinical_doc_conservative_home_exercise_checkbox)
        self.clinical_doc_conservative_nsaids_checkbox = self._make_checkbox("NSAIDs", True)
        clinical_doc_form.addRow("", self.clinical_doc_conservative_nsaids_checkbox)
        self.clinical_doc_conservative_non_opioid_checkbox = self._make_checkbox("Non-opioid analgesics", True)
        clinical_doc_form.addRow("", self.clinical_doc_conservative_non_opioid_checkbox)
        self.clinical_doc_conservative_activity_checkbox = self._make_checkbox("Activity modification", True)
        clinical_doc_form.addRow("", self.clinical_doc_conservative_activity_checkbox)
        self.clinical_doc_conservative_esi_checkbox = self._make_checkbox("Epidural steroid injections", True)
        clinical_doc_form.addRow("", self.clinical_doc_conservative_esi_checkbox)
        self.clinical_doc_conservative_other_checkbox = self._make_checkbox("Other interventional procedures", False)
        clinical_doc_form.addRow("", self.clinical_doc_conservative_other_checkbox)
        self.clinical_doc_conservative_duration_input = self._make_line_edit("Duration of conservative treatment")
        clinical_doc_form.addRow("Conservative Care Duration", self.clinical_doc_conservative_duration_input)
        self.clinical_doc_esi_response_combo = self._make_choice_combo(["Temporary relief", "Partial relief", "No relief"])
        clinical_doc_form.addRow("Response to Prior ESI", self.clinical_doc_esi_response_combo)

        self.clinical_doc_mri_date_input = self._make_line_edit("MM/DD/YYYY")
        clinical_doc_form.addRow("MRI Date", self.clinical_doc_mri_date_input)
        self.clinical_doc_imaging_annular_checkbox = self._make_checkbox("Annular tear", True)
        clinical_doc_form.addRow("Imaging Findings", self.clinical_doc_imaging_annular_checkbox)
        self.clinical_doc_imaging_degeneration_checkbox = self._make_checkbox("Disc degeneration", True)
        clinical_doc_form.addRow("", self.clinical_doc_imaging_degeneration_checkbox)
        self.clinical_doc_imaging_protrusion_checkbox = self._make_checkbox("Disc protrusion", False)
        clinical_doc_form.addRow("", self.clinical_doc_imaging_protrusion_checkbox)
        self.clinical_doc_imaging_displacement_checkbox = self._make_checkbox("Disc displacement", False)
        clinical_doc_form.addRow("", self.clinical_doc_imaging_displacement_checkbox)
        self.clinical_doc_affected_levels_input = self._make_line_edit("Affected levels")
        clinical_doc_form.addRow("Affected Levels", self.clinical_doc_affected_levels_input)

        self.clinical_doc_primary_diagnosis_input = self._make_line_edit("Primary diagnosis (ICD-10)")
        clinical_doc_form.addRow("Primary Diagnosis", self.clinical_doc_primary_diagnosis_input)
        self.clinical_doc_secondary_diagnosis_input = self._make_line_edit("Secondary diagnosis")
        clinical_doc_form.addRow("Secondary Diagnosis", self.clinical_doc_secondary_diagnosis_input)
        self.clinical_doc_assessment_summary_input = QTextEdit()
        self.clinical_doc_assessment_summary_input.setMinimumHeight(70)
        self.clinical_doc_assessment_summary_input.setPlaceholderText("Summarize the assessment")
        self.clinical_doc_assessment_summary_input.setPlainText("")
        self.clinical_doc_assessment_summary_input.textChanged.connect(self.refresh_preview)
        clinical_doc_form.addRow("Assessment", self.clinical_doc_assessment_summary_input)

        self.clinical_doc_treatment_plan_intro_input = QTextEdit()
        self.clinical_doc_treatment_plan_intro_input.setMinimumHeight(70)
        self.clinical_doc_treatment_plan_intro_input.setPlaceholderText("Summarize the treatment plan")
        self.clinical_doc_treatment_plan_intro_input.setPlainText("")
        self.clinical_doc_treatment_plan_intro_input.textChanged.connect(self.refresh_preview)
        clinical_doc_form.addRow("Treatment Plan", self.clinical_doc_treatment_plan_intro_input)
        self.clinical_doc_plan_diagnostic_checkbox = self._make_checkbox("Diagnostic confirmation if necessary", True)
        clinical_doc_form.addRow("Plan Includes", self.clinical_doc_plan_diagnostic_checkbox)
        self.clinical_doc_plan_intervention_checkbox = self._make_checkbox("Intradiscal annular intervention if indicated", True)
        clinical_doc_form.addRow("", self.clinical_doc_plan_intervention_checkbox)
        self.clinical_doc_plan_follow_up_checkbox = self._make_checkbox("Standard post-procedure follow-up", True)
        clinical_doc_form.addRow("", self.clinical_doc_plan_follow_up_checkbox)
        self.clinical_doc_plan_exclusion_input = QTextEdit()
        self.clinical_doc_plan_exclusion_input.setMinimumHeight(70)
        self.clinical_doc_plan_exclusion_input.setPlaceholderText("Optional plan limitation or exclusion note")
        self.clinical_doc_plan_exclusion_input.setPlainText(
            "This plan does not request open-ended medication management or unrelated pain treatment outside the documented lumbar condition."
        )
        self.clinical_doc_plan_exclusion_input.textChanged.connect(self.refresh_preview)
        clinical_doc_form.addRow("Plan Exclusion", self.clinical_doc_plan_exclusion_input)

        self.clinical_doc_physician_narrative_input = QTextEdit()
        self.clinical_doc_physician_narrative_input.setMinimumHeight(110)
        self.clinical_doc_physician_narrative_input.setPlaceholderText("Document the physician narrative for this patient")
        self.clinical_doc_physician_narrative_input.setPlainText("")
        self.clinical_doc_physician_narrative_input.textChanged.connect(self.refresh_preview)
        clinical_doc_form.addRow("Physician Narrative", self.clinical_doc_physician_narrative_input)

        workspace_layout.addWidget(self.clinical_doc_group)

        self.va10172_group = QGroupBox("VA Form 10-10172")
        self.va10172_group.setStyleSheet("QGroupBox { font-weight:700; color:#E5E7EB; }")
        va10172_form = QFormLayout(self.va10172_group)
        self._configure_form_layout(va10172_form)

        template_row = QWidget()
        template_layout = QHBoxLayout(template_row)
        template_layout.setContentsMargins(0, 0, 0, 0)
        template_layout.setSpacing(8)
        self.va10172_template_value = QLabel(self.config.get("va_form_10172_template_path") or "No template selected")
        self.va10172_template_value.setWordWrap(True)
        self.va10172_template_value.setStyleSheet("color:#8FA6C1;")
        choose_va10172_template_button = QPushButton("Choose PDF Template")
        choose_va10172_template_button.clicked.connect(self.choose_va10172_template)
        template_layout.addWidget(self.va10172_template_value, stretch=1)
        template_layout.addWidget(choose_va10172_template_button)
        va10172_form.addRow("Template PDF", template_row)

        self.va10172_facility_address_input = QTextEdit()
        self.va10172_facility_address_input.setMinimumHeight(70)
        self.va10172_facility_address_input.textChanged.connect(self.refresh_preview)
        va10172_form.addRow("VA Facility & Address", self.va10172_facility_address_input)

        self.va10172_provider_office_address_input = QTextEdit()
        self.va10172_provider_office_address_input.setMinimumHeight(80)
        self.va10172_provider_office_address_input.textChanged.connect(self.refresh_preview)
        va10172_form.addRow("Ordering Provider Office Name & Address", self.va10172_provider_office_address_input)

        self.va10172_is_ihs_combo = self._make_choice_combo(["No", "Yes"])
        va10172_form.addRow("Indian Health Services / THP Provider", self.va10172_is_ihs_combo)
        self.va10172_provider_phone_input = self._make_line_edit("Provider phone")
        va10172_form.addRow("Ordering Provider Phone", self.va10172_provider_phone_input)
        self.va10172_provider_fax_input = self._make_line_edit("Provider fax")
        va10172_form.addRow("Ordering Provider Fax", self.va10172_provider_fax_input)
        self.va10172_provider_email_input = self._make_line_edit("Secure email")
        va10172_form.addRow("Ordering Provider Secure Email", self.va10172_provider_email_input)

        self.va10172_48h_combo = self._make_choice_combo(["No", "Yes"])
        va10172_form.addRow("Care Needed Within 48 Hours", self.va10172_48h_combo)
        self.va10172_continuation_combo = self._make_choice_combo(["No", "Yes"])
        va10172_form.addRow("Continuation of Care", self.va10172_continuation_combo)
        self.va10172_referral_specialty_combo = self._make_choice_combo(["No", "Yes"])
        va10172_form.addRow("Referral to Another Specialty", self.va10172_referral_specialty_combo)
        self.va10172_referral_specialty_text_input = self._make_line_edit("Specialty name")
        va10172_form.addRow("Specialty", self.va10172_referral_specialty_text_input)

        self.va10172_diagnosis_description_input = QTextEdit()
        self.va10172_diagnosis_description_input.setMinimumHeight(70)
        self.va10172_diagnosis_description_input.textChanged.connect(self.refresh_preview)
        va10172_form.addRow("Diagnosis Description", self.va10172_diagnosis_description_input)
        self.va10172_requested_cpt_input = self._make_line_edit("Requested CPT/HCPCS code")
        va10172_form.addRow("Requested CPT/HCPCS Code", self.va10172_requested_cpt_input)
        self.va10172_requested_cpt_description_input = self._make_line_edit("CPT/HCPCS description")
        va10172_form.addRow("Description CPT/HCPCS Code", self.va10172_requested_cpt_description_input)

        self.va10172_geriatric_option_combo = self._make_choice_combo(
            [
                "None",
                "Community Nursing Home",
                "Home Infusion",
                "Hospice/Palliative Care",
                "Skilled Home Health Care",
                "Community Adult Day Health Care",
                "Home Homemaker/Home Health Aide",
                "Respite",
            ]
        )
        va10172_form.addRow("Geriatric / Extended Care", self.va10172_geriatric_option_combo)

        self.va10172_reason_for_request_input = QTextEdit()
        self.va10172_reason_for_request_input.setMinimumHeight(130)
        self.va10172_reason_for_request_input.textChanged.connect(self.refresh_preview)
        va10172_form.addRow("Reason for Request", self.va10172_reason_for_request_input)

        self.va10172_provider_name_printed_input = self._make_line_edit("CCN ordering provider printed name")
        va10172_form.addRow("CCN Ordering Provider Name (Printed)", self.va10172_provider_name_printed_input)
        self.va10172_provider_npi_input = self._make_line_edit("CCN ordering provider NPI")
        va10172_form.addRow("CCN Ordering Provider NPI", self.va10172_provider_npi_input)
        va_signature_widget, self.va10172_signature_text_input = self._make_signature_field("va10172_signature_text", "CCN ordering provider signature")
        va10172_form.addRow("CCN Ordering Provider Signature", va_signature_widget)
        self.va10172_today_date_input = self._make_line_edit("MM/DD/YYYY")
        va10172_form.addRow("Today's Date", self.va10172_today_date_input)

        workspace_layout.addWidget(self.va10172_group)

        action_row = QGridLayout()
        action_row.setHorizontalSpacing(10)
        action_row.setVerticalSpacing(10)
        self.new_packet_button = QPushButton("New Packet")
        self._style_builder_button(self.new_packet_button)
        self.new_packet_button.clicked.connect(self.start_new_packet)
        self.save_to_library_button = QPushButton("Save To Library")
        self._style_builder_button(self.save_to_library_button)
        self.save_to_library_button.clicked.connect(self.save_draft_json)
        self.export_word_button = QPushButton("Word")
        self._style_builder_button(self.export_word_button)
        self.export_word_button.clicked.connect(self.export_word)
        self.export_pdf_button = QPushButton("PDF")
        self._style_builder_button(self.export_pdf_button)
        self.export_pdf_button.clicked.connect(self.export_pdf)
        self.export_both_button = QPushButton("Word + PDF")
        self._style_builder_button(self.export_both_button)
        self.export_both_button.clicked.connect(self.export_both)
        self.open_folder_button = QPushButton("Open Folder")
        self._style_builder_button(self.open_folder_button)
        self.open_folder_button.clicked.connect(self.open_export_folder)
        action_row.addWidget(self.new_packet_button, 0, 0)
        action_row.addWidget(self.save_to_library_button, 0, 1)
        action_row.addWidget(self.export_word_button, 0, 2)
        action_row.addWidget(self.export_pdf_button, 1, 0)
        action_row.addWidget(self.export_both_button, 1, 1)
        action_row.addWidget(self.open_folder_button, 1, 2)
        workspace_layout.addLayout(action_row)

        bundle_row = QGridLayout()
        bundle_row.setHorizontalSpacing(10)
        bundle_row.setVerticalSpacing(10)
        bundle_label = QLabel("Bundle Export")
        bundle_label.setStyleSheet("font-weight:700; color:#C8D8E8;")
        self.bundle_format_combo = QComboBox()
        self.bundle_format_combo.addItems(["Word", "PDF", "Both"])
        self.bundle_format_combo.setCurrentText("Both")
        self.export_referral_button = QPushButton("Referral Request")
        self._style_builder_button(self.export_referral_button)
        self.export_referral_button.clicked.connect(self.export_referral_request_bundle)
        self.export_packet_button = QPushButton("Patient Packet")
        self._style_builder_button(self.export_packet_button)
        self.export_packet_button.clicked.connect(self.export_patient_packet_bundle)
        bundle_row.addWidget(bundle_label, 0, 0)
        bundle_row.addWidget(self.bundle_format_combo, 0, 1)
        bundle_row.addWidget(self.export_referral_button, 0, 2)
        bundle_row.addWidget(self.export_packet_button, 0, 3)
        bundle_row.setColumnStretch(4, 1)
        workspace_layout.addLayout(bundle_row)
        workspace_layout.addStretch(1)
        self._export_controls = [
            self.export_word_button,
            self.export_pdf_button,
            self.export_both_button,
            self.export_referral_button,
            self.export_packet_button,
            self.bundle_format_combo,
            choose_folder_button,
            self.open_folder_button,
        ]

        preview_panel = QWidget()
        preview_layout = QVBoxLayout(preview_panel)
        preview_layout.setContentsMargins(8, 0, 0, 0)
        preview_layout.setSpacing(6)

        self.preview_tabs = QTabWidget()
        self.preview_tabs.setDocumentMode(True)

        preview_tab = QWidget()
        preview_tab_layout = QVBoxLayout(preview_tab)
        preview_tab_layout.setContentsMargins(0, 0, 0, 0)
        preview_tab_layout.setSpacing(8)
        self.preview_status_label = QLabel()
        self.preview_status_label.setWordWrap(True)
        self.preview_status_label.setStyleSheet("color:#8FA6C1; font-size:12px;")
        preview_tab_layout.addWidget(self.preview_status_label)
        preview_action_row = QHBoxLayout()
        preview_action_row.setContentsMargins(0, 0, 0, 0)
        preview_action_row.setSpacing(8)
        preview_action_row.addStretch(1)
        self.exact_preview_button = QPushButton("Refresh Exact Preview")
        self._style_builder_button(self.exact_preview_button)
        self.exact_preview_button.setMaximumWidth(180)
        self.exact_preview_button.clicked.connect(self.refresh_exact_preview)
        preview_action_row.addWidget(self.exact_preview_button, 0, Qt.AlignRight)
        preview_tab_layout.addLayout(preview_action_row)
        self.preview_stack = QStackedWidget()
        self.preview_view = QTextEdit()
        self.preview_view.setReadOnly(True)
        self.preview_view.setMinimumWidth(300)
        self.preview_view.document().setDocumentMargin(0)
        self.preview_view.setStyleSheet(
            "background:#FFFFFF; border:1px solid #D6DCE5; border-radius:10px; color:#111827; padding:0px;"
        )
        self.preview_pdf_document = QPdfDocument(self)
        self.preview_pdf_view = QPdfView()
        self.preview_pdf_view.setDocument(self.preview_pdf_document)
        self.preview_pdf_view.setPageMode(QPdfView.PageMode.MultiPage)
        self.preview_pdf_view.setZoomMode(QPdfView.ZoomMode.FitToWidth)
        self.preview_pdf_view.setStyleSheet(
            "background:#09121B; border:1px solid #243446; border-radius:10px;"
        )
        self.preview_stack.addWidget(self.preview_view)
        self.preview_stack.addWidget(self.preview_pdf_view)
        preview_tab_layout.addWidget(self.preview_stack, stretch=1)

        lab_tab = QWidget()
        lab_tab_layout = QVBoxLayout(lab_tab)
        lab_tab_layout.setContentsMargins(0, 0, 0, 0)
        lab_tab_layout.setSpacing(8)
        self.lab_view = QTextEdit()
        self.lab_view.setReadOnly(True)
        self.lab_view.setMinimumWidth(300)
        self.lab_view.setStyleSheet(
            "background:#09121B; border:1px solid #243446; border-radius:10px; color:#E8F1FC; padding:0px;"
        )
        lab_tab_layout.addWidget(self.lab_view, stretch=1)

        wording_tab = QWidget()
        wording_tab_layout = QVBoxLayout(wording_tab)
        wording_tab_layout.setContentsMargins(0, 0, 0, 0)
        wording_tab_layout.setSpacing(8)
        wording_help = QLabel(
            "Wording Assist turns rough facts into provider-style language for denial-sensitive fields. Review the original wording, cycle the phrasing when you want a different version, then accept, edit, or intentionally keep the original before final export."
        )
        wording_help.setWordWrap(True)
        wording_help.setStyleSheet("color:#8FA6C1;")
        wording_tab_layout.addWidget(wording_help)
        wording_splitter = QSplitter(Qt.Horizontal)
        wording_splitter.setChildrenCollapsible(False)
        wording_splitter.setHandleWidth(10)
        self.wording_splitter = wording_splitter
        self.wording_assist_list = QListWidget()
        self.wording_assist_list.setMinimumWidth(0)
        self.wording_assist_list.setMaximumWidth(230)
        self.wording_assist_list.setSizePolicy(QSizePolicy.Fixed, QSizePolicy.Expanding)
        self.wording_assist_list.setStyleSheet(
            "QListWidget { background:#09121B; border:1px solid #243446; border-radius:10px; padding:8px; }"
            "QListWidget::item { color:#E8F1FC; padding:8px 10px; margin:2px 0; border-radius:8px; }"
            "QListWidget::item:selected { background:#1D2A3A; }"
        )
        self.wording_assist_list.currentItemChanged.connect(self._on_wording_entry_changed)
        wording_splitter.addWidget(self.wording_assist_list)

        wording_detail = QWidget()
        wording_detail.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
        wording_detail.setMinimumWidth(0)
        wording_detail_layout = QVBoxLayout(wording_detail)
        wording_detail_layout.setContentsMargins(0, 0, 18, 0)
        wording_detail_layout.setSpacing(8)
        self.wording_field_title_label = QLabel("No wording-assist fields on this form.")
        self.wording_field_title_label.setWordWrap(True)
        self.wording_field_title_label.setStyleSheet("color:#EAF2FF; font-size:15px; font-weight:700;")
        wording_detail_layout.addWidget(self.wording_field_title_label)
        self.wording_status_label = QLabel("")
        self.wording_status_label.setWordWrap(True)
        wording_detail_layout.addWidget(self.wording_status_label)
        self.wording_followup_label = QLabel("")
        self.wording_followup_label.setWordWrap(True)
        self.wording_followup_label.setStyleSheet("color:#FCD7A4;")
        wording_detail_layout.addWidget(self.wording_followup_label)
        self.wording_risk_label = QLabel("")
        self.wording_risk_label.setWordWrap(True)
        self.wording_risk_label.setStyleSheet("color:#F4C7C3;")
        wording_detail_layout.addWidget(self.wording_risk_label)
        original_label = QLabel("Original Wording")
        original_label.setStyleSheet("color:#C8D8E8; font-weight:700;")
        wording_detail_layout.addWidget(original_label)
        self.wording_original_view = AutoSizingPlainTextEdit(150)
        self.wording_original_view.setReadOnly(True)
        self.wording_original_view.document().setDocumentMargin(10)
        self.wording_original_view.setLineWrapMode(QPlainTextEdit.WidgetWidth)
        self.wording_original_view.setWordWrapMode(QTextOption.WrapAtWordBoundaryOrAnywhere)
        self.wording_original_view.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.wording_original_view.setCenterOnScroll(False)
        self.wording_original_view.setStyleSheet(
            "background:#09121B; border:1px solid #243446; border-radius:10px; color:#E8F1FC; padding:12px 10px 18px 10px; font-size:13px;"
        )
        wording_detail_layout.addWidget(self.wording_original_view)
        suggestion_label = QLabel("Suggested Professional Wording")
        suggestion_label.setStyleSheet("color:#C8D8E8; font-weight:700;")
        wording_detail_layout.addWidget(suggestion_label)
        self.wording_suggestion_input = AutoSizingPlainTextEdit(220)
        self.wording_suggestion_input.document().setDocumentMargin(10)
        self.wording_suggestion_input.setLineWrapMode(QPlainTextEdit.WidgetWidth)
        self.wording_suggestion_input.setWordWrapMode(QTextOption.WrapAtWordBoundaryOrAnywhere)
        self.wording_suggestion_input.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.wording_suggestion_input.setCenterOnScroll(False)
        self.wording_suggestion_input.setStyleSheet(
            "background:#0F1823; border:1px solid #243446; border-radius:10px; color:#EAF2FF; padding:12px 10px 18px 10px; font-size:13px;"
        )
        wording_detail_layout.addWidget(self.wording_suggestion_input, stretch=1)
        wording_button_row = QGridLayout()
        wording_button_row.setContentsMargins(0, 0, 0, 0)
        wording_button_row.setHorizontalSpacing(8)
        wording_button_row.setVerticalSpacing(8)
        self.wording_accept_button = QPushButton("Accept Suggestion")
        self._style_builder_button(self.wording_accept_button)
        self.wording_accept_button.clicked.connect(self.accept_wording_suggestion)
        self.wording_accept_button.setMaximumWidth(180)
        wording_button_row.addWidget(self.wording_accept_button, 0, 0)
        self.wording_edit_button = QPushButton("Apply Edited Suggestion")
        self._style_builder_button(self.wording_edit_button)
        self.wording_edit_button.clicked.connect(self.apply_edited_wording_suggestion)
        self.wording_edit_button.setMaximumWidth(180)
        wording_button_row.addWidget(self.wording_edit_button, 0, 1)
        self.wording_keep_button = QPushButton("Keep Original")
        self._style_builder_button(self.wording_keep_button)
        self.wording_keep_button.clicked.connect(self.keep_original_wording)
        self.wording_keep_button.setMaximumWidth(160)
        wording_button_row.addWidget(self.wording_keep_button, 0, 2)
        self.wording_cycle_button = QPushButton("Cycle Phrasing")
        self._style_builder_button(self.wording_cycle_button)
        self.wording_cycle_button.clicked.connect(self.cycle_wording_suggestion)
        self.wording_cycle_button.setMaximumWidth(160)
        wording_button_row.addWidget(self.wording_cycle_button, 1, 0)
        self.wording_reset_button = QPushButton("Reset Review")
        self._style_builder_button(self.wording_reset_button)
        self.wording_reset_button.clicked.connect(self.reset_wording_review)
        self.wording_reset_button.setMaximumWidth(150)
        wording_button_row.addWidget(self.wording_reset_button, 1, 1)
        wording_button_row.setColumnStretch(0, 1)
        wording_button_row.setColumnStretch(1, 1)
        wording_button_row.setColumnStretch(2, 1)
        wording_detail_layout.addLayout(wording_button_row)
        wording_detail_scroll = QScrollArea()
        wording_detail_scroll.setWidgetResizable(True)
        wording_detail_scroll.setFrameShape(QFrame.NoFrame)
        wording_detail_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        wording_detail_scroll.setWidget(wording_detail)
        wording_splitter.addWidget(wording_detail_scroll)
        wording_splitter.setStretchFactor(0, 0)
        wording_splitter.setStretchFactor(1, 1)
        wording_splitter.setSizes([180, 720])
        wording_tab_layout.addWidget(wording_splitter, stretch=1)

        library_tab = QWidget()
        library_tab_layout = QVBoxLayout(library_tab)
        library_tab_layout.setContentsMargins(0, 0, 0, 0)
        library_tab_layout.setSpacing(8)
        library_help = QLabel(
            "Saved packet drafts live here. Click the packet name to reopen it, open its latest PDF, or remove the draft from the library."
        )
        library_help.setWordWrap(True)
        library_help.setStyleSheet("color:#8FA6C1;")
        library_tab_layout.addWidget(library_help)
        self.library_list = QListWidget()
        self.library_list.setStyleSheet(
            "QListWidget { background:#09121B; border:1px solid #243446; border-radius:10px; padding:8px; }"
            "QListWidget::item { border:0px; margin:0px; padding:0px; }"
        )
        library_tab_layout.addWidget(self.library_list, stretch=1)

        self.preview_tabs.addTab(preview_tab, "Export Preview")
        self.preview_tabs.addTab(wording_tab, "Wording Assist")
        self.preview_tabs.addTab(lab_tab, "Lab")
        self.preview_tabs.addTab(library_tab, "Library")
        self.preview_tabs.tabBar().setUsesScrollButtons(True)
        preview_layout.addWidget(self.preview_tabs, stretch=1)
        splitter.addWidget(preview_panel)
        splitter.setStretchFactor(0, 1)
        splitter.setStretchFactor(1, 1)
        splitter.setSizes([760, 760])

        self.packet_title_input.setText("Community Care Referral Request")
        self._builder_ready = True
        self.on_profile_changed(self.profile_combo.currentText())
        self.refresh_library()
        self._update_current_packet_status()
        self.refresh_preview()
        QTimer.singleShot(0, self._rebalance_builder_splitter)

    def _make_line_edit(self, placeholder):
        widget = QLineEdit()
        widget.setPlaceholderText(placeholder)
        widget.textChanged.connect(self.refresh_preview)
        return widget

    def _make_signature_field(self, field_name, placeholder):
        line_edit = self._make_line_edit(placeholder)
        line_edit.setPlaceholderText(f"{placeholder} (double-click to draw)")
        line_edit.mouseDoubleClickEvent = lambda event, key=field_name: self._open_signature_capture(key)
        button = QPushButton("Draw")
        button.setFixedHeight(32)
        button.setMaximumWidth(84)
        button.setStyleSheet(
            "QPushButton {"
            "background-color:#1D2A3A; color:#FFFFFF; border:1px solid #34506B; border-radius:8px; "
            "padding:4px 10px; font-size:12px; font-weight:600; }"
            "QPushButton:hover { background-color:#223246; }"
        )
        button.clicked.connect(lambda _=False, key=field_name: self._open_signature_capture(key))
        container = QWidget()
        layout = QHBoxLayout(container)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(8)
        layout.addWidget(line_edit, stretch=1)
        layout.addWidget(button, 0)
        return container, line_edit

    def _apply_signature_field_state(self, field_name):
        binding = PACKET_WIDGET_BINDINGS.get(field_name)
        if not binding:
            return
        widget = getattr(self, binding[0], None)
        if widget is None:
            return
        has_signature = bool(self._signature_images.get(SIGNATURE_IMAGE_FIELDS.get(field_name, field_name)))
        tooltip = "Double-click to draw a signature." if not has_signature else "Drawn signature captured. Double-click or use Draw to replace it."
        widget.setToolTip(tooltip)
        if has_signature:
            widget.setStyleSheet("QLineEdit { border:1px solid #2C8B57; border-radius:8px; padding:6px; }")
        else:
            widget.setStyleSheet("")

    def _open_signature_capture(self, field_name):
        binding = PACKET_WIDGET_BINDINGS.get(field_name)
        if not binding:
            return
        widget = getattr(self, binding[0], None)
        if widget is None:
            return
        image_field = SIGNATURE_IMAGE_FIELDS.get(field_name, field_name)
        role_label = SIGNATURE_ROLE_LABELS.get(field_name, "signature")
        dialog = SignatureCaptureDialog(
            label_text=f"Draw the {role_label} below. Use Reset to clear it, then Accept to apply it to the current packet.",
            existing_base64=self._signature_images.get(image_field, ""),
            parent=self,
        )
        if dialog.exec() != QDialog.Accepted:
            return
        encoded = dialog.signature_base64()
        self._signature_images[image_field] = encoded
        if encoded and not widget.text().strip():
            fallback_name = ""
            if field_name == "patient_signature_name":
                fallback_name = self.patient_name_input.text().strip()
            elif field_name == "office_staff_signature":
                fallback_name = self.office_staff_name_input.text().strip()
            elif field_name == "va10172_signature_text":
                fallback_name = self.va10172_provider_name_printed_input.text().strip() or self.provider_input.text().strip()
            if fallback_name:
                widget.setText(fallback_name)
        self._apply_signature_field_state(field_name)
        self.refresh_preview()

    def _configure_form_layout(self, form_layout):
        form_layout.setLabelAlignment(Qt.AlignLeft)
        form_layout.setFormAlignment(Qt.AlignLeft | Qt.AlignTop)
        form_layout.setFieldGrowthPolicy(QFormLayout.AllNonFixedFieldsGrow)
        form_layout.setRowWrapPolicy(QFormLayout.WrapAllRows)
        form_layout.setHorizontalSpacing(14)
        form_layout.setVerticalSpacing(8)

    def _style_builder_button(self, button):
        button.setFixedHeight(38)
        button.setMaximumWidth(190)
        button.setSizePolicy(QSizePolicy.Maximum, QSizePolicy.Fixed)
        button.setStyleSheet(
            "QPushButton {"
            "background-color:#1D2A3A; color:#FFFFFF; border:1px solid #34506B; border-radius:8px; "
            "padding:4px 12px; font-size:13px; font-weight:600; }"
            "QPushButton:hover { background-color:#223246; }"
        )

    def _style_library_button(self, button, destructive=False):
        background = "#1D2A3A"
        border = "#34506B"
        hover = "#223246"
        if destructive:
            background = "#4A1F1F"
            border = "#934141"
            hover = "#5A2626"
        button.setFixedHeight(30)
        button.setMaximumWidth(110)
        button.setSizePolicy(QSizePolicy.Maximum, QSizePolicy.Fixed)
        button.setStyleSheet(
            "QPushButton {"
            f"background-color:{background}; color:#FFFFFF; border:1px solid {border}; border-radius:8px; "
            "padding:4px 10px; font-size:12px; font-weight:600; }"
            f"QPushButton:hover {{ background-color:{hover}; }}"
        )

    def _update_wording_assist_list_width(self):
        if not hasattr(self, "wording_assist_list"):
            return
        metrics = QFontMetrics(self.wording_assist_list.font())
        longest = 0
        for index in range(self.wording_assist_list.count()):
            item = self.wording_assist_list.item(index)
            if item:
                longest = max(longest, metrics.horizontalAdvance(item.text()))
        content_width = longest + 54
        target_width = max(156, min(220, content_width))
        self.wording_assist_list.setFixedWidth(target_width)
        splitter = getattr(self, "wording_splitter", None)
        if splitter is not None:
            total_width = max(splitter.size().width(), 820)
            right_width = max(500, total_width - target_width - splitter.handleWidth() - 18)
            splitter.setSizes([target_width, right_width])

    def _set_payload_field_value(self, field_name, value):
        binding = PACKET_WIDGET_BINDINGS.get(field_name)
        if not binding:
            return False
        attr_name, kind = binding
        widget = getattr(self, attr_name, None)
        if widget is None:
            return False
        self._set_bound_widget_value(widget, value, kind)
        return True

    def _wording_status_markup(self, entry):
        bg, border, text = WORDING_STATUS_COLORS.get(entry.get("status_key"), WORDING_STATUS_COLORS["needs_review"])
        return (
            f"<span style='display:inline-block; padding:4px 10px; border-radius:10px; "
            f"background:{bg}; color:{text}; border:1px solid {border}; font-weight:700;'>"
            f"{html.escape(entry.get('status_label') or 'Needs Review')}</span>"
        )

    def _refresh_wording_assist(self):
        payload = self.collect_payload() if self._builder_ready else dict(self.packet_payload or {})
        entries = build_wording_assist_entries(payload)
        self._wording_entries = {entry["key"]: entry for entry in entries}
        selected_key = self._current_wording_entry_key if self._current_wording_entry_key in self._wording_entries else ""
        with QSignalBlocker(self.wording_assist_list):
            self.wording_assist_list.clear()
            for entry in entries:
                item = QListWidgetItem(entry["label"])
                item.setData(Qt.UserRole, entry["key"])
                bg, _, fg = WORDING_STATUS_COLORS.get(entry["status_key"], WORDING_STATUS_COLORS["needs_review"])
                item.setBackground(QColor(bg))
                item.setForeground(QColor(fg))
                item.setToolTip(entry["status_label"])
                self.wording_assist_list.addItem(item)
                if entry["key"] == selected_key:
                    self.wording_assist_list.setCurrentItem(item)
            if not selected_key and self.wording_assist_list.count() > 0:
                self.wording_assist_list.setCurrentRow(0)
        self._update_wording_assist_list_width()
        current_item = self.wording_assist_list.currentItem()
        if current_item:
            self._on_wording_entry_changed(current_item, None)
        else:
            self._current_wording_entry_key = ""
            self.wording_field_title_label.setText("No wording-assist fields on this form.")
            self.wording_status_label.setText("")
            self.wording_followup_label.setText("")
            self.wording_risk_label.setText("")
            self.wording_original_view.setPlainText("")
            self.wording_suggestion_input.setPlainText("")
            self.wording_accept_button.setEnabled(False)
            self.wording_edit_button.setEnabled(False)
            self.wording_keep_button.setEnabled(False)
            self.wording_reset_button.setEnabled(False)

    def _on_wording_entry_changed(self, current, previous):
        key = current.data(Qt.UserRole) if current else ""
        entry = dict(self._wording_entries.get(key) or {})
        self._current_wording_entry_key = key or ""
        if not entry:
            self.wording_field_title_label.setText("No wording-assist fields on this form.")
            self.wording_status_label.setText("")
            self.wording_followup_label.setText("")
            self.wording_risk_label.setText("")
            self.wording_original_view.setPlainText("")
            self.wording_suggestion_input.setPlainText("")
            self.wording_accept_button.setEnabled(False)
            self.wording_edit_button.setEnabled(False)
            self.wording_keep_button.setEnabled(False)
            self.wording_cycle_button.setEnabled(False)
            self.wording_reset_button.setEnabled(False)
            return
        self.wording_field_title_label.setText(entry["label"])
        self.wording_status_label.setText(self._wording_status_markup(entry))
        if entry["missing_facts"]:
            self.wording_followup_label.setText(
                "Follow-up needed before professional wording can be approved:\n- "
                + "\n- ".join(entry["missing_facts"])
            )
        else:
            scenario_lines = ["All supporting facts needed for wording assist are present."]
            if entry.get("scenario_label"):
                scenario_lines.append(f"Matched blueprint: {entry['scenario_label']}")
            if entry.get("scenario_use_when"):
                scenario_lines.append(f"Use when: {entry['scenario_use_when']}")
            self.wording_followup_label.setText("\n".join(scenario_lines))
        if entry["risk_messages"]:
            self.wording_risk_label.setText("Wording checks:\n- " + "\n- ".join(entry["risk_messages"]))
        else:
            self.wording_risk_label.setText("No wording-risk flags detected in the current text.")
        self.wording_original_view.setPlainText(entry["raw_text"])
        self.wording_suggestion_input.setPlainText(entry["suggestion_text"])
        allow_suggestion = bool(entry["suggestion_text"]) and not entry["missing_facts"]
        self.wording_accept_button.setEnabled(allow_suggestion)
        self.wording_edit_button.setEnabled(True)
        self.wording_keep_button.setEnabled(bool(entry["raw_text"]))
        self.wording_cycle_button.setEnabled(allow_suggestion)
        self.wording_reset_button.setEnabled(bool((self._wording_assist_state or {}).get(entry["key"])))

    def _store_wording_review(self, entry_key, decision, approved_text):
        payload = self.collect_payload()
        entries = {entry["key"]: entry for entry in build_wording_assist_entries(payload)}
        entry = entries.get(entry_key)
        if not entry:
            return
        state = dict(self._wording_assist_state or {})
        prior_state = dict(state.get(entry_key) or {})
        state[entry_key] = {
            "decision": decision,
            "approved_text": sanitize_packet_builder_text(approved_text),
            "source_fingerprint": entry["source_fingerprint"],
            "reviewed_at": datetime.now().isoformat(timespec="seconds"),
            "cycle_index": prior_state.get("cycle_index", entry.get("cycle_index", 0)),
        }
        self._wording_assist_state = state

    def cycle_wording_suggestion(self):
        entry = dict(self._wording_entries.get(self._current_wording_entry_key) or {})
        if not entry or entry["missing_facts"]:
            return
        state = dict(self._wording_assist_state or {})
        current_state = dict(state.get(entry["key"]) or {})
        try:
            cycle_index = max(0, int(current_state.get("cycle_index") or 0)) + 1
        except Exception:
            cycle_index = 1
        current_state["cycle_index"] = cycle_index
        current_state["reviewed_at"] = datetime.now().isoformat(timespec="seconds")
        state[entry["key"]] = current_state
        self._wording_assist_state = state
        self.refresh_preview()
        self._select_wording_entry(entry["key"])

    def accept_wording_suggestion(self):
        entry = dict(self._wording_entries.get(self._current_wording_entry_key) or {})
        suggestion = sanitize_packet_builder_text(self.wording_suggestion_input.toPlainText())
        if not entry or not suggestion:
            return
        self._set_payload_field_value(entry["field_name"], suggestion)
        self._store_wording_review(entry["key"], "accepted", suggestion)
        self.refresh_preview()
        self._select_wording_entry(entry["key"])

    def apply_edited_wording_suggestion(self):
        entry = dict(self._wording_entries.get(self._current_wording_entry_key) or {})
        edited_text = sanitize_packet_builder_text(self.wording_suggestion_input.toPlainText())
        if not entry:
            return
        if not edited_text:
            QMessageBox.information(self, "Edited Suggestion Required", "Enter the edited wording you want to approve for this field.")
            return
        self._set_payload_field_value(entry["field_name"], edited_text)
        self._store_wording_review(entry["key"], "edited", edited_text)
        self.refresh_preview()
        self._select_wording_entry(entry["key"])

    def keep_original_wording(self):
        entry = dict(self._wording_entries.get(self._current_wording_entry_key) or {})
        if not entry or not entry["raw_text"]:
            return
        self._store_wording_review(entry["key"], "keep_original", entry["raw_text"])
        self.refresh_preview()
        self._select_wording_entry(entry["key"])

    def reset_wording_review(self):
        entry_key = self._current_wording_entry_key
        if not entry_key:
            return
        state = dict(self._wording_assist_state or {})
        if entry_key in state:
            state.pop(entry_key, None)
            self._wording_assist_state = state
        self.refresh_preview()
        self._select_wording_entry(entry_key)

    def _select_wording_entry(self, entry_key):
        if not entry_key:
            return
        for index in range(self.wording_assist_list.count()):
            item = self.wording_assist_list.item(index)
            if item.data(Qt.UserRole) == entry_key:
                self.wording_assist_list.setCurrentItem(item)
                break

    def _set_bound_widget_value(self, widget, value, kind):
        with QSignalBlocker(widget):
            if kind == "text":
                widget.setText(str(value or ""))
            elif kind == "toPlainText":
                widget.setPlainText(str(value or ""))
            elif kind == "currentText":
                text = str(value or "")
                index = widget.findText(text)
                if index >= 0:
                    widget.setCurrentIndex(index)
                elif text:
                    widget.setCurrentText(text)
            elif kind == "isChecked":
                widget.setChecked(bool(value))

    def _apply_payload_to_widgets(self, payload, base_filename=None):
        packet = normalize_packet_builder_payload(payload)
        self._builder_ready = False
        try:
            self._wording_assist_state = dict(packet.get("wording_assist_state") or {})
            self._signature_images = {
                image_field: str(packet.get(image_field) or "").strip()
                for image_field in SIGNATURE_IMAGE_FIELDS.values()
            }
            for field_name, (attr_name, kind) in PACKET_WIDGET_BINDINGS.items():
                widget = getattr(self, attr_name, None)
                if widget is None:
                    continue
                self._set_bound_widget_value(widget, packet.get(field_name), kind)
            for field_name in SIGNATURE_IMAGE_FIELDS:
                self._apply_signature_field_state(field_name)
            if hasattr(self, "base_filename_input"):
                with QSignalBlocker(self.base_filename_input):
                    self.base_filename_input.setText(sanitize_builder_filename(base_filename or "truecore_packet"))
            template_path = str(packet.get("va10172_template_path") or self.config.get("va_form_10172_template_path") or "").strip()
            if template_path:
                self.config = dict(self.save_config_callback({"va_form_10172_template_path": template_path}) or {})
            if hasattr(self, "va10172_template_value"):
                self.va10172_template_value.setText(
                    str(self.config.get("va_form_10172_template_path") or _default_va_form_10172_template_path() or "")
                )
            self._shared_sync_state.clear()
        finally:
            self._builder_ready = True
        self.on_profile_changed(packet.get("packet_profile") or self.profile_combo.currentText())
        self.refresh_preview()

    def _current_display_name(self, payload=None, base_filename=None):
        working_payload = dict(payload or self.collect_payload())
        return packet_library_display_name(working_payload, base_filename or self._resolve_base_filename())

    def _build_current_draft_record(self):
        payload = self.collect_payload()
        now = datetime.now().isoformat(timespec="seconds")
        existing = dict(self.current_draft_record or {})
        artifacts = dict(existing.get("artifacts") or {})
        return {
            "draft_id": self.current_draft_id or existing.get("draft_id") or generate_packet_library_draft_id(),
            "display_name": self._current_display_name(payload),
            "base_filename": self._resolve_base_filename(),
            "packet_profile": str(payload.get("packet_profile") or "").strip(),
            "saved_at": existing.get("saved_at") or now,
            "updated_at": now,
            "artifacts": artifacts,
            "payload": payload,
        }

    def _save_current_draft_record(self, silent=False):
        record = save_packet_library_record(self._build_current_draft_record())
        self.current_draft_id = record.get("draft_id")
        self.current_draft_record = dict(record)
        self.refresh_library()
        self._update_current_packet_status()
        if not silent:
            QMessageBox.information(
                self,
                "Saved To Packet Library",
                f"{record.get('display_name') or 'Packet draft'} was saved to the Packet Library.",
            )
        return record

    def _update_current_draft_artifacts(self, updates):
        if not updates:
            return
        record = self.current_draft_record if self.current_draft_id else self._save_current_draft_record(silent=True)
        artifacts = dict((record or {}).get("artifacts") or {})
        artifacts.update({key: value for key, value in dict(updates or {}).items() if value})
        record = dict(record or {})
        record["artifacts"] = artifacts
        record = save_packet_library_record(record)
        self.current_draft_record = dict(record)
        self.refresh_library()
        self._update_current_packet_status()

    def _preferred_pdf_path_for_record(self, record):
        artifacts = dict((record or {}).get("artifacts") or {})
        candidates = [
            artifacts.get("compiled_patient_packet_pdf"),
            artifacts.get("single_doc_pdf"),
            artifacts.get("referral_request_pdf"),
        ]
        for bundle_key in ("patient_packet_bundle_dir", "referral_request_bundle_dir"):
            bundle_dir = str(artifacts.get(bundle_key) or "").strip()
            if bundle_dir and os.path.isdir(bundle_dir):
                for name in sorted(os.listdir(bundle_dir)):
                    if str(name).lower().endswith(".pdf"):
                        candidates.append(os.path.join(bundle_dir, name))
        for candidate in candidates:
            normalized = str(candidate or "").strip()
            if normalized and os.path.exists(normalized):
                return normalized
        return ""

    def _artifact_folder_for_record(self, record):
        artifacts = dict((record or {}).get("artifacts") or {})
        for key in ("patient_packet_bundle_dir", "referral_request_bundle_dir"):
            bundle_dir = str(artifacts.get(key) or "").strip()
            if bundle_dir and os.path.isdir(bundle_dir):
                return bundle_dir
        preferred_pdf = self._preferred_pdf_path_for_record(record)
        if preferred_pdf:
            return os.path.dirname(preferred_pdf)
        single_doc = str(artifacts.get("single_doc_docx") or artifacts.get("single_doc_pdf") or "").strip()
        if single_doc and os.path.exists(single_doc):
            return os.path.dirname(single_doc)
        return str(self.config.get("packet_builder_export_dir") or "").strip()

    def _format_library_timestamp(self, value):
        try:
            return datetime.fromisoformat(str(value)).strftime("%b %d, %Y %I:%M %p")
        except Exception:
            return "Unknown"

    def _update_current_packet_status(self):
        payload = self.collect_payload() if self._builder_ready else dict(self.packet_payload or {})
        display_name = self._current_display_name(payload)
        completion = build_packet_library_completion(payload)
        metrics = build_packet_library_production_metrics(payload)
        saved_state = "Saved Draft" if self.current_draft_id else "Unsaved Draft"
        readiness_label = {
            "ready": "Ready",
            "requires_review": "Review Required",
            "hold": "Hold",
        }.get(metrics.get("readiness"), "Hold")
        self.current_packet_status_label.setText(
            f"{saved_state}: {display_name} | {completion.get('status_label')} | "
            f"Prod Score {int(round(metrics.get('score') or 0))} | {readiness_label}"
        )

    def _build_library_row(self, record):
        payload = dict(record.get("payload") or {})
        completion = build_packet_library_completion(payload)
        metrics = build_packet_library_production_metrics(payload)
        status_palette = packet_profile_status_palette(completion.get("status_key"))
        readiness = str(metrics.get("readiness") or "hold")
        readiness_colors = {
            "ready": ("#173B29", "#2C8B57", "#EAFBF1"),
            "requires_review": ("#4A3410", "#C7922E", "#FFF6DB"),
            "hold": ("#441A1A", "#B55050", "#FFECEC"),
        }
        ready_bg, ready_border, ready_text = readiness_colors.get(readiness, readiness_colors["hold"])
        is_current = record.get("draft_id") == self.current_draft_id

        frame = QFrame()
        frame.setStyleSheet(
            "QFrame {"
            f"background:#0F1823; border:1px solid {'#4A89C7' if is_current else '#243446'}; border-radius:12px; "
            "padding:10px; }"
        )
        layout = QVBoxLayout(frame)
        layout.setContentsMargins(12, 12, 12, 12)
        layout.setSpacing(10)

        top_row = QHBoxLayout()
        top_row.setContentsMargins(0, 0, 0, 0)
        top_row.setSpacing(8)
        title_button = QPushButton(str(record.get("display_name") or "Untitled Packet"))
        title_button.setCursor(Qt.PointingHandCursor)
        title_button.setFlat(True)
        title_button.setStyleSheet(
            "QPushButton { color:#EAF2FF; font-size:15px; font-weight:700; text-align:left; border:0px; padding:0px; }"
            "QPushButton:hover { color:#8BC8FF; }"
        )
        title_button.clicked.connect(lambda _=False, draft_id=record.get("draft_id"): self.load_library_draft(draft_id))
        top_row.addWidget(title_button, stretch=1)

        status_badge = QLabel(completion.get("status_label") or "In Progress")
        status_badge.setStyleSheet(
            "padding:4px 10px; border-radius:10px; font-weight:700; "
            f"background:{status_palette['combo_background']}; color:{status_palette['combo_text']}; "
            f"border:1px solid {status_palette['combo_border']};"
        )
        top_row.addWidget(status_badge, 0, Qt.AlignRight)

        readiness_badge = QLabel(
            {
                "ready": "Ready",
                "requires_review": "Review Required",
                "hold": "Hold",
            }.get(readiness, "Hold")
        )
        readiness_badge.setStyleSheet(
            "padding:4px 10px; border-radius:10px; font-weight:700; "
            f"background:{ready_bg}; color:{ready_text}; border:1px solid {ready_border};"
        )
        top_row.addWidget(readiness_badge, 0, Qt.AlignRight)
        layout.addLayout(top_row)

        meta_text = (
            f"Updated {self._format_library_timestamp(record.get('updated_at'))} | "
            f"{'Patient Packet' if completion.get('group_name') == 'patient_packet' else 'Referral Request'} | "
            f"Prod Score {int(round(metrics.get('score') or 0))}"
        )
        if is_current:
            meta_text += " | Current Draft"
        meta_label = QLabel(meta_text)
        meta_label.setWordWrap(True)
        meta_label.setStyleSheet("color:#8FA6C1;")
        layout.addWidget(meta_label)

        blocker = str(metrics.get("main_blocker") or "").strip()
        summary_label = QLabel(
            blocker
            or "This draft does not currently show a blocking paperwork issue."
        )
        summary_label.setWordWrap(True)
        summary_label.setStyleSheet("color:#DDE8F5;")
        layout.addWidget(summary_label)

        action_row = QHBoxLayout()
        action_row.setContentsMargins(0, 0, 0, 0)
        action_row.setSpacing(8)
        resume_button = QPushButton("Resume")
        self._style_library_button(resume_button)
        resume_button.clicked.connect(lambda _=False, draft_id=record.get("draft_id"): self.load_library_draft(draft_id))
        pdf_button = QPushButton("Open PDF")
        self._style_library_button(pdf_button)
        pdf_button.clicked.connect(lambda _=False, draft_id=record.get("draft_id"): self.open_library_pdf(draft_id))
        folder_button = QPushButton("Open Folder")
        self._style_library_button(folder_button)
        folder_button.clicked.connect(lambda _=False, draft_id=record.get("draft_id"): self.open_library_folder(draft_id))
        delete_button = QPushButton("Delete")
        self._style_library_button(delete_button, destructive=True)
        delete_button.clicked.connect(lambda _=False, draft_id=record.get("draft_id"): self.delete_library_draft(draft_id))
        for button in (resume_button, pdf_button, folder_button, delete_button):
            action_row.addWidget(button)
        action_row.addStretch(1)
        layout.addLayout(action_row)
        return frame

    def refresh_library(self):
        if not hasattr(self, "library_list"):
            return
        self.library_list.clear()
        records = list_packet_library_records()
        if not records:
            empty_item = QListWidgetItem()
            empty_widget = QLabel(
                "No saved packet drafts yet. Use Save To Library after you start a packet, and it will appear here."
            )
            empty_widget.setWordWrap(True)
            empty_widget.setStyleSheet("color:#8FA6C1; padding:14px;")
            empty_item.setSizeHint(empty_widget.sizeHint())
            self.library_list.addItem(empty_item)
            self.library_list.setItemWidget(empty_item, empty_widget)
            return
        for record in records:
            item = QListWidgetItem()
            widget = self._build_library_row(record)
            item.setSizeHint(widget.sizeHint())
            self.library_list.addItem(item)
            self.library_list.setItemWidget(item, widget)

    def load_library_draft(self, draft_id):
        record = next((item for item in list_packet_library_records() if item.get("draft_id") == draft_id), None)
        if not record:
            QMessageBox.information(self, "Draft Not Found", "That packet draft could not be found in the library.")
            self.refresh_library()
            return
        self.current_draft_id = record.get("draft_id")
        self.current_draft_record = dict(record)
        self._apply_payload_to_widgets(record.get("payload") or {}, record.get("base_filename"))
        self.preview_tabs.setCurrentIndex(0)
        if hasattr(self, "packet_builder_left_tabs"):
            self.packet_builder_left_tabs.setCurrentWidget(self.packet_workspace_page)
        self.refresh_library()
        self._update_current_packet_status()

    def open_library_pdf(self, draft_id):
        record = next((item for item in list_packet_library_records() if item.get("draft_id") == draft_id), None)
        pdf_path = self._preferred_pdf_path_for_record(record or {})
        if pdf_path:
            os.startfile(pdf_path)
            return
        QMessageBox.information(
            self,
            "No PDF Yet",
            "This draft does not have a saved PDF yet. Export a PDF first, then reopen it from the library.",
        )

    def open_library_folder(self, draft_id):
        record = next((item for item in list_packet_library_records() if item.get("draft_id") == draft_id), None)
        folder_path = self._artifact_folder_for_record(record or {})
        if folder_path and os.path.isdir(folder_path):
            os.startfile(folder_path)
            return
        QMessageBox.information(self, "No Export Folder Yet", "This draft does not have an export folder yet.")

    def delete_library_draft(self, draft_id):
        reply = QMessageBox.question(
            self,
            "Delete Saved Draft",
            "Remove this saved draft from the Packet Library? Exported files on disk will be left alone.",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No,
        )
        if reply != QMessageBox.Yes:
            return
        delete_packet_library_record(draft_id)
        if draft_id == self.current_draft_id:
            self.current_draft_id = None
            self.current_draft_record = {}
        self.refresh_library()
        self._update_current_packet_status()

    def start_new_packet(self):
        reply = QMessageBox.question(
            self,
            "Start New Packet",
            "Clear the current working packet and start a new one?",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No,
        )
        if reply != QMessageBox.Yes:
            return
        self.current_draft_id = None
        self.current_draft_record = {}
        self._signature_images = {}
        self._apply_payload_to_widgets(default_packet_builder_payload(), "truecore_packet")
        self.preview_tabs.setCurrentIndex(0)
        if hasattr(self, "packet_builder_left_tabs"):
            self.packet_builder_left_tabs.setCurrentWidget(self.packet_intake_page)
        self.refresh_library()
        self._update_current_packet_status()

    def _rebalance_builder_splitter(self):
        if not hasattr(self, "builder_splitter"):
            return
        total = max(0, self.builder_splitter.width())
        if total <= 0:
            return
        half = total // 2
        self.builder_splitter.setSizes([half, max(half, total - half)])

    def _make_choice_combo(self, options):
        widget = QComboBox()
        widget.addItems(list(options or []))
        widget.currentTextChanged.connect(self.refresh_preview)
        return widget

    def _make_checkbox(self, label, checked=False):
        widget = QCheckBox(label)
        widget.setChecked(bool(checked))
        widget.stateChanged.connect(self.refresh_preview)
        return widget

    def _apply_profile_combo_status_style(self, status_key):
        palette = packet_profile_status_palette(status_key)
        self.profile_combo.setStyleSheet(
            "QComboBox {"
            f"background-color:{palette['combo_background']}; color:{palette['combo_text']}; "
            f"border:1px solid {palette['combo_border']}; border-radius:8px; padding:5px 10px; min-height:28px;"
            "}"
            "QComboBox::drop-down { border:0px; width:28px; }"
            "QComboBox QAbstractItemView {"
            "background:#0F1823; color:#E8F1FC; border:1px solid #243446; selection-background-color:#223246;"
            "}"
        )

    def _update_profile_statuses(self, base_payload):
        model = self.profile_combo.model()
        status_map = {}
        for index, profile_name in enumerate(PACKET_BUILDER_PROFILES):
            profile_payload = build_profile_export_payload(base_payload, profile_name)
            profile_report = build_packet_lab_report(profile_payload)
            status_key, status_label = classify_packet_lab_completion(profile_report)
            status_map[profile_name] = status_key
            palette = packet_profile_status_palette(status_key)
            model.setData(model.index(index, 0), palette["item_background"], Qt.BackgroundRole)
            model.setData(model.index(index, 0), palette["item_foreground"], Qt.ForegroundRole)
            model.setData(
                model.index(index, 0),
                f"{status_label} | Shared: {profile_report['shared_complete']}/{len(profile_report['shared_checks'])} | "
                f"Form: {profile_report['current_complete']}/{len(profile_report['current_form_checks'])} | "
                f"Inconsistencies: {len(profile_report.get('inconsistency_messages') or [])}",
                Qt.ToolTipRole,
            )
        self._profile_status_map = status_map
        current_status = status_map.get(self.profile_combo.currentText().strip(), "not_started")
        self._apply_profile_combo_status_style(current_status)

    def _sync_line_edit_from_shared(self, state_key, target_widget, value, fallback_values=None):
        normalized = str(value or "").strip()
        previous = str(self._shared_sync_state.get(state_key) or "").strip()
        current = target_widget.text().strip()
        if not normalized:
            if previous and current == previous:
                with QSignalBlocker(target_widget):
                    target_widget.clear()
            self._shared_sync_state[state_key] = ""
            return
        if _can_replace_with_shared_value(current, [previous, *(fallback_values or [])]):
            if current != normalized:
                with QSignalBlocker(target_widget):
                    target_widget.setText(normalized)
            self._shared_sync_state[state_key] = normalized
            return
        self._shared_sync_state[state_key] = previous

    def _sync_text_edit_from_shared(self, state_key, target_widget, value, fallback_values=None):
        normalized = str(value or "").strip()
        previous = str(self._shared_sync_state.get(state_key) or "").strip()
        current = target_widget.toPlainText().strip()
        if not normalized:
            if previous and current == previous:
                with QSignalBlocker(target_widget):
                    target_widget.clear()
            self._shared_sync_state[state_key] = ""
            return
        if _can_replace_with_shared_value(current, [previous, *(fallback_values or [])]):
            if current != normalized:
                with QSignalBlocker(target_widget):
                    target_widget.setPlainText(normalized)
            self._shared_sync_state[state_key] = normalized
            return
        self._shared_sync_state[state_key] = previous

    def _propagate_shared_fields(self):
        defaults = default_packet_builder_payload()
        auth_number = self.auth_input.text().strip()
        ordering_doctor = self.ordering_doctor_input.text().strip()
        provider_name = self.provider_input.text().strip()
        facility = self.facility_input.text().strip()
        diagnosis = self.diagnosis_input.text().strip()
        secondary_diagnosis = self.secondary_diagnosis_input.text().strip()
        icd_codes = self.icd_codes_input.text().strip()
        first_icd_code = _first_icd_code(icd_codes)
        ssn = self.ssn_input.text().strip()
        last_four = re.sub(r"\D", "", ssn)[-4:] if ssn else ""
        provider_credentials = self.master_provider_credentials_input.text().strip()
        provider_specialty = self.master_provider_specialty_input.text().strip()
        provider_npi = self.master_provider_npi_input.text().strip()
        practice_name = self.master_practice_name_input.text().strip()
        provider_phone = self.master_provider_phone_input.text().strip()
        provider_fax = self.master_provider_fax_input.text().strip()
        provider_email = self.master_provider_email_input.text().strip()
        provider_address = self.master_provider_address_input.toPlainText().strip()
        provider_address_one_line = ", ".join(part.strip() for part in provider_address.splitlines() if part.strip())
        requested_cpt = self.master_requested_cpt_code_input.text().strip()
        requested_service = self.requested_service_input.text().strip()
        mri_date = self.master_mri_date_input.text().strip()
        mri_findings = self.master_mri_findings_input.text().strip()
        affected_levels = self.master_affected_levels_input.text().strip()

        self._sync_line_edit_from_shared("auth:lmn", self.lomn_va_claim_number_input, auth_number, [defaults.get("lmn_va_claim_number")])
        self._sync_line_edit_from_shared("auth:consult", self.consult_va_claim_number_input, auth_number, [defaults.get("consult_va_claim_number")])
        self._sync_line_edit_from_shared("icd:cover", self.primary_diagnosis_code_input, first_icd_code, [defaults.get("primary_diagnosis_code")])
        self._sync_line_edit_from_shared("icd:seoc", self.episode_icd_code_input, first_icd_code, [defaults.get("episode_icd_code")])
        self._sync_line_edit_from_shared("facility:seoc", self.va_medical_center_input, facility, [defaults.get("va_medical_center_name")])
        self._sync_line_edit_from_shared("doctor:consult", self.consult_referring_va_provider_input, ordering_doctor, [defaults.get("consult_referring_va_provider")])
        self._sync_line_edit_from_shared(
            "provider:va10172",
            self.va10172_provider_name_printed_input,
            provider_name or ordering_doctor,
            [defaults.get("va10172_ordering_provider_name_printed")],
        )
        self._sync_line_edit_from_shared("diagnosis:seoc", self.episode_diagnosis_input, diagnosis, [defaults.get("episode_diagnosis")])
        self._sync_line_edit_from_shared("diagnosis:lomn", self.lomn_primary_diagnosis_input, diagnosis, [defaults.get("lmn_primary_diagnosis")])
        self._sync_line_edit_from_shared("diagnosis:consult", self.consult_primary_diagnosis_input, diagnosis, [defaults.get("consult_primary_diagnosis")])
        self._sync_line_edit_from_shared("diagnosis:clinical", self.clinical_doc_primary_diagnosis_input, diagnosis, [defaults.get("clinical_doc_primary_diagnosis")])
        self._sync_line_edit_from_shared("diagnosis:secondary_lomn", self.lomn_secondary_diagnosis_input, secondary_diagnosis, [defaults.get("lmn_secondary_diagnosis")])
        self._sync_line_edit_from_shared("diagnosis:secondary_consult", self.consult_secondary_diagnosis_input, secondary_diagnosis, [defaults.get("consult_secondary_diagnosis")])
        self._sync_line_edit_from_shared("diagnosis:secondary_clinical", self.clinical_doc_secondary_diagnosis_input, secondary_diagnosis, [defaults.get("clinical_doc_secondary_diagnosis")])
        self._sync_text_edit_from_shared("diagnosis:va10172", self.va10172_diagnosis_description_input, diagnosis, [defaults.get("va10172_diagnosis_description")])
        self._sync_line_edit_from_shared("ssn:last4", self.last_four_ssn_input, last_four)
        self._sync_line_edit_from_shared("provider:credentials", self.provider_credentials_input, provider_credentials, [defaults.get("provider_credentials")])
        self._sync_line_edit_from_shared("provider:specialty", self.provider_specialty_input, provider_specialty, [defaults.get("provider_specialty")])
        self._sync_line_edit_from_shared("provider:npi", self.provider_npi_input, provider_npi, [defaults.get("provider_npi")])
        self._sync_line_edit_from_shared("provider:npi_va10172", self.va10172_provider_npi_input, provider_npi, [defaults.get("va10172_ordering_provider_npi")])
        self._sync_line_edit_from_shared("provider:practice", self.practice_name_input, practice_name, [defaults.get("practice_name")])
        self._sync_line_edit_from_shared("provider:phone", self.provider_phone_input, provider_phone, [defaults.get("provider_phone")])
        self._sync_line_edit_from_shared("provider:phone_va10172", self.va10172_provider_phone_input, provider_phone, [defaults.get("va10172_ordering_provider_phone")])
        self._sync_line_edit_from_shared("provider:fax", self.provider_fax_input, provider_fax, [defaults.get("provider_fax")])
        self._sync_line_edit_from_shared("provider:fax_va10172", self.va10172_provider_fax_input, provider_fax, [defaults.get("va10172_ordering_provider_fax")])
        self._sync_line_edit_from_shared("provider:email", self.provider_email_input, provider_email, [defaults.get("provider_email")])
        self._sync_line_edit_from_shared("provider:email_va10172", self.va10172_provider_email_input, provider_email, [defaults.get("va10172_ordering_provider_secure_email")])
        self._sync_text_edit_from_shared("provider:address", self.va10172_provider_office_address_input, provider_address, [defaults.get("va10172_ordering_provider_office_address")])
        self._sync_line_edit_from_shared("provider:address_consult", self.provider_address_input, provider_address_one_line)
        self._sync_line_edit_from_shared("request:cpt", self.va10172_requested_cpt_input, requested_cpt, [defaults.get("va10172_requested_cpt_hcpcs_code")])
        self._sync_line_edit_from_shared("request:description", self.va10172_requested_cpt_description_input, requested_service, [defaults.get("va10172_description_cpt_hcpcs_code")])
        self._sync_line_edit_from_shared("mri:date_lomn", self.lomn_mri_date_input, mri_date, [defaults.get("lmn_mri_date")])
        self._sync_line_edit_from_shared("mri:date_consult", self.consult_mri_date_input, mri_date, [defaults.get("consult_mri_date")])
        self._sync_line_edit_from_shared("mri:date_clinical", self.clinical_doc_mri_date_input, mri_date, [defaults.get("clinical_doc_mri_date")])
        self._sync_line_edit_from_shared("mri:findings_lomn", self.lomn_mri_findings_input, mri_findings, [defaults.get("lmn_mri_findings")])
        self._sync_line_edit_from_shared("mri:findings_consult", self.consult_mri_findings_input, mri_findings, [defaults.get("consult_mri_findings")])
        self._sync_line_edit_from_shared("levels:clinical", self.clinical_doc_affected_levels_input, affected_levels, [defaults.get("clinical_doc_affected_levels")])

    def collect_payload(self):
        return normalize_packet_builder_payload({
            "packet_title": self.packet_title_input.text().strip(),
            "packet_profile": self.profile_combo.currentText().strip(),
            "patient_name": self.patient_name_input.text().strip(),
            "date_of_birth": self.dob_input.text().strip(),
            "authorization_number": self.auth_input.text().strip(),
            "va_icn": self.icn_input.text().strip(),
            "ordering_doctor": self.ordering_doctor_input.text().strip(),
            "provider": self.provider_input.text().strip(),
            "facility": self.facility_input.text().strip(),
            "community_facility": self.community_facility_input.text().strip(),
            "requested_service": self.requested_service_input.text().strip(),
            "diagnosis": self.diagnosis_input.text().strip(),
            "secondary_diagnosis": self.secondary_diagnosis_input.text().strip(),
            "icd_codes": self.icd_codes_input.text().strip(),
            "master_requested_cpt_code": self.master_requested_cpt_code_input.text().strip(),
            "master_provider_credentials": self.master_provider_credentials_input.text().strip(),
            "master_provider_specialty": self.master_provider_specialty_input.text().strip(),
            "master_provider_npi": self.master_provider_npi_input.text().strip(),
            "master_practice_name": self.master_practice_name_input.text().strip(),
            "master_provider_phone": self.master_provider_phone_input.text().strip(),
            "master_provider_fax": self.master_provider_fax_input.text().strip(),
            "master_provider_email": self.master_provider_email_input.text().strip(),
            "master_provider_address": self.master_provider_address_input.toPlainText().strip(),
            "master_mri_date": self.master_mri_date_input.text().strip(),
            "master_mri_findings": self.master_mri_findings_input.text().strip(),
            "master_affected_levels": self.master_affected_levels_input.text().strip(),
            "clinical_summary": self.clinical_summary_input.toPlainText().strip(),
            "packet_notes": self.packet_notes_input.toPlainText().strip(),
            "scenario_pathology_pattern": self.scenario_pathology_pattern_input.text().strip(),
            "scenario_conservative_duration": self.scenario_conservative_duration_input.text().strip(),
            "scenario_prior_esi_response": self.scenario_prior_esi_response_input.text().strip(),
            "scenario_functional_emphasis": self.scenario_functional_emphasis_input.text().strip(),
            "scenario_request_framing": self.scenario_request_framing_input.text().strip(),
            "scenario_symptom_pattern": self.scenario_symptom_pattern_input.text().strip(),
            "scenario_conservative_modalities": self.scenario_conservative_modalities_input.text().strip(),
            "scenario_review_concern": self.scenario_review_concern_input.text().strip(),
            "scenario_treatment_goals": self.scenario_treatment_goals_input.text().strip(),
            "referral_subtitle": self.referral_subtitle_input.text().strip(),
            "pcp_request_text": self.pcp_request_input.text().strip(),
            "referral_entry_text": self.referral_entry_input.text().strip(),
            "areas_of_concern": self.areas_of_concern_input.text().strip(),
            "group_npi": self.group_npi_input.text().strip(),
            "fax_number": self.fax_number_input.text().strip(),
            "liaison_contact_info": self.liaison_contact_input.toPlainText().strip(),
            "consent_form_title": self.consent_form_title_input.text().strip(),
            "street_address": self.street_address_input.text().strip(),
            "city": self.city_input.text().strip(),
            "state": self.state_input.text().strip(),
            "zip_code": self.zip_code_input.text().strip(),
            "home_phone": self.home_phone_input.text().strip(),
            "mobile_phone": self.mobile_phone_input.text().strip(),
            "work_phone": self.work_phone_input.text().strip(),
            "email_address": self.email_address_input.text().strip(),
            "ssn": self.ssn_input.text().strip(),
            "drivers_license": self.drivers_license_input.text().strip(),
            "drivers_license_state": self.drivers_license_state_input.text().strip(),
            "appointment_confirmation_method": self.appointment_confirmation_combo.currentText().strip(),
            "filed_for_disability": self.filed_for_disability_combo.currentText().strip(),
            "condition_work_related": self.condition_work_related_combo.currentText().strip(),
            "condition_due_to_accident": self.condition_due_to_accident_combo.currentText().strip(),
            "yes_response_explanation": self.yes_response_explanation_input.toPlainText().strip(),
            "has_attorney": self.has_attorney_combo.currentText().strip(),
            "attorney_name": self.attorney_name_input.text().strip(),
            "attorney_phone": self.attorney_phone_input.text().strip(),
            "emergency_contact_name": self.emergency_contact_name_input.text().strip(),
            "emergency_contact_relationship": self.emergency_contact_relationship_input.text().strip(),
            "emergency_contact_phone": self.emergency_contact_phone_input.text().strip(),
            "primary_insurance_carrier": self.primary_insurance_carrier_input.text().strip(),
            "primary_insurance_id": self.primary_insurance_id_input.text().strip(),
            "primary_insurance_phone": self.primary_insurance_phone_input.text().strip(),
            "secondary_insurance_carrier": self.secondary_insurance_carrier_input.text().strip(),
            "secondary_insurance_id": self.secondary_insurance_id_input.text().strip(),
            "secondary_insurance_phone": self.secondary_insurance_phone_input.text().strip(),
            "pcp_pcm_name": self.pcp_pcm_name_input.text().strip(),
            "pcp_pcm_phone": self.pcp_pcm_phone_input.text().strip(),
            "pcp_pcm_fax": self.pcp_pcm_fax_input.text().strip(),
            "consent_provider_name": self.consent_provider_name_input.text().strip(),
            "consent_initials": self.consent_initials_input.text().strip(),
            "minor_doctor_name": self.minor_doctor_name_input.text().strip(),
            "minor_consent_initials": self.minor_consent_initials_input.text().strip(),
            "service_authorization_name": self.service_authorization_name_input.text().strip(),
            "patient_signature_name": self.patient_signature_name_input.text().strip(),
            "patient_signature_image": str(self._signature_images.get("patient_signature_image") or "").strip(),
            "patient_signature_date": self.patient_signature_date_input.text().strip(),
            "submission_cover_title": self.submission_cover_title_input.text().strip(),
            "submission_date": self.submission_date_input.text().strip(),
            "primary_diagnosis_code": self.primary_diagnosis_code_input.text().strip(),
            "included_virtual_consent_form": self.included_virtual_consent_checkbox.isChecked(),
            "included_va_form_10_10172": self.included_va_form_checkbox.isChecked(),
            "included_seoc_request": self.included_seoc_checkbox.isChecked(),
            "included_consult_request": self.included_consult_checkbox.isChecked(),
            "included_lomn": self.included_lomn_checkbox.isChecked(),
            "included_clinical_notes": self.included_clinical_notes_checkbox.isChecked(),
            "included_mri_report": self.included_mri_checkbox.isChecked(),
            "submitting_office": self.submitting_office_input.text().strip(),
            "office_staff_name": self.office_staff_name_input.text().strip(),
            "office_staff_signature": self.office_staff_signature_input.text().strip(),
            "office_staff_signature_image": str(self._signature_images.get("office_staff_signature_image") or "").strip(),
            "date_reviewed": self.date_reviewed_input.text().strip(),
            "seoc_request_date": self.seoc_request_date_input.text().strip(),
            "va_medical_center_name": self.va_medical_center_input.text().strip(),
            "last_four_ssn": self.last_four_ssn_input.text().strip(),
            "episode_diagnosis": self.episode_diagnosis_input.text().strip(),
            "episode_icd_code": self.episode_icd_code_input.text().strip(),
            "seoc_scope_text": self.seoc_scope_text_input.toPlainText().strip(),
            "estimated_duration_text": self.estimated_duration_input.text().strip(),
            "clinical_objectives": self.clinical_objectives_input.toPlainText().strip(),
            "seoc_continuity_text": self.seoc_continuity_text_input.toPlainText().strip(),
            "provider_credentials": self.provider_credentials_input.text().strip(),
            "provider_specialty": self.provider_specialty_input.text().strip(),
            "provider_npi": self.provider_npi_input.text().strip(),
            "practice_name": self.practice_name_input.text().strip(),
            "provider_phone": self.provider_phone_input.text().strip(),
            "provider_fax": self.provider_fax_input.text().strip(),
            "seoc_include_preprocedure_eval": self.seoc_preprocedure_checkbox.isChecked(),
            "seoc_include_annulargram": self.seoc_annulargram_checkbox.isChecked(),
            "seoc_include_fibrin_injection": self.seoc_fibrin_checkbox.isChecked(),
            "seoc_include_follow_up": self.seoc_follow_up_checkbox.isChecked(),
            "lmn_request_date": self.lomn_request_date_input.text().strip(),
            "lmn_va_claim_number": self.lomn_va_claim_number_input.text().strip(),
            "lmn_primary_diagnosis": self.lomn_primary_diagnosis_input.text().strip(),
            "lmn_secondary_diagnosis": self.lomn_secondary_diagnosis_input.text().strip(),
            "lmn_clinical_summary": self.lomn_clinical_summary_input.toPlainText().strip(),
            "lmn_mri_date": self.lomn_mri_date_input.text().strip(),
            "lmn_mri_findings": self.lomn_mri_findings_input.text().strip(),
            "lmn_conservative_duration": self.lomn_conservative_duration_input.text().strip(),
            "lmn_include_physical_therapy": self.lomn_physical_therapy_checkbox.isChecked(),
            "lmn_include_nsaids": self.lomn_nsaids_checkbox.isChecked(),
            "lmn_include_activity_modification": self.lomn_activity_modification_checkbox.isChecked(),
            "lmn_include_home_exercise": self.lomn_home_exercise_checkbox.isChecked(),
            "lmn_include_epidural_steroid_injections": self.lomn_esi_checkbox.isChecked(),
            "lmn_medical_necessity_statement": self.lomn_medical_necessity_input.toPlainText().strip(),
            "lmn_indication_reduce_pain": self.lomn_reduce_pain_checkbox.isChecked(),
            "lmn_indication_improve_function": self.lomn_improve_function_checkbox.isChecked(),
            "lmn_indication_prevent_degeneration": self.lomn_prevent_degeneration_checkbox.isChecked(),
            "lmn_indication_reduce_opioid_reliance": self.lomn_reduce_opioids_checkbox.isChecked(),
            "lmn_indication_prevent_surgery": self.lomn_prevent_surgery_checkbox.isChecked(),
            "lmn_risk_statement": self.lomn_risk_statement_input.toPlainText().strip(),
            "lmn_reasonable_necessary_statement": self.lomn_reasonable_statement_input.toPlainText().strip(),
            "lmn_contact_statement": self.lomn_contact_statement_input.toPlainText().strip(),
            "consult_request_date": self.consult_request_date_input.text().strip(),
            "consult_va_claim_number": self.consult_va_claim_number_input.text().strip(),
            "consult_referring_va_provider": self.consult_referring_va_provider_input.text().strip(),
            "consult_reason_text": self.consult_reason_input.toPlainText().strip(),
            "consult_primary_diagnosis": self.consult_primary_diagnosis_input.text().strip(),
            "consult_secondary_diagnosis": self.consult_secondary_diagnosis_input.text().strip(),
            "consult_symptom_axial_pain": self.consult_symptom_axial_pain_checkbox.isChecked(),
            "consult_symptom_activity_exacerbation": self.consult_symptom_activity_checkbox.isChecked(),
            "consult_symptom_reduced_tolerance": self.consult_symptom_tolerance_checkbox.isChecked(),
            "consult_symptom_functional_impairment": self.consult_symptom_function_checkbox.isChecked(),
            "consult_mri_date": self.consult_mri_date_input.text().strip(),
            "consult_mri_findings": self.consult_mri_findings_input.text().strip(),
            "consult_conservative_duration": self.consult_conservative_duration_input.text().strip(),
            "consult_include_physical_therapy": self.consult_physical_therapy_checkbox.isChecked(),
            "consult_include_nsaids": self.consult_nsaids_checkbox.isChecked(),
            "consult_include_activity_modification": self.consult_activity_mod_checkbox.isChecked(),
            "consult_include_home_exercise": self.consult_home_exercise_checkbox.isChecked(),
            "consult_include_interventional_history": self.consult_interventional_history_checkbox.isChecked(),
            "consult_include_pain_management_consultation": self.consult_include_pm_consult_checkbox.isChecked(),
            "consult_include_procedural_planning": self.consult_include_planning_checkbox.isChecked(),
            "consult_include_annulargram": self.consult_include_annulargram_checkbox.isChecked(),
            "consult_include_fibrin_injection": self.consult_include_fibrin_checkbox.isChecked(),
            "consult_include_follow_up": self.consult_include_follow_up_checkbox.isChecked(),
            "consult_fibrin_levels": self.consult_fibrin_levels_input.text().strip(),
            "consult_scope_exclusion_text": self.consult_scope_exclusion_input.toPlainText().strip(),
            "consult_medical_rationale_text": self.consult_medical_rationale_input.toPlainText().strip(),
            "consult_goal_pain_reduction": self.consult_goal_pain_checkbox.isChecked(),
            "consult_goal_functional_improvement": self.consult_goal_function_checkbox.isChecked(),
            "consult_goal_reduce_analgesics": self.consult_goal_analgesics_checkbox.isChecked(),
            "consult_goal_prevent_surgery": self.consult_goal_surgery_checkbox.isChecked(),
            "consult_risk_without_treatment": self.consult_risk_without_treatment_input.toPlainText().strip(),
            "consult_duration_scope_text": self.consult_duration_scope_input.toPlainText().strip(),
            "consult_contact_statement": self.consult_contact_statement_input.toPlainText().strip(),
            "provider_address": self.provider_address_input.text().strip(),
            "provider_email": self.provider_email_input.text().strip(),
            "clinical_doc_title": self.clinical_doc_title_input.text().strip(),
            "clinical_doc_chief_complaint": self.clinical_doc_chief_complaint_input.text().strip(),
            "clinical_doc_duration_gt_3m": self.clinical_doc_duration_3m_checkbox.isChecked(),
            "clinical_doc_duration_gt_6m": self.clinical_doc_duration_6m_checkbox.isChecked(),
            "clinical_doc_duration_gt_12m": self.clinical_doc_duration_12m_checkbox.isChecked(),
            "clinical_doc_exact_duration": self.clinical_doc_exact_duration_input.text().strip(),
            "clinical_doc_pain_axial": self.clinical_doc_pain_axial_checkbox.isChecked(),
            "clinical_doc_pain_discogenic": self.clinical_doc_pain_discogenic_checkbox.isChecked(),
            "clinical_doc_pain_activity_exacerbation": self.clinical_doc_pain_activity_checkbox.isChecked(),
            "clinical_doc_pain_sitting_intolerance": self.clinical_doc_pain_sitting_checkbox.isChecked(),
            "clinical_doc_pain_standing_intolerance": self.clinical_doc_pain_standing_checkbox.isChecked(),
            "clinical_doc_pain_bending_lifting": self.clinical_doc_pain_bending_checkbox.isChecked(),
            "clinical_doc_pain_severity": self.clinical_doc_pain_severity_input.text().strip(),
            "clinical_doc_limit_occupational": self.clinical_doc_limit_occupational_checkbox.isChecked(),
            "clinical_doc_limit_prolonged_sitting": self.clinical_doc_limit_sitting_checkbox.isChecked(),
            "clinical_doc_limit_prolonged_standing": self.clinical_doc_limit_standing_checkbox.isChecked(),
            "clinical_doc_limit_ambulation": self.clinical_doc_limit_ambulation_checkbox.isChecked(),
            "clinical_doc_limit_household": self.clinical_doc_limit_household_checkbox.isChecked(),
            "clinical_doc_limit_sleep": self.clinical_doc_limit_sleep_checkbox.isChecked(),
            "clinical_doc_functional_impact": self.clinical_doc_functional_impact_input.toPlainText().strip(),
            "clinical_doc_conservative_pt": self.clinical_doc_conservative_pt_checkbox.isChecked(),
            "clinical_doc_conservative_home_exercise": self.clinical_doc_conservative_home_exercise_checkbox.isChecked(),
            "clinical_doc_conservative_nsaids": self.clinical_doc_conservative_nsaids_checkbox.isChecked(),
            "clinical_doc_conservative_non_opioid": self.clinical_doc_conservative_non_opioid_checkbox.isChecked(),
            "clinical_doc_conservative_activity_modification": self.clinical_doc_conservative_activity_checkbox.isChecked(),
            "clinical_doc_conservative_esi": self.clinical_doc_conservative_esi_checkbox.isChecked(),
            "clinical_doc_conservative_other_interventional": self.clinical_doc_conservative_other_checkbox.isChecked(),
            "clinical_doc_conservative_duration": self.clinical_doc_conservative_duration_input.text().strip(),
            "clinical_doc_esi_response": self.clinical_doc_esi_response_combo.currentText().strip(),
            "clinical_doc_mri_date": self.clinical_doc_mri_date_input.text().strip(),
            "clinical_doc_imaging_annular_tear": self.clinical_doc_imaging_annular_checkbox.isChecked(),
            "clinical_doc_imaging_disc_degeneration": self.clinical_doc_imaging_degeneration_checkbox.isChecked(),
            "clinical_doc_imaging_disc_protrusion": self.clinical_doc_imaging_protrusion_checkbox.isChecked(),
            "clinical_doc_imaging_disc_displacement": self.clinical_doc_imaging_displacement_checkbox.isChecked(),
            "clinical_doc_affected_levels": self.clinical_doc_affected_levels_input.text().strip(),
            "clinical_doc_primary_diagnosis": self.clinical_doc_primary_diagnosis_input.text().strip(),
            "clinical_doc_secondary_diagnosis": self.clinical_doc_secondary_diagnosis_input.text().strip(),
            "clinical_doc_assessment_summary": self.clinical_doc_assessment_summary_input.toPlainText().strip(),
            "clinical_doc_treatment_plan_intro": self.clinical_doc_treatment_plan_intro_input.toPlainText().strip(),
            "clinical_doc_plan_diagnostic_confirmation": self.clinical_doc_plan_diagnostic_checkbox.isChecked(),
            "clinical_doc_plan_intradiscal_intervention": self.clinical_doc_plan_intervention_checkbox.isChecked(),
            "clinical_doc_plan_follow_up": self.clinical_doc_plan_follow_up_checkbox.isChecked(),
            "clinical_doc_plan_exclusion": self.clinical_doc_plan_exclusion_input.toPlainText().strip(),
            "clinical_doc_physician_narrative": self.clinical_doc_physician_narrative_input.toPlainText().strip(),
            "va10172_template_path": self.config.get("va_form_10172_template_path") or "",
            "va10172_va_facility_address": self.va10172_facility_address_input.toPlainText().strip(),
            "va10172_ordering_provider_office_address": self.va10172_provider_office_address_input.toPlainText().strip(),
            "va10172_is_ihs_provider": self.va10172_is_ihs_combo.currentText().strip(),
            "va10172_ordering_provider_phone": self.va10172_provider_phone_input.text().strip(),
            "va10172_ordering_provider_fax": self.va10172_provider_fax_input.text().strip(),
            "va10172_ordering_provider_secure_email": self.va10172_provider_email_input.text().strip(),
            "va10172_care_needed_within_48_hours": self.va10172_48h_combo.currentText().strip(),
            "va10172_is_continuation_of_care": self.va10172_continuation_combo.currentText().strip(),
            "va10172_referral_to_specialty": self.va10172_referral_specialty_combo.currentText().strip(),
            "va10172_referral_specialty_text": self.va10172_referral_specialty_text_input.text().strip(),
            "va10172_diagnosis_description": self.va10172_diagnosis_description_input.toPlainText().strip(),
            "va10172_requested_cpt_hcpcs_code": self.va10172_requested_cpt_input.text().strip(),
            "va10172_description_cpt_hcpcs_code": self.va10172_requested_cpt_description_input.text().strip(),
            "va10172_geriatric_care_option": self.va10172_geriatric_option_combo.currentText().strip(),
            "va10172_reason_for_request": self.va10172_reason_for_request_input.toPlainText().strip(),
            "va10172_ordering_provider_name_printed": self.va10172_provider_name_printed_input.text().strip(),
            "va10172_ordering_provider_npi": self.va10172_provider_npi_input.text().strip(),
            "va10172_signature_text": self.va10172_signature_text_input.text().strip(),
            "va10172_signature_image": str(self._signature_images.get("va10172_signature_image") or "").strip(),
            "va10172_today_date": self.va10172_today_date_input.text().strip(),
            "wording_assist_state": dict(self._wording_assist_state or {}),
        })

    def _set_preview_status(self, text):
        if hasattr(self, "preview_status_label"):
            self.preview_status_label.setText(str(text or "").strip())

    def _preview_render_key_for_payload(self, payload):
        packet = apply_packet_builder_shared_field_sync(default_packet_builder_payload() | dict(payload or {}))
        serialized = json.dumps(packet, sort_keys=True, ensure_ascii=False, default=str)
        return hashlib.sha1(serialized.encode("utf-8")).hexdigest()

    def _set_exact_preview_button_state(self, busy=False):
        if not hasattr(self, "exact_preview_button"):
            return
        preview_available = bool(self.word_executable_path) or is_va_10172_profile((self.packet_payload or {}).get("packet_profile"))
        if busy:
            self.exact_preview_button.setEnabled(False)
            self.exact_preview_button.setText("Rendering Exact Preview...")
            return
        self.exact_preview_button.setEnabled(preview_available)
        self.exact_preview_button.setText("Refresh Exact Preview" if preview_available else "Exact Preview Unavailable")

    def _show_export_html_preview(self, payload, status_text=""):
        self.preview_view.setHtml(build_packet_builder_preview_html(payload))
        if hasattr(self, "preview_stack"):
            self.preview_stack.setCurrentWidget(self.preview_view)
        self._set_preview_status(status_text)

    def _load_exact_preview_pdf(self, render_key, pdf_path):
        normalized = str(pdf_path or "").strip()
        if not normalized or not os.path.exists(normalized):
            return False
        try:
            self.preview_pdf_document.close()
        except Exception:
            pass
        self.preview_pdf_document.load(normalized)
        if hasattr(self, "preview_stack"):
            self.preview_stack.setCurrentWidget(self.preview_pdf_view)
        self._preview_loaded_key = render_key
        self._set_preview_status("Exact preview matches the current exported document layout and pagination.")
        return True

    def refresh_exact_preview(self):
        payload = dict(self.packet_payload or {})
        if not payload:
            payload = self.collect_payload()
        if not self.word_executable_path and not is_va_10172_profile(payload.get("packet_profile")):
            self._show_export_html_preview(
                payload,
                "Fast preview shown. Exact Word-faithful preview is unavailable because Microsoft Word was not detected.",
            )
            self._set_exact_preview_button_state(False)
            return
        if self._preview_render_thread:
            self._set_preview_status("Exact preview is already rendering. Please let it finish.")
            return
        render_key = self._preview_render_key_for_payload(payload)
        self._preview_latest_requested_key = render_key
        os.makedirs(self._preview_cache_dir, exist_ok=True)
        cached_pdf = os.path.join(self._preview_cache_dir, f"{render_key}.pdf")
        if render_key == self._preview_loaded_key and os.path.exists(cached_pdf):
            self._set_preview_status("Exact preview matches the current exported document layout and pagination.")
            self._set_exact_preview_button_state(False)
            return
        if os.path.exists(cached_pdf):
            self._load_exact_preview_pdf(render_key, cached_pdf)
            self._set_exact_preview_button_state(False)
            return
        self._set_exact_preview_button_state(True)
        self._set_preview_status("Rendering exact preview from the exported document. You can keep editing while this finishes.")
        self._start_exact_preview_worker(render_key, payload)

    def _start_exact_preview_worker(self, render_key, payload):
        thread = QThread(self)
        worker = ExactPreviewRenderWorker(self, render_key, payload, self._preview_cache_dir)
        worker.moveToThread(thread)
        thread.started.connect(worker.run)
        worker.finished.connect(self._on_exact_preview_render_finished)
        worker.failed.connect(self._on_exact_preview_render_failed)
        worker.finished.connect(thread.quit)
        worker.failed.connect(thread.quit)
        thread.finished.connect(worker.deleteLater)
        thread.finished.connect(lambda: self._on_exact_preview_thread_finished(thread))
        self._preview_render_thread = thread
        self._preview_render_worker = worker
        thread.start()

    def _on_exact_preview_render_finished(self, render_key, pdf_path):
        self._set_exact_preview_button_state(False)
        if render_key == self._preview_current_key:
            if not self._load_exact_preview_pdf(render_key, pdf_path):
                self._show_export_html_preview(
                    self.packet_payload,
                    "The exact preview finished rendering, but the PDF could not be loaded. Showing the fast export preview instead.",
                )

    def _on_exact_preview_render_failed(self, render_key, error_text):
        self._set_exact_preview_button_state(False)
        if render_key == self._preview_current_key:
            self._show_export_html_preview(
                self.packet_payload,
                "Exact preview render failed, so Packet Studio is showing the fast export preview. "
                f"Details: {error_text}",
            )

    def _on_exact_preview_thread_finished(self, thread):
        if self._preview_render_thread is thread:
            self._preview_render_thread = None
            self._preview_render_worker = None
        self._set_exact_preview_button_state(False)

    def refresh_preview(self):
        if not self._builder_ready:
            return
        self._preview_refresh_timer.start()

    def _refresh_preview_now(self):
        if not self._builder_ready:
            return
        self._propagate_shared_fields()
        self.packet_payload = self.collect_payload()
        self._preview_current_key = self._preview_render_key_for_payload(self.packet_payload)
        self._update_profile_statuses(self.packet_payload)
        self._refresh_wording_assist()
        if self._preview_render_thread:
            status_text = "Fast preview shown while exact preview finishes in the background."
        elif self.word_executable_path or is_va_10172_profile(self.packet_payload.get("packet_profile")):
            status_text = "Fast preview shown. Use Refresh Exact Preview for the Word-faithful view."
        else:
            status_text = "Fast preview shown. Exact preview is unavailable because Microsoft Word was not detected."
        self._show_export_html_preview(
            self.packet_payload,
            status_text,
        )
        self.lab_view.setHtml(build_packet_lab_html(self.packet_payload))
        self._update_current_packet_status()
        self._set_exact_preview_button_state(False)

    def on_profile_changed(self, profile_name):
        is_referral = is_referral_request_profile(profile_name)
        is_consent = is_virtual_consent_profile(profile_name)
        is_cover_sheet = is_submission_cover_profile(profile_name)
        is_seoc = is_seoc_request_profile(profile_name)
        is_lomn = is_lomn_profile(profile_name)
        is_consult = is_consult_request_profile(profile_name)
        is_clinical_doc = is_clinical_documentation_profile(profile_name)
        is_va_10172 = is_va_10172_profile(profile_name)
        self.referral_group.setVisible(is_referral)
        self.virtual_consent_group.setVisible(is_consent)
        self.submission_cover_group.setVisible(is_cover_sheet)
        self.seoc_request_group.setVisible(is_seoc)
        self.lomn_group.setVisible(is_lomn)
        self.consult_request_group.setVisible(is_consult)
        self.clinical_doc_group.setVisible(is_clinical_doc)
        self.va10172_group.setVisible(is_va_10172)
        self.export_context_label.setText(describe_packet_export_context(profile_name))
        current_title = self.packet_title_input.text().strip()
        known_defaults = {
            default_title_for_profile(name)
            for name in PACKET_BUILDER_PROFILES
            if default_title_for_profile(name)
        }
        default_title = default_title_for_profile(profile_name)
        if default_title and (not current_title or current_title in known_defaults):
            self.packet_title_input.setText(default_title)
        current_status = getattr(self, "_profile_status_map", {}).get(profile_name, "not_started")
        self._apply_profile_combo_status_style(current_status)
        self.refresh_preview()

    def choose_export_dir(self):
        current_dir = self.config.get("packet_builder_export_dir") or _default_export_dir()
        selected = QFileDialog.getExistingDirectory(self, "Choose Packet Builder Export Folder", current_dir)
        if not selected:
            return
        self.config = dict(self.save_config_callback({"packet_builder_export_dir": selected}) or {})
        self.export_dir_value.setText(selected)
        self.refresh_preview()

    def choose_va10172_template(self):
        current_path = self.config.get("va_form_10172_template_path") or _default_va_form_10172_template_path() or _default_export_dir()
        selected, _ = QFileDialog.getOpenFileName(
            self,
            "Choose VA Form 10-10172 PDF Template",
            current_path,
            "PDF Files (*.pdf)",
        )
        if not selected:
            return
        self.config = dict(self.save_config_callback({"va_form_10172_template_path": selected}) or {})
        self.va10172_template_value.setText(selected)
        self.refresh_preview()

    def _resolve_export_dir(self):
        export_dir = str(self.config.get("packet_builder_export_dir") or "").strip()
        if export_dir and os.path.isdir(export_dir):
            return export_dir
        self.choose_export_dir()
        return str(self.config.get("packet_builder_export_dir") or "").strip()

    def _resolve_va10172_template_path(self):
        template_path = str(self.config.get("va_form_10172_template_path") or "").strip()
        if template_path and os.path.exists(template_path):
            return template_path
        template_path = _default_va_form_10172_template_path()
        if template_path and os.path.exists(template_path):
            self.config = dict(self.save_config_callback({"va_form_10172_template_path": template_path}) or {})
            if hasattr(self, "va10172_template_value"):
                self.va10172_template_value.setText(template_path)
            return template_path
        return ""

    def _resolve_base_filename(self):
        raw_text = self.base_filename_input.text()
        typed = sanitize_builder_filename(raw_text)
        if typed != raw_text:
            self.base_filename_input.setText(typed)
        return typed or "truecore_packet"

    def _ensure_export_targets(self, extensions):
        export_dir = self._resolve_export_dir()
        if not export_dir:
            QMessageBox.warning(
                self,
                "Export Folder Required",
                "Choose an export folder before saving packet drafts.",
            )
            return None

        base_name = self._resolve_base_filename()
        if not base_name:
            base_name = "truecore_packet"
        return {
            ext: os.path.join(export_dir, f"{base_name}.{ext}")
            for ext in extensions
        }

    def _ensure_bundle_export_dir(self, group_name):
        export_dir = self._resolve_export_dir()
        if not export_dir:
            QMessageBox.warning(
                self,
                "Export Folder Required",
                "Choose an export folder before exporting referral requests or patient packets.",
            )
            return None
        bundle_dir = os.path.join(export_dir, bundle_folder_name(self._resolve_base_filename(), group_name))
        os.makedirs(bundle_dir, exist_ok=True)
        return bundle_dir

    def _selected_bundle_format(self):
        return self.bundle_format_combo.currentText().strip() or "Both"

    def _section_rows(self):
        return build_packet_builder_sections(self.collect_payload())

    def save_draft_json(self):
        self._save_current_draft_record()
        self.preview_tabs.setCurrentIndex(3)

    def _set_export_busy_state(self, busy=False, status_text=""):
        for control in self._export_controls or []:
            if control is not None:
                control.setEnabled(not busy)
        if busy:
            self._set_preview_status(status_text or "Exporting packet files in the background.")
        elif status_text:
            self._set_preview_status(status_text)

    def _start_export_request(self, request):
        if self._export_thread:
            QMessageBox.information(
                self,
                "Export In Progress",
                "A packet export is already running. Let it finish before starting another one.",
            )
            return False

        action_label = str(request.get("action_label") or "packet").strip().lower()
        self._set_export_busy_state(True, f"Exporting {action_label} in the background. You can keep reviewing the packet while it runs.")
        thread = QThread(self)
        worker = PacketExportWorker(self, request)
        worker.moveToThread(thread)
        thread.started.connect(worker.run)
        worker.finished.connect(self._on_export_worker_finished)
        worker.failed.connect(self._on_export_worker_failed)
        worker.finished.connect(thread.quit)
        worker.failed.connect(thread.quit)
        thread.finished.connect(worker.deleteLater)
        thread.finished.connect(lambda: self._on_export_thread_finished(thread))
        self._export_thread = thread
        self._export_worker = worker
        thread.start()
        return True

    def _on_export_worker_finished(self, result):
        result = dict(result or {})
        artifact_updates = dict(result.get("artifact_updates") or {})
        if artifact_updates:
            self._update_current_draft_artifacts(artifact_updates)
        title = str(result.get("title") or "Export Complete").strip()
        message = str(result.get("message") or "Packet export completed.").strip()
        self._set_export_busy_state(False, "Export complete. Packet Studio is ready for the next action.")
        QMessageBox.information(self, title, message)

    def _on_export_worker_failed(self, error_text):
        self._set_export_busy_state(False, "Packet export failed. Review the error details and try again.")
        QMessageBox.warning(self, "Export Failed", str(error_text or "Packet export failed."))

    def _on_export_thread_finished(self, thread):
        if self._export_thread is thread:
            self._export_thread = None
            self._export_worker = None
        self._set_export_busy_state(False)

    def _run_export_request(self, request):
        request = dict(request or {})
        mode = str(request.get("mode") or "").strip().lower()

        if mode == "single":
            output_map = dict(request.get("output_map") or {})
            for extension, target_path in output_map.items():
                normalized_extension = str(extension or "").strip().lower()
                if normalized_extension == "docx":
                    self._write_word_doc_for_payload(target_path, request.get("payload"))
                elif normalized_extension == "pdf":
                    self._write_pdf_doc_for_payload(target_path, request.get("payload"))
            artifact_updates = dict(request.get("artifact_updates") or {})
            return {
                "title": request.get("success_title") or "Export Complete",
                "message": request.get("success_message") or "Packet export completed.",
                "artifact_updates": artifact_updates,
            }

        if mode == "bundle":
            written_files = []
            bundle_dir = str(request.get("bundle_dir") or "").strip()
            for item in request.get("documents") or []:
                payload = dict(item.get("payload") or {})
                target_path = str(item.get("target_path") or "").strip()
                extension = str(item.get("extension") or "").strip().lower()
                if extension == "docx":
                    self._write_word_doc_for_payload(target_path, payload)
                elif extension == "pdf":
                    self._write_pdf_doc_for_payload(target_path, payload)
                else:
                    continue
                written_files.append(target_path)
            notice = [
                f"Bundle folder:\n{bundle_dir}",
                "",
                f"Exported {len(written_files)} file(s).",
            ]
            if request.get("group_name") == "patient_packet":
                notice.extend(
                    [
                        "",
                        "Note: VA Form 10-10172 always includes the real filled PDF in the patient packet export.",
                    ]
                )
            return {
                "title": request.get("success_title") or "Bundle Export Complete",
                "message": "\n".join(notice),
                "artifact_updates": dict(request.get("artifact_updates") or {}),
            }

        if mode == "compiled_pdf":
            temp_paths = []
            target_path = str(request.get("target_path") or "").strip()
            with tempfile.TemporaryDirectory() as temp_dir:
                for item in request.get("documents") or []:
                    payload = dict(item.get("payload") or {})
                    temp_path = os.path.join(
                        temp_dir,
                        f"{bundle_member_filename(payload.get('packet_profile'), 'patient_packet')}.pdf",
                    )
                    self._write_pdf_doc_for_payload(temp_path, payload)
                    temp_paths.append(temp_path)
                # Keep compiled packet generation lossless. The prior continuous-page-number
                # overlay path corrupted merged PDFs, so for now we preserve the original
                # rendered pages exactly and revisit packet-wide numbering separately.
                self._merge_pdf_files(temp_paths, target_path, add_continuous_page_numbers=False)
            return {
                "title": request.get("success_title") or "Patient Packet Export Complete",
                "message": request.get("success_message") or f"Compiled patient packet PDF saved to:\n{target_path}",
                "artifact_updates": dict(request.get("artifact_updates") or {}),
            }

        raise RuntimeError("Unknown export request.")

    def _confirm_export_warning_if_needed(self, payload, group_name=""):
        warning = build_export_warning_context(payload, group_name=group_name)
        if not warning:
            return True
        choice = QMessageBox.question(
            self,
            str(warning.get("title") or "Export Warning"),
            str(warning.get("message") or "This export looks incomplete. Continue?"),
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No,
        )
        return choice == QMessageBox.Yes

    def _confirm_wording_review_if_needed(self, payload, group_name=""):
        blockers = build_wording_export_blockers(payload, group_name=group_name)
        if not blockers:
            return True
        self.preview_tabs.setCurrentIndex(1)
        blocker_lines = "\n".join(f"- {item}" for item in blockers[:12])
        if len(blockers) > 12:
            blocker_lines += f"\n- ...and {len(blockers) - 12} more"
        QMessageBox.information(
            self,
            "Wording Review Required",
            "Before final export, Packet Studio needs wording approval on the denial-sensitive fields below.\n\n"
            + blocker_lines
            + "\n\nOpen the Wording Assist tab, review the original wording, and either accept the suggestion, edit it, or intentionally keep the original once the supporting facts are complete.",
        )
        return False

    def export_word(self):
        targets = self._ensure_export_targets(["docx"])
        if not targets:
            return
        payload = self.collect_payload()
        if not self._confirm_export_warning_if_needed(payload):
            return
        if not self._confirm_wording_review_if_needed(payload):
            return
        self._start_export_request(
            {
                "mode": "single",
                "action_label": "word export",
                "payload": payload,
                "output_map": {"docx": targets["docx"]},
                "artifact_updates": {"single_doc_docx": targets["docx"]},
                "success_title": "Word Export Complete",
                "success_message": f"Word packet draft saved to:\n{targets['docx']}",
            }
        )

    def export_pdf(self):
        targets = self._ensure_export_targets(["pdf"])
        if not targets:
            return
        payload = self.collect_payload()
        if not self._confirm_export_warning_if_needed(payload):
            return
        if not self._confirm_wording_review_if_needed(payload):
            return
        self._start_export_request(
            {
                "mode": "single",
                "action_label": "pdf export",
                "payload": payload,
                "output_map": {"pdf": targets["pdf"]},
                "artifact_updates": {"single_doc_pdf": targets["pdf"]},
                "success_title": "PDF Export Complete",
                "success_message": f"PDF packet draft saved to:\n{targets['pdf']}",
            }
        )

    def export_both(self):
        targets = self._ensure_export_targets(["docx", "pdf"])
        if not targets:
            return
        payload = self.collect_payload()
        if not self._confirm_export_warning_if_needed(payload):
            return
        if not self._confirm_wording_review_if_needed(payload):
            return
        self._start_export_request(
            {
                "mode": "single",
                "action_label": "word and pdf export",
                "payload": payload,
                "output_map": {"docx": targets["docx"], "pdf": targets["pdf"]},
                "artifact_updates": {"single_doc_docx": targets["docx"], "single_doc_pdf": targets["pdf"]},
                "success_title": "Packet Exports Complete",
                "success_message": f"Word draft:\n{targets['docx']}\n\nPDF draft:\n{targets['pdf']}",
            }
        )

    def export_referral_request_bundle(self):
        self._export_bundle_group("referral_request", "Referral Request Export Complete")

    def export_patient_packet_bundle(self):
        if normalize_bundle_export_format(self._selected_bundle_format()) == "pdf":
            self._export_compiled_patient_packet_pdf()
            return
        self._export_bundle_group("patient_packet", "Patient Packet Export Complete")

    def open_export_folder(self):
        export_dir = self._resolve_export_dir()
        if export_dir and os.path.isdir(export_dir):
            os.startfile(export_dir)
        else:
            QMessageBox.information(self, "No Export Folder Yet", "Choose an export folder first.")

    def _export_bundle_group(self, group_name, success_title):
        bundle_dir = self._ensure_bundle_export_dir(group_name)
        if not bundle_dir:
            return
        base_payload = self.collect_payload()
        if not self._confirm_export_warning_if_needed(base_payload, group_name=group_name):
            return
        if not self._confirm_wording_review_if_needed(base_payload, group_name=group_name):
            return
        plan = build_bundle_export_plan(self._resolve_base_filename(), group_name, self._selected_bundle_format())
        documents = []
        for item in plan.get("documents") or []:
            payload = build_profile_export_payload(base_payload, item.get("profile_name"))
            target_path = os.path.join(bundle_dir, item.get("filename") or "packet_document")
            documents.append(
                {
                    "payload": payload,
                    "target_path": target_path,
                    "extension": str(item.get("extension") or "").strip().lower(),
                }
            )
        if group_name == "patient_packet":
            artifact_key = "patient_packet_bundle_dir"
        else:
            artifact_key = "referral_request_bundle_dir"
        artifact_updates = {artifact_key: bundle_dir}
        if group_name == "referral_request":
            referral_pdf = os.path.join(bundle_dir, "Community_Care_Referral_Request.pdf")
            artifact_updates["referral_request_pdf"] = referral_pdf
        self._start_export_request(
            {
                "mode": "bundle",
                "action_label": f"{group_name.replace('_', ' ')} bundle",
                "group_name": group_name,
                "bundle_dir": bundle_dir,
                "documents": documents,
                "artifact_updates": artifact_updates,
                "success_title": success_title,
            }
        )

    def _export_compiled_patient_packet_pdf(self):
        export_dir = self._resolve_export_dir()
        if not export_dir:
            QMessageBox.warning(
                self,
                "Export Folder Required",
                "Choose an export folder before exporting the compiled patient packet PDF.",
            )
            return

        base_payload = self.collect_payload()
        if not self._confirm_export_warning_if_needed(base_payload, group_name="patient_packet"):
            return
        if not self._confirm_wording_review_if_needed(base_payload, group_name="patient_packet"):
            return
        target_path = os.path.join(export_dir, compiled_packet_filename(self._resolve_base_filename(), "patient_packet"))
        documents = [
            {"payload": build_profile_export_payload(base_payload, profile_name)}
            for profile_name in PATIENT_PACKET_PROFILE_ORDER
        ]
        self._start_export_request(
            {
                "mode": "compiled_pdf",
                "action_label": "compiled patient packet pdf",
                "documents": documents,
                "target_path": target_path,
                "artifact_updates": {"compiled_patient_packet_pdf": target_path},
                "success_title": "Patient Packet Export Complete",
                "success_message": f"Compiled patient packet PDF saved to:\n{target_path}",
            }
        )

    def _merge_pdf_files(self, source_paths, target_path, add_continuous_page_numbers=False):
        merged = PdfWriter()
        for source_path in source_paths or []:
            reader = PdfReader(source_path)
            for page in reader.pages:
                merged.add_page(page)
        if add_continuous_page_numbers and merged.pages:
            self._overlay_compiled_pdf_page_numbers(merged)
        with open(target_path, "wb") as handle:
            merged.write(handle)

    def _overlay_compiled_pdf_page_numbers(self, writer):
        total_pages = len(writer.pages or [])
        if total_pages <= 0:
            return

        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_handle:
            overlay_path = temp_handle.name

        pdf_writer = QPdfWriter(overlay_path)
        pdf_writer.setResolution(72)
        pdf_writer.setPageMargins(QMarginsF(0, 0, 0, 0))
        pdf_writer.setPageSize(QPageSize(QSizeF(612, 792), QPageSize.Point))

        painter = QPainter(pdf_writer)
        for index in range(total_pages):
            if index > 0:
                pdf_writer.newPage()
            painter.fillRect(QRectF(196, 756, 220, 22), QColor("#FFFFFF"))
            painter.setPen(QColor("#374151"))
            font = QFont("Arial", 9)
            painter.setFont(font)
            painter.drawText(
                QRectF(170, 755, 272, 18),
                Qt.AlignCenter,
                f"Page {index + 1} of {total_pages}",
            )
        painter.end()

        overlay_reader = PdfReader(overlay_path)
        for index, page in enumerate(writer.pages):
            if index < len(overlay_reader.pages):
                page.merge_page(overlay_reader.pages[index])

        try:
            os.remove(overlay_path)
        except OSError:
            pass

    def _load_pdf_widget_maps(self, template_path):
        reader = PdfReader(template_path)
        page_maps = []
        for page in reader.pages:
            page_map = {"fields": {}, "groups": {}}
            annots = page.get("/Annots")
            if hasattr(annots, "get_object"):
                annots = annots.get_object()
            for annot_ref in annots or []:
                annot = annot_ref.get_object()
                if annot.get("/Subtype") != "/Widget":
                    continue
                rect = annot.get("/Rect")
                rect = [float(value) for value in rect] if rect else None
                field_name = annot.get("/T")
                field_type = annot.get("/FT")
                parent = annot.get("/Parent")
                parent_name = None
                if parent:
                    parent_name = str(parent.get_object().get("/T") or "")
                if field_name and field_type in {"/Tx", "/Sig"} and rect:
                    page_map["fields"][str(field_name)] = {"rect": rect, "type": str(field_type)}
                    continue
                if parent_name and rect:
                    appearance = annot.get("/AP", {}).get("/N")
                    option_names = []
                    if appearance:
                        appearance_obj = appearance.get_object() if hasattr(appearance, "get_object") else appearance
                        if hasattr(appearance_obj, "keys"):
                            option_names = [str(key) for key in appearance_obj.keys()]
                    page_map["groups"].setdefault(parent_name, []).append(
                        {"rect": rect, "option": option_names[0] if option_names else ""}
                    )
            for group_name, widgets in page_map["groups"].items():
                page_map["groups"][group_name] = sorted(
                    widgets,
                    key=lambda item: (-item["rect"][1], item["rect"][0]),
                )
            page_maps.append(page_map)
        return page_maps

    def _pdf_rect_to_qt_rect(self, rect, page_height=792.0):
        x0, y0, x1, y1 = rect
        return QRectF(float(x0), float(page_height) - float(y1), float(x1 - x0), float(y1 - y0))

    def _draw_text_in_pdf_rect(self, painter, rect, text, font_size=9, bold=False, multiline=False):
        if not rect or not str(text or "").strip():
            return
        painter.save()
        font = QFont("Arial", font_size)
        font.setBold(bool(bold))
        painter.setFont(font)
        flags = Qt.AlignLeft | Qt.AlignTop
        if multiline:
            flags |= Qt.TextWordWrap
        else:
            flags = Qt.AlignLeft | Qt.AlignVCenter
        painter.drawText(self._pdf_rect_to_qt_rect(rect), flags, str(text).strip())
        painter.restore()

    def _draw_signature_image_in_pdf_rect(self, painter, rect, payload, field_name):
        image_bytes = _signature_image_bytes(payload, field_name)
        if not rect or not image_bytes:
            return False
        pixmap = QPixmap()
        if not pixmap.loadFromData(image_bytes, "PNG"):
            return False
        target_rect = self._pdf_rect_to_qt_rect(rect)
        scaled = pixmap.scaled(
            int(target_rect.width()),
            int(target_rect.height()),
            Qt.KeepAspectRatio,
            Qt.SmoothTransformation,
        )
        painter.save()
        x = target_rect.x()
        y = target_rect.y() + max(0.0, (target_rect.height() - scaled.height()) / 2.0)
        painter.drawPixmap(int(x), int(y), scaled)
        painter.restore()
        return True

    def _draw_checkbox_mark(self, painter, rect):
        if not rect:
            return
        painter.save()
        font = QFont("Arial", 12)
        font.setBold(True)
        painter.setFont(font)
        painter.drawText(self._pdf_rect_to_qt_rect(rect), Qt.AlignCenter, "X")
        painter.restore()

    def _draw_group_choice(self, painter, page_map, group_name, choice_label, option_map):
        widgets = page_map["groups"].get(group_name) or []
        option_name = option_map.get(str(choice_label or "").strip(), "")
        if not option_name:
            return
        for widget in widgets:
            if widget.get("option") == option_name:
                self._draw_checkbox_mark(painter, widget.get("rect"))
                return

    def _draw_va10172_page_overlay(self, painter, page_map, payload, page_index):
        patient_name = payload.get("patient_name") or ""
        dob = payload.get("date_of_birth") or ""
        auth_number = payload.get("authorization_number") or ""
        va_facility_address = payload.get("va10172_va_facility_address") or ""
        provider_office_address = payload.get("va10172_ordering_provider_office_address") or ""
        provider_phone = payload.get("va10172_ordering_provider_phone") or ""
        provider_fax = payload.get("va10172_ordering_provider_fax") or ""
        provider_email = payload.get("va10172_ordering_provider_secure_email") or ""
        provider_name_printed = payload.get("va10172_ordering_provider_name_printed") or payload.get("provider") or payload.get("ordering_doctor") or ""
        provider_npi = payload.get("va10172_ordering_provider_npi") or payload.get("provider_npi") or ""
        today_date = payload.get("va10172_today_date") or ""
        signature_text = payload.get("va10172_signature_text") or ""

        if page_index == 0:
            field_values = {
                "VETERANSNAME[0]": patient_name,
                "DOB[0]": dob,
                "VAFACILITYADDRESS[0]": va_facility_address,
                "VAAUTHORIZATIONNUMBER[0]": auth_number,
                "ORDERINGPROVIDEROFFICENAMEADDRESS[0]": provider_office_address,
                "ORDERINGPROVIDERPHONENUMBER[0]": provider_phone,
                "ORDERINGPROVIDERFAXNUMBER[0]": provider_fax,
                "ORDERINGPROVIDERSECUREEMAILADDRESS[0]": provider_email,
                "SPECIALTY[0]": payload.get("va10172_referral_specialty_text") or "",
                "DIAGNOSISCODES[0]": payload.get("icd_codes") or "",
                "DIAGNOSISDESCRIPTION[0]": payload.get("va10172_diagnosis_description") or payload.get("diagnosis") or "",
                "REQUESTEDCPTHCPCSCODE[0]": payload.get("va10172_requested_cpt_hcpcs_code") or "",
                "DESCRIPTIONCPTHCPCSCODE[0]": payload.get("va10172_description_cpt_hcpcs_code") or "",
                "TextField1[0]": payload.get("va10172_reason_for_request") or "",
                "ORDERINGPROVIDERSNAMEPRINTED[0]": provider_name_printed,
                "ORDERINGPROVIDERSNPI[0]": provider_npi,
                "Date[0]": today_date,
                "SignatureField11[0]": signature_text,
            }
            multiline_fields = {
                "VAFACILITYADDRESS[0]",
                "ORDERINGPROVIDEROFFICENAMEADDRESS[0]",
                "DIAGNOSISDESCRIPTION[0]",
                "DESCRIPTIONCPTHCPCSCODE[0]",
                "TextField1[0]",
            }
            small_font_fields = {
                "VAFACILITYADDRESS[0]": 8,
                "ORDERINGPROVIDEROFFICENAMEADDRESS[0]": 8,
                "ORDERINGPROVIDERSECUREEMAILADDRESS[0]": 8,
                "SPECIALTY[0]": 8,
                "DIAGNOSISCODES[0]": 8,
                "DIAGNOSISDESCRIPTION[0]": 7,
                "REQUESTEDCPTHCPCSCODE[0]": 8,
                "DESCRIPTIONCPTHCPCSCODE[0]": 7,
                "TextField1[0]": 8,
            }
            for field_name, value in field_values.items():
                field_info = page_map["fields"].get(field_name) or {}
                if field_name == "SignatureField11[0]" and self._draw_signature_image_in_pdf_rect(
                    painter, field_info.get("rect"), payload, "va10172_signature_text"
                ):
                    continue
                self._draw_text_in_pdf_rect(
                    painter,
                    field_info.get("rect"),
                    value,
                    font_size=small_font_fields.get(field_name, 9),
                    bold=False,
                    multiline=field_name in multiline_fields,
                )

            yes_no_map = {"No": "/0", "Yes": "/1"}
            self._draw_group_choice(painter, page_map, "HISTHP[0]", payload.get("va10172_is_ihs_provider"), yes_no_map)
            self._draw_group_choice(painter, page_map, "RadioButtonList[0]", payload.get("va10172_care_needed_within_48_hours"), yes_no_map)
            self._draw_group_choice(painter, page_map, "RadioButtonList[1]", payload.get("va10172_is_continuation_of_care"), yes_no_map)
            self._draw_group_choice(painter, page_map, "RadioButtonList[2]", payload.get("va10172_referral_to_specialty"), yes_no_map)

            geriatric_map = {
                "Community Nursing Home": "/0",
                "Home Infusion": "/1",
                "Hospice/Palliative Care": "/2",
                "Skilled Home Health Care": "/3",
                "Community Adult Day Health Care": "/4",
                "Home Homemaker/Home Health Aide": "/5",
                "Respite": "/6",
            }
            self._draw_group_choice(
                painter,
                page_map,
                "RadioButtonList[3]",
                payload.get("va10172_geriatric_care_option"),
                geriatric_map,
            )
            return

        if page_index == 1:
            field_values = {
                "VETERANSNAME[1]": patient_name,
                "DOB[1]": dob,
                "VAFACILITYADDRESS[1]": va_facility_address,
                "VAAUTHORIZATIONNUMBER[1]": auth_number,
                "ORDERINGPROVIDEROFFICENAMEADDRESS[1]": provider_office_address,
                "ORDERINGPROVIDERPHONENUMBER[1]": provider_phone,
                "ORDERINGPROVIDERFAXNUMBER[1]": provider_fax,
                "ORDERINGPROVIDERSECUREEMAILADDRESS[1]": provider_email,
                "ORDERINGPROVIDERSNAMEPRINTED2[0]": provider_name_printed,
                "ORDERINGPROVIDERSNPI2[0]": provider_npi,
                "Date2[0]": today_date,
                "SignatureField2[0]": signature_text,
            }
            for field_name, value in field_values.items():
                field_info = page_map["fields"].get(field_name) or {}
                if field_name == "SignatureField2[0]" and self._draw_signature_image_in_pdf_rect(
                    painter, field_info.get("rect"), payload, "va10172_signature_text"
                ):
                    continue
                self._draw_text_in_pdf_rect(
                    painter,
                    field_info.get("rect"),
                    value,
                    font_size=8 if field_name in {"VAFACILITYADDRESS[1]", "ORDERINGPROVIDEROFFICENAMEADDRESS[1]"} else 9,
                    bold=False,
                    multiline=field_name in {"VAFACILITYADDRESS[1]", "ORDERINGPROVIDEROFFICENAMEADDRESS[1]"},
                )

    def _write_va_10172_pdf(self, path, payload):
        template_path = self._resolve_va10172_template_path()
        if not template_path or not os.path.exists(template_path):
            raise FileNotFoundError("VA Form 10-10172 template PDF not found.")

        page_maps = self._load_pdf_widget_maps(template_path)
        app = QGuiApplication.instance() or QGuiApplication([])
        del app
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_handle:
            overlay_path = temp_handle.name

        pdf_writer = QPdfWriter(overlay_path)
        pdf_writer.setResolution(72)
        pdf_writer.setPageMargins(QMarginsF(0, 0, 0, 0))
        pdf_writer.setPageSize(QPageSize(QSizeF(612, 792), QPageSize.Point))

        painter = QPainter(pdf_writer)
        for index, page_map in enumerate(page_maps):
            if index > 0:
                pdf_writer.newPage()
            self._draw_va10172_page_overlay(painter, page_map, payload, index)
        painter.end()

        base_reader = PdfReader(template_path)
        overlay_reader = PdfReader(overlay_path)
        merged = PdfWriter()
        for index, base_page in enumerate(base_reader.pages):
            if index < len(overlay_reader.pages):
                base_page.merge_page(overlay_reader.pages[index])
            merged.add_page(base_page)
        with open(path, "wb") as handle:
            merged.write(handle)

        try:
            os.remove(overlay_path)
        except OSError:
            pass

    def _write_word_doc(self, path):
        payload = self.collect_payload()
        self._write_word_doc_for_payload(path, payload)

    def _write_pdf_doc(self, path):
        payload = self.collect_payload()
        self._write_pdf_doc_for_payload(path, payload)

    def _write_word_doc_for_payload(self, path, payload):
        if is_referral_request_profile(payload.get("packet_profile")):
            self._write_referral_request_doc(path, payload)
            return
        if is_virtual_consent_profile(payload.get("packet_profile")):
            self._write_virtual_consent_doc_v2(path, payload)
            return
        if is_submission_cover_profile(payload.get("packet_profile")):
            self._write_submission_cover_doc_v2(path, payload)
            return
        if is_seoc_request_profile(payload.get("packet_profile")):
            self._write_seoc_request_doc_v2(path, payload)
            return
        if is_lomn_profile(payload.get("packet_profile")):
            self._write_lomn_doc_v2(path, payload)
            return
        if is_consult_request_profile(payload.get("packet_profile")):
            self._write_consult_request_doc_v2(path, payload)
            return
        if is_clinical_documentation_profile(payload.get("packet_profile")):
            self._write_clinical_documentation_doc_v2(path, payload)
            return
        if is_va_10172_profile(payload.get("packet_profile")):
            self._write_va_10172_summary_doc(path, payload)
            return
        document = Document()
        document.add_heading(payload.get("packet_title") or "TrueCore Packet Draft", level=0)
        document.add_paragraph(f"Packet Profile: {payload.get('packet_profile') or 'Pending'}")
        for section_title, rows in build_packet_builder_sections(payload):
            document.add_heading(section_title, level=1)
            table = document.add_table(rows=0, cols=2)
            table.style = "Table Grid"
            for label, value in rows:
                cells = table.add_row().cells
                cells[0].text = label
                cells[1].text = str(value or "Pending")
        document.save(path)

    def _write_pdf_doc_for_payload(self, path, payload):
        if is_va_10172_profile(payload.get("packet_profile")):
            self._write_va_10172_pdf(path, payload)
            return
        if find_word_executable():
            with tempfile.TemporaryDirectory() as temp_dir:
                docx_name = sanitize_builder_filename(payload.get("packet_title") or payload.get("packet_profile") or "packet_preview")
                docx_path = os.path.join(temp_dir, f"{docx_name}.docx")
                self._write_word_doc_for_payload(docx_path, payload)
                convert_docx_to_pdf_via_word(docx_path, path)
            return
        document = QTextDocument()
        document.setHtml(build_packet_builder_export_html(payload))
        printer = QPrinter(QPrinter.HighResolution)
        printer.setOutputFormat(QPrinter.PdfFormat)
        printer.setOutputFileName(path)
        printer.setPageMargins(QMarginsF(14, 14, 14, 14), QPageLayout.Millimeter)
        document.print_(printer)

    def _write_referral_request_doc(self, path, payload):
        document = Document()
        title = document.add_paragraph()
        title.alignment = 1
        title_run = title.add_run(payload.get("packet_title") or "Community Care Referral Request")
        title_run.bold = True
        title_run.font.size = Pt(20)

        subtitle = document.add_paragraph()
        subtitle.alignment = 1
        subtitle_run = subtitle.add_run(payload.get("referral_subtitle") or "Spine & Pain Evaluation")
        subtitle_run.bold = True

        intro = document.add_paragraph()
        intro.alignment = 1
        intro.add_run(
            "This document supports your request for a Community Care Referral for specialty spine and pain evaluation. "
            "Please share this with your Primary Care VA provider during your appointment."
        )

        emphasis = document.add_paragraph()
        emphasis_run = emphasis.add_run("You do not need to request a specific procedure to be referred.")
        emphasis_run.bold = True
        document.add_paragraph(
            "We provide a comprehensive evaluation in a non-surgical spine care setting and determine the most appropriate treatment based on your condition."
        )
        document.add_paragraph(
            "Community Care allows you to receive specialty care outside the VA when clinically appropriate."
        )

        document.add_heading("What to Do Next", level=1)
        steps = [
            "Schedule a visit with your Primary Care VA Provider (PCP).",
            "Make this request clearly:",
            f"“{payload.get('pcp_request_text') or 'Pending'}”",
            "Referral should be entered as:",
            payload.get("referral_entry_text") or "Pending",
            "Include all relevant diagnoses and areas of concern:",
            payload.get("areas_of_concern") or "Pending",
        ]
        for item in steps:
            document.add_paragraph(item, style="List Bullet")

        document.add_heading("Important", level=1)
        for item in [
            "This request is for:",
            "“Specialty evaluation and consultation – not a request for a specific procedure.”",
        ]:
            document.add_paragraph(item, style="List Bullet")

        document.add_heading("Referral Routing", level=1)
        document.add_paragraph(f"Group NPI: {payload.get('group_npi') or 'Pending'}")
        document.add_paragraph(f"Fax: {payload.get('fax_number') or 'Pending'}")
        document.add_paragraph("This ensures proper routing through the HSRM system.")

        document.add_heading("Once Authorized", level=1)
        for item in [
            "Our team will coordinate your evaluation",
            "Imaging and medical history will be reviewed",
            "A personalized treatment plan will be developed",
        ]:
            document.add_paragraph(item, style="List Bullet")

        document.add_heading("Send Records", level=1)
        document.add_paragraph("Please ensure all prior imaging and relevant medical records are included.")

        document.add_heading("Need Help?", level=1)
        document.add_paragraph("Our Veteran liaison team is here to assist you.")
        document.add_paragraph(payload.get("liaison_contact_info") or "Pending")

        document.save(path)

    def _write_virtual_consent_doc(self, path, payload):
        document = Document()

        title = document.add_paragraph()
        title.alignment = 1
        title_run = title.add_run(payload.get("packet_title") or payload.get("consent_form_title") or "TELEHEALTH VIRTUAL CONSENT FORM")
        title_run.bold = True
        title_run.font.size = Pt(18)

        document.add_heading("Demographic Information", level=1)
        demographic_rows = [
            ("Full Name", payload.get("patient_name") or "Pending"),
            ("Date of Birth", payload.get("date_of_birth") or "Pending"),
            ("Street Address", payload.get("street_address") or "Pending"),
            ("City", payload.get("city") or "Pending"),
            ("State", payload.get("state") or "Pending"),
            ("Zip", payload.get("zip_code") or "Pending"),
            ("Home Phone", payload.get("home_phone") or "Pending"),
            ("Mobile Phone", payload.get("mobile_phone") or "Pending"),
            ("Work Phone", payload.get("work_phone") or "Pending"),
            ("Email Address", payload.get("email_address") or "Pending"),
            ("SSN", payload.get("ssn") or "Pending"),
            ("Driver License", payload.get("drivers_license") or "Pending"),
            ("DL State", payload.get("drivers_license_state") or "Pending"),
            ("Appointment Confirmation Method", payload.get("appointment_confirmation_method") or "Pending"),
        ]
        self._write_key_value_table(document, demographic_rows)

        document.add_heading("Eligibility And Legal Questions", level=1)
        legal_rows = [
            ("Filed for Disability", payload.get("filed_for_disability") or "Pending"),
            ("Condition Work Related", payload.get("condition_work_related") or "Pending"),
            ("Condition Due to Accident", payload.get("condition_due_to_accident") or "Pending"),
            ("Yes Response Explanation", payload.get("yes_response_explanation") or "Pending"),
            ("Has Attorney", payload.get("has_attorney") or "Pending"),
            ("Attorney Name", payload.get("attorney_name") or "Pending"),
            ("Attorney Phone", payload.get("attorney_phone") or "Pending"),
            ("Emergency Contact", payload.get("emergency_contact_name") or "Pending"),
            ("Relationship", payload.get("emergency_contact_relationship") or "Pending"),
            ("Emergency Contact Phone", payload.get("emergency_contact_phone") or "Pending"),
        ]
        self._write_key_value_table(document, legal_rows)

        document.add_heading("Insurance Information", level=1)
        insurance_rows = [
            ("Primary Insurance Carrier", payload.get("primary_insurance_carrier") or "Pending"),
            ("Primary Insurance ID", payload.get("primary_insurance_id") or "Pending"),
            ("Primary Insurance Phone", payload.get("primary_insurance_phone") or "Pending"),
            ("Secondary Insurance Carrier", payload.get("secondary_insurance_carrier") or "Pending"),
            ("Secondary Insurance ID", payload.get("secondary_insurance_id") or "Pending"),
            ("Secondary Insurance Phone", payload.get("secondary_insurance_phone") or "Pending"),
            ("PCP / PCM", payload.get("pcp_pcm_name") or "Pending"),
            ("PCP / PCM Phone", payload.get("pcp_pcm_phone") or "Pending"),
            ("PCP / PCM Fax", payload.get("pcp_pcm_fax") or "Pending"),
        ]
        self._write_key_value_table(document, insurance_rows)

        document.add_heading("Consent for Medical Care and Treatment", level=1)
        document.add_paragraph(
            f"I, {payload.get('patient_name') or 'Pending'}, hereby agree and give my consent for "
            f"{payload.get('consent_provider_name') or 'Pending'} to furnish medical care and treatment considered necessary "
            f"and proper in evaluating or treating my physical condition. Initials: {payload.get('consent_initials') or 'Pending'}"
        )
        document.add_paragraph("FOR MINORS ONLY CONSENT FOR CARE:")
        document.add_paragraph(
            f"As parent and/or legal guardian, I authorize {payload.get('minor_doctor_name') or 'Pending'} to treat the minor patient "
            f"named in the attached form while I am not present. Initials: {payload.get('minor_consent_initials') or 'Pending'}"
        )

        document.add_heading("Telehealth Consent", level=1)
        for bullet in [
            "I understand that through an interactive video connection, telehealth might be used to perform my consultation and that we may use telehealth to connect while working together.",
            "I understand the benefits and risks involved with telehealth technology.",
            "I do not need to travel to the consultation location.",
            "I have access to a specialist through this consultation.",
            "There may be interruptions, unauthorized access and technical difficulties and the telemedicine consult/visit can be discontinued if needed.",
            "My healthcare information may be shared with other individuals for scheduling and billing purposes.",
            "Reasonable steps will be taken to maintain confidentiality.",
            "I understand the initial virtual meet-and-greet discussion is primarily intended to determine any and all treatment I might pursue, in addition to the best approach for scheduling these treatments.",
            "I have read this document in its entirety and understand the risks and benefits of telehealth sessions.",
        ]:
            document.add_paragraph(bullet, style="List Bullet")

        document.add_heading("Final Authorization", level=1)
        document.add_paragraph(
            f"By signing below, I agree that all the above information is correct, and that I authorize "
            f"{payload.get('service_authorization_name') or 'Pending'} to provide me with medical services and to furnish my physician, insurance company or attorney information concerning my injury and/or treatment."
        )
        self._add_signature_paragraph(
            document,
            "Veteran / Patient Signature (Parent/Guardian if applicable)",
            payload.get("patient_signature_name") or payload.get("patient_name") or "Pending",
            "patient_signature_name",
            payload,
        )
        document.add_paragraph(f"Date: {payload.get('patient_signature_date') or 'Pending'}")
        document.save(path)

    def _write_virtual_consent_doc_v2(self, path, payload):
        document = Document()
        section = document.sections[0]
        section.top_margin = Pt(48)
        section.bottom_margin = Pt(48)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)

        normal_style = document.styles["Normal"]
        normal_style.font.name = "Calibri"
        normal_style.font.size = Pt(11)

        title = document.add_paragraph()
        title.alignment = 1
        title_run = title.add_run(payload.get("packet_title") or payload.get("consent_form_title") or "TELEHEALTH VIRTUAL CONSENT FORM")
        title_run.bold = True
        title_run.font.size = Pt(16)
        title.paragraph_format.space_after = Pt(12)

        def fill_text(value, blank="_____________"):
            text = str(value or "").strip()
            return text or blank

        def add_line(text, bold=False, space_after=Pt(2)):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = space_after
            run = paragraph.add_run(text)
            run.bold = bold
            return paragraph

        def add_tabbed_line(parts, tab_stops, bold=False, space_after=Pt(2)):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = space_after
            for stop in tab_stops:
                paragraph.paragraph_format.tab_stops.add_tab_stop(Pt(stop), WD_TAB_ALIGNMENT.LEFT)
            run = paragraph.add_run("\t".join(parts))
            run.bold = bold
            return paragraph

        def add_two_line_row(labels, values, tab_stops):
            add_tabbed_line(labels, tab_stops, bold=True, space_after=Pt(0))
            add_tabbed_line(values, tab_stops, space_after=Pt(5))

        add_line("Demographic Information:", bold=True, space_after=Pt(4))
        add_two_line_row(
            ["Full Name (Shown on your Ins Card):", "Date of Birth:", "Phone:"],
            [
                fill_text(payload.get("patient_name"), "____________________________"),
                fill_text(payload.get("date_of_birth")),
                fill_text(payload.get("mobile_phone") or payload.get("phone")),
            ],
            [270, 405],
        )
        add_two_line_row(
            ["Street Address:", "City:", "State:", "Zip:"],
            [
                fill_text(payload.get("street_address"), "______________________"),
                fill_text(payload.get("city"), "_____"),
                fill_text(payload.get("state"), "___"),
                fill_text(payload.get("zip_code"), "_________"),
            ],
            [180, 280, 360],
        )
        add_two_line_row(
            ["Home Phone:", "Mobile Phone:", "Work Phone:"],
            [
                fill_text(payload.get("home_phone")),
                fill_text(payload.get("mobile_phone") or payload.get("phone"), "______________"),
                fill_text(payload.get("work_phone"), "_____________"),
            ],
            [170, 320],
        )
        add_two_line_row(
            ["Email Address:", "SSN:", "DL:", "State:"],
            [
                fill_text(payload.get("email_address"), "____________________"),
                fill_text(payload.get("ssn"), "______________"),
                fill_text(payload.get("drivers_license"), "_____________"),
                fill_text(payload.get("drivers_license_state"), "__________"),
            ],
            [180, 270, 360],
        )

        appointment = str(payload.get("appointment_confirmation_method") or "").strip().lower()
        add_line(
            "Appointment Confirmation Method:\t"
            + export_checkbox_marker(appointment == "phone")
            + " Phone\t"
            + export_checkbox_marker(appointment == "text")
            + " Text\t"
            + export_checkbox_marker(appointment == "email")
            + " Email",
            bold=True,
            space_after=Pt(6),
        )

        def yes_no_line(label, value):
            selected = str(value or "").strip().lower()
            add_line(
                f"{label}\t"
                + export_checkbox_marker(selected == "yes")
                + " Yes\t"
                + export_checkbox_marker(selected == "no")
                + " No",
                bold=True,
                space_after=Pt(4),
            )

        yes_no_line("Have you filed for Disability?", payload.get("filed_for_disability"))
        yes_no_line("Is your condition work related?", payload.get("condition_work_related"))
        yes_no_line("Is your condition due to an accident?", payload.get("condition_due_to_accident"))
        add_line("Please explain any responses you answered 'Yes' to:", bold=True, space_after=Pt(0))
        add_line(fill_text(payload.get("yes_response_explanation"), "________________________________________________________"), space_after=Pt(6))
        yes_no_line("Do you have an Attorney?", payload.get("has_attorney"))
        add_two_line_row(
            ["Attorney Full Name:", "Attorney Phone Number:"],
            [
                fill_text(payload.get("attorney_name"), "________________________"),
                fill_text(payload.get("attorney_phone"), "______________________"),
            ],
            [240],
        )
        add_two_line_row(
            ["Emergency Contact:", "Relationship:", "Phone:"],
            [
                fill_text(payload.get("emergency_contact_name"), "__________________"),
                fill_text(payload.get("emergency_contact_relationship")),
                fill_text(payload.get("emergency_contact_phone"), "_______________"),
            ],
            [190, 310],
        )

        add_line(
            "Insurance Information - If your condition is due to a work injury or other Personal Injury/Accident, there is a separate form you will need to complete. See accident information form.",
            bold=True,
            space_after=Pt(6),
        )
        add_two_line_row(
            ["Primary Insurance Carrier:", "ID Number:", "Insurance Phone:"],
            [
                fill_text(payload.get("primary_insurance_carrier"), "______________________"),
                fill_text(payload.get("primary_insurance_id"), "____________"),
                fill_text(payload.get("primary_insurance_phone"), "________________"),
            ],
            [250, 365],
        )
        add_two_line_row(
            ["Secondary Insurance Carrier:", "ID Number:", "Insurance Phone:"],
            [
                fill_text(payload.get("secondary_insurance_carrier"), "_______________________"),
                fill_text(payload.get("secondary_insurance_id"), "___________"),
                fill_text(payload.get("secondary_insurance_phone"), "_______________"),
            ],
            [250, 365],
        )
        add_two_line_row(
            ["PCP/PCM:", "Phone:", "Fax:"],
            [
                fill_text(payload.get("pcp_pcm_name"), "_______________________"),
                fill_text(payload.get("pcp_pcm_phone"), "___________"),
                fill_text(payload.get("pcp_pcm_fax"), "_______________"),
            ],
            [170, 300],
        )

        document.add_page_break()
        add_line("CONSENT FOR MEDICAL CARE AND TREATMENT", bold=True, space_after=Pt(8))
        add_line(
            "I, "
            + fill_text(payload.get("patient_name"), "________________")
            + " hereby agree and give my consent for "
            + fill_text(payload.get("consent_provider_name"), "_______________")
            + " to furnish medical care and treatment considered necessary and proper in evaluating or treating my physical condition. "
            + fill_text(payload.get("consent_initials"), "______")
            + " (initial)",
            space_after=Pt(8),
        )
        add_line("FOR MINORS ONLY CONSENT FOR CARE:", bold=True, space_after=Pt(4))
        add_line(
            "As parent and/or legal guardian, I authorize "
            + fill_text(payload.get("minor_doctor_name"), "_________________")
            + " (Doctors Name) to treat the minor patient named in the attached form while I am not present. "
            + fill_text(payload.get("minor_consent_initials"), "_____")
            + " (Parent/Guardian initial)",
            space_after=Pt(10),
        )
        add_line("TELEHEALTH CONSENT", bold=True, space_after=Pt(6))

        telehealth_lines = [
            "I understand that through an interactive video connection, telehealth might be used to perform my consultation and that we may use telehealth to connect while working together.",
            "I understand the benefits and risks involved with telehealth technology:",
            "I do not need to travel to the consultation location.",
            "I have access to a specialist through this consultation.",
            "There may be interruptions, unauthorized access and technical difficulties and my health care provider(s) or myself can discontinue the telemedicine consult/visit if it is felt that the videoconferencing connections are not adequate for the situation.",
            "My healthcare information may be shared with other individuals for scheduling and billing purposes.",
            "I also understand that other individuals may need to use the telehealth platform and that reasonable steps will be taken to maintain confidentiality of the information obtained.",
            "I understand the initial virtual meet and greet discussion is primarily intended to determine any and all treatment I might pursue, in addition to the best approach for scheduling these treatments. The virtual discussion will be followed by individual cost estimates (if applicable).",
            "I have read this document in its entirety and understand the risks and benefits of telehealth consultation/post op visits and have had my questions explained. I hereby consent to participate in telehealth sessions under the conditions described in this document.",
        ]
        for line in telehealth_lines:
            add_line(line, space_after=Pt(4))

        add_line(
            "By signing below, I agree that all the above information is correct, and that I authorize "
            + fill_text(payload.get("service_authorization_name"), "_____________")
            + " to provide me with medical services and to furnish my physician, insurance company or attorney information concerning my injury and/or treatment.",
            space_after=Pt(8),
        )
        self._add_signature_paragraph(
            document,
            "Veteran / Patient Signature (Parent/Guardian if applicable)",
            payload.get("patient_signature_name") or payload.get("patient_name"),
            "patient_signature_name",
            payload,
        )
        add_line(f"Date: {fill_text(payload.get('patient_signature_date'), '___________')}", space_after=Pt(2))
        document.save(path)

    def _write_key_value_table(self, document, rows):
        table = document.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        for label, value in rows:
            cells = table.add_row().cells
            cells[0].text = str(label or "")
            cells[1].text = str(value or "Pending")

    def _add_signature_paragraph(self, document, label, text_value, image_field_name, payload, left_indent=None):
        image_bytes = _signature_image_bytes(payload, image_field_name)
        paragraph = document.add_paragraph()
        if left_indent is not None:
            paragraph.paragraph_format.left_indent = left_indent
        paragraph.paragraph_format.space_after = Pt(2 if image_bytes else 4)
        label_run = paragraph.add_run(f"{label}: ")
        label_run.bold = True
        if image_bytes:
            signature_paragraph = document.add_paragraph()
            if left_indent is not None:
                signature_paragraph.paragraph_format.left_indent = left_indent + Pt(18)
            signature_paragraph.paragraph_format.space_after = Pt(4)
            run = signature_paragraph.add_run()
            run.add_picture(io.BytesIO(image_bytes), width=Inches(2.15))
        else:
            value_text = str(text_value or "").strip() or "_____________"
            value_run = paragraph.add_run(value_text)
            if str(text_value or "").strip():
                value_run.underline = True

    def _write_va_10172_summary_doc(self, path, payload):
        document = Document()
        title = document.add_paragraph()
        title.alignment = 1
        title_run = title.add_run(payload.get("packet_title") or "VA Form 10-10172")
        title_run.bold = True
        title_run.font.size = Pt(18)

        document.add_paragraph(
            "This Word export is a drafting summary for the VA Form 10-10172 profile. "
            "Use the PDF export to generate the actual filled VA form."
        )

        rows = [
            ("Veteran Legal Full Name", payload.get("patient_name") or "Pending"),
            ("Date of Birth", payload.get("date_of_birth") or "Pending"),
            ("VA Authorization Number", payload.get("authorization_number") or "Pending"),
            ("VA Facility & Address", payload.get("va10172_va_facility_address") or "Pending"),
            ("Ordering Provider Office Name & Address", payload.get("va10172_ordering_provider_office_address") or "Pending"),
            ("Indian Health Services / THP Provider", payload.get("va10172_is_ihs_provider") or "Pending"),
            ("Ordering Provider Phone", payload.get("va10172_ordering_provider_phone") or "Pending"),
            ("Ordering Provider Fax", payload.get("va10172_ordering_provider_fax") or "Pending"),
            ("Ordering Provider Secure Email", payload.get("va10172_ordering_provider_secure_email") or "Pending"),
            ("Care Needed Within 48 Hours", payload.get("va10172_care_needed_within_48_hours") or "Pending"),
            ("Continuation of Care", payload.get("va10172_is_continuation_of_care") or "Pending"),
            ("Referral to Another Specialty", payload.get("va10172_referral_to_specialty") or "Pending"),
            ("Specialty", payload.get("va10172_referral_specialty_text") or "Pending"),
            ("Diagnosis Codes", payload.get("icd_codes") or "Pending"),
            ("Diagnosis Description", payload.get("va10172_diagnosis_description") or "Pending"),
            ("Requested CPT/HCPCS Code", payload.get("va10172_requested_cpt_hcpcs_code") or "Pending"),
            ("Description CPT/HCPCS Code", payload.get("va10172_description_cpt_hcpcs_code") or "Pending"),
            ("Geriatric / Extended Care", payload.get("va10172_geriatric_care_option") or "Pending"),
            ("Reason for Request", payload.get("va10172_reason_for_request") or "Pending"),
            ("CCN Ordering Provider Name (Printed)", payload.get("va10172_ordering_provider_name_printed") or "Pending"),
            ("CCN Ordering Provider NPI", payload.get("va10172_ordering_provider_npi") or "Pending"),
            ("CCN Ordering Provider Signature", payload.get("va10172_signature_text") or "Pending"),
            ("Today's Date", payload.get("va10172_today_date") or "Pending"),
            ("Template PDF Path", payload.get("va10172_template_path") or "Pending"),
        ]
        self._write_key_value_table(document, rows)
        document.save(path)

    def _write_submission_cover_doc(self, path, payload):
        document = Document()
        section = document.sections[0]
        section.top_margin = Pt(54)
        section.bottom_margin = Pt(54)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)

        normal_style = document.styles["Normal"]
        normal_style.font.name = "Calibri"
        normal_style.font.size = Pt(12)

        title = document.add_paragraph()
        title.alignment = 1
        title_run = title.add_run(payload.get("packet_title") or payload.get("submission_cover_title") or "VA Submission Cover Sheet")
        title_run.bold = True
        title_run.font.size = Pt(18)
        title.paragraph_format.space_after = Pt(18)

        def add_fill_line(label, value):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.add_run(f"{label}: ").bold = True
            value_text = str(value or "").strip()
            value_run = paragraph.add_run(value_text or "_____________")
            if value_text:
                value_run.underline = True

        add_fill_line("Patient Name", payload.get("patient_name"))
        add_fill_line("DOB", payload.get("date_of_birth"))
        add_fill_line("VA Facility", payload.get("facility"))
        add_fill_line("Ordering Provider", payload.get("ordering_doctor") or payload.get("provider"))
        add_fill_line("Date of Submission", payload.get("submission_date"))
        add_fill_line("Primary Diagnosis Code", payload.get("primary_diagnosis_code") or payload.get("icd_codes"))

        checklist_heading = document.add_paragraph()
        checklist_heading.paragraph_format.space_before = Pt(12)
        checklist_heading.paragraph_format.space_after = Pt(6)
        checklist_heading.add_run("Documents Included (check all):").bold = True
        checklist_rows = [
            ("Virtual Consent Form completed and signed", payload.get("included_virtual_consent_form")),
            ("VA Form 10-10172 completed and signed", payload.get("included_va_form_10_10172")),
            ("SEOC request signed by CCN ordering provider", payload.get("included_seoc_request")),
            ("Consultation & Treatment Request completed", payload.get("included_consult_request")),
            ("Letter of Medical Necessity completed", payload.get("included_lomn")),
            ("Clinical Notes included", payload.get("included_clinical_notes")),
            ("MRI Report included", payload.get("included_mri_report")),
        ]
        for label, checked in checklist_rows:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(2)
            marker = "[X]" if checked else "[ ]"
            paragraph.add_run(f"{marker} {label}")

        review_heading = document.add_paragraph()
        review_heading.paragraph_format.space_before = Pt(12)
        review_heading.paragraph_format.space_after = Pt(6)
        review_heading.add_run("Submitting Office Review").bold = True
        add_fill_line("Submitting Office", payload.get("submitting_office"))
        add_fill_line("Office Staff Name", payload.get("office_staff_name"))
        signature_image = _signature_image_bytes(payload, "office_staff_signature")
        if signature_image:
            self._add_signature_paragraph(document, "Program User / Office Staff Signature", payload.get("office_staff_signature"), "office_staff_signature", payload, left_indent=body_indent)
        else:
            add_fill_line("Program User / Office Staff Signature", payload.get("office_staff_signature"))
        add_fill_line("Date Reviewed", payload.get("date_reviewed"))
        document.save(path)

    def _write_submission_cover_doc_v2(self, path, payload):
        document = Document()
        section = document.sections[0]
        section.top_margin = Pt(48)
        section.bottom_margin = Pt(48)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)

        normal_style = document.styles["Normal"]
        normal_style.font.name = "Calibri"
        normal_style.font.size = Pt(11)
        body_indent = Pt(18)

        title = document.add_paragraph()
        title.alignment = 1
        title_run = title.add_run(payload.get("packet_title") or payload.get("submission_cover_title") or "VA Submission Cover Sheet")
        title_run.bold = True
        title_run.font.size = Pt(18)
        title.paragraph_format.space_after = Pt(14)

        def add_fill_line(label, value):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = body_indent
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.add_run(f"{label}: ").bold = True
            value_text = str(value or "").strip()
            value_run = paragraph.add_run(value_text or "_____________")
            if value_text:
                value_run.underline = True

        add_fill_line("Patient Name", payload.get("patient_name"))
        add_fill_line("DOB", payload.get("date_of_birth"))
        add_fill_line("VA Facility", payload.get("facility"))
        add_fill_line("Ordering Provider", payload.get("ordering_doctor") or payload.get("provider"))
        add_fill_line("Date of Submission", payload.get("submission_date"))
        add_fill_line("Primary Diagnosis Code", payload.get("primary_diagnosis_code") or payload.get("icd_codes"))

        checklist_heading = document.add_paragraph()
        checklist_heading.paragraph_format.left_indent = body_indent
        checklist_heading.paragraph_format.space_before = Pt(10)
        checklist_heading.paragraph_format.space_after = Pt(5)
        checklist_heading.add_run("Documents Included (check all):").bold = True

        checklist_rows = [
            ("Virtual Consent Form completed and signed", payload.get("included_virtual_consent_form")),
            ("VA Form 10-10172 completed and signed", payload.get("included_va_form_10_10172")),
            ("SEOC request signed by CCN ordering provider", payload.get("included_seoc_request")),
            ("Consultation & Treatment Request completed", payload.get("included_consult_request")),
            ("Letter of Medical Necessity completed", payload.get("included_lomn")),
            ("Clinical Notes included", payload.get("included_clinical_notes")),
            ("MRI Report included", payload.get("included_mri_report")),
        ]
        for label, checked in checklist_rows:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = body_indent
            paragraph.paragraph_format.space_after = Pt(1)
            marker = export_checkbox_marker(checked)
            paragraph.add_run(f"{marker} {label}")

        spacer = document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)
        add_fill_line("Submitting Office", payload.get("submitting_office"))
        add_fill_line("Office Staff Name", payload.get("office_staff_name"))
        signature_image = _signature_image_bytes(payload, "office_staff_signature")
        if signature_image:
            self._add_signature_paragraph(document, "Program User / Office Staff Signature", payload.get("office_staff_signature"), "office_staff_signature", payload)
        else:
            add_fill_line("Program User / Office Staff Signature", payload.get("office_staff_signature"))
        add_fill_line("Date Reviewed", payload.get("date_reviewed"))
        document.save(path)

    def _write_seoc_request_doc(self, path, payload):
        document = Document()
        document.add_paragraph(payload.get("seoc_request_date") or "Pending")
        document.add_paragraph("Department of Veterans Affairs\nCommunity Care Office\n" + (payload.get("va_medical_center_name") or "Pending"))

        heading = document.add_paragraph()
        heading_run = heading.add_run("RE: Single Episode of Care (SEOC) Request")
        heading_run.bold = True
        document.add_paragraph(f"Veteran Name: {payload.get('patient_name') or 'Pending'}")
        document.add_paragraph(f"DOB: {payload.get('date_of_birth') or 'Pending'}")
        document.add_paragraph(f"Last Four SSN: {payload.get('last_four_ssn') or 'Pending'}")

        document.add_paragraph(
            "This request is for authorization of a defined, time-limited Single Episode of Care (SEOC) for treatment of lumbar disc pathology."
        )
        document.add_heading("Episode Diagnosis", level=1)
        diagnosis_line = payload.get("episode_diagnosis") or "Pending"
        if payload.get("episode_icd_code"):
            diagnosis_line = f"{diagnosis_line} – {payload.get('episode_icd_code')}"
        document.add_paragraph(diagnosis_line)

        document.add_heading("Scope of Requested Episode", level=1)
        document.add_paragraph(
            "This SEOC includes only the following services directly related to treatment of the above diagnosis:"
        )
        scope_items = [
            (payload.get("seoc_include_preprocedure_eval"), "Pre-procedure evaluation and procedural planning"),
            (payload.get("seoc_include_annulargram"), "Diagnostic Annulargram"),
            (payload.get("seoc_include_fibrin_injection"), "Inter Annular Fibrin Injections if indicated"),
            (payload.get("seoc_include_follow_up"), "Standard post-procedure follow-up visit(s) within the global postoperative period"),
        ]
        for enabled, label in scope_items:
            if enabled:
                document.add_paragraph(label, style="List Bullet")
        document.add_paragraph(
            payload.get("seoc_scope_text")
            or "No additional pain management services, long-term medication management, or unrelated spine care are requested under this episode."
        )

        document.add_heading("Estimated Duration of Episode", level=1)
        document.add_paragraph(
            "The episode is expected to begin upon authorization and conclude within approximately "
            + (payload.get("estimated_duration_text") or "Pending")
        )

        document.add_heading("Clinical Objective", level=1)
        objective_lines = [line.strip() for line in str(payload.get("clinical_objectives") or "").splitlines() if line.strip()]
        for line in objective_lines or ["Pending clinical objective"]:
            document.add_paragraph(line, style="List Bullet")

        document.add_heading("Continuity of Care", level=1)
        document.add_paragraph(
            payload.get("seoc_continuity_text")
            or "Upon completion of the episode, a summary of treatment rendered and clinical outcome will be forwarded to the referring VA provider. Any need for further care beyond this defined episode will require separate evaluation and authorization."
        )
        document.add_paragraph(
            "This request represents a discrete, procedure-based intervention and meets criteria for authorization under the Single Episode of Care (SEOC) framework."
        )

        document.add_paragraph("Sincerely,")
        document.add_paragraph(
            "\n".join(
                [
                    f"{payload.get('provider') or payload.get('ordering_doctor') or 'Pending'}"
                    + (f", {payload.get('provider_credentials')}" if payload.get("provider_credentials") else ""),
                    payload.get("provider_specialty") or "Pending",
                    f"NPI: {payload.get('provider_npi') or 'Pending'}",
                    payload.get("practice_name") or "Pending",
                    payload.get("provider_phone") or "Pending",
                    payload.get("provider_fax") or "Pending",
                ]
            )
        )
        document.save(path)

    def _write_lomn_doc(self, path, payload):
        document = Document()

        title = document.add_paragraph()
        title.alignment = 1
        title_run = title.add_run(payload.get("packet_title") or "LETTER OF MEDICAL NECESSITY")
        title_run.bold = True
        title_run.font.size = Pt(18)

        subtitle = document.add_paragraph()
        subtitle.alignment = 1
        subtitle_run = subtitle.add_run("(Pain Management - Interventional Spine)")
        subtitle_run.bold = True

        document.add_paragraph(payload.get("lmn_request_date") or "Pending")
        document.add_paragraph(
            "Department of Veterans Affairs\nCommunity Care Office\n"
            + (payload.get("va_medical_center_name") or "Pending")
            + ("\n" + (payload.get("facility") or "") if payload.get("facility") else "")
        )

        heading = document.add_paragraph()
        heading_run = heading.add_run("RE: Letter of Medical Necessity")
        heading_run.bold = True
        document.add_paragraph(f"Veteran Name: {payload.get('patient_name') or 'Pending'}")
        document.add_paragraph(f"DOB: {payload.get('date_of_birth') or 'Pending'}")
        document.add_paragraph(f"Last Four SSN: {payload.get('last_four_ssn') or 'Pending'}")
        document.add_paragraph(f"VA Claim Number: {payload.get('lmn_va_claim_number') or 'If Applicable'}")

        document.add_paragraph("To Whom It May Concern:")
        document.add_paragraph(
            "I am writing to formally document the medical necessity of interventional spine treatment for the above-named Veteran."
        )

        document.add_heading("Diagnosis", level=1)
        document.add_paragraph(f"Primary: {payload.get('lmn_primary_diagnosis') or 'Pending'}")
        document.add_paragraph(f"Secondary: {payload.get('lmn_secondary_diagnosis') or 'Pending'}")

        document.add_heading("Clinical Summary", level=1)
        document.add_paragraph(payload.get("lmn_clinical_summary") or "Pending")
        document.add_paragraph(f"MRI Date: {payload.get('lmn_mri_date') or 'Pending'}")
        document.add_paragraph(f"MRI Findings: {payload.get('lmn_mri_findings') or 'Pending'}")
        document.add_paragraph("The Veteran has completed extensive conservative management, including:")

        conservative_items = [
            (payload.get("lmn_include_physical_therapy"), "Structured physical therapy"),
            (payload.get("lmn_include_nsaids"), "NSAIDs and non-opioid analgesics"),
            (payload.get("lmn_include_activity_modification"), "Activity modification"),
            (payload.get("lmn_include_home_exercise"), "Home exercise program"),
            (payload.get("lmn_include_epidural_steroid_injections"), "Epidural Steroid injections (if applicable)"),
        ]
        for enabled, label in conservative_items:
            if enabled:
                document.add_paragraph(label, style="List Bullet")
        document.add_paragraph(
            "Despite appropriate and guideline-based conservative treatment for greater than "
            + (payload.get("lmn_conservative_duration") or "Pending")
            + ", the Veteran continues to experience persistent pain and functional limitation."
        )

        document.add_heading("Medical Necessity", level=1)
        document.add_paragraph(payload.get("lmn_medical_necessity_statement") or "Pending")
        document.add_paragraph("This intervention is indicated to:")

        indication_items = [
            (payload.get("lmn_indication_reduce_pain"), "Reduce pain severity"),
            (payload.get("lmn_indication_improve_function"), "Improve functional capacity"),
            (payload.get("lmn_indication_prevent_degeneration"), "Prevent further disc degeneration"),
            (payload.get("lmn_indication_reduce_opioid_reliance"), "Decrease reliance on long-term opioid therapy"),
            (payload.get("lmn_indication_prevent_surgery"), "Potentially prevent need for more invasive surgical intervention"),
        ]
        for enabled, label in indication_items:
            if enabled:
                document.add_paragraph(label, style="List Bullet")

        document.add_paragraph(payload.get("lmn_risk_statement") or "Pending")
        document.add_paragraph(payload.get("lmn_reasonable_necessary_statement") or "Pending")
        document.add_paragraph(payload.get("lmn_contact_statement") or "Pending")

        document.add_paragraph("Sincerely,")
        document.add_paragraph(
            "\n".join(
                [
                    f"{payload.get('provider') or payload.get('ordering_doctor') or 'Pending'}"
                    + (f", {payload.get('provider_credentials')}" if payload.get("provider_credentials") else ""),
                    payload.get("provider_specialty") or "Pending",
                    f"NPI: {payload.get('provider_npi') or 'Pending'}",
                    payload.get("practice_name") or "Pending",
                    payload.get("provider_phone") or "Pending",
                    payload.get("provider_fax") or "Pending",
                ]
            )
        )
        document.save(path)

    def _write_consult_request_doc(self, path, payload):
        document = Document()

        title = document.add_paragraph()
        title.alignment = 1
        title_run = title.add_run(payload.get("packet_title") or "CONSULTATION AND TREATMENT REQUEST")
        title_run.bold = True
        title_run.font.size = Pt(18)

        subtitle = document.add_paragraph()
        subtitle.alignment = 1
        subtitle_run = subtitle.add_run("(Pain Management - Interventional Spine)")
        subtitle_run.bold = True

        document.add_paragraph(payload.get("consult_request_date") or "Pending")
        document.add_paragraph(
            "Department of Veterans Affairs\nCommunity Care Office\n"
            + (payload.get("va_medical_center_name") or "Pending")
            + ("\n" + (payload.get("facility") or "") if payload.get("facility") else "")
        )

        heading = document.add_paragraph()
        heading_run = heading.add_run("RE: Consultation and Treatment Request")
        heading_run.bold = True
        document.add_paragraph(f"Veteran Name: {payload.get('patient_name') or 'Pending'}")
        document.add_paragraph(f"DOB: {payload.get('date_of_birth') or 'Pending'}")
        document.add_paragraph(f"Last Four SSN: {payload.get('last_four_ssn') or 'Pending'}")
        document.add_paragraph(f"VA Claim Number: {payload.get('consult_va_claim_number') or 'If Applicable'}")
        document.add_paragraph(f"Referring VA Provider: {payload.get('consult_referring_va_provider') or 'If Applicable'}")

        document.add_heading("Reason for Consultation", level=1)
        document.add_paragraph(payload.get("consult_reason_text") or "Pending")

        document.add_heading("Diagnoses", level=1)
        document.add_paragraph(f"Primary: {payload.get('consult_primary_diagnosis') or 'Pending'}")
        document.add_paragraph(f"Secondary: {payload.get('consult_secondary_diagnosis') or 'Pending'}")

        document.add_heading("Clinical Summary", level=1)
        document.add_paragraph(
            "The Veteran presents with persistent, function-limiting low back pain refractory to conservative management. Symptoms include:"
        )
        symptom_items = [
            (payload.get("consult_symptom_axial_pain"), "Chronic axial lumbar pain"),
            (payload.get("consult_symptom_activity_exacerbation"), "Activity-related exacerbation (sitting, standing, bending)"),
            (payload.get("consult_symptom_reduced_tolerance"), "Reduced tolerance for prolonged positioning"),
            (payload.get("consult_symptom_functional_impairment"), "Functional impairment affecting occupational and daily activities"),
        ]
        for enabled, label in symptom_items:
            if enabled:
                document.add_paragraph(label, style="List Bullet")

        document.add_paragraph(f"MRI Date: {payload.get('consult_mri_date') or 'Pending'}")
        document.add_paragraph(f"MRI Findings: {payload.get('consult_mri_findings') or 'Pending'}")
        document.add_paragraph("These findings are reviewed alongside the clinical presentation and prior treatment history.")
        document.add_paragraph("Conservative management has included:")
        conservative_items = [
            (payload.get("consult_include_physical_therapy"), "Physical therapy"),
            (payload.get("consult_include_nsaids"), "NSAIDs and non-opioid analgesics"),
            (payload.get("consult_include_activity_modification"), "Activity modification"),
            (payload.get("consult_include_home_exercise"), "Home exercise program"),
            (payload.get("consult_include_interventional_history"), "Epidural steroid injections and/or other interventional procedures (if applicable)"),
        ]
        for enabled, label in conservative_items:
            if enabled:
                document.add_paragraph(label, style="List Bullet")
        document.add_paragraph(
            "Despite appropriate treatment for greater than "
            + (payload.get("consult_conservative_duration") or "Pending")
            + ", the Veteran continues to experience persistent pain and impaired function."
        )

        document.add_heading("Requested Services", level=1)
        document.add_paragraph("Authorization is requested for:")
        service_items = [
            (payload.get("consult_include_pain_management_consultation"), "Comprehensive pain management consultation"),
            (payload.get("consult_include_procedural_planning"), "Diagnostic confirmation and procedural planning"),
            (payload.get("consult_include_annulargram"), "Diagnostic Annulargram"),
            (
                payload.get("consult_include_fibrin_injection"),
                "Intraannular Fibrin injection"
                + (f" at {payload.get('consult_fibrin_levels')}" if payload.get("consult_fibrin_levels") else "")
                + " if indicated",
            ),
            (payload.get("consult_include_follow_up"), "Routine post-procedure follow-up visit(s)"),
        ]
        for enabled, label in service_items:
            if enabled:
                document.add_paragraph(label, style="List Bullet")
        document.add_paragraph(payload.get("consult_scope_exclusion_text") or "Pending")

        document.add_heading("Medical Rationale", level=1)
        document.add_paragraph(payload.get("consult_medical_rationale_text") or "Pending")
        document.add_paragraph("Clinical goals include:")
        goal_items = [
            (payload.get("consult_goal_pain_reduction"), "Pain reduction"),
            (payload.get("consult_goal_functional_improvement"), "Functional improvement"),
            (payload.get("consult_goal_reduce_analgesics"), "Decreased reliance on chronic analgesic therapy"),
            (payload.get("consult_goal_prevent_surgery"), "Prevention of progression requiring more invasive surgical intervention"),
        ]
        for enabled, label in goal_items:
            if enabled:
                document.add_paragraph(label, style="List Bullet")
        document.add_paragraph(payload.get("consult_risk_without_treatment") or "Pending")

        document.add_heading("Duration and Scope of Care", level=1)
        document.add_paragraph(payload.get("consult_duration_scope_text") or "Pending")
        document.add_paragraph(payload.get("consult_contact_statement") or "Pending")

        document.add_paragraph("Sincerely,")
        document.add_paragraph(
            "\n".join(
                [
                    f"{payload.get('provider') or payload.get('ordering_doctor') or 'Pending'}"
                    + (f", {payload.get('provider_credentials')}" if payload.get("provider_credentials") else ""),
                    payload.get("provider_specialty") or "Pending",
                    f"NPI: {payload.get('provider_npi') or 'Pending'}",
                    payload.get("practice_name") or "Pending",
                    payload.get("provider_address") or "Pending",
                    payload.get("provider_phone") or "Pending",
                    payload.get("provider_fax") or "Pending",
                    payload.get("provider_email") or "Pending",
                ]
            )
        )
        document.save(path)

    def _write_clinical_documentation_doc(self, path, payload):
        document = Document()

        title = document.add_paragraph()
        title.alignment = 1
        title_run = title.add_run(payload.get("packet_title") or payload.get("clinical_doc_title") or "Clinical Documentation Template")
        title_run.bold = True
        title_run.font.size = Pt(18)

        document.add_heading("I. Chief Complaint", level=1)
        document.add_paragraph(payload.get("clinical_doc_chief_complaint") or "Pending")

        document.add_heading("II. History of Present Illness", level=1)
        document.add_paragraph("Duration of Symptoms:")
        duration_items = [
            (payload.get("clinical_doc_duration_gt_3m"), "> 3 months"),
            (payload.get("clinical_doc_duration_gt_6m"), "> 6 months"),
            (payload.get("clinical_doc_duration_gt_12m"), "> 12 months"),
        ]
        for enabled, label in duration_items:
            if enabled:
                document.add_paragraph(label, style="List Bullet")
        document.add_paragraph(f"Exact duration: {payload.get('clinical_doc_exact_duration') or 'Pending'}")

        document.add_paragraph("Pain Characteristics:")
        pain_items = [
            (payload.get("clinical_doc_pain_axial"), "Axial lumbar pain"),
            (payload.get("clinical_doc_pain_discogenic"), "Discogenic pattern"),
            (payload.get("clinical_doc_pain_activity_exacerbation"), "Activity-related exacerbation"),
            (payload.get("clinical_doc_pain_sitting_intolerance"), "Sitting intolerance"),
            (payload.get("clinical_doc_pain_standing_intolerance"), "Standing intolerance"),
            (payload.get("clinical_doc_pain_bending_lifting"), "Bending/lifting provocation"),
        ]
        for enabled, label in pain_items:
            if enabled:
                document.add_paragraph(label, style="List Bullet")
        document.add_paragraph(f"Pain severity (0-10): {payload.get('clinical_doc_pain_severity') or 'Pending'}")

        document.add_heading("III. Functional Impairment", level=1)
        document.add_paragraph("Patient reports limitation in:")
        limitation_items = [
            (payload.get("clinical_doc_limit_occupational"), "Occupational duties"),
            (payload.get("clinical_doc_limit_prolonged_sitting"), "Prolonged sitting"),
            (payload.get("clinical_doc_limit_prolonged_standing"), "Prolonged standing"),
            (payload.get("clinical_doc_limit_ambulation"), "Ambulation tolerance"),
            (payload.get("clinical_doc_limit_household"), "Household activities"),
            (payload.get("clinical_doc_limit_sleep"), "Sleep disturbance"),
        ]
        for enabled, label in limitation_items:
            if enabled:
                document.add_paragraph(label, style="List Bullet")
        document.add_paragraph(
            "Describe specific functional impact: "
            + (payload.get("clinical_doc_functional_impact") or "Pending")
        )

        document.add_heading("IV. Conservative Therapy History", level=1)
        document.add_paragraph("The patient has completed and/or attempted:")
        conservative_items = [
            (payload.get("clinical_doc_conservative_pt"), "Physical therapy"),
            (payload.get("clinical_doc_conservative_home_exercise"), "Home exercise program"),
            (payload.get("clinical_doc_conservative_nsaids"), "NSAIDs"),
            (payload.get("clinical_doc_conservative_non_opioid"), "Non-opioid analgesics"),
            (payload.get("clinical_doc_conservative_activity_modification"), "Activity modification"),
            (payload.get("clinical_doc_conservative_esi"), "Epidural steroid injections"),
            (payload.get("clinical_doc_conservative_other_interventional"), "Other interventional procedures"),
        ]
        for enabled, label in conservative_items:
            if enabled:
                document.add_paragraph(label, style="List Bullet")
        document.add_paragraph(
            "Duration of conservative treatment: "
            + (payload.get("clinical_doc_conservative_duration") or "Pending")
        )
        document.add_paragraph(
            "Response to prior ESI (if applicable): "
            + (payload.get("clinical_doc_esi_response") or "Pending")
        )

        document.add_heading("V. Imaging Findings", level=1)
        document.add_paragraph(f"MRI Date: {payload.get('clinical_doc_mri_date') or 'Pending'}")
        imaging_items = [
            (payload.get("clinical_doc_imaging_annular_tear"), "Annular tear"),
            (payload.get("clinical_doc_imaging_disc_degeneration"), "Disc degeneration"),
            (payload.get("clinical_doc_imaging_disc_protrusion"), "Disc protrusion"),
            (payload.get("clinical_doc_imaging_disc_displacement"), "Disc displacement"),
        ]
        document.add_paragraph("Findings:")
        for enabled, label in imaging_items:
            if enabled:
                document.add_paragraph(label, style="List Bullet")
        document.add_paragraph(f"Affected Levels: {payload.get('clinical_doc_affected_levels') or 'Pending'}")
        document.add_paragraph("Imaging correlates with clinical presentation.")

        document.add_heading("VI. Assessment", level=1)
        document.add_paragraph(f"Primary: {payload.get('clinical_doc_primary_diagnosis') or 'Pending'}")
        document.add_paragraph(f"Secondary: {payload.get('clinical_doc_secondary_diagnosis') or 'Pending'}")
        document.add_paragraph(payload.get("clinical_doc_assessment_summary") or "Pending")

        document.add_heading("VII. Treatment Plan", level=1)
        document.add_paragraph(payload.get("clinical_doc_treatment_plan_intro") or "Pending")
        document.add_paragraph("Plan includes:")
        plan_items = [
            (payload.get("clinical_doc_plan_diagnostic_confirmation"), "Diagnostic confirmation if necessary"),
            (payload.get("clinical_doc_plan_intradiscal_intervention"), "Intradiscal annular intervention if indicated"),
            (payload.get("clinical_doc_plan_follow_up"), "Standard post-procedure follow-up"),
        ]
        for enabled, label in plan_items:
            if enabled:
                document.add_paragraph(label, style="List Bullet")
        document.add_paragraph(payload.get("clinical_doc_plan_exclusion") or "Pending")

        document.add_heading("Physician Narrative Paragraph", level=1)
        document.add_paragraph(payload.get("clinical_doc_physician_narrative") or "Pending")

        document.save(path)

    def _write_seoc_request_doc_v2(self, path, payload):
        document = Document()
        section = document.sections[0]
        section.top_margin = Pt(46)
        section.bottom_margin = Pt(46)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)

        normal_style = document.styles["Normal"]
        normal_style.font.name = "Calibri"
        normal_style.font.size = Pt(10.5)
        body_indent = Pt(18)
        bullet_indent = Pt(30)

        def fill_text(value, blank="______________"):
            text = str(value or "").strip()
            return text or blank

        def add_line(text, bold=False, space_after=Pt(4), alignment=None, left_indent=None):
            paragraph = document.add_paragraph()
            if alignment is not None:
                paragraph.alignment = alignment
            if left_indent is not None:
                paragraph.paragraph_format.left_indent = left_indent
            paragraph.paragraph_format.space_after = space_after
            run = paragraph.add_run(text)
            run.bold = bold
            return paragraph

        def add_bullets(items, fallback_text):
            selected = [label for enabled, label in items if enabled]
            if not selected:
                selected = [fallback_text]
            for label in selected:
                add_line("- " + label, space_after=Pt(1), left_indent=bullet_indent)

        diagnosis_line = fill_text(payload.get("episode_diagnosis"), "________________________________________")

        objective_lines = [line.strip() for line in str(payload.get("clinical_objectives") or "").splitlines() if line.strip()]
        if not objective_lines:
            objective_lines = ["________________________________________"]
        scope_summary = str(payload.get("seoc_scope_text") or "").strip()
        continuity_text = str(payload.get("seoc_continuity_text") or "").strip()

        add_line(fill_text(payload.get("seoc_request_date"), "[Date]"))
        add_line(
            "Department of Veterans Affairs\nCommunity Care Office\n"
            + fill_text(payload.get("va_medical_center_name"), "[VA Medical Center Name]"),
            space_after=Pt(10),
        )
        add_line("RE: Single Episode of Care (SEOC) Request", bold=True, space_after=Pt(2))
        add_line(f"Veteran Name: {fill_text(payload.get('patient_name'), '[Full Name]')}", space_after=Pt(2))
        add_line(f"DOB: {fill_text(payload.get('date_of_birth'), '[MM/DD/YYYY]')}", space_after=Pt(2))
        add_line(f"Last Four SSN: {fill_text(payload.get('last_four_ssn'), '[XXXX]')}", space_after=Pt(8))
        add_line(
            "This request is for authorization of a defined, time-limited Single Episode of Care (SEOC) related to the diagnosis or condition listed below.",
            space_after=Pt(6),
            left_indent=body_indent,
        )
        add_line("Episode Diagnosis / Condition:", bold=True, space_after=Pt(2))
        add_line(f"Diagnosis: {diagnosis_line}", space_after=Pt(2), left_indent=body_indent)
        add_line(f"ICD-10 Code: {fill_text(payload.get('episode_icd_code'), '____________')}", space_after=Pt(6), left_indent=body_indent)
        add_line("Scope of Requested Episode:", bold=True, space_after=Pt(2))
        add_line(
            "This SEOC includes only the following services directly related to treatment of the above diagnosis:",
            space_after=Pt(3),
            left_indent=body_indent,
        )
        add_bullets(
            [
                (payload.get("seoc_include_preprocedure_eval"), "Pre-procedure evaluation and procedural planning"),
                (payload.get("seoc_include_annulargram"), "Diagnostic Annulargram"),
                (payload.get("seoc_include_fibrin_injection"), "Inter Annular Fibrin Injections if indicated"),
                (payload.get("seoc_include_follow_up"), "Standard post-procedure follow-up visit(s) within the global postoperative period"),
            ],
            "Pending scope items",
        )
        add_line(
            scope_summary or "No additional pain management services, long-term medication management, or unrelated spine care are requested under this episode.",
            space_after=Pt(6),
            left_indent=body_indent,
        )
        add_line("Estimated Duration of Episode:", bold=True, space_after=Pt(2))
        add_line(
            "Expected episode duration: " + fill_text(payload.get("estimated_duration_text"), "__________________"),
            space_after=Pt(6),
            left_indent=body_indent,
        )
        add_line("Clinical Objective:", bold=True, space_after=Pt(2))
        for line in objective_lines:
            add_line("- " + line, space_after=Pt(1), left_indent=bullet_indent)
        add_line("Continuity of Care:", bold=True, space_after=Pt(2))
        add_line(
            continuity_text or "Upon completion of the episode, a summary of treatment rendered and clinical outcome will be forwarded to the referring VA provider. Any need for further care beyond this defined episode will require separate evaluation and authorization.",
            space_after=Pt(3),
            left_indent=body_indent,
        )
        add_line("Sincerely,", space_after=Pt(8))
        provider_header = fill_text(payload.get("provider") or payload.get("ordering_doctor"), "[Provider Name]")
        credentials = str(payload.get("provider_credentials") or "").strip()
        if credentials:
            provider_header += f", {credentials}"
        add_line(provider_header, space_after=Pt(2))
        add_line(
            f"{fill_text(payload.get('provider_specialty'), '[Specialty]')} | NPI Number: {fill_text(payload.get('provider_npi'), '[NPI Number]')}",
            space_after=Pt(2),
        )
        add_line(fill_text(payload.get("practice_name") or payload.get("community_facility"), "[Practice Name]"), space_after=Pt(2))
        add_line(
            "Phone: "
            + fill_text(payload.get("provider_phone"), "[Phone]")
            + "    Fax: "
            + fill_text(payload.get("provider_fax"), "[Fax]"),
            space_after=Pt(0),
        )
        document.save(path)

    def _write_lomn_doc_v2(self, path, payload):
        document = Document()
        section = document.sections[0]
        section.top_margin = Pt(46)
        section.bottom_margin = Pt(46)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)

        normal_style = document.styles["Normal"]
        normal_style.font.name = "Calibri"
        normal_style.font.size = Pt(10.5)
        body_indent = Pt(18)
        bullet_indent = Pt(30)

        def fill_text(value, blank="______________"):
            text = str(value or "").strip()
            return text or blank

        def add_line(text, bold=False, space_after=Pt(4), alignment=None, left_indent=None):
            paragraph = document.add_paragraph()
            if alignment is not None:
                paragraph.alignment = alignment
            if left_indent is not None:
                paragraph.paragraph_format.left_indent = left_indent
            paragraph.paragraph_format.space_after = space_after
            run = paragraph.add_run(text)
            run.bold = bold
            return paragraph

        def add_bullets(items, fallback_text):
            selected = [label for enabled, label in items if enabled]
            if not selected:
                selected = [fallback_text]
            for label in selected:
                add_line("- " + label, space_after=Pt(1), left_indent=bullet_indent)

        provider_header = fill_text(payload.get("provider") or payload.get("ordering_doctor"), "[Provider Name]")
        credentials = str(payload.get("provider_credentials") or "").strip()
        if credentials:
            provider_header += f", {credentials}"
        practice_name = fill_text(payload.get("practice_name") or payload.get("community_facility"), "[Practice Name]")
        provider_address = str(payload.get("provider_address") or "").strip()
        provider_phone = fill_text(payload.get("provider_phone"), "[Phone]")
        provider_fax = fill_text(payload.get("provider_fax"), "[Fax]")

        title = payload.get("packet_title") or "LETTER OF MEDICAL NECESSITY"
        add_line(title, bold=True, alignment=1, space_after=Pt(2))
        add_line("(Pain Management - Interventional Spine)", bold=True, alignment=1, space_after=Pt(6))
        add_line(practice_name, alignment=1, space_after=Pt(1))
        add_line(provider_header, alignment=1, space_after=Pt(1))
        if provider_address:
            add_line(provider_address, alignment=1, space_after=Pt(1))
        add_line(f"Phone: {provider_phone} | Fax: {provider_fax}", alignment=1, space_after=Pt(8))

        add_line(fill_text(payload.get("lmn_request_date"), "[Date]"))
        va_block = [
            "Department of Veterans Affairs",
            "Community Care Office",
            fill_text(payload.get("va_medical_center_name"), "[VA Medical Center Name]"),
        ]
        facility = str(payload.get("facility") or "").strip()
        if facility and facility != va_block[-1]:
            va_block.append(facility)
        add_line("\n".join(va_block), space_after=Pt(10))

        add_line("RE: Letter of Medical Necessity", bold=True, space_after=Pt(2))
        add_line(f"Veteran Name: {fill_text(payload.get('patient_name'), '[Full Name]')}", space_after=Pt(2))
        add_line(f"DOB: {fill_text(payload.get('date_of_birth'), '[MM/DD/YYYY]')}", space_after=Pt(2))
        add_line(f"Last Four SSN: {fill_text(payload.get('last_four_ssn'), '[XXXX]')}", space_after=Pt(2))
        add_line(f"VA Claim Number: {fill_text(payload.get('lmn_va_claim_number'), '[If Applicable]')}", space_after=Pt(8))

        add_line("To Whom It May Concern:", bold=True, space_after=Pt(4))
        add_line(
            "I am writing to formally document the medical necessity of interventional spine treatment for the above-named Veteran.",
            space_after=Pt(6),
            left_indent=body_indent,
        )
        add_line("Clinical Diagnoses:", bold=True, space_after=Pt(2))
        add_line(f"Primary: {fill_text(payload.get('lmn_primary_diagnosis'), '______________________________')}", space_after=Pt(2), left_indent=body_indent)
        add_line(f"Secondary: {fill_text(payload.get('lmn_secondary_diagnosis'), '______________________________')}", space_after=Pt(6), left_indent=body_indent)

        add_line("Clinical Basis:", bold=True, space_after=Pt(2))
        add_line(
            str(payload.get("lmn_clinical_summary") or "").strip()
            or "____________________________________________________________",
            space_after=Pt(3),
            left_indent=body_indent,
        )
        add_line("Imaging Support:", bold=True, space_after=Pt(2))
        add_line(f"MRI Date: {fill_text(payload.get('lmn_mri_date'), '________________')}", space_after=Pt(2), left_indent=body_indent)
        add_line(f"MRI Findings: {fill_text(payload.get('lmn_mri_findings'), '________________________________________')}", space_after=Pt(3), left_indent=body_indent)
        add_line("The Veteran has completed extensive conservative management, including:", space_after=Pt(3), left_indent=body_indent)
        add_bullets(
            [
                (payload.get("lmn_include_physical_therapy"), "Structured physical therapy"),
                (payload.get("lmn_include_nsaids"), "NSAIDs and non-opioid analgesics"),
                (payload.get("lmn_include_activity_modification"), "Activity modification"),
                (payload.get("lmn_include_home_exercise"), "Home exercise program"),
                (payload.get("lmn_include_epidural_steroid_injections"), "Epidural Steroid injections (if applicable)"),
            ],
            "________________________________________",
        )
        add_line(
            "Despite appropriate and guideline-based conservative treatment for greater than "
            + fill_text(payload.get("lmn_conservative_duration"), "________________")
            + ", the Veteran continues to experience persistent pain and functional limitation.",
            space_after=Pt(6),
            left_indent=body_indent,
        )

        add_line("Basis for Medical Necessity:", bold=True, space_after=Pt(2))
        add_line(
            str(payload.get("lmn_medical_necessity_statement") or "").strip()
            or "____________________________________________________________",
            space_after=Pt(3),
            left_indent=body_indent,
        )
        add_line("Requested Treatment Objectives:", bold=True, space_after=Pt(2))
        add_line("This intervention is indicated to:", space_after=Pt(3), left_indent=body_indent)
        add_bullets(
            [
                (payload.get("lmn_indication_reduce_pain"), "Reduce pain severity"),
                (payload.get("lmn_indication_improve_function"), "Improve functional capacity"),
                (payload.get("lmn_indication_prevent_degeneration"), "Prevent further disc degeneration"),
                (payload.get("lmn_indication_reduce_opioid_reliance"), "Decrease reliance on long-term opioid therapy"),
                (payload.get("lmn_indication_prevent_surgery"), "Potentially prevent need for more invasive surgical intervention"),
            ],
            "________________________________________",
        )
        add_line("Risk if Treatment Is Delayed or Denied:", bold=True, space_after=Pt(2))
        add_line(
            str(payload.get("lmn_risk_statement") or "").strip()
            or "____________________________________________________________",
            space_after=Pt(3),
            left_indent=body_indent,
        )
        add_line("Reasonable and Necessary Determination:", bold=True, space_after=Pt(2))
        add_line(
            str(payload.get("lmn_reasonable_necessary_statement") or "").strip()
            or "____________________________________________________________",
            space_after=Pt(3),
            left_indent=body_indent,
        )
        add_line("Provider Contact Statement:", bold=True, space_after=Pt(2))
        add_line(
            str(payload.get("lmn_contact_statement") or "").strip()
            or "____________________________________________________________",
            space_after=Pt(10),
            left_indent=body_indent,
        )
        add_line("Sincerely,", space_after=Pt(8))
        add_line(provider_header, space_after=Pt(2))
        add_line(
            f"{fill_text(payload.get('provider_specialty'), '[Specialty - Pain Management / Interventional Spine]')} | NPI Number: {fill_text(payload.get('provider_npi'), '[NPI Number]')}",
            space_after=Pt(2),
        )
        add_line(practice_name, space_after=Pt(2))
        add_line(f"Phone: {provider_phone} | Fax: {provider_fax}", space_after=Pt(0))
        document.save(path)

    def _write_consult_request_doc_v2(self, path, payload):
        document = Document()
        section = document.sections[0]
        section.top_margin = Pt(46)
        section.bottom_margin = Pt(46)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)

        normal_style = document.styles["Normal"]
        normal_style.font.name = "Calibri"
        normal_style.font.size = Pt(10.5)
        body_indent = Pt(18)
        bullet_indent = Pt(30)

        def fill_text(value, blank="______________"):
            text = str(value or "").strip()
            return text or blank

        def add_line(text, bold=False, space_after=Pt(4), alignment=None, left_indent=None):
            paragraph = document.add_paragraph()
            if alignment is not None:
                paragraph.alignment = alignment
            if left_indent is not None:
                paragraph.paragraph_format.left_indent = left_indent
            paragraph.paragraph_format.space_after = space_after
            run = paragraph.add_run(text)
            run.bold = bold
            return paragraph

        def add_bullets(items, fallback_text):
            selected = [label for enabled, label in items if enabled]
            if not selected:
                selected = [fallback_text]
            for label in selected:
                add_line("- " + label, space_after=Pt(1), left_indent=bullet_indent)

        title = payload.get("packet_title") or "CONSULTATION AND TREATMENT REQUEST"
        provider_header = fill_text(payload.get("provider") or payload.get("ordering_doctor"), "[Provider Name]")
        credentials = str(payload.get("provider_credentials") or "").strip()
        if credentials:
            provider_header += f", {credentials}"
        practice_name = fill_text(payload.get("practice_name") or payload.get("community_facility"), "[Practice Name]")
        provider_address = fill_text(payload.get("provider_address"), "[Address]")
        provider_phone = fill_text(payload.get("provider_phone"), "[Phone]")
        provider_fax = fill_text(payload.get("provider_fax"), "[Fax]")
        provider_email = fill_text(payload.get("provider_email"), "[Email]")

        add_line(title, bold=True, alignment=1, space_after=Pt(2))
        add_line("(Pain Management - Interventional Spine)", bold=True, alignment=1, space_after=Pt(6))
        add_line(practice_name, alignment=1, space_after=Pt(1))
        add_line(provider_header, alignment=1, space_after=Pt(1))
        add_line(provider_address, alignment=1, space_after=Pt(1))
        add_line(f"Phone: {provider_phone} | Fax: {provider_fax}", alignment=1, space_after=Pt(1))
        add_line(provider_email, alignment=1, space_after=Pt(6))

        add_line(fill_text(payload.get("consult_request_date"), "[Date]"))
        va_block = [
            "Department of Veterans Affairs",
            "Community Care Office",
            fill_text(payload.get("va_medical_center_name"), "[VA Medical Center Name]"),
        ]
        facility = str(payload.get("facility") or "").strip()
        if facility and facility != va_block[-1]:
            va_block.append(facility)
        add_line("\n".join(va_block), space_after=Pt(10))

        add_line("RE: Consultation and Treatment Request", bold=True, space_after=Pt(2))
        add_line(f"Veteran Name: {fill_text(payload.get('patient_name'), '[Full Name]')}", space_after=Pt(2))
        add_line(f"DOB: {fill_text(payload.get('date_of_birth'), '[MM/DD/YYYY]')}", space_after=Pt(2))
        add_line(f"Last Four SSN: {fill_text(payload.get('last_four_ssn'), '[XXXX]')}", space_after=Pt(2))
        add_line(f"VA Claim Number: {fill_text(payload.get('consult_va_claim_number'), '[If Applicable]')}", space_after=Pt(2))
        add_line(f"Referring VA Provider: {fill_text(payload.get('consult_referring_va_provider'), '[Name, if applicable]')}", space_after=Pt(8))

        add_line("Reason for Consultation", bold=True, space_after=Pt(2))
        add_line(
            str(payload.get("consult_reason_text") or "").strip()
            or "____________________________________________________________",
            space_after=Pt(6),
            left_indent=body_indent,
        )
        add_line("Diagnoses", bold=True, space_after=Pt(2))
        add_line(f"Primary: {fill_text(payload.get('consult_primary_diagnosis'), '______________________________')}", space_after=Pt(2), left_indent=body_indent)
        add_line(f"Secondary: {fill_text(payload.get('consult_secondary_diagnosis'), '______________________________')}", space_after=Pt(6), left_indent=body_indent)

        add_line("Clinical Summary", bold=True, space_after=Pt(2))
        add_line(
            "Symptoms and functional concerns documented for this consultation include:",
            space_after=Pt(3),
            left_indent=body_indent,
        )
        add_bullets(
            [
                (payload.get("consult_symptom_axial_pain"), "Chronic axial lumbar pain"),
                (payload.get("consult_symptom_activity_exacerbation"), "Activity-related exacerbation (sitting, standing, bending)"),
                (payload.get("consult_symptom_reduced_tolerance"), "Reduced tolerance for prolonged positioning"),
                (payload.get("consult_symptom_functional_impairment"), "Functional impairment affecting occupational and daily activities"),
            ],
            "________________________________________",
        )
        add_line("Imaging Review", bold=True, space_after=Pt(2))
        add_line(f"MRI Date: {fill_text(payload.get('consult_mri_date'), '________________')}", space_after=Pt(2), left_indent=body_indent)
        add_line(f"MRI Findings: {fill_text(payload.get('consult_mri_findings'), '________________________________________')}", space_after=Pt(3), left_indent=body_indent)
        add_line("Conservative management has included:", space_after=Pt(3), left_indent=body_indent)
        add_bullets(
            [
                (payload.get("consult_include_physical_therapy"), "Physical therapy"),
                (payload.get("consult_include_nsaids"), "NSAIDs and non-opioid analgesics"),
                (payload.get("consult_include_activity_modification"), "Activity modification"),
                (payload.get("consult_include_home_exercise"), "Home exercise program"),
                (payload.get("consult_include_interventional_history"), "Epidural steroid injections and/or other interventional procedures (if applicable)"),
            ],
            "________________________________________",
        )
        add_line(
            "Despite appropriate treatment for greater than "
            + fill_text(payload.get("consult_conservative_duration"), "________________")
            + ", the Veteran continues to experience persistent pain and impaired function.",
            space_after=Pt(6),
            left_indent=body_indent,
        )

        add_line("Requested Services", bold=True, space_after=Pt(2))
        add_line("Authorization is requested for:", space_after=Pt(3), left_indent=body_indent)
        fibrin_line = "Intraannular Fibrin injection"
        if str(payload.get("consult_fibrin_levels") or "").strip():
            fibrin_line += f" at {payload.get('consult_fibrin_levels').strip()}"
        fibrin_line += " if indicated"
        add_bullets(
            [
                (payload.get("consult_include_pain_management_consultation"), "Comprehensive pain management consultation"),
                (payload.get("consult_include_procedural_planning"), "Diagnostic confirmation and procedural planning"),
                (payload.get("consult_include_annulargram"), "Diagnostic Annulargram"),
                (payload.get("consult_include_fibrin_injection"), fibrin_line),
                (payload.get("consult_include_follow_up"), "Routine post-procedure follow-up visit(s)"),
            ],
            "________________________________________",
        )
        add_line("Scope of Request", bold=True, space_after=Pt(2))
        add_line(
            str(payload.get("consult_scope_exclusion_text") or "").strip()
            or "____________________________________________________________",
            space_after=Pt(6),
            left_indent=body_indent,
        )

        add_line("Medical Rationale", bold=True, space_after=Pt(2))
        add_line(
            str(payload.get("consult_medical_rationale_text") or "").strip()
            or "____________________________________________________________",
            space_after=Pt(3),
            left_indent=body_indent,
        )
        add_line("Clinical goals include:", space_after=Pt(3), left_indent=body_indent)
        add_bullets(
            [
                (payload.get("consult_goal_pain_reduction"), "Pain reduction"),
                (payload.get("consult_goal_functional_improvement"), "Functional improvement"),
                (payload.get("consult_goal_reduce_analgesics"), "Decreased reliance on chronic analgesic therapy"),
                (payload.get("consult_goal_prevent_surgery"), "Prevention of progression requiring more invasive surgical intervention"),
            ],
            "________________________________________",
        )
        add_line("Risk Without Requested Treatment", bold=True, space_after=Pt(2))
        add_line(
            str(payload.get("consult_risk_without_treatment") or "").strip()
            or "____________________________________________________________",
            space_after=Pt(6),
            left_indent=body_indent,
        )

        add_line("Duration and Scope of Care", bold=True, space_after=Pt(2))
        add_line(
            str(payload.get("consult_duration_scope_text") or "").strip()
            or "____________________________________________________________",
            space_after=Pt(3),
            left_indent=body_indent,
        )
        add_line("Provider Contact Statement", bold=True, space_after=Pt(2))
        add_line(
            str(payload.get("consult_contact_statement") or "").strip()
            or "____________________________________________________________",
            space_after=Pt(10),
            left_indent=body_indent,
        )
        add_line("Sincerely,", space_after=Pt(8))
        add_line(provider_header, space_after=Pt(2))
        add_line(
            f"{fill_text(payload.get('provider_specialty'), '[Specialty - Pain Management / Interventional Spine]')} | NPI Number: {fill_text(payload.get('provider_npi'), '[NPI Number]')}",
            space_after=Pt(2),
        )
        add_line(practice_name, space_after=Pt(2))
        add_line(provider_address, space_after=Pt(2))
        add_line(f"Phone: {provider_phone} | Fax: {provider_fax} | Email: {provider_email}", space_after=Pt(0))
        document.save(path)

    def _write_clinical_documentation_doc_v2(self, path, payload):
        document = Document()
        section = document.sections[0]
        section.top_margin = Pt(46)
        section.bottom_margin = Pt(46)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)

        normal_style = document.styles["Normal"]
        normal_style.font.name = "Calibri"
        normal_style.font.size = Pt(10.5)
        body_indent = Pt(18)
        checkbox_indent = Pt(30)

        title = document.add_paragraph()
        title.alignment = 1
        title_run = title.add_run(payload.get("packet_title") or payload.get("clinical_doc_title") or "Clinical Documentation Template")
        title_run.bold = True
        title_run.font.size = Pt(16)
        title.paragraph_format.space_after = Pt(10)

        def fill_text(value, blank="_____________"):
            text = str(value or "").strip()
            return text or blank

        def add_line(text, bold=False, space_after=Pt(4), left_indent=None):
            paragraph = document.add_paragraph()
            if left_indent is not None:
                paragraph.paragraph_format.left_indent = left_indent
            paragraph.paragraph_format.space_after = space_after
            run = paragraph.add_run(text)
            run.bold = bold
            return paragraph

        def marker(enabled):
            return "[X]" if enabled else "[ ]"

        def add_checkbox_lines(items):
            for enabled, label in items:
                add_line(f"{marker(enabled)} {label}", space_after=Pt(1), left_indent=checkbox_indent)

        add_line("I. Chief Complaint", bold=True)
        add_line(fill_text(payload.get("clinical_doc_chief_complaint"), "____________________________"), left_indent=body_indent)

        add_line("II. History of Present Illness", bold=True)
        add_line("Duration of Symptoms:", space_after=Pt(2), left_indent=body_indent)
        add_checkbox_lines(
            [
                (payload.get("clinical_doc_duration_gt_3m"), "> 3 months"),
                (payload.get("clinical_doc_duration_gt_6m"), "> 6 months"),
                (payload.get("clinical_doc_duration_gt_12m"), "> 12 months"),
            ]
        )
        add_line(f"Exact duration: {fill_text(payload.get('clinical_doc_exact_duration'), '______________________')}", left_indent=body_indent)
        add_line("Pain Characteristics:", space_after=Pt(2), left_indent=body_indent)
        add_checkbox_lines(
            [
                (payload.get("clinical_doc_pain_axial"), "Axial lumbar pain"),
                (payload.get("clinical_doc_pain_discogenic"), "Discogenic pattern"),
                (payload.get("clinical_doc_pain_activity_exacerbation"), "Activity-related exacerbation"),
                (payload.get("clinical_doc_pain_sitting_intolerance"), "Sitting intolerance"),
                (payload.get("clinical_doc_pain_standing_intolerance"), "Standing intolerance"),
                (payload.get("clinical_doc_pain_bending_lifting"), "Bending/lifting provocation"),
            ]
        )
        add_line(f"Pain severity (0-10): {fill_text(payload.get('clinical_doc_pain_severity'), '______')}", left_indent=body_indent)

        add_line("III. Functional Impairment", bold=True)
        add_line("Patient reports limitation in:", left_indent=body_indent)
        add_checkbox_lines(
            [
                (payload.get("clinical_doc_limit_occupational"), "Occupational duties"),
                (payload.get("clinical_doc_limit_prolonged_sitting"), "Prolonged sitting"),
                (payload.get("clinical_doc_limit_prolonged_standing"), "Prolonged standing"),
                (payload.get("clinical_doc_limit_ambulation"), "Ambulation tolerance"),
                (payload.get("clinical_doc_limit_household"), "Household activities"),
                (payload.get("clinical_doc_limit_sleep"), "Sleep disturbance"),
            ]
        )
        add_line("Describe specific functional impact:", left_indent=body_indent)
        add_line(fill_text(payload.get("clinical_doc_functional_impact"), "____________________________"), left_indent=body_indent)

        add_line("IV. Conservative Therapy History", bold=True)
        add_line("The patient has completed and/or attempted:", left_indent=body_indent)
        add_checkbox_lines(
            [
                (payload.get("clinical_doc_conservative_pt"), "Physical therapy"),
                (payload.get("clinical_doc_conservative_home_exercise"), "Home exercise program"),
                (payload.get("clinical_doc_conservative_nsaids"), "NSAIDs"),
                (payload.get("clinical_doc_conservative_non_opioid"), "Non-opioid analgesics"),
                (payload.get("clinical_doc_conservative_activity_modification"), "Activity modification"),
                (payload.get("clinical_doc_conservative_esi"), "Epidural steroid injections"),
                (payload.get("clinical_doc_conservative_other_interventional"), "Other interventional procedures"),
            ]
        )
        add_line(f"Duration of conservative treatment: {fill_text(payload.get('clinical_doc_conservative_duration'), '____________________')}", left_indent=body_indent)
        add_line("Response to prior ESI (if applicable):", space_after=Pt(2), left_indent=body_indent)
        esi_response = str(payload.get("clinical_doc_esi_response") or "").strip().lower()
        add_checkbox_lines(
            [
                (esi_response == "temporary relief", "Temporary relief"),
                (esi_response == "partial relief", "Partial relief"),
                (esi_response == "no relief", "No relief"),
            ]
        )

        add_line("V. Imaging Findings", bold=True)
        add_line(f"MRI Date: {fill_text(payload.get('clinical_doc_mri_date'), '____________________')}", left_indent=body_indent)
        add_line("Findings:", left_indent=body_indent)
        add_checkbox_lines(
            [
                (payload.get("clinical_doc_imaging_annular_tear"), "Annular tear"),
                (payload.get("clinical_doc_imaging_disc_degeneration"), "Disc degeneration"),
                (payload.get("clinical_doc_imaging_disc_protrusion"), "Disc protrusion"),
                (payload.get("clinical_doc_imaging_disc_displacement"), "Disc displacement"),
            ]
        )
        add_line(f"Affected Levels: {fill_text(payload.get('clinical_doc_affected_levels'), '____________________')}", left_indent=body_indent)
        add_line("Imaging correlates with clinical presentation.", left_indent=body_indent)

        add_line("VI. Assessment", bold=True)
        add_line(f"Primary Diagnosis / ICD-10: {fill_text(payload.get('clinical_doc_primary_diagnosis'), '__________________________')}", left_indent=body_indent)
        add_line(f"Secondary Diagnosis / ICD-10: {fill_text(payload.get('clinical_doc_secondary_diagnosis'), '__________________________')}", left_indent=body_indent)
        add_line(
            fill_text(
                payload.get("clinical_doc_assessment_summary"),
                "____________________________________________________________",
            ),
            left_indent=body_indent,
        )

        add_line("VII. Treatment Plan", bold=True)
        add_line(
            fill_text(
                payload.get("clinical_doc_treatment_plan_intro"),
                "____________________________________________________________",
            ),
            left_indent=body_indent,
        )
        selected_plan_items = []
        if payload.get("clinical_doc_plan_diagnostic_confirmation"):
            selected_plan_items.append("Diagnostic confirmation if necessary")
        if payload.get("clinical_doc_plan_intradiscal_intervention"):
            selected_plan_items.append("Intradiscal annular intervention if indicated")
        if payload.get("clinical_doc_plan_follow_up"):
            selected_plan_items.append("Standard post-procedure follow-up")
        if not selected_plan_items:
            selected_plan_items = ["________________________________________"]
        add_line("Plan includes:", left_indent=body_indent)
        for item in selected_plan_items:
            add_line("- " + item, space_after=Pt(1), left_indent=checkbox_indent)
        add_line(
            fill_text(
                payload.get("clinical_doc_plan_exclusion"),
                "____________________________________________________________",
            ),
            left_indent=body_indent,
        )

        add_line("Physician Narrative Paragraph", bold=True, space_after=Pt(4))
        add_line(
            fill_text(
                payload.get("clinical_doc_physician_narrative"),
                "____________________________________________________________",
            ),
            left_indent=body_indent,
        )
        document.save(path)


class DevToolsDialog(QDialog):
    def __init__(self, machine_role, primary_update_channel, reference_update_channel, developer_tools_enabled, private_dev_channel, parent=None):
        super().__init__(parent)
        self.machine_role = machine_role
        self.primary_update_channel = primary_update_channel
        self.reference_update_channel = reference_update_channel
        self.developer_tools_enabled = developer_tools_enabled
        self.private_dev_channel = dict(private_dev_channel or {})
        self.config = load_dev_tools_config()

        self.setWindowTitle("Developer Tools")
        self.setMinimumSize(980, 720)
        self.setWindowFlags(
            Qt.Window
            | Qt.WindowCloseButtonHint
            | Qt.WindowMinimizeButtonHint
            | Qt.WindowMaximizeButtonHint
            | Qt.WindowSystemMenuHint
        )
        self.setStyleSheet(
            """
            QDialog { background-color: #0D1520; color: #E5E7EB; }
            QLabel { color: #E5E7EB; }
            QScrollArea, QScrollArea > QWidget > QWidget {
                background-color: #0F1823;
                border: none;
            }
            QTextEdit, QListWidget, QLineEdit, QComboBox {
                background-color: #111A25;
                border: 1px solid #253243;
                border-radius: 8px;
                color: #DCE6F2;
                padding: 8px;
            }
            QComboBox QAbstractItemView {
                background-color: #111A25;
                color: #DCE6F2;
                border: 1px solid #34506B;
                selection-background-color: #223246;
                selection-color: #FFFFFF;
                outline: 0;
            }
            QTextEdit:focus, QListWidget:focus, QLineEdit:focus, QComboBox:focus {
                border: 1px solid #69BCFF;
            }
            QPushButton {
                background-color: #1D2A3A;
                color: #FFFFFF;
                border: 1px solid #34506B;
                border-radius: 6px;
                padding: 8px 14px;
            }
            QPushButton:hover { background-color: #223246; }
            QTabWidget::pane {
                border: 1px solid #243446;
                background: #0F1823;
                border-radius: 10px;
            }
            QTabBar::tab {
                background: #111A25;
                color: #C9D7E5;
                padding: 10px 16px;
                margin-right: 4px;
                border-top-left-radius: 8px;
                border-top-right-radius: 8px;
            }
            QTabBar::tab:selected {
                background: #1D2A3A;
                color: #FFFFFF;
                font-weight: 700;
            }
            QGroupBox {
                background-color: #0F1823;
                border: 1px solid #243446;
                border-radius: 10px;
                margin-top: 10px;
                padding-top: 14px;
            }
            QGroupBox:title {
                subcontrol-origin: margin;
                left: 12px;
                padding: 0 4px;
            }
            QSplitter::handle {
                background: #243446;
            }
            QSplitter::handle:horizontal {
                width: 6px;
            }
            QSplitter::handle:vertical {
                height: 6px;
            }
            """
        )

        layout = QVBoxLayout(self)
        layout.setContentsMargins(18, 18, 18, 18)
        layout.setSpacing(12)

        title = QLabel("Developer Tools")
        title.setStyleSheet("font-size:22px; font-weight:700; color:#69BCFF;")
        layout.addWidget(title)

        tabs = QTabWidget()
        tabs.addTab(self._build_overview_tab(), "Overview")
        tabs.addTab(LegacyGalleryTab(self.config, self.save_config, self), "Legacy UI Gallery")
        tabs.addTab(PacketBuilderTab(self.config, self.save_config, self), "Packet Builder Studio")
        layout.addWidget(tabs, stretch=1)

        close_button = QPushButton("Close")
        close_button.clicked.connect(self.accept)
        layout.addWidget(close_button, alignment=Qt.AlignRight)
        self._apply_initial_geometry(parent)

    def keyPressEvent(self, event):
        if event.key() == Qt.Key_Escape:
            self.showMinimized()
            event.accept()
            return
        super().keyPressEvent(event)

    def save_config(self, changes):
        self.config = dict(update_dev_tools_config(changes or {}) or {})
        return self.config

    def _apply_initial_geometry(self, parent):
        screen = None
        if parent and parent.windowHandle():
            screen = parent.windowHandle().screen()
        if not screen:
            screen = QGuiApplication.primaryScreen()
        if not screen:
            self.resize(1360, 900)
            return
        available = screen.availableGeometry()
        width = min(max(1180, int(available.width() * 0.88)), 1700)
        height = min(max(820, int(available.height() * 0.88)), 1100)
        x = available.x() + max(0, (available.width() - width) // 2)
        y = available.y() + max(0, (available.height() - height) // 2)
        self.setGeometry(x, y, width, height)

    def _build_overview_tab(self):
        tab = QWidget()
        layout = QVBoxLayout(tab)
        layout.setContentsMargins(12, 12, 12, 12)
        layout.setSpacing(12)

        intro = QLabel(
            "This install is marked as a developer machine. Offices will never see this area. Use it to keep visual references nearby and begin packet creation work without exposing half-built tools to production users."
        )
        intro.setWordWrap(True)
        intro.setStyleSheet("color:#BFD0E3;")
        layout.addWidget(intro)

        summary = QTextEdit()
        summary.setReadOnly(True)
        summary.setStyleSheet(
            "background:#0F1823; border:1px solid #243446; border-radius:12px; color:#DCE6F2; padding:12px;"
        )
        summary.setHtml(
            f"""
            <div style="font-family:'Segoe UI'; font-size:14px; color:#DCE6F2; line-height:1.55;">
              <div style="font-size:18px; font-weight:700; color:#69BCFF; margin-bottom:10px;">Developer Install Status</div>
              <table style="width:100%; border-collapse:collapse; margin-bottom:18px;">
                <tr><td style="padding:6px 0; color:#8FA6C1; width:38%;">Machine role</td><td style="padding:6px 0; color:#F2F5F9; font-weight:600;">{html.escape(str(self.machine_role or 'unknown').title())}</td></tr>
                <tr><td style="padding:6px 0; color:#8FA6C1;">Primary update channel</td><td style="padding:6px 0; color:#F2F5F9; font-weight:600;">{html.escape(str(self.primary_update_channel or 'unknown').title())}</td></tr>
                <tr><td style="padding:6px 0; color:#8FA6C1;">Production reference visible</td><td style="padding:6px 0; color:#F2F5F9; font-weight:600;">{'Yes' if self.reference_update_channel else 'No'}</td></tr>
                <tr><td style="padding:6px 0; color:#8FA6C1;">Developer tools enabled</td><td style="padding:6px 0; color:#F2F5F9; font-weight:600;">{'Yes' if self.developer_tools_enabled else 'No'}</td></tr>
                <tr><td style="padding:6px 0; color:#8FA6C1;">Private dev repo</td><td style="padding:6px 0; color:#F2F5F9; font-weight:600;">{html.escape((f"{self.private_dev_channel.get('owner', '')}/{self.private_dev_channel.get('repo', '')}".strip("/") or "Not configured"))}</td></tr>
                <tr><td style="padding:6px 0; color:#8FA6C1;">Private dev updates enabled</td><td style="padding:6px 0; color:#F2F5F9; font-weight:600;">{'Yes' if self.private_dev_channel.get('enabled') else 'No'}</td></tr>
              </table>

              <div style="font-size:17px; font-weight:700; color:#69BCFF; margin:0 0 8px 0;">What This Version Gives You</div>
              <ul style="margin:0 0 18px 18px; padding:0;">
                <li style="margin-bottom:6px;">Legacy launcher and GUI previews for visual reference only</li>
                <li style="margin-bottom:6px;">Packet Builder Studio with live preview</li>
                <li style="margin-bottom:6px;">Single-form export to Word, PDF, or both</li>
                <li style="margin-bottom:6px;">Referral-request and patient-packet export flows</li>
                <li style="margin-bottom:6px;">Reusable export-folder selection</li>
              </ul>

              <div style="font-size:17px; font-weight:700; color:#69BCFF; margin:0 0 8px 0;">Next Dev-Only Ideas</div>
              <ul style="margin:0 0 0 18px; padding:0;">
                <li style="margin-bottom:6px;">Packet Lab</li>
                <li style="margin-bottom:6px;">Extraction Trace</li>
                <li style="margin-bottom:6px;">Threshold tuning</li>
                <li style="margin-bottom:6px;">Regression helpers</li>
              </ul>
            </div>
            """
        )
        layout.addWidget(summary, stretch=1)
        return tab
